import streamlit as st
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from matplotlib.patches import FancyArrowPatch, Circle, Rectangle, FancyBboxPatch, Polygon
from matplotlib.lines import Line2D
import numpy as np
import plotly.graph_objects as go
import json
from io import BytesIO

# ============================================================
# 3D SHAPE LIBRARY — Realistic geometric primitives
# (Used by both single-department and full-hospital 3D views)
# ============================================================
def box_mesh_global(x0, y0, z0, dx, dy, dz, color, opacity=1.0, name=''):
    """Closed rectangular box."""
    x = [x0, x0+dx, x0+dx, x0, x0, x0+dx, x0+dx, x0]
    y = [y0, y0, y0+dy, y0+dy, y0, y0, y0+dy, y0+dy]
    z = [z0, z0, z0, z0, z0+dz, z0+dz, z0+dz, z0+dz]
    i = [7,0,0,0,4,4,6,6,4,0,3,2]
    j = [3,4,1,2,5,6,5,2,0,1,6,3]
    k = [0,7,2,3,6,7,1,1,5,5,7,6]
    return go.Mesh3d(x=x, y=y, z=z, i=i, j=j, k=k, color=color,
        opacity=opacity, name=name, hovertext=name,
        hoverinfo='text', flatshading=True)

def cylinder_mesh(cx, cy, z0, radius, height, axis='z', n=20,
                  color='gray', opacity=1.0, name=''):
    """Closed cylinder with end caps. axis: 'x', 'y', or 'z'."""
    theta = np.linspace(0, 2*np.pi, n, endpoint=False)
    if axis == 'z':
        a_circ = radius*np.cos(theta); b_circ = radius*np.sin(theta)
        x = np.concatenate([cx + a_circ, cx + a_circ, [cx, cx]])
        y = np.concatenate([cy + b_circ, cy + b_circ, [cy, cy]])
        z = np.concatenate([np.full(n, z0), np.full(n, z0+height),
                            [z0, z0+height]])
    elif axis == 'y':
        a_circ = radius*np.cos(theta); b_circ = radius*np.sin(theta)
        x = np.concatenate([cx + a_circ, cx + a_circ, [cx, cx]])
        z = np.concatenate([z0 + b_circ, z0 + b_circ, [z0, z0]])
        y = np.concatenate([np.full(n, cy), np.full(n, cy+height),
                            [cy, cy+height]])
    else:  # axis == 'x'
        a_circ = radius*np.cos(theta); b_circ = radius*np.sin(theta)
        y = np.concatenate([cy + a_circ, cy + a_circ, [cy, cy]])
        z = np.concatenate([z0 + b_circ, z0 + b_circ, [z0, z0]])
        x = np.concatenate([np.full(n, cx), np.full(n, cx+height),
                            [cx, cx+height]])
    I, J, K = [], [], []
    for i in range(n):
        ni = (i+1) % n
        I += [i, ni];  J += [ni, ni+n];  K += [i+n, i+n]
    for i in range(n):
        ni = (i+1) % n
        I.append(2*n); J.append(ni); K.append(i)
    for i in range(n):
        ni = (i+1) % n
        I.append(2*n+1); J.append(i+n); K.append(ni+n)
    return go.Mesh3d(x=x, y=y, z=z, i=I, j=J, k=K, color=color,
        opacity=opacity, name=name, hovertext=name,
        hoverinfo='text', flatshading=True)

def torus_mesh(cx, cy, cz, R=1.0, r=0.3, axis='y', n_major=24, n_minor=12,
               color='gray', opacity=1.0, name=''):
    """Torus (donut) shape. axis = which direction the donut hole faces."""
    u = np.linspace(0, 2*np.pi, n_major, endpoint=False)
    v = np.linspace(0, 2*np.pi, n_minor, endpoint=False)
    U, V = np.meshgrid(u, v)
    if axis == 'y':
        X = ((R + r*np.cos(V)) * np.cos(U)).flatten() + cx
        Z = ((R + r*np.cos(V)) * np.sin(U)).flatten() + cz
        Y = (r*np.sin(V)).flatten() + cy
    elif axis == 'z':
        X = ((R + r*np.cos(V)) * np.cos(U)).flatten() + cx
        Y = ((R + r*np.cos(V)) * np.sin(U)).flatten() + cy
        Z = (r*np.sin(V)).flatten() + cz
    else:  # axis == 'x'
        Y = ((R + r*np.cos(V)) * np.cos(U)).flatten() + cy
        Z = ((R + r*np.cos(V)) * np.sin(U)).flatten() + cz
        X = (r*np.sin(V)).flatten() + cx
    I, J, K = [], [], []
    for i in range(n_minor):
        ni = (i+1) % n_minor
        for j in range(n_major):
            nj = (j+1) % n_major
            a = i*n_major + j;  b = i*n_major + nj
            c = ni*n_major + j; d = ni*n_major + nj
            I += [a, b]; J += [b, d]; K += [c, c]
    return go.Mesh3d(x=X, y=Y, z=Z, i=I, j=J, k=K, color=color,
        opacity=opacity, name=name, hovertext=name,
        hoverinfo='text', flatshading=True)

def cone_mesh(cx, cy, z0, radius, height, n=12, color='green',
              opacity=1.0, name=''):
    """Closed cone (apex pointing up)."""
    theta = np.linspace(0, 2*np.pi, n, endpoint=False)
    bx = cx + radius*np.cos(theta)
    by = cy + radius*np.sin(theta)
    x = np.concatenate([bx, [cx, cx]])  # base ring + apex + base-center
    y = np.concatenate([by, [cy, cy]])
    z = np.concatenate([np.full(n, z0), [z0+height, z0]])
    apex = n; base_center = n+1
    I, J, K = [], [], []
    for i in range(n):
        ni = (i+1) % n
        I.append(i); J.append(ni); K.append(apex)
        I.append(base_center); J.append(ni); K.append(i)
    return go.Mesh3d(x=x, y=y, z=z, i=I, j=J, k=K, color=color,
        opacity=opacity, name=name, hovertext=name,
        hoverinfo='text', flatshading=True)

def tree_3d(cx, cy, base_z=0, total_height=2.5,
            trunk_color='#5d4037', leaf_color='#2e7d32'):
    """Returns a list of meshes representing a tree (trunk + cone leaves)."""
    trunk_h = total_height * 0.35
    trunk = cylinder_mesh(cx, cy, base_z, 0.12, trunk_h, 'z', 8,
                          trunk_color, 1.0, 'Tree Trunk')
    leaves = cone_mesh(cx, cy, base_z + trunk_h, 0.6,
                       total_height - trunk_h, 14, leaf_color, 1.0, 'Tree')
    return [trunk, leaves]

def chair_3d(cx, cy, z0, seat_w=0.5, seat_d=0.5, seat_h=0.45,
             back_h=0.5, color='#90a4ae', name='Chair'):
    """Realistic chair: seat + backrest + 4 legs."""
    traces = []
    # Seat
    traces.append(box_mesh_global(cx-seat_w/2, cy-seat_d/2, z0+seat_h-0.05,
                                   seat_w, seat_d, 0.05, color, 1.0, name))
    # Backrest
    traces.append(box_mesh_global(cx-seat_w/2, cy+seat_d/2-0.05, z0+seat_h,
                                   seat_w, 0.05, back_h, color, 1.0, name+' back'))
    # 4 legs
    leg_t = 0.04
    for lx, ly in [(-seat_w/2, -seat_d/2), (seat_w/2-leg_t, -seat_d/2),
                    (-seat_w/2, seat_d/2-leg_t), (seat_w/2-leg_t, seat_d/2-leg_t)]:
        traces.append(box_mesh_global(cx+lx, cy+ly, z0,
                                       leg_t, leg_t, seat_h-0.05,
                                       '#37474f', 1.0, name+' leg'))
    return traces

def recliner_3d(cx, cy, z0, color='#1976d2', name='Recliner'):
    """Reclining medical chair (dialysis/chemo)."""
    traces = []
    # Base/legs platform
    traces.append(box_mesh_global(cx-0.4, cy-0.8, z0,
                                   0.8, 1.6, 0.35, '#455a64', 1.0, name+' base'))
    # Seat cushion
    traces.append(box_mesh_global(cx-0.35, cy-0.3, z0+0.35,
                                   0.7, 0.6, 0.15, color, 1.0, name+' seat'))
    # Leg rest (extended forward)
    traces.append(box_mesh_global(cx-0.35, cy-0.85, z0+0.35,
                                   0.7, 0.55, 0.1, color, 1.0, name+' legrest'))
    # Backrest (reclined slightly back)
    traces.append(box_mesh_global(cx-0.35, cy+0.3, z0+0.5,
                                   0.7, 0.15, 0.7, color, 1.0, name+' backrest'))
    # Headrest
    traces.append(box_mesh_global(cx-0.25, cy+0.4, z0+1.15,
                                   0.5, 0.1, 0.2, color, 1.0, name+' headrest'))
    # Arm rests
    for ax in [-0.4, 0.3]:
        traces.append(box_mesh_global(cx+ax, cy-0.3, z0+0.5,
                                       0.1, 0.6, 0.1, '#546e7a', 1.0,
                                       name+' arm'))
    return traces

def incubator_3d(cx, cy, z0, color='#ec407a', name='Incubator'):
    """NICU baby incubator with transparent dome."""
    traces = []
    # Stand with wheels
    traces.append(box_mesh_global(cx-0.45, cy-0.55, z0,
                                   0.9, 1.1, 0.7, '#37474f', 1.0, name+' stand'))
    # Bed inside
    traces.append(box_mesh_global(cx-0.35, cy-0.45, z0+0.7,
                                   0.7, 0.9, 0.08, '#fff8e1', 1.0, name+' bed'))
    # Transparent acrylic dome (top)
    traces.append(box_mesh_global(cx-0.4, cy-0.5, z0+0.78,
                                   0.8, 1.0, 0.45, '#e1f5fe', 0.35,
                                   name+' dome'))
    # Pink accent strip
    traces.append(box_mesh_global(cx-0.45, cy-0.55, z0+0.65,
                                   0.9, 1.1, 0.05, color, 1.0,
                                   name+' accent'))
    return traces

def cabinet_3d(cx, cy, z0, w=1.4, d=0.4, h=1.8, color='#a5d6a7',
               n_shelves=4, name='Cabinet'):
    """Tall shelving unit (pharmacy/lab storage)."""
    traces = []
    # Frame
    traces.append(box_mesh_global(cx-w/2, cy-d/2, z0, w, d, h, color, 1.0, name))
    # Visible shelf dividers (slightly darker stripes)
    shelf_h = h / n_shelves
    for s in range(1, n_shelves):
        traces.append(box_mesh_global(cx-w/2+0.02, cy-d/2-0.02, z0+s*shelf_h,
                                       w-0.04, 0.02, 0.04, '#2e7d32', 1.0,
                                       name+' shelf'))
    return traces

def hospital_bed_3d(cx, cy, z0, color='#1976d2', size=1.0, name='Bed'):
    """Realistic hospital bed: frame + mattress + headboard + pillow + wheels.
    `size` scales the bed (use 0.7 for pediatric, 1.2 for maternity)."""
    bw, bl = 0.9 * size, 2.0 * size
    traces = []
    # Frame
    traces.append(box_mesh_global(cx-bw/2, cy-bl/2, z0+0.35,
                                   bw, bl, 0.15, color, 1.0, name+' frame'))
    # Mattress
    traces.append(box_mesh_global(cx-bw/2+0.05, cy-bl/2+0.05, z0+0.5,
                                   bw-0.1, bl-0.1, 0.15, '#eceff1', 1.0,
                                   name+' mattress'))
    # Headboard
    traces.append(box_mesh_global(cx-bw/2, cy-bl/2, z0+0.35,
                                   bw, 0.08, 0.6, color, 1.0, name+' headboard'))
    # Pillow
    traces.append(box_mesh_global(cx-bw/2+0.15, cy-bl/2+0.15, z0+0.65,
                                   bw-0.3, 0.4, 0.08, 'white', 1.0,
                                   name+' pillow'))
    # 4 wheels
    for wx, wy in [(-bw/2+0.05, -bl/2+0.1), (bw/2-0.1, -bl/2+0.1),
                    (-bw/2+0.05, bl/2-0.15), (bw/2-0.1, bl/2-0.15)]:
        traces.append(cylinder_mesh(cx+wx+0.04, cy+wy+0.04, z0+0.075,
                                     0.08, 0.15, 'y', 8, '#212121', 1.0,
                                     name+' wheel'))
    return traces

def mri_scanner_3d(cx, cy, z0, color='#4a148c', name='MRI Scanner'):
    """MRI: large vertical torus + sliding patient table."""
    traces = []
    # The donut (torus standing vertical; patient slides through the hole)
    traces.append(torus_mesh(cx, cy, z0+1.0, R=0.9, r=0.45, axis='y',
                              n_major=28, n_minor=14, color=color,
                              opacity=0.95, name=name+' bore'))
    # Inner tube (the bore that patient lies in)
    traces.append(cylinder_mesh(cx, cy-0.4, z0+1.0, 0.45, 0.8, 'y', 16,
                                 '#1a1a2e', 0.6, name+' bore-tube'))
    # Sliding patient table
    traces.append(box_mesh_global(cx-0.3, cy-1.8, z0+0.7,
                                   0.6, 1.4, 0.1, '#eceff1', 1.0,
                                   name+' table'))
    # Table base
    traces.append(box_mesh_global(cx-0.25, cy-1.7, z0,
                                   0.5, 1.2, 0.7, '#37474f', 1.0,
                                   name+' table-base'))
    return traces

def ct_scanner_3d(cx, cy, z0, color='#6a1b9a', name='CT Scanner'):
    """CT: flatter ring + table."""
    traces = []
    # The CT ring (flatter than MRI)
    traces.append(torus_mesh(cx, cy, z0+0.9, R=0.7, r=0.25, axis='y',
                              n_major=24, n_minor=12, color=color,
                              opacity=1.0, name=name+' ring'))
    # Cylindrical housing
    traces.append(cylinder_mesh(cx, cy, z0+0.9, 0.95, 0.5, 'y', 16,
                                 '#fafafa', 0.4, name+' housing'))
    # Inner bore (dark)
    traces.append(cylinder_mesh(cx, cy, z0+0.9, 0.4, 0.5, 'y', 16,
                                 '#212121', 0.8, name+' bore'))
    # Patient table
    traces.append(box_mesh_global(cx-0.3, cy-1.6, z0+0.7,
                                   0.6, 1.2, 0.1, '#eceff1', 1.0,
                                   name+' table'))
    traces.append(box_mesh_global(cx-0.25, cy-1.5, z0,
                                   0.5, 1.0, 0.7, '#37474f', 1.0,
                                   name+' base'))
    return traces

def autoclave_3d(cx, cy, z0, color='#e64a19', name='Autoclave'):
    """Horizontal cylindrical autoclave on stand."""
    traces = []
    # Stand
    traces.append(box_mesh_global(cx-0.45, cy-0.55, z0,
                                   0.9, 1.1, 0.7, '#546e7a', 1.0, name+' stand'))
    # Horizontal cylinder body
    traces.append(cylinder_mesh(cx, cy-0.5, z0+1.1, 0.4, 1.0, 'y', 20,
                                 color, 1.0, name+' chamber'))
    # Round door (slightly thicker disc on front)
    traces.append(cylinder_mesh(cx, cy-0.55, z0+1.1, 0.42, 0.08, 'y', 20,
                                 '#bf360c', 1.0, name+' door'))
    # Wheel handle
    traces.append(torus_mesh(cx, cy-0.62, z0+1.1, R=0.15, r=0.03, axis='y',
                              n_major=16, n_minor=6,
                              color='#212121', opacity=1.0,
                              name=name+' handle'))
    return traces

def washer_3d(cx, cy, z0, color='#0288d1', name='Washer'):
    """Front-loading washing machine with porthole door."""
    traces = []
    # Body
    traces.append(box_mesh_global(cx-0.4, cy-0.4, z0,
                                   0.8, 0.8, 1.0, '#eceff1', 1.0, name))
    # Color accent top
    traces.append(box_mesh_global(cx-0.4, cy-0.4, z0+0.85,
                                   0.8, 0.8, 0.05, color, 1.0, name+' top'))
    # Round porthole door
    traces.append(cylinder_mesh(cx, cy-0.41, z0+0.5, 0.25, 0.05, 'y', 16,
                                 '#212121', 1.0, name+' door-frame'))
    traces.append(cylinder_mesh(cx, cy-0.43, z0+0.5, 0.22, 0.02, 'y', 16,
                                 '#4fc3f7', 0.6, name+' door-glass'))
    # Control panel strip
    traces.append(box_mesh_global(cx-0.35, cy-0.41, z0+0.82,
                                   0.7, 0.02, 0.1, '#1a237e', 1.0,
                                   name+' controls'))
    return traces

def blood_fridge_3d(cx, cy, z0, color='#c62828', name='Blood Fridge'):
    """Tall vertical blood storage refrigerator."""
    traces = []
    # Main body
    traces.append(box_mesh_global(cx-0.4, cy-0.4, z0,
                                   0.8, 0.8, 1.9, '#eceff1', 1.0, name))
    # Door (slightly recessed)
    traces.append(box_mesh_global(cx-0.35, cy-0.42, z0+0.1,
                                   0.7, 0.02, 1.7, '#fafafa', 1.0, name+' door'))
    # Red accent top
    traces.append(box_mesh_global(cx-0.4, cy-0.4, z0+1.85,
                                   0.8, 0.8, 0.08, color, 1.0, name+' top'))
    # Handle
    traces.append(box_mesh_global(cx+0.2, cy-0.45, z0+0.9,
                                   0.1, 0.03, 0.3, '#37474f', 1.0,
                                   name+' handle'))
    # Window with red glow
    traces.append(box_mesh_global(cx-0.2, cy-0.43, z0+1.0,
                                   0.4, 0.01, 0.5, color, 0.5, name+' window'))
    return traces

def lab_bench_3d(cx, cy, z0, color='#00695c', name='Lab Bench'):
    """Lab workbench with microscope and equipment."""
    traces = []
    # Bench top
    traces.append(box_mesh_global(cx-0.8, cy-0.3, z0+0.85,
                                   1.6, 0.6, 0.05, '#b2dfdb', 1.0, name+' top'))
    # Legs/cabinets below
    traces.append(box_mesh_global(cx-0.8, cy-0.3, z0,
                                   1.6, 0.6, 0.85, '#80cbc4', 1.0,
                                   name+' cabinet'))
    # Overhead shelf with light
    traces.append(box_mesh_global(cx-0.8, cy-0.3, z0+1.6,
                                   1.6, 0.3, 0.05, color, 1.0, name+' shelf'))
    # Microscope (cylinder + small box)
    traces.append(cylinder_mesh(cx-0.5, cy, z0+0.9, 0.06, 0.3, 'z', 12,
                                 '#37474f', 1.0, name+' microscope'))
    traces.append(box_mesh_global(cx-0.55, cy-0.05, z0+0.9,
                                   0.1, 0.1, 0.08, '#212121', 1.0, 'eyepiece'))
    # Test tube rack
    traces.append(box_mesh_global(cx, cy-0.1, z0+0.9,
                                   0.3, 0.15, 0.05, '#5d4037', 1.0, 'rack'))
    for tx in [0.05, 0.15, 0.25]:
        traces.append(cylinder_mesh(cx-0.15+tx, cy-0.025, z0+0.92,
                                     0.025, 0.18, 'z', 8,
                                     '#80cbc4', 0.8, 'tube'))
    return traces

def pharmacy_counter_3d(cx, cy, z0, color='#388e3c', name='Pharmacy'):
    """Pharmacy dispensing counter + tall shelving behind."""
    traces = []
    # Counter
    traces.append(box_mesh_global(cx-0.8, cy-0.5, z0,
                                   1.6, 0.5, 0.95, '#a5d6a7', 1.0,
                                   name+' counter'))
    # Counter top
    traces.append(box_mesh_global(cx-0.85, cy-0.55, z0+0.95,
                                   1.7, 0.55, 0.05, color, 1.0,
                                   name+' top'))
    # Shelving behind
    for s in range(4):
        traces.append(box_mesh_global(cx-0.8, cy+0.3, z0+0.4 + s*0.4,
                                       1.6, 0.2, 0.05, color, 1.0,
                                       name+' shelf'))
    # Shelf back panel
    traces.append(box_mesh_global(cx-0.8, cy+0.5, z0,
                                   1.6, 0.05, 2.0, '#c8e6c9', 1.0,
                                   name+' back-panel'))
    return traces

def dining_set_3d(cx, cy, z0, color='#fbc02d', name='Dining'):
    """Round table with 4 chairs around it."""
    traces = []
    # Round table top
    traces.append(cylinder_mesh(cx, cy, z0+0.7, 0.55, 0.04, 'z', 20,
                                 '#fff8e1', 1.0, name+' table-top'))
    # Table pedestal
    traces.append(cylinder_mesh(cx, cy, z0, 0.1, 0.7, 'z', 12,
                                 '#5d4037', 1.0, name+' pedestal'))
    # Table base
    traces.append(cylinder_mesh(cx, cy, z0, 0.4, 0.05, 'z', 16,
                                 '#5d4037', 1.0, name+' base'))
    # 4 chairs around
    chair_dist = 0.95
    for ang in [0, np.pi/2, np.pi, 3*np.pi/2]:
        chx = cx + chair_dist*np.cos(ang)
        chy = cy + chair_dist*np.sin(ang)
        traces.extend(chair_3d(chx, chy, z0, 0.4, 0.4, 0.4, 0.4, color,
                                name+' chair'))
    return traces

def desk_3d(cx, cy, z0, color='#757575', name='Desk'):
    """Office desk with chair and monitor."""
    traces = []
    # Desk surface
    traces.append(box_mesh_global(cx-0.7, cy-0.4, z0+0.7,
                                   1.4, 0.7, 0.05, '#bcaaa4', 1.0,
                                   name+' top'))
    # Side cabinet
    traces.append(box_mesh_global(cx+0.4, cy-0.4, z0,
                                   0.3, 0.7, 0.7, '#a1887b', 1.0,
                                   name+' cabinet'))
    # 2 legs
    for lx in [-0.65, -0.05]:
        traces.append(box_mesh_global(cx+lx, cy-0.35, z0,
                                       0.05, 0.6, 0.7, '#8d6e63', 1.0,
                                       name+' leg'))
    # Monitor (screen on back of desk)
    traces.append(box_mesh_global(cx-0.25, cy+0.2, z0+0.85,
                                   0.5, 0.05, 0.35, '#212121', 1.0,
                                   name+' monitor'))
    traces.append(box_mesh_global(cx-0.2, cy+0.18, z0+0.9,
                                   0.4, 0.01, 0.25, '#4fc3f7', 1.0,
                                   name+' screen'))
    # Office chair in front
    traces.extend(chair_3d(cx, cy-0.7, z0, 0.5, 0.5, 0.45, 0.55, color,
                            name+' chair'))
    return traces

def parallel_bars_3d(cx, cy, z0, color='#26a69a', name='PT Bars'):
    """Physical therapy parallel bars + exercise mat."""
    traces = []
    # Mat on floor
    traces.append(box_mesh_global(cx-0.8, cy-1.0, z0,
                                   1.6, 2.0, 0.05, '#b2dfdb', 1.0, name+' mat'))
    # 4 vertical posts
    for px, py in [(-0.7, -0.9), (0.7, -0.9), (-0.7, 0.9), (0.7, 0.9)]:
        traces.append(cylinder_mesh(cx+px, cy+py, z0+0.05,
                                     0.04, 1.0, 'z', 8, color, 1.0,
                                     name+' post'))
    # 2 horizontal bars
    for bx in [-0.7, 0.7]:
        traces.append(cylinder_mesh(cx+bx, cy-0.9, z0+1.05,
                                     0.04, 1.8, 'y', 8, color, 1.0,
                                     name+' bar'))
    return traces

def maternity_bed_3d(cx, cy, z0, color='#f06292', name='Delivery Bed'):
    """Wider birthing bed + baby warmer beside it."""
    traces = []
    # Larger bed (scale 1.2)
    traces.extend(hospital_bed_3d(cx, cy, z0, color, size=1.15, name=name))
    # Baby warmer next to bed (small heated unit)
    warm_x = cx + 1.0
    traces.append(box_mesh_global(warm_x-0.3, cy-0.4, z0,
                                   0.6, 0.8, 0.9, '#fff59d', 1.0,
                                   name+' warmer-stand'))
    traces.append(box_mesh_global(warm_x-0.35, cy-0.45, z0+0.9,
                                   0.7, 0.9, 0.05, '#fce4ec', 1.0,
                                   name+' bassinet'))
    # Heat lamp above
    traces.append(cylinder_mesh(warm_x, cy, z0+1.5, 0.15, 0.2, 'z', 12,
                                 '#ff6f00', 1.0, name+' heat-lamp'))
    return traces

# ============================================================
# HOSPITAL BUILDING BLOCKS (for full hospital 3D view)
# ============================================================
class MeshAccumulator:
    """Accumulates many shape vertices into ONE big Mesh3d.
    Critical for performance: 500 small meshes → laggy.
    Combined into 10 grouped meshes → smooth."""
    def __init__(self, color, name='Group', opacity=1.0):
        self.x, self.y, self.z = [], [], []
        self.i, self.j, self.k = [], [], []
        self.color = color
        self.name = name
        self.opacity = opacity
        self.n_verts = 0

    def add_box(self, x0, y0, z0, dx, dy, dz):
        b = self.n_verts
        self.x.extend([x0, x0+dx, x0+dx, x0, x0, x0+dx, x0+dx, x0])
        self.y.extend([y0, y0, y0+dy, y0+dy, y0, y0, y0+dy, y0+dy])
        self.z.extend([z0, z0, z0, z0, z0+dz, z0+dz, z0+dz, z0+dz])
        ti = [7,0,0,0,4,4,6,6,4,0,3,2]
        tj = [3,4,1,2,5,6,5,2,0,1,6,3]
        tk = [0,7,2,3,6,7,1,1,5,5,7,6]
        self.i.extend([b+v for v in ti])
        self.j.extend([b+v for v in tj])
        self.k.extend([b+v for v in tk])
        self.n_verts += 8

    def add_cylinder(self, cx, cy, z0, radius, height, n=12):
        """Add a vertical cylinder."""
        b = self.n_verts
        theta = np.linspace(0, 2*np.pi, n, endpoint=False)
        cos_t = np.cos(theta) * radius
        sin_t = np.sin(theta) * radius
        # Bottom + top rings + 2 cap centers
        for ct, st in zip(cos_t, sin_t):
            self.x.append(cx + ct); self.y.append(cy + st); self.z.append(z0)
        for ct, st in zip(cos_t, sin_t):
            self.x.append(cx + ct); self.y.append(cy + st); self.z.append(z0+height)
        self.x.extend([cx, cx]); self.y.extend([cy, cy])
        self.z.extend([z0, z0+height])
        # Side quads (2 triangles each)
        for k_ in range(n):
            nk = (k_+1) % n
            self.i.append(b+k_); self.j.append(b+nk); self.k.append(b+k_+n)
            self.i.append(b+nk); self.j.append(b+nk+n); self.k.append(b+k_+n)
        # Bottom cap
        for k_ in range(n):
            nk = (k_+1) % n
            self.i.append(b+2*n); self.j.append(b+nk); self.k.append(b+k_)
        # Top cap
        for k_ in range(n):
            nk = (k_+1) % n
            self.i.append(b+2*n+1); self.j.append(b+k_+n); self.k.append(b+nk+n)
        self.n_verts += 2*n + 2

    def add_cone(self, cx, cy, z0, radius, height, n=10):
        """Add a cone (apex up). Good for tree leaves."""
        b = self.n_verts
        theta = np.linspace(0, 2*np.pi, n, endpoint=False)
        for ang in theta:
            self.x.append(cx + radius*np.cos(ang))
            self.y.append(cy + radius*np.sin(ang))
            self.z.append(z0)
        # apex + base center
        self.x.extend([cx, cx])
        self.y.extend([cy, cy])
        self.z.extend([z0+height, z0])
        apex_idx = b + n
        base_c = b + n + 1
        for k_ in range(n):
            nk = (k_+1) % n
            self.i.append(b+k_); self.j.append(b+nk); self.k.append(apex_idx)
            self.i.append(base_c); self.j.append(b+nk); self.k.append(b+k_)
        self.n_verts += n + 2

    def to_mesh3d(self):
        if not self.x:
            return None
        return go.Mesh3d(
            x=self.x, y=self.y, z=self.z,
            i=self.i, j=self.j, k=self.k,
            color=self.color, opacity=self.opacity,
            name=self.name, hovertext=self.name,
            hoverinfo='text', flatshading=True)


def building_block_3d(x, y, w, d, h, color, name='Building',
                      add_solar=False, add_green_roof=False):
    """A complete hospital wing: walls + roof + optional solar/green roof."""
    traces = []
    # Main building mass
    traces.append(box_mesh_global(x, y, 0, w, d, h, color, 0.92, name))
    # Roof slab (slight overhang)
    traces.append(box_mesh_global(x-0.2, y-0.2, h, w+0.4, d+0.4, 0.15,
                                   '#37474f', 1.0, name+' roof'))
    # Solar panels on top
    if add_solar:
        # Cover ~70% of roof with panels
        panel_w, panel_d = 0.9, 1.2
        nx = max(1, int((w-0.4) / panel_w))
        ny = max(1, int((d-0.4) / panel_d))
        for i in range(nx):
            for j in range(ny):
                px = x + 0.2 + i*panel_w
                py = y + 0.2 + j*panel_d
                traces.append(box_mesh_global(px, py, h+0.18,
                                               panel_w*0.85, panel_d*0.85, 0.04,
                                               '#1565c0', 1.0,
                                               name+' solar'))
    # Green roof carpet
    if add_green_roof and not add_solar:
        traces.append(box_mesh_global(x+0.1, y+0.1, h+0.16,
                                       w-0.2, d-0.2, 0.06, '#66bb6a', 1.0,
                                       name+' green-roof'))
    # Windows pattern on front facade (small dark rectangles)
    win_size = 0.4
    floors = max(1, int(h / 1.2))
    win_count = max(2, int(w / 1.0))
    for f in range(floors):
        for wi in range(win_count):
            wx = x + 0.3 + wi * (w - 0.6) / max(1, win_count - 1)
            wy = y - 0.01  # slightly in front of facade
            wz = 0.5 + f * 1.2
            traces.append(box_mesh_global(wx-win_size/2, wy, wz,
                                           win_size, 0.02, 0.5,
                                           '#4fc3f7', 0.75,
                                           name+' window'))
    return traces


# ════════════════════════════════════════════════════════════════
# PDF REPORT GENERATOR
# Uses matplotlib's PdfPages for a self-contained, dependency-free
# report (no need for reportlab/fpdf — matplotlib ships with Streamlit).
# ════════════════════════════════════════════════════════════════
from matplotlib.backends.backend_pdf import PdfPages
from datetime import datetime as _dt

def generate_pdf_report(config_type, config, metrics, blueprint_fig=None):
    """Build a multi-page PDF report and return it as bytes.

    config_type: 'single' or 'hospital'
    config: the single_config or hospital_config dict
    metrics: dict of computed numbers (green_score, energy_saved, etc.)
    blueprint_fig: optional matplotlib figure to embed
    """
    buf = BytesIO()
    with PdfPages(buf) as pdf:
        # ────── PAGE 1: COVER ──────
        fig = plt.figure(figsize=(8.27, 11.69))  # A4 portrait
        fig.patch.set_facecolor('#f1f8e9')
        ax = fig.add_subplot(111)
        ax.axis('off')

        # Logo + title
        ax.text(0.5, 0.85, "🍃", fontsize=80, ha='center',
                color='#2e7d32', transform=ax.transAxes)
        ax.text(0.5, 0.74, "VIRIDIS", fontsize=32, ha='center',
                weight='bold', color='#1b5e20', transform=ax.transAxes,
                family='serif')
        ax.text(0.5, 0.69, "Green Hospital Design Report",
                fontsize=14, ha='center', color='#388e3c',
                transform=ax.transAxes, style='italic')

        # Decorative line
        ax.plot([0.2, 0.8], [0.64, 0.64], color='#2e7d32',
                linewidth=2, transform=ax.transAxes)

        # Subject
        if config_type == 'single':
            subject = config.get('facility_type', 'Medical Facility')
            subtitle = f"Single Department Design"
        else:
            subject = config.get('hospital_scale', 'Hospital')
            n_depts = len(config.get('selected_depts', {}))
            subtitle = f"Full Hospital — {n_depts} Departments"

        ax.text(0.5, 0.55, subject, fontsize=20, ha='center',
                weight='bold', color='#1b5e20', transform=ax.transAxes)
        ax.text(0.5, 0.51, subtitle, fontsize=12, ha='center',
                color='#558b2f', transform=ax.transAxes, style='italic')

        # Key metrics preview
        date_str = _dt.now().strftime("%B %d, %Y")
        ax.text(0.5, 0.42, f"Generated: {date_str}",
                fontsize=10, ha='center', color='#558b2f',
                transform=ax.transAxes)

        # Metrics summary box
        if config_type == 'single':
            metrics_text = (
                f"Green Sustainability Score:  {metrics.get('green_score', 0):.0f}%\n"
                f"Monthly Energy Saved:  {metrics.get('energy_saved', 0):.0f} kWh\n"
                f"Monthly Cost Savings:  ${metrics.get('money_saved', 0):.0f}\n"
            )
        else:
            metrics_text = (
                f"Annual CO₂ Reduction:  {metrics.get('co2_saved', 0)*12/1000:.1f} tons\n"
                f"Annual Energy Saved:  {metrics.get('energy_saved', 0)*12/1000:.0f} MWh\n"
                f"Annual Cost Savings:  ${metrics.get('money_saved', 0)*12:,.0f}\n"
                f"Trees Equivalent:  {int(metrics.get('trees_equiv', 0)*12):,}\n"
            )

        ax.text(0.5, 0.28, "KEY ENVIRONMENTAL METRICS", fontsize=10,
                ha='center', weight='bold', color='#1b5e20',
                transform=ax.transAxes, style='italic')
        ax.text(0.5, 0.18, metrics_text, fontsize=11, ha='center',
                color='#263238', transform=ax.transAxes, family='monospace',
                bbox=dict(boxstyle='round,pad=1.2', facecolor='white',
                          edgecolor='#2e7d32', linewidth=1.5))

        # Footer
        ax.text(0.5, 0.05, "Automated Eco-Friendly Hospital Design",
                fontsize=8, ha='center', color='#558b2f',
                transform=ax.transAxes, style='italic')
        pdf.savefig(fig, bbox_inches='tight')
        plt.close(fig)

        # ────── PAGE 2: DESIGN PARAMETERS ──────
        fig = plt.figure(figsize=(8.27, 11.69))
        fig.patch.set_facecolor('#ffffff')
        ax = fig.add_subplot(111)
        ax.axis('off')
        ax.text(0.5, 0.95, "Design Parameters", fontsize=20, ha='center',
                weight='bold', color='#1b5e20', transform=ax.transAxes)
        ax.plot([0.1, 0.9], [0.92, 0.92], color='#2e7d32',
                linewidth=1.5, transform=ax.transAxes)

        # Parameters table
        if config_type == 'single':
            params = [
                ("Facility Type", config.get('facility_type', '—')),
                ("Total Area", f"{config.get('total_area', 0)} m²"),
                ("Number of Rooms", config.get('rooms_count', 0)),
                ("Units per Room", config.get('units_per_room', 0)),
                ("Climate Zone", config.get('climate', '—')),
                ("Lighting", config.get('lighting_type', '—')),
                ("Flooring Material", config.get('flooring', '—')),
                ("Wall Material", config.get('wall_material', '—')),
                ("Ceiling Material", config.get('ceiling_material', '—')),
                ("Insulation", config.get('insulation_material', '—')),
                ("Interior Paint", config.get('paint_material', '—')),
                ("Waste Management", config.get('waste_mgmt', '—')),
                ("Equipment Class", config.get('equip_efficiency', '—')),
                ("Power Source", config.get('power_source', '—')),
                ("Scrub Sink", config.get('scrub_sink', '—')),
                ("Control Panel", config.get('control_panel', '—')),
                ("Glass Partitions", "Yes" if config.get('partition_type') == "Antibacterial Glass Walls" else "No"),
                ("Laminar Flow Ceiling", "Yes" if config.get('laminar_ceiling') else "No"),
            ]
        else:
            sd = config.get('selected_depts', {})
            params = [
                ("Hospital Scale", config.get('hospital_scale', '—')),
                ("Total Departments", len(sd)),
                ("Building Mode", config.get('building_mode', '—')),
                ("Solar PV Farm", "Yes" if config.get('h_solar') else "No"),
                ("Healing Garden", "Yes" if config.get('h_garden') else "No"),
                ("Green Roof", "Yes" if config.get('h_green_roof') else "No"),
                ("Rainwater Harvesting", "Yes" if config.get('h_rainwater') else "No"),
                ("Greywater Recycling", "Yes" if config.get('h_greywater') else "No"),
                ("EV Charging", "Yes" if config.get('h_ev_chargers') else "No"),
                ("Smart Grid", "Yes" if config.get('h_smart_grid') else "No"),
                ("Parking Garage", "Yes" if config.get('h_garage') else "No"),
                ("Road Network", "Yes" if config.get('h_roads') else "No"),
                ("Outdoor Seating", "Yes" if config.get('h_outdoor_seating') else "No"),
            ]

        y_pos = 0.85
        for label, value in params:
            ax.text(0.1, y_pos, f"{label}:", fontsize=10,
                    ha='left', color='#37474f', weight='bold',
                    transform=ax.transAxes)
            ax.text(0.6, y_pos, str(value), fontsize=10,
                    ha='left', color='#1b5e20',
                    transform=ax.transAxes)
            y_pos -= 0.045
            if y_pos < 0.05: break
        pdf.savefig(fig, bbox_inches='tight')
        plt.close(fig)

        # ────── PAGE 3: BLUEPRINT (if provided) ──────
        if blueprint_fig is not None:
            pdf.savefig(blueprint_fig, bbox_inches='tight')

        # ────── PAGE 4: ENVIRONMENTAL IMPACT ──────
        fig = plt.figure(figsize=(8.27, 11.69))
        fig.patch.set_facecolor('#ffffff')

        # Top: title
        ax_t = fig.add_axes([0.1, 0.92, 0.8, 0.05])
        ax_t.axis('off')
        ax_t.text(0.5, 0.5, "Environmental Impact Analysis",
                  fontsize=18, ha='center', weight='bold',
                  color='#1b5e20', transform=ax_t.transAxes)

        # Bar chart (energy)
        ax_bar = fig.add_axes([0.15, 0.55, 0.7, 0.3])
        if config_type == 'single':
            normal = metrics.get('normal_env', 1000)
            green = metrics.get('green_env', 700)
        else:
            normal = metrics.get('energy_normal', 50000) * 12 / 1000
            green = metrics.get('energy_green', 30000) * 12 / 1000
        ax_bar.bar(['Traditional', 'Viridis'], [normal, green],
                   color=['#ef5350', '#66bb6a'], width=0.5,
                   edgecolor='black', linewidth=1.2)
        ax_bar.set_ylabel("Energy (kWh/mo)" if config_type == 'single'
                          else "Energy (MWh/year)", fontsize=11)
        ax_bar.set_title("Energy Comparison", fontsize=13, weight='bold',
                         color='#1b5e20')
        ax_bar.grid(axis='y', alpha=0.3)
        for i, v in enumerate([normal, green]):
            ax_bar.text(i, v + max(normal, green)*0.02, f"{v:.0f}",
                        ha='center', fontsize=10, weight='bold')

        # Bottom: impact narrative
        ax_n = fig.add_axes([0.1, 0.05, 0.8, 0.45])
        ax_n.axis('off')
        if config_type == 'hospital':
            co2_yr = metrics.get('co2_saved', 0) * 12 / 1000
            money_yr = metrics.get('money_saved', 0) * 12
            trees_yr = int(metrics.get('trees_equiv', 0) * 12)
            water_yr = metrics.get('water_saved', 0) * 12 / 1000
            narrative = (
                "ANNUAL SUSTAINABILITY ACHIEVEMENTS\n\n"
                f"• CO₂ emissions avoided:  {co2_yr:.1f} tons/year\n"
                f"• Trees equivalent:         {trees_yr:,} trees\n"
                f"• Energy saved:             {metrics.get('energy_saved', 0)*12/1000:.0f} MWh/year\n"
                f"• Money saved:              ${money_yr:,.0f}/year\n"
                f"• Water saved:              {water_yr:.0f} m³/year\n\n"
                "These results are equivalent to:\n"
                f"  ▸ Removing {int(co2_yr / 4.6):,} cars from the road for a year\n"
                f"  ▸ Powering {int(metrics.get('energy_saved', 0)*12/3500):,} homes annually\n"
                f"  ▸ Drinking water for {int(water_yr * 1000 / 180):,} people per year"
            )
        else:
            narrative = (
                "MONTHLY SUSTAINABILITY METRICS\n\n"
                f"• Green Score:           {metrics.get('green_score', 0):.0f}%\n"
                f"• Energy saved:          {metrics.get('energy_saved', 0):.0f} kWh/month\n"
                f"• Savings percentage:    {metrics.get('savings_pct', 0)*100:.1f}%\n"
                f"• Monthly cost savings:  ${metrics.get('money_saved', 0):.0f}\n\n"
                "Equivalent annual impact:\n"
                f"  ▸ ${metrics.get('money_saved', 0)*12:,.0f} saved annually\n"
                f"  ▸ {metrics.get('energy_saved', 0)*12/1000:.1f} MWh/year reduction"
            )
        ax_n.text(0.05, 0.95, narrative, fontsize=11, ha='left',
                  va='top', family='monospace', color='#37474f',
                  transform=ax_n.transAxes,
                  bbox=dict(boxstyle='round,pad=0.8', facecolor='#f1f8e9',
                            edgecolor='#2e7d32', linewidth=1.5))

        pdf.savefig(fig, bbox_inches='tight')
        plt.close(fig)

        # Set PDF metadata
        info = pdf.infodict()
        info['Title'] = f'Viridis Report - {subject}'
        info['Author'] = 'Viridis Green Hospital Planner'
        info['Subject'] = 'Hospital sustainability and design report'
        info['CreationDate'] = _dt.now()

    buf.seek(0)
    return buf.getvalue()


# ════════════════════════════════════════════════════════════════
# LUX HEATMAP GENERATOR
# Estimates illumination level across a room based on window position,
# ceiling lights, and time of day. Produces a matplotlib heatmap.
# ════════════════════════════════════════════════════════════════
def generate_lux_heatmap(width, height, facility_type, lighting_type,
                          required_lux):
    """Generate a 2D lux heatmap of the room."""
    # Build grid
    resolution = 50
    x = np.linspace(0, width, resolution)
    y = np.linspace(0, height, resolution)
    X, Y = np.meshgrid(x, y)

    # Base illumination from ceiling LEDs (uniform contribution)
    led_positions = []
    for lx in np.linspace(width*0.25, width*0.75, 3):
        for ly in np.linspace(height*0.3, height*0.8, 2):
            led_positions.append((lx, ly))

    # Each LED produces ~600 lux directly below, falling off with distance²
    led_lux = np.zeros_like(X)
    for lx, ly in led_positions:
        dist_sq = (X - lx)**2 + (Y - ly)**2 + 0.5
        led_lux += 600 / (1 + dist_sq * 0.3)

    # Natural light from window (on right wall if lighting includes natural)
    natural_lux = np.zeros_like(X)
    if lighting_type in ["Natural", "Mixed"]:
        # Window centered on right wall, illuminates inward
        win_y = height / 2
        for ix in range(resolution):
            for iy in range(resolution):
                dist_from_window = abs(width - X[iy, ix])
                vertical_offset = abs(win_y - Y[iy, ix])
                # Brighter near window, falls off
                natural_lux[iy, ix] = 1500 * np.exp(-dist_from_window * 0.4) \
                                       * np.exp(-vertical_offset * 0.3)

    total_lux = led_lux + natural_lux

    # Plot heatmap
    fig, ax = plt.subplots(figsize=(9, 6))
    levels = [0, 100, 200, 300, 500, 750, 1000, 1500, 2000, 3000]
    cmap = plt.cm.viridis
    cs = ax.contourf(X, Y, total_lux, levels=levels, cmap=cmap, alpha=0.85)
    cbar = plt.colorbar(cs, ax=ax, label='Illuminance (lux)', shrink=0.8)
    cbar.ax.tick_params(labelsize=8)

    # Show LED positions
    for lx, ly in led_positions:
        ax.plot(lx, ly, '*', markersize=14, color='#fff59d',
                markeredgecolor='#f57f17', markeredgewidth=1.5,
                label='LED Panel' if (lx, ly) == led_positions[0] else None)

    # Show window
    if lighting_type in ["Natural", "Mixed"]:
        ax.plot([width, width], [height/2 - 1.2, height/2 + 1.2],
                color='#0277bd', linewidth=6, label='Window')

    # Required lux line indicator
    ax.contour(X, Y, total_lux, levels=[required_lux], colors='red',
               linewidths=2, linestyles='dashed')

    # Walls outline
    ax.add_patch(plt.Rectangle((0, 0), width, height, fill=False,
                                edgecolor='black', linewidth=2))

    ax.set_xlabel('Width (m)', fontsize=10)
    ax.set_ylabel('Length (m)', fontsize=10)
    ax.set_title(f"Lux Heatmap — {facility_type}\n"
                 f"Required: {required_lux} lux  •  "
                 f"Lighting: {lighting_type}",
                 fontsize=11, weight='bold')
    ax.set_aspect('equal')
    ax.legend(loc='upper left', fontsize=8, framealpha=0.9)
    ax.grid(alpha=0.2)
    plt.tight_layout()
    return fig, total_lux


# ════════════════════════════════════════════════════════════════
# ════════════════════════════════════════════════════════════════
#         CUSTOM DEPARTMENT LAYOUTS (architecturally realistic)
#
#  Each layout function draws ALL the furniture/equipment for that
#  specific department in an architecturally correct arrangement
#  (beds on walls, service bars on walls, nursing stations centered,
#   etc.) instead of the generic "units in middle" pattern.
# ════════════════════════════════════════════════════════════════
# ════════════════════════════════════════════════════════════════

import matplotlib.patches as _mpat
from matplotlib.patches import (Circle as _MCircle, Rectangle as _MRect,
                                 FancyBboxPatch as _MFBox)


# ────────────────────────────────────────────────────────────
# 🍽️ CAFETERIA & KITCHEN (2D)
# Service bar on back wall + salad bar + cashier + dining tables
# ────────────────────────────────────────────────────────────
def layout_cafeteria_2d(ax, width, height, color, n_seats=40):
    """Cafeteria with proper service-line architecture.
    n_seats drives number of dining tables (4 seats per table)."""
    # ─── SERVICE BAR along back wall ───
    bar_y = height - 1.5
    bar_h = 1.0
    # Main service counter (long rectangle)
    ax.add_patch(_MFBox((width*0.1, bar_y), width*0.6, bar_h,
        boxstyle="round,pad=0.05", facecolor='#a5d6a7',
        edgecolor=color, linewidth=2, zorder=6))
    # Sneeze guard / glass on top
    ax.plot([width*0.1, width*0.7], [bar_y + bar_h + 0.1,
            bar_y + bar_h + 0.1], color='#4fc3f7',
            linewidth=2, alpha=0.6, zorder=7)

    # ─── Service stations scale with size: 3 for small, 4 for medium, 5 for big ───
    n_stations = 3 if n_seats <= 24 else 4 if n_seats <= 36 else 5
    station_w = (width * 0.55) / n_stations
    station_options = [
        ("🥗 Salad\nBar", '#81c784'),
        ("🍲 Hot\nFood", '#ff8a65'),
        ("🥤 Beverages", '#64b5f6'),
        ("🥖 Bakery", '#bcaaa4'),
        ("🍰 Desserts", '#f48fb1'),
    ]
    for i in range(n_stations):
        lbl, scol = station_options[i % len(station_options)]
        sx = width*0.12 + i * station_w
        # Heat lamp / display top
        ax.add_patch(_MRect((sx, bar_y + bar_h - 0.2), station_w*0.85, 0.2,
            facecolor=scol, edgecolor='black', linewidth=0.5, zorder=7))
        # Station icon/label
        ax.text(sx + station_w*0.42, bar_y + bar_h/2 + 0.1, lbl,
            color='white', fontsize=6, ha='center', va='center',
            weight='bold', zorder=8)

    # Bar label
    ax.text(width*0.4, bar_y + bar_h + 0.4,
        f"SERVICE LINE ({n_stations} stations)",
        color=color, fontsize=8, ha='center', weight='bold',
        style='italic', zorder=7)

    # ─── TRAY PICKUP at start of line ───
    ax.add_patch(_MRect((width*0.05, bar_y + 0.1), 0.15, bar_h - 0.2,
        facecolor='#bcaaa4', edgecolor='black', linewidth=0.5, zorder=6))
    ax.text(width*0.05 + 0.075, bar_y - 0.2, "Trays",
        color='#5d4037', fontsize=5, ha='center', weight='bold', zorder=6)

    # ─── CASHIER STATIONS (more for larger cafeterias) ───
    n_cashiers = 1 if n_seats <= 20 else 2 if n_seats <= 32 else 3
    for ci in range(n_cashiers):
        cash_x = width*0.72 + ci * 0.6
        ax.add_patch(_MFBox((cash_x, bar_y + 0.1), 0.5, bar_h - 0.2,
            boxstyle="round,pad=0.03", facecolor='#ffd54f',
            edgecolor='#f57c00', linewidth=1.5, zorder=6))
        ax.text(cash_x + 0.25, bar_y + bar_h/2, "💳",
            fontsize=10, ha='center', va='center', zorder=7)
        ax.text(cash_x + 0.25, bar_y - 0.2, f"Cashier {ci+1}" if n_cashiers > 1 else "Cashier",
            color='#ef6c00', fontsize=5, ha='center', weight='bold', zorder=6)

    # ─── DINING TABLES (round, 4 seats each) ───
    n_tables = max(1, n_seats // 4)
    # Try to fit n_tables in a grid; calculate optimal cols/rows
    cols = max(2, int(width / 2.2))
    rows = max(1, (n_tables + cols - 1) // cols)
    table_x_step = (width - 1.5) / cols
    table_y_step = (height - 3.5) / max(1, rows)

    t_idx = 0
    for r in range(rows):
        for c in range(cols):
            if t_idx >= n_tables:
                break
            tx = 0.75 + c * table_x_step + table_x_step/2
            ty = 0.8 + r * table_y_step + table_y_step/2
            if ty > height - 2.2:  # don't overlap service bar
                continue
            # Table top
            ax.add_patch(_MCircle((tx, ty), 0.4, facecolor='#fff8e1',
                edgecolor='#5d4037', linewidth=1, zorder=5))
            # 4 chairs around
            for ang in [0, np.pi/2, np.pi, 3*np.pi/2]:
                cx = tx + 0.6*np.cos(ang)
                cy = ty + 0.6*np.sin(ang)
                ax.add_patch(_MCircle((cx, cy), 0.15,
                    facecolor=color, edgecolor='black',
                    linewidth=0.3, zorder=5))
            t_idx += 1
        if t_idx >= n_tables:
            break

    actual_seats = t_idx * 4
    ax.text(width/2, 0.4,
        f"DINING AREA ({t_idx} tables · {actual_seats} seats)",
        color='#5d4037', fontsize=7, ha='center',
        weight='bold', style='italic', zorder=5)

    # ─── TRASH/RECYCLING at exit ───
    ax.add_patch(_MRect((width*0.92, bar_y - 0.3), 0.4, 0.5,
        facecolor='#9e9e9e', edgecolor='black', linewidth=0.5, zorder=6))
    ax.text(width*0.92 + 0.2, bar_y - 0.5, "🗑️",
        fontsize=6, ha='center', zorder=7)


# ────────────────────────────────────────────────────────────
# 🍽️ CAFETERIA & KITCHEN (3D)
# ────────────────────────────────────────────────────────────
def layout_cafeteria_3d(fig, width, height, color, wall_height, n_seats=40):
    """Cafeteria 3D: service bar, stations, cashier, dining tables.
    n_seats drives number of tables and stations."""
    # Service bar
    bar_y = height - 1.5
    fig.add_trace(box_mesh_global(width*0.1, bar_y, 0,
        width*0.6, 1.0, 0.95, '#a5d6a7', 1.0, 'Service Counter'))
    # Counter top
    fig.add_trace(box_mesh_global(width*0.1, bar_y, 0.95,
        width*0.6, 1.0, 0.05, color, 1.0, 'Counter Top'))

    # Service stations scale with size
    n_stations = 3 if n_seats <= 24 else 4 if n_seats <= 36 else 5
    station_w = (width * 0.55) / n_stations
    station_options = [
        ('#81c784', 'Salad Bar', False),
        ('#ff8a65', 'Hot Food', True),  # has heat lamp
        ('#64b5f6', 'Beverages', False),
        ('#bcaaa4', 'Bakery', False),
        ('#f48fb1', 'Desserts', False),
    ]
    for i in range(n_stations):
        scol, nm, has_lamp = station_options[i % len(station_options)]
        sx = width*0.12 + i * station_w
        fig.add_trace(box_mesh_global(sx, bar_y + 0.1, 1.0,
            station_w*0.85, 0.8, 0.3, scol, 1.0, nm))
        fig.add_trace(box_mesh_global(sx, bar_y - 0.05, 1.3,
            station_w*0.85, 0.05, 0.5, '#81d4fa', 0.4, nm + ' Glass'))
        if has_lamp:
            fig.add_trace(cylinder_mesh(sx + station_w*0.42,
                bar_y + 0.5, 2.0, 0.1, 0.15, 'z', 12,
                '#ff6f00', 1.0, 'Heat Lamp'))

    # Tray rack
    fig.add_trace(box_mesh_global(width*0.05, bar_y + 0.1, 0,
        0.15, 0.8, 1.1, '#bcaaa4', 1.0, 'Tray Rack'))

    # Cashier stations (dynamic count)
    n_cashiers = 1 if n_seats <= 20 else 2 if n_seats <= 32 else 3
    for ci in range(n_cashiers):
        cash_x = width*0.72 + ci * 0.6
        fig.add_trace(box_mesh_global(cash_x, bar_y + 0.1, 0,
            0.5, 0.8, 1.0, '#ffd54f', 1.0, f'Cashier {ci+1}'))
        fig.add_trace(box_mesh_global(cash_x + 0.1, bar_y + 0.2, 1.0,
            0.3, 0.4, 0.2, '#212121', 1.0, 'POS Terminal'))
        fig.add_trace(box_mesh_global(cash_x + 0.13, bar_y + 0.18, 1.05,
            0.24, 0.02, 0.15, '#4fc3f7', 1.0, 'POS Screen'))

    # Dining tables + chairs (round)
    n_tables = max(1, n_seats // 4)
    cols = max(2, int(width / 2.2))
    rows = max(1, (n_tables + cols - 1) // cols)
    table_x_step = (width - 1.5) / cols
    table_y_step = (height - 3.5) / max(1, rows)
    t_idx = 0
    for r in range(rows):
        for c in range(cols):
            if t_idx >= n_tables:
                break
            tx = 0.75 + c * table_x_step + table_x_step/2
            ty = 0.8 + r * table_y_step + table_y_step/2
            if ty > height - 2.2: continue
            # Table top
            fig.add_trace(cylinder_mesh(tx, ty, 0.7, 0.4, 0.04, 'z', 16,
                '#fff8e1', 1.0, 'Table'))
            # Pedestal
            fig.add_trace(cylinder_mesh(tx, ty, 0, 0.08, 0.7, 'z', 8,
                '#5d4037', 1.0, 'Pedestal'))
            # 4 chairs
            for ang in [0, np.pi/2, np.pi, 3*np.pi/2]:
                cx = tx + 0.65*np.cos(ang)
                cy = ty + 0.65*np.sin(ang)
                fig.add_trace(cylinder_mesh(cx, cy, 0.45, 0.18, 0.05, 'z', 10,
                    color, 1.0, 'Chair Seat'))
                back_ang = ang - np.pi
                fig.add_trace(box_mesh_global(
                    cx - 0.05 + 0.15*np.cos(back_ang),
                    cy - 0.05 + 0.15*np.sin(back_ang),
                    0.5, 0.1, 0.04, 0.4, color, 1.0, 'Chair Back'))
            t_idx += 1
        if t_idx >= n_tables:
            break


# ────────────────────────────────────────────────────────────
# 🛏️ PATIENT WARD / ICU / ER (2D)
# Beds on walls (head against wall), nursing station center
# ────────────────────────────────────────────────────────────
def layout_ward_2d(ax, width, height, color, n_beds,
                   has_monitor=True, has_iv=True, label_prefix="Bed",
                   glass_partitions=False):
    """Hospital ward with beds along walls and central nursing station.
    If glass_partitions=True, draws antibacterial glass walls between
    adjacent beds for infection control."""
    # Distribute beds: half along top wall, half along bottom wall
    # (heads against wall, foot pointing inward)
    bed_w, bed_l = 0.9, 2.0

    # How many fit per wall
    max_per_wall = max(1, int((width - 1.5) / 1.5))
    top_beds = (n_beds + 1) // 2
    bottom_beds = n_beds - top_beds
    top_beds = min(top_beds, max_per_wall)
    bottom_beds = min(bottom_beds, max_per_wall)

    bed_positions_drawn = []

    # ─── TOP WALL beds (head against top wall, foot pointing down) ───
    if top_beds > 0:
        x_step = (width - 1) / top_beds
        for i in range(top_beds):
            bx = 0.5 + x_step * (i + 0.5)
            by = height - 1.5  # bed center near top
            # Bed body
            ax.add_patch(_MFBox((bx - bed_w/2, by - bed_l/2 + 0.2),
                bed_w, bed_l - 0.2, boxstyle="round,pad=0.02",
                facecolor='#eceff1', edgecolor=color, linewidth=1.5, zorder=6))
            # Headboard against TOP wall
            ax.add_patch(_MRect((bx - bed_w/2, by + bed_l/2 - 0.25),
                bed_w, 0.25, facecolor=color, zorder=6))
            # Pillow (top, near head)
            ax.add_patch(_MFBox((bx - bed_w/2 + 0.1, by + bed_l/2 - 0.55),
                bed_w - 0.2, 0.3, boxstyle="round,pad=0.02",
                facecolor='white', edgecolor='#bdbdbd', linewidth=0.4, zorder=7))
            # Bed label
            ax.text(bx, by - bed_l/2 - 0.3, f"{label_prefix} {i+1}",
                color=color, fontsize=6, ha='center', weight='bold', zorder=7)
            # Vital monitor on wall above
            if has_monitor:
                ax.add_patch(_MRect((bx - 0.2, height - 0.2),
                    0.4, 0.15, facecolor='#1a237e', zorder=7))
                ax.text(bx, height - 0.12, "M", color='lime',
                    fontsize=4, ha='center', va='center', weight='bold', zorder=8)
            # IV pole beside bed
            if has_iv:
                ax.add_patch(_MCircle((bx - bed_w/2 - 0.2, by + 0.3),
                    0.07, facecolor='#9e9e9e', zorder=7))
            # Side table for visitor
            ax.add_patch(_MRect((bx + bed_w/2 + 0.05, by - 0.2),
                0.3, 0.4, facecolor='#bcaaa4', edgecolor='black',
                linewidth=0.3, zorder=6))
            bed_positions_drawn.append((bx, by))

    # ─── BOTTOM WALL beds (head against bottom wall, foot pointing up) ───
    if bottom_beds > 0:
        x_step = (width - 1) / bottom_beds
        for i in range(bottom_beds):
            bx = 0.5 + x_step * (i + 0.5)
            by = 1.5  # bed center near bottom
            # Bed body
            ax.add_patch(_MFBox((bx - bed_w/2, by - bed_l/2),
                bed_w, bed_l - 0.2, boxstyle="round,pad=0.02",
                facecolor='#eceff1', edgecolor=color, linewidth=1.5, zorder=6))
            # Headboard against BOTTOM wall
            ax.add_patch(_MRect((bx - bed_w/2, by - bed_l/2),
                bed_w, 0.25, facecolor=color, zorder=6))
            # Pillow
            ax.add_patch(_MFBox((bx - bed_w/2 + 0.1, by - bed_l/2 + 0.25),
                bed_w - 0.2, 0.3, boxstyle="round,pad=0.02",
                facecolor='white', edgecolor='#bdbdbd', linewidth=0.4, zorder=7))
            ax.text(bx, by + bed_l/2 + 0.1, f"{label_prefix} {top_beds + i + 1}",
                color=color, fontsize=6, ha='center', weight='bold', zorder=7)
            if has_monitor:
                ax.add_patch(_MRect((bx - 0.2, 0.05),
                    0.4, 0.15, facecolor='#1a237e', zorder=7))
                ax.text(bx, 0.12, "M", color='lime',
                    fontsize=4, ha='center', va='center', weight='bold', zorder=8)
            if has_iv:
                ax.add_patch(_MCircle((bx - bed_w/2 - 0.2, by - 0.3),
                    0.07, facecolor='#9e9e9e', zorder=7))
            ax.add_patch(_MRect((bx + bed_w/2 + 0.05, by - 0.2),
                0.3, 0.4, facecolor='#bcaaa4', edgecolor='black',
                linewidth=0.3, zorder=6))
            bed_positions_drawn.append((bx, by))

    # ─── CENTRAL NURSING STATION ───
    ns_w, ns_h = min(2.5, width*0.4), 1.0
    ns_x = width/2 - ns_w/2
    ns_y = height/2 - ns_h/2
    ax.add_patch(_MFBox((ns_x, ns_y), ns_w, ns_h,
        boxstyle="round,pad=0.05", facecolor='#37474f',
        edgecolor=color, linewidth=2, zorder=6))
    # Counter top (lighter)
    ax.add_patch(_MFBox((ns_x + 0.05, ns_y + 0.05), ns_w - 0.1, ns_h - 0.1,
        boxstyle="round,pad=0.03", facecolor='#cfd8dc',
        edgecolor='none', zorder=7))
    ax.text(ns_x + ns_w/2, ns_y + ns_h/2, "🩺 NURSING\nSTATION",
        color='#37474f', fontsize=7, ha='center', va='center',
        weight='bold', zorder=8)
    # 2 staff chairs at nursing station
    for cx_off in [-0.4, 0.4]:
        ax.add_patch(_MCircle((ns_x + ns_w/2 + cx_off, ns_y - 0.2),
            0.15, facecolor=color, zorder=6))

    # ─── HAND-WASH SINK in center ───
    ax.add_patch(_MRect((width/2 - 0.2, height - 0.5),
        0.4, 0.2, facecolor='#90a4ae', edgecolor='#37474f',
        linewidth=0.5, zorder=6))
    ax.text(width/2, height - 0.7, "🚿 Sink",
        color='#37474f', fontsize=5, ha='center', zorder=6)

    # ─── GLASS PARTITIONS between adjacent beds (if enabled) ───
    if glass_partitions:
        # Top wall partitions (vertical glass strips between top beds)
        if top_beds > 1:
            x_step = (width - 1) / top_beds
            for i in range(top_beds - 1):
                # Partition between bed i and bed i+1
                bx1 = 0.5 + x_step * (i + 0.5)
                bx2 = 0.5 + x_step * (i + 1.5)
                part_x = (bx1 + bx2) / 2
                # Glass wall (vertical strip from top wall extending down)
                ax.plot([part_x, part_x],
                        [height - bed_l - 0.3, height - 0.3],
                        color='#4fc3f7', linewidth=3.5, alpha=0.65,
                        solid_capstyle='round', zorder=5)
                # Glass shimmer (lighter inner line)
                ax.plot([part_x - 0.04, part_x - 0.04],
                        [height - bed_l - 0.2, height - 0.4],
                        color='#b3e5fc', linewidth=1, alpha=0.8, zorder=6)
                ax.text(part_x, height - bed_l - 0.45,
                        "║Glass║", color="#0277bd", fontsize=4,
                        ha="center", style='italic', weight='bold', zorder=6)
        # Bottom wall partitions
        if bottom_beds > 1:
            x_step = (width - 1) / bottom_beds
            for i in range(bottom_beds - 1):
                bx1 = 0.5 + x_step * (i + 0.5)
                bx2 = 0.5 + x_step * (i + 1.5)
                part_x = (bx1 + bx2) / 2
                ax.plot([part_x, part_x],
                        [0.3, bed_l + 0.3],
                        color='#4fc3f7', linewidth=3.5, alpha=0.65,
                        solid_capstyle='round', zorder=5)
                ax.plot([part_x - 0.04, part_x - 0.04],
                        [0.4, bed_l + 0.2],
                        color='#b3e5fc', linewidth=1, alpha=0.8, zorder=6)
                ax.text(part_x, bed_l + 0.45,
                        "║Glass║", color="#0277bd", fontsize=4,
                        ha="center", style='italic', weight='bold', zorder=6)

    return bed_positions_drawn


# ────────────────────────────────────────────────────────────
# 🛏️ PATIENT WARD / ICU / ER (3D)
# ────────────────────────────────────────────────────────────
def layout_ward_3d(fig, width, height, color, wall_height, n_beds,
                   has_monitor=True, has_iv=True, glass_partitions=False):
    """Ward 3D: beds along walls, central nursing station, optional
    glass partitions between beds (transparent walls)."""
    bed_w, bed_l = 0.9, 2.0
    max_per_wall = max(1, int((width - 1.5) / 1.5))
    top_beds = min((n_beds + 1) // 2, max_per_wall)
    bottom_beds = min(n_beds - top_beds, max_per_wall)

    # Top wall beds (head against top wall, foot pointing inward)
    if top_beds > 0:
        x_step = (width - 1) / top_beds
        for i in range(top_beds):
            bx = 0.5 + x_step * (i + 0.5)
            by = height - 1.3
            # Reuse the existing hospital bed mesh
            for t in hospital_bed_3d(bx, by, 0, color, 1.0, f'Bed Top {i+1}'):
                fig.add_trace(t)
            # Monitor on top wall above bed
            if has_monitor:
                fig.add_trace(box_mesh_global(bx - 0.2,
                    height - 0.15, 1.6, 0.4, 0.05, 0.3,
                    '#0d47a1', 1.0, f'Monitor T{i+1}'))
            # IV pole
            if has_iv:
                fig.add_trace(cylinder_mesh(bx - bed_w/2 - 0.25,
                    by + 0.3, 0, 0.025, 1.6, 'z', 8,
                    '#9e9e9e', 1.0, 'IV Pole'))
                fig.add_trace(box_mesh_global(bx - bed_w/2 - 0.33,
                    by + 0.24, 1.4, 0.16, 0.12, 0.2,
                    '#e1f5fe', 0.7, 'IV Bag'))
            # Side table
            fig.add_trace(box_mesh_global(bx + bed_w/2 + 0.1,
                by - 0.2, 0, 0.3, 0.4, 0.7,
                '#bcaaa4', 1.0, f'Side Table T{i+1}'))

    # Bottom wall beds
    if bottom_beds > 0:
        x_step = (width - 1) / bottom_beds
        for i in range(bottom_beds):
            bx = 0.5 + x_step * (i + 0.5)
            by = 1.3
            for t in hospital_bed_3d(bx, by, 0, color, 1.0, f'Bed Bot {i+1}'):
                fig.add_trace(t)
            if has_monitor:
                fig.add_trace(box_mesh_global(bx - 0.2,
                    0.1, 1.6, 0.4, 0.05, 0.3,
                    '#0d47a1', 1.0, f'Monitor B{i+1}'))
            if has_iv:
                fig.add_trace(cylinder_mesh(bx - bed_w/2 - 0.25,
                    by - 0.3, 0, 0.025, 1.6, 'z', 8,
                    '#9e9e9e', 1.0, 'IV Pole'))
                fig.add_trace(box_mesh_global(bx - bed_w/2 - 0.33,
                    by - 0.36, 1.4, 0.16, 0.12, 0.2,
                    '#e1f5fe', 0.7, 'IV Bag'))
            fig.add_trace(box_mesh_global(bx + bed_w/2 + 0.1,
                by - 0.2, 0, 0.3, 0.4, 0.7,
                '#bcaaa4', 1.0, f'Side Table B{i+1}'))

    # Central Nursing Station
    ns_w = min(2.5, width*0.4)
    ns_d = 1.0
    ns_x = width/2 - ns_w/2
    ns_y = height/2 - ns_d/2
    # Counter body
    fig.add_trace(box_mesh_global(ns_x, ns_y, 0, ns_w, ns_d, 0.95,
        '#37474f', 1.0, 'Nursing Station'))
    # Counter top
    fig.add_trace(box_mesh_global(ns_x - 0.05, ns_y - 0.05, 0.95,
        ns_w + 0.1, ns_d + 0.1, 0.05,
        '#cfd8dc', 1.0, 'Counter Top'))
    # Computer monitors on counter
    for cm_off in [-0.5, 0, 0.5]:
        fig.add_trace(box_mesh_global(ns_x + ns_w/2 + cm_off - 0.15,
            ns_y + ns_d - 0.1, 1.0, 0.3, 0.04, 0.25,
            '#212121', 1.0, 'Monitor'))
        fig.add_trace(box_mesh_global(ns_x + ns_w/2 + cm_off - 0.13,
            ns_y + ns_d - 0.12, 1.03, 0.26, 0.01, 0.19,
            '#4fc3f7', 1.0, 'Screen'))
    # 2 office chairs
    for cx_off in [-0.5, 0.5]:
        fig.add_trace(cylinder_mesh(ns_x + ns_w/2 + cx_off,
            ns_y - 0.3, 0.45, 0.2, 0.05, 'z', 12,
            color, 1.0, 'Office Chair'))

    # ═══ GLASS PARTITIONS between beds (3D transparent walls) ═══
    if glass_partitions:
        bed_l = 2.0
        glass_color = '#81d4fa'
        glass_opacity = 0.3
        glass_thick = 0.05  # very thin wall

        # Top wall partitions
        if top_beds > 1:
            x_step = (width - 1) / top_beds
            for i in range(top_beds - 1):
                bx1 = 0.5 + x_step * (i + 0.5)
                bx2 = 0.5 + x_step * (i + 1.5)
                part_x = (bx1 + bx2) / 2
                # Glass wall: from top wall extending inward, full height
                fig.add_trace(box_mesh_global(
                    part_x - glass_thick/2, height - bed_l - 0.3, 0,
                    glass_thick, bed_l + 0.1, wall_height * 0.85,
                    glass_color, glass_opacity,
                    f'Glass Partition Top {i+1}'))
                # Add a metal frame around the edges for realism
                # Top edge
                fig.add_trace(box_mesh_global(
                    part_x - glass_thick*1.5, height - bed_l - 0.3,
                    wall_height * 0.85 - 0.05,
                    glass_thick*3, bed_l + 0.1, 0.05,
                    '#cfd8dc', 1.0, 'Glass Frame Top'))
        # Bottom wall partitions
        if bottom_beds > 1:
            x_step = (width - 1) / bottom_beds
            for i in range(bottom_beds - 1):
                bx1 = 0.5 + x_step * (i + 0.5)
                bx2 = 0.5 + x_step * (i + 1.5)
                part_x = (bx1 + bx2) / 2
                fig.add_trace(box_mesh_global(
                    part_x - glass_thick/2, 0.2, 0,
                    glass_thick, bed_l + 0.1, wall_height * 0.85,
                    glass_color, glass_opacity,
                    f'Glass Partition Bot {i+1}'))
                fig.add_trace(box_mesh_global(
                    part_x - glass_thick*1.5, 0.2,
                    wall_height * 0.85 - 0.05,
                    glass_thick*3, bed_l + 0.1, 0.05,
                    '#cfd8dc', 1.0, 'Glass Frame Bot'))


# ────────────────────────────────────────────────────────────
# 💊 PHARMACY (2D) - Dispensing counter + shelves + waiting
# ────────────────────────────────────────────────────────────
def layout_pharmacy_2d(ax, width, height, color, n_counters=3):
    """Pharmacy with dispensing window, rear shelves, waiting bench.
    n_counters drives the number of dispensing windows and shelving units."""
    # Clamp to reasonable range
    n_counters = max(1, min(n_counters, 6))

    # ─── DISPENSING COUNTER (back wall) ───
    counter_y = height - 1.3
    ax.add_patch(_MFBox((width*0.1, counter_y), width*0.8, 0.5,
        boxstyle="round,pad=0.05", facecolor='#a5d6a7',
        edgecolor=color, linewidth=2, zorder=6))
    # Dispensing windows (n_counters cutouts)
    win_total_w = width * 0.7
    win_w = win_total_w / n_counters * 0.6
    win_gap = win_total_w / n_counters * 0.4
    for i in range(n_counters):
        wx = width*0.15 + i * (win_w + win_gap)
        ax.add_patch(_MRect((wx, counter_y + 0.2), win_w, 0.25,
            facecolor='#fff', edgecolor=color, linewidth=1, zorder=7))
        ax.text(wx + win_w/2, counter_y + 0.32, f"#{i+1}",
            color=color, fontsize=6, ha='center', weight='bold', zorder=8)
    ax.text(width/2, counter_y + 0.6,
        f"DISPENSING COUNTER ({n_counters} windows)",
        color=color, fontsize=8, ha='center', weight='bold',
        style='italic', zorder=7)

    # ─── DRUG STORAGE SHELVES (scales with counters) ───
    shelf_y = height - 0.5
    n_shelves = max(3, n_counters + 1)
    shelf_total_w = width * 0.8
    shelf_w = shelf_total_w / n_shelves - 0.1
    for i in range(n_shelves):
        sx = width*0.1 + i * (shelf_w + 0.1) + 0.05
        # Shelf unit (tall narrow rectangle)
        ax.add_patch(_MRect((sx, shelf_y - 0.05), shelf_w, 0.15,
            facecolor=color, alpha=0.6, edgecolor='black',
            linewidth=0.5, zorder=5))
        # Pill icon
        ax.text(sx + shelf_w/2, shelf_y + 0.02, "💊" * 3,
            fontsize=5, ha='center', va='center', zorder=6)

    # ─── REFRIGERATED SECTION (cold storage, left side) ───
    ax.add_patch(_MFBox((width*0.05, height*0.4), 0.6, 1.5,
        boxstyle="round,pad=0.03", facecolor='#bbdefb',
        edgecolor='#0277bd', linewidth=2, zorder=6))
    ax.text(width*0.05 + 0.3, height*0.4 + 0.75, "❄️\nCold\nStorage",
        color='#0277bd', fontsize=5, ha='center', va='center',
        weight='bold', zorder=7)

    # ─── COMPOUNDING AREA (right side, scales with size) ───
    compounding_w = 1.0 if n_counters <= 3 else 1.4
    ax.add_patch(_MFBox((width - compounding_w - 0.22, height*0.4),
        compounding_w, 1.5,
        boxstyle="round,pad=0.03", facecolor='#fff9c4',
        edgecolor='#f57f17', linewidth=2, zorder=6))
    ax.text(width - compounding_w/2 - 0.22, height*0.4 + 0.75,
        "🧪\nCompounding\nArea",
        color='#e65100', fontsize=5, ha='center', va='center',
        weight='bold', zorder=7)

    # ─── WAITING BENCH (longer for more counters) ───
    n_bench_seats = max(4, n_counters * 2)
    bench_y = 0.4
    bench_total_w = width * 0.6
    seat_w = bench_total_w / n_bench_seats * 0.85
    for i in range(n_bench_seats):
        bx = width*0.2 + i * (bench_total_w / n_bench_seats)
        ax.add_patch(_MRect((bx, bench_y), seat_w, 0.3,
            facecolor=color, alpha=0.7, edgecolor='black',
            linewidth=0.3, zorder=5))
    ax.text(width/2, bench_y - 0.2,
        f"Patient Waiting ({n_bench_seats} seats)",
        color='#5d4037', fontsize=6, ha='center',
        weight='bold', style='italic', zorder=5)

    # ─── CONSULTATION CUBICLE ───
    cube_x = width*0.05
    cube_y = 0.4
    ax.add_patch(_MFBox((cube_x, cube_y), 1.2, 1.2,
        boxstyle="round,pad=0.03", facecolor='none',
        edgecolor=color, linewidth=1.5, linestyle='--', zorder=5))
    ax.text(cube_x + 0.6, cube_y + 0.6, "Private\nConsult",
        color=color, fontsize=5, ha='center', va='center',
        weight='bold', zorder=6)


# ────────────────────────────────────────────────────────────
# 💊 PHARMACY (3D)
# ────────────────────────────────────────────────────────────
def layout_pharmacy_3d(fig, width, height, color, wall_height, n_counters=3):
    """Pharmacy 3D with dispensing counter + shelves.
    n_counters drives dispensing windows + shelf units."""
    n_counters = max(1, min(n_counters, 6))

    # Main counter (long, low)
    counter_y = height - 1.3
    fig.add_trace(box_mesh_global(width*0.1, counter_y, 0,
        width*0.8, 0.5, 1.1, '#a5d6a7', 1.0, 'Dispensing Counter'))
    # Counter top
    fig.add_trace(box_mesh_global(width*0.1, counter_y, 1.1,
        width*0.8, 0.5, 0.05, color, 1.0, 'Counter Top'))

    # Dispensing windows
    win_total_w = width * 0.7
    win_w = win_total_w / n_counters * 0.6
    win_gap = win_total_w / n_counters * 0.4
    for i in range(n_counters):
        wx = width*0.15 + i * (win_w + win_gap)
        fig.add_trace(box_mesh_global(wx, counter_y - 0.02, 1.1,
            win_w, 0.04, 1.0, '#cfd8dc', 0.8, f'Window Frame #{i+1}'))
        fig.add_trace(box_mesh_global(wx + 0.05, counter_y - 0.04, 1.3,
            win_w - 0.1, 0.02, 0.6, '#4fc3f7', 0.4, f'Window Glass #{i+1}'))

    # Tall shelving units behind counter (n_shelves scales)
    n_shelves = max(3, n_counters + 1)
    shelf_total_w = width * 0.8
    shelf_w = shelf_total_w / n_shelves - 0.1
    for i in range(n_shelves):
        sx = width*0.1 + i * (shelf_w + 0.1) + 0.05
        # Shelf body
        fig.add_trace(box_mesh_global(sx, height - 0.4, 0,
            shelf_w, 0.3, wall_height - 0.3, color, 0.95,
            f'Drug Shelf {i+1}'))
        # Shelf dividers
        for sz in [1.0, 1.6, 2.2]:
            fig.add_trace(box_mesh_global(sx + 0.02, height - 0.42, sz,
                shelf_w - 0.04, 0.34, 0.03, '#2e7d32', 1.0, 'Shelf'))

    # Cold storage refrigerator
    fig.add_trace(box_mesh_global(width*0.05, height*0.4, 0,
        0.6, 1.5, 1.8, '#bbdefb', 1.0, 'Cold Storage'))
    fig.add_trace(box_mesh_global(width*0.05 - 0.02, height*0.4 + 0.1, 0.2,
        0.04, 1.3, 1.4, '#01579b', 1.0, 'Fridge Door'))

    # Compounding bench (scales)
    compounding_w = 1.0 if n_counters <= 3 else 1.4
    fig.add_trace(box_mesh_global(width - compounding_w - 0.22, height*0.4, 0,
        compounding_w, 1.5, 0.9, '#fff9c4', 1.0, 'Compounding Bench'))
    fig.add_trace(cylinder_mesh(width - compounding_w/2 - 0.22,
        height*0.4 + 0.5, 0.9, 0.1, 0.4, 'z', 12,
        '#37474f', 1.0, 'Mortar Mixer'))

    # Waiting bench (dynamic seats)
    n_bench_seats = max(4, n_counters * 2)
    bench_total_w = width * 0.6
    seat_w = bench_total_w / n_bench_seats * 0.85
    for i in range(n_bench_seats):
        bx = width*0.2 + i * (bench_total_w / n_bench_seats)
        # Seat
        fig.add_trace(box_mesh_global(bx, 0.4, 0.45,
            seat_w, 0.4, 0.05, color, 1.0, f'Waiting Seat {i+1}'))
        # Back
        fig.add_trace(box_mesh_global(bx, 0.75, 0.5,
            seat_w, 0.05, 0.5, color, 1.0, 'Back'))
        # Legs
        for lx in [0.05, seat_w - 0.1]:
            fig.add_trace(box_mesh_global(bx + lx, 0.45, 0,
                0.04, 0.04, 0.45, '#37474f', 1.0, 'Leg'))


# ────────────────────────────────────────────────────────────
# 🪑 RECEPTION & WAITING (2D) - Info desk + seating arrangement
# ────────────────────────────────────────────────────────────
def layout_reception_2d(ax, width, height, color, n_seats=30):
    """Reception with info desk, staff workstation, organized seating.
    n_seats drives the waiting area size + staff count."""
    n_seats = max(4, min(n_seats, 60))

    # ─── INFO DESK (curved, near entrance) ───
    # Larger desk for more seats (more staff)
    n_staff = 2 if n_seats <= 20 else 3 if n_seats <= 40 else 4
    desk_w = min(max(2.5, n_staff * 1.1), width*0.6)
    desk_x = width/2 - desk_w/2
    desk_y = 0.5
    ax.add_patch(_MFBox((desk_x, desk_y), desk_w, 0.8,
        boxstyle="round,pad=0.08", facecolor='#a1887b',
        edgecolor='#5d4037', linewidth=2, zorder=6))
    # Counter top
    ax.add_patch(_MFBox((desk_x + 0.05, desk_y + 0.05),
        desk_w - 0.1, 0.7, boxstyle="round,pad=0.04",
        facecolor='#d7ccc8', edgecolor='none', zorder=7))
    ax.text(desk_x + desk_w/2, desk_y + 0.4,
        f"ℹ️ INFORMATION ({n_staff} staff)",
        color='#5d4037', fontsize=7, ha='center', va='center',
        weight='bold', zorder=8)

    # Staff workstations behind desk
    for i in range(n_staff):
        wx = desk_x + 0.4 + i * (desk_w - 0.8) / max(1, n_staff - 1) if n_staff > 1 else desk_x + desk_w/2
        ax.add_patch(_MCircle((wx, desk_y - 0.3), 0.18,
            facecolor=color, edgecolor='black', linewidth=0.3, zorder=6))
        # Computer monitor on desk
        ax.add_patch(_MRect((wx - 0.15, desk_y + 0.55), 0.3, 0.05,
            facecolor='#1a237e', zorder=7))

    # ─── WAITING SEATING ARRANGEMENT (dynamic rows × cols) ───
    # Calculate how many rows and columns from n_seats
    chairs_per_row = max(4, int(width / 0.6))
    rows_needed = max(1, (n_seats + chairs_per_row - 1) // chairs_per_row)
    rows = min(rows_needed, max(1, int((height - 2.5) / 1.0)))  # don't overflow
    seat_w = 0.5
    seat_gap = 0.1

    seats_drawn = 0
    for r in range(rows):
        if seats_drawn >= n_seats:
            break
        row_y = height - 1.5 - r * 1.0
        # Fit as many seats as possible in this row
        seats_in_row = min(chairs_per_row, n_seats - seats_drawn)
        total_w = seats_in_row * (seat_w + seat_gap)
        start_x = (width - total_w) / 2
        for c in range(seats_in_row):
            cx = start_x + c * (seat_w + seat_gap)
            # Chair seat
            ax.add_patch(_MFBox((cx, row_y), seat_w, 0.5,
                boxstyle="round,pad=0.03", facecolor=color,
                edgecolor='black', linewidth=0.3, zorder=5))
            # Backrest line
            ax.add_patch(_MRect((cx, row_y + 0.45), seat_w, 0.08,
                facecolor='#37474f', zorder=6))
            seats_drawn += 1

    ax.text(width/2, height - 0.8 + 0.6,
        f"WAITING AREA ({seats_drawn} seats / {rows} rows)",
        color=color, fontsize=7, ha='center', weight='bold',
        style='italic', zorder=7)

    # ─── WALL-MOUNTED TV (back wall, possibly multiple for big lobbies) ───
    n_tvs = 1 if n_seats <= 30 else 2
    for ti in range(n_tvs):
        tv_x = width/2 - 0.6 + (ti - (n_tvs-1)/2) * 2.0 if n_tvs > 1 else width/2 - 0.6
        ax.add_patch(_MRect((tv_x, height - 0.2),
            1.2, 0.15, facecolor='#212121', edgecolor='#37474f',
            linewidth=1, zorder=7))
        ax.add_patch(_MRect((tv_x + 0.05, height - 0.18),
            1.1, 0.11, facecolor='#1976d2', alpha=0.7, zorder=8))

    # ─── INFO KIOSKS (1-2 depending on size) ───
    n_kiosks = 1 if n_seats <= 30 else 2
    for ki in range(n_kiosks):
        kx = width*0.05 if ki == 0 else width*0.92
        ax.add_patch(_MFBox((kx, 0.5), 0.4, 0.8,
            boxstyle="round,pad=0.03", facecolor='#1a237e',
            edgecolor='black', linewidth=1, zorder=6))
        ax.add_patch(_MRect((kx + 0.05, 0.7), 0.3, 0.5,
            facecolor='#4fc3f7', zorder=7))
        ax.text(kx + 0.2, 0.3, "🖥️ Kiosk",
            color='#1a237e', fontsize=5, ha='center',
            weight='bold', zorder=6)


# ────────────────────────────────────────────────────────────
# 🪑 RECEPTION & WAITING (3D)
# ────────────────────────────────────────────────────────────
def layout_reception_3d(fig, width, height, color, wall_height, n_seats=30):
    """Reception 3D with info desk + seating + TV. Scales with n_seats."""
    n_seats = max(4, min(n_seats, 60))
    n_staff = 2 if n_seats <= 20 else 3 if n_seats <= 40 else 4
    desk_w = min(max(2.5, n_staff * 1.1), width*0.6)
    desk_x = width/2 - desk_w/2
    desk_y = 0.5
    fig.add_trace(box_mesh_global(desk_x, desk_y, 0,
        desk_w, 0.8, 1.0, '#a1887b', 1.0, 'Information Desk'))
    # Counter top
    fig.add_trace(box_mesh_global(desk_x - 0.05, desk_y - 0.05, 1.0,
        desk_w + 0.1, 0.9, 0.05, '#d7ccc8', 1.0, 'Counter Top'))

    # Staff seats + monitors
    for i in range(n_staff):
        wx = desk_x + 0.4 + i * (desk_w - 0.8) / max(1, n_staff - 1) if n_staff > 1 else desk_x + desk_w/2
        fig.add_trace(cylinder_mesh(wx, desk_y - 0.4, 0.45, 0.2, 0.05,
            'z', 12, color, 1.0, f'Staff Chair {i+1}'))
        fig.add_trace(box_mesh_global(wx - 0.15, desk_y + 0.6, 1.05,
            0.3, 0.04, 0.3, '#212121', 1.0, f'Monitor {i+1}'))
        fig.add_trace(box_mesh_global(wx - 0.13, desk_y + 0.58, 1.08,
            0.26, 0.02, 0.22, '#4fc3f7', 1.0, 'Screen'))

    # Rows of waiting chairs (dynamic)
    chairs_per_row = max(4, int(width / 0.7))
    rows_needed = max(1, (n_seats + chairs_per_row - 1) // chairs_per_row)
    rows = min(rows_needed, max(1, int((height - 2.5) / 1.0)))
    seat_w, seat_d, seat_gap = 0.5, 0.5, 0.1
    seats_drawn = 0
    for r in range(rows):
        if seats_drawn >= n_seats:
            break
        row_y = height - 1.5 - r * 1.0
        seats_in_row = min(chairs_per_row, n_seats - seats_drawn)
        total_w = seats_in_row * (seat_w + seat_gap)
        start_x = (width - total_w) / 2
        for c in range(seats_in_row):
            cx = start_x + c * (seat_w + seat_gap)
            # Seat
            fig.add_trace(box_mesh_global(cx, row_y, 0.45,
                seat_w, seat_d, 0.05, color, 1.0, f'Seat'))
            # Backrest
            fig.add_trace(box_mesh_global(cx, row_y + seat_d - 0.05, 0.5,
                seat_w, 0.05, 0.5, color, 1.0, 'Back'))
            # 4 legs
            for lx, ly in [(0.05, 0.05), (seat_w-0.09, 0.05),
                            (0.05, seat_d-0.09), (seat_w-0.09, seat_d-0.09)]:
                fig.add_trace(box_mesh_global(cx + lx, row_y + ly, 0,
                    0.04, 0.04, 0.45, '#37474f', 1.0, 'Leg'))
            seats_drawn += 1

    # Wall-mounted TVs (1 or 2)
    n_tvs = 1 if n_seats <= 30 else 2
    for ti in range(n_tvs):
        tv_x = width/2 - 0.6 + (ti - (n_tvs-1)/2) * 2.0 if n_tvs > 1 else width/2 - 0.6
        fig.add_trace(box_mesh_global(tv_x, height - 0.1, 2.0,
            1.2, 0.05, 0.7, '#212121', 1.0, 'TV Frame'))
        fig.add_trace(box_mesh_global(tv_x + 0.05, height - 0.12, 2.05,
            1.1, 0.02, 0.6, '#1976d2', 0.9, 'TV Screen'))

    # Info kiosks (1-2)
    n_kiosks = 1 if n_seats <= 30 else 2
    for ki in range(n_kiosks):
        kx = width*0.05 if ki == 0 else width*0.92
        fig.add_trace(box_mesh_global(kx, 0.5, 0,
            0.4, 0.4, 1.4, '#1a237e', 1.0, 'Info Kiosk'))
        fig.add_trace(box_mesh_global(kx + 0.02, 0.48, 0.5,
            0.36, 0.02, 0.7, '#4fc3f7', 1.0, 'Touchscreen'))


# ════════════════════════════════════════════════════════════════
# EXTENDED EQUIPMENT LIBRARY (14 pieces)
# Each piece is a smart placement: tries different positions based
# on what's needed, with realistic 2D + 3D rendering.
# ════════════════════════════════════════════════════════════════

EXTENDED_EQUIPMENT = [
    {'key': 'eq_defib',     'icon': '⚡', 'label_en': 'Defibrillator',
     'label_ar': 'مزيل الرجفان', 'power_w': 50,
     'default_for': ['Operating Room (OR)', 'Intensive Care Unit (ICU)',
                     'Emergency Room (ER)', 'NICU (Neonatal ICU)']},
    {'key': 'eq_ventilator','icon': '💨', 'label_en': 'Ventilator',
     'label_ar': 'جهاز التنفس الصناعي', 'power_w': 250,
     'default_for': ['Intensive Care Unit (ICU)', 'Operating Room (OR)',
                     'Emergency Room (ER)']},
    {'key': 'eq_ultrasound','icon': '🩻', 'label_en': 'Ultrasound Machine',
     'label_ar': 'جهاز السونار', 'power_w': 400,
     'default_for': ['Radiology Department', 'Maternity & Delivery',
                     'Emergency Room (ER)']},
    {'key': 'eq_crashcart', 'icon': '🚑', 'label_en': 'Crash Cart',
     'label_ar': 'عربة الإنقاذ', 'power_w': 30,
     'default_for': ['Operating Room (OR)', 'Intensive Care Unit (ICU)',
                     'Emergency Room (ER)']},
    {'key': 'eq_surgical_light','icon': '💡','label_en': 'Surgical Light',
     'label_ar': 'إضاءة الجراحة', 'power_w': 200,
     'default_for': ['Operating Room (OR)']},
    {'key': 'eq_warmer',    'icon': '🌡️', 'label_en': 'Patient Warmer',
     'label_ar': 'مدفأة المرضى', 'power_w': 800,
     'default_for': ['Operating Room (OR)', 'NICU (Neonatal ICU)']},
    {'key': 'eq_infusion',  'icon': '💧', 'label_en': 'Infusion Pump',
     'label_ar': 'مضخة الحقن', 'power_w': 25,
     'default_for': ['Intensive Care Unit (ICU)', 'Oncology']},
    {'key': 'eq_endoscopy', 'icon': '📺', 'label_en': 'Endoscopy Tower',
     'label_ar': 'برج المنظار', 'power_w': 600,
     'default_for': ['Operating Room (OR)']},
    {'key': 'eq_ecg',       'icon': '〽️', 'label_en': 'ECG Machine',
     'label_ar': 'جهاز رسم القلب', 'power_w': 80,
     'default_for': ['Intensive Care Unit (ICU)', 'Emergency Room (ER)',
                     'Outpatient Clinic']},
    {'key': 'eq_suction',   'icon': '🌀', 'label_en': 'Suction Machine',
     'label_ar': 'جهاز الشفط', 'power_w': 150,
     'default_for': ['Operating Room (OR)', 'Intensive Care Unit (ICU)']},
    {'key': 'eq_xray',      'icon': '☢️', 'label_en': 'Mobile X-Ray',
     'label_ar': 'الأشعة المتحركة', 'power_w': 1200,
     'default_for': ['Intensive Care Unit (ICU)', 'Emergency Room (ER)']},
    {'key': 'eq_music',     'icon': '🎵', 'label_en': 'Music Therapy System',
     'label_ar': 'نظام العلاج بالموسيقى', 'power_w': 20,
     'default_for': ['NICU (Neonatal ICU)', 'Pediatric Ward']},
    {'key': 'eq_bairhugger','icon': '🌬️', 'label_en': 'Bair Hugger (warming)',
     'label_ar': 'جهاز التدفئة الجراحية', 'power_w': 700,
     'default_for': ['Operating Room (OR)']},
    {'key': 'eq_comm_board','icon': '📋', 'label_en': 'Patient Communication Board',
     'label_ar': 'لوحة التواصل مع المرضى', 'power_w': 5,
     'default_for': ['Intensive Care Unit (ICU)', 'General Patient Ward']},
    {'key': 'eq_fire_ext','icon': '🧯', 'label_en': 'Fire Extinguisher',
     'label_ar': 'طفاية حريق', 'power_w': 0,
     # Required by NFPA 10 in virtually every facility space.
     # '_all_' sentinel = default ON for every department.
     'default_for': '_all_'},
]


# Equipment applicability: defines if equipment is even RELEVANT (selectable)
# for a given facility, beyond just the "default_for" auto-select.
# A piece is applicable if the facility is similar to defaults.
_EQ_APPLICABLE_GROUPS = {
    # Surgical equipment — relevant for surgical/critical zones
    'eq_surgical_light': {'Operating Room (OR)', 'Maternity & Delivery'},
    'eq_anesthesia':     {'Operating Room (OR)', 'Maternity & Delivery'},
    'eq_endoscopy':      {'Operating Room (OR)', 'Outpatient Clinic'},
    'eq_bairhugger':     {'Operating Room (OR)', 'NICU (Neonatal ICU)',
                          'Maternity & Delivery'},
    # Critical care / monitoring — wide clinical applicability
    'eq_defib':          {'Operating Room (OR)', 'Intensive Care Unit (ICU)',
                          'Emergency Room (ER)', 'NICU (Neonatal ICU)',
                          'Maternity & Delivery', 'Dialysis Unit',
                          'Outpatient Clinic', 'Dental Clinic',
                          'Pediatric Ward', 'Oncology'},
    'eq_ventilator':     {'Operating Room (OR)', 'Intensive Care Unit (ICU)',
                          'Emergency Room (ER)', 'NICU (Neonatal ICU)',
                          'Maternity & Delivery'},
    'eq_crashcart':      {'Operating Room (OR)', 'Intensive Care Unit (ICU)',
                          'Emergency Room (ER)', 'NICU (Neonatal ICU)',
                          'Maternity & Delivery', 'Dialysis Unit',
                          'Pediatric Ward', 'General Patient Ward',
                          'Oncology', 'Outpatient Clinic'},
    'eq_ecg':            {'Operating Room (OR)', 'Intensive Care Unit (ICU)',
                          'Emergency Room (ER)', 'NICU (Neonatal ICU)',
                          'Outpatient Clinic', 'General Patient Ward',
                          'Pediatric Ward', 'Maternity & Delivery'},
    'eq_suction':        {'Operating Room (OR)', 'Intensive Care Unit (ICU)',
                          'Emergency Room (ER)', 'NICU (Neonatal ICU)',
                          'Maternity & Delivery', 'Dental Clinic'},
    'eq_warmer':         {'Operating Room (OR)', 'NICU (Neonatal ICU)',
                          'Maternity & Delivery'},
    'eq_infusion':       {'Operating Room (OR)', 'Intensive Care Unit (ICU)',
                          'Emergency Room (ER)', 'NICU (Neonatal ICU)',
                          'Oncology', 'General Patient Ward',
                          'Pediatric Ward', 'Dialysis Unit',
                          'Maternity & Delivery'},
    # Imaging
    'eq_ultrasound':     {'Operating Room (OR)', 'Intensive Care Unit (ICU)',
                          'Emergency Room (ER)', 'Radiology Department',
                          'Maternity & Delivery', 'Outpatient Clinic'},
    'eq_xray':           {'Intensive Care Unit (ICU)', 'Emergency Room (ER)',
                          'Radiology Department', 'NICU (Neonatal ICU)'},
    # Patient-comfort
    'eq_music':          {'NICU (Neonatal ICU)', 'Pediatric Ward',
                          'Maternity & Delivery', 'Psychiatric Ward',
                          'Oncology'},
    'eq_comm_board':     {'Intensive Care Unit (ICU)', 'General Patient Ward',
                          'NICU (Neonatal ICU)', 'Pediatric Ward',
                          'Maternity & Delivery', 'Oncology'},
    # Fire extinguishers — universal applicability (NFPA 10 in every space)
    'eq_fire_ext':       '_all_',
}


def _is_eq_applicable(facility_type, eq):
    """Return True if equipment is relevant (selectable) for this facility."""
    # Universal items (e.g. fire extinguishers) — always applicable
    if eq.get('default_for') == '_all_':
        return True
    allowed = _EQ_APPLICABLE_GROUPS.get(eq['key'])
    if allowed == '_all_':
        return True
    if allowed is None:
        # If no rule defined, fall back to "default_for" only
        return facility_type in eq.get('default_for', [])
    return facility_type in allowed


def _is_eq_default_on(facility_type, eq):
    """Return True if the equipment should default to ON for this facility."""
    if eq.get('default_for') == '_all_':
        return True
    return facility_type in eq.get('default_for', [])


# ────────────────────────────────────────────────────────────
# 2D drawing for each equipment piece — small footprint icons
# placed at supplied (x, y) position
# ────────────────────────────────────────────────────────────
def draw_equipment_2d(ax, x, y, eq_key, color):
    """Draw a small 2D representation of equipment at (x, y)."""
    eq = next((e for e in EXTENDED_EQUIPMENT if e['key'] == eq_key), None)
    if not eq: return

    if eq_key == 'eq_defib':
        # Red box with lightning bolt
        ax.add_patch(_MFBox((x-0.2, y-0.15), 0.4, 0.3,
            boxstyle="round,pad=0.02", facecolor='#e53935',
            edgecolor='black', linewidth=0.8, zorder=7))
        ax.text(x, y, "⚡", fontsize=10, ha='center', va='center', zorder=8)
    elif eq_key == 'eq_ventilator':
        # Tall blue box on wheels
        ax.add_patch(_MFBox((x-0.18, y-0.25), 0.36, 0.5,
            boxstyle="round,pad=0.02", facecolor='#1976d2',
            edgecolor='black', linewidth=0.8, zorder=7))
        ax.add_patch(_MCircle((x-0.12, y-0.28), 0.04,
            facecolor='black', zorder=8))
        ax.add_patch(_MCircle((x+0.12, y-0.28), 0.04,
            facecolor='black', zorder=8))
        ax.text(x, y, "VENT", fontsize=5, color='white',
            ha='center', va='center', weight='bold', zorder=9)
    elif eq_key == 'eq_ultrasound':
        # Gray cart with monitor
        ax.add_patch(_MFBox((x-0.22, y-0.18), 0.44, 0.36,
            boxstyle="round,pad=0.02", facecolor='#546e7a',
            edgecolor='black', linewidth=0.8, zorder=7))
        ax.add_patch(_MRect((x-0.15, y+0.05), 0.3, 0.12,
            facecolor='#212121', zorder=8))
        ax.text(x, y-0.05, "ECHO", fontsize=4, color='#b0bec5',
            ha='center', va='center', weight='bold', zorder=9)
    elif eq_key == 'eq_crashcart':
        # Red rolling cart with drawers
        ax.add_patch(_MFBox((x-0.2, y-0.15), 0.4, 0.3,
            boxstyle="round,pad=0.02", facecolor='#d32f2f',
            edgecolor='black', linewidth=0.8, zorder=7))
        for dy in [-0.08, 0.0, 0.08]:
            ax.plot([x-0.18, x+0.18], [y+dy, y+dy],
                color='white', linewidth=0.5, zorder=8)
        ax.text(x, y+0.18, "Crash", fontsize=4, color='#d32f2f',
            ha='center', weight='bold', zorder=8)
    elif eq_key == 'eq_surgical_light':
        # Big yellow circle (ceiling-mounted from above)
        ax.add_patch(_MCircle((x, y), 0.25, facecolor='#fff59d',
            edgecolor='#f57f17', linewidth=1.5, zorder=8, alpha=0.85))
        # Inner LEDs
        for ang_l in [0, np.pi/2, np.pi, 3*np.pi/2]:
            ax.add_patch(_MCircle(
                (x + 0.12*np.cos(ang_l), y + 0.12*np.sin(ang_l)),
                0.03, facecolor='#fff', edgecolor='#f57f17',
                linewidth=0.5, zorder=9))
        ax.text(x, y-0.35, "💡 Surgical Light",
            color='#f57f17', fontsize=4, ha='center', weight='bold', zorder=8)
    elif eq_key == 'eq_warmer':
        # Orange rectangle (radiant warmer)
        ax.add_patch(_MFBox((x-0.2, y-0.12), 0.4, 0.24,
            boxstyle="round,pad=0.02", facecolor='#ff9800',
            edgecolor='black', linewidth=0.8, zorder=7))
        ax.text(x, y, "🌡️", fontsize=8, ha='center', va='center', zorder=8)
    elif eq_key == 'eq_infusion':
        # Small blue box with IV symbol
        ax.add_patch(_MFBox((x-0.1, y-0.12), 0.2, 0.24,
            boxstyle="round,pad=0.02", facecolor='#42a5f5',
            edgecolor='black', linewidth=0.8, zorder=7))
        ax.text(x, y, "💧", fontsize=7, ha='center', va='center', zorder=8)
    elif eq_key == 'eq_endoscopy':
        # Large dark cart with stacked monitors
        ax.add_patch(_MFBox((x-0.2, y-0.25), 0.4, 0.5,
            boxstyle="round,pad=0.02", facecolor='#37474f',
            edgecolor='black', linewidth=0.8, zorder=7))
        for dy in [0.1, -0.05]:
            ax.add_patch(_MRect((x-0.15, y+dy), 0.3, 0.08,
                facecolor='#0277bd', zorder=8))
        ax.text(x, y-0.18, "Endo", fontsize=4, color='white',
            ha='center', weight='bold', zorder=9)
    elif eq_key == 'eq_ecg':
        # Cream-colored compact unit
        ax.add_patch(_MFBox((x-0.18, y-0.12), 0.36, 0.24,
            boxstyle="round,pad=0.02", facecolor='#fff8e1',
            edgecolor='#5d4037', linewidth=0.8, zorder=7))
        # ECG waveform line
        wave_x = np.linspace(x-0.15, x+0.15, 30)
        wave_y = y + 0.04 * np.sin(wave_x * 50) * np.exp(-((wave_x-x)*8)**2)
        ax.plot(wave_x, wave_y, color='#43a047', linewidth=0.8, zorder=8)
        ax.text(x, y-0.18, "ECG", fontsize=4, color='#5d4037',
            ha='center', weight='bold', zorder=8)
    elif eq_key == 'eq_suction':
        # Small yellow canister
        ax.add_patch(_MFBox((x-0.1, y-0.15), 0.2, 0.3,
            boxstyle="round,pad=0.02", facecolor='#fdd835',
            edgecolor='black', linewidth=0.8, zorder=7))
        ax.text(x, y, "🌀", fontsize=8, ha='center', va='center', zorder=8)
    elif eq_key == 'eq_xray':
        # C-arm shape (gray with arc)
        ax.add_patch(_MFBox((x-0.25, y-0.1), 0.5, 0.2,
            boxstyle="round,pad=0.02", facecolor='#90a4ae',
            edgecolor='black', linewidth=0.8, zorder=7))
        ax.text(x, y, "X", fontsize=10, color='white',
            ha='center', va='center', weight='bold', zorder=8)
        ax.text(x, y-0.18, "Mobile X-Ray", fontsize=4, color='#37474f',
            ha='center', weight='bold', zorder=8)
    elif eq_key == 'eq_music':
        # Pink small speaker
        ax.add_patch(_MCircle((x, y), 0.12, facecolor='#f48fb1',
            edgecolor='black', linewidth=0.8, zorder=7))
        ax.text(x, y, "🎵", fontsize=8, ha='center', va='center', zorder=8)
    elif eq_key == 'eq_bairhugger':
        # Light blue square (forced air warming unit)
        ax.add_patch(_MFBox((x-0.15, y-0.12), 0.3, 0.24,
            boxstyle="round,pad=0.02", facecolor='#80deea',
            edgecolor='black', linewidth=0.8, zorder=7))
        ax.text(x, y, "🌬️", fontsize=7, ha='center', va='center', zorder=8)
    elif eq_key == 'eq_comm_board':
        # White board on wall
        ax.add_patch(_MFBox((x-0.25, y-0.05), 0.5, 0.1,
            boxstyle="round,pad=0.01", facecolor='white',
            edgecolor='#37474f', linewidth=0.8, zorder=7))
        ax.text(x, y, "📋 PT-101 | Dr. Smith", fontsize=4,
            color='#37474f', ha='center', va='center', weight='bold', zorder=8)
    elif eq_key == 'eq_fire_ext':
        # Small red fire extinguisher (cylinder + black neck)
        # Body
        ax.add_patch(_MCircle((x, y), 0.13, facecolor='#c62828',
            edgecolor='black', linewidth=0.8, zorder=8))
        # Black neck/valve at top
        ax.add_patch(_MRect((x-0.04, y+0.10), 0.08, 0.06,
            facecolor='#212121', edgecolor='black', linewidth=0.5, zorder=9))
        # White label band
        ax.add_patch(_MRect((x-0.10, y-0.03), 0.20, 0.05,
            facecolor='white', edgecolor='#c62828', linewidth=0.5, zorder=9))
        ax.text(x, y, "🧯", fontsize=7, ha='center', va='center', zorder=10)
        # Floor mounting bracket indicator
        ax.text(x, y-0.22, "FE", color='#c62828', fontsize=4,
            ha='center', weight='bold', zorder=8)


# ────────────────────────────────────────────────────────────
# 3D drawing for each equipment piece
# ────────────────────────────────────────────────────────────
def draw_equipment_3d(fig, x, y, eq_key, wall_height=3.0):
    """Draw a 3D representation of equipment at (x, y) with floor at z=0."""
    eq = next((e for e in EXTENDED_EQUIPMENT if e['key'] == eq_key), None)
    if not eq: return

    if eq_key == 'eq_defib':
        # Red wall-mounted box
        fig.add_trace(box_mesh_global(x-0.2, y, 1.0,
            0.4, 0.15, 0.3, '#e53935', 1.0, 'Defibrillator'))
        # Lightning bolt sticker
        fig.add_trace(box_mesh_global(x-0.05, y-0.02, 1.1,
            0.1, 0.01, 0.15, '#fff59d', 1.0, 'Symbol'))
    elif eq_key == 'eq_ventilator':
        # Tall blue cart
        fig.add_trace(box_mesh_global(x-0.18, y-0.15, 0,
            0.36, 0.3, 1.2, '#1976d2', 1.0, 'Ventilator'))
        # Monitor on top
        fig.add_trace(box_mesh_global(x-0.15, y-0.13, 1.2,
            0.3, 0.04, 0.25, '#212121', 1.0, 'Vent Screen'))
        # Wheels (4 small)
        for wx, wy in [(-0.15, -0.13), (0.11, -0.13), (-0.15, 0.11), (0.11, 0.11)]:
            fig.add_trace(cylinder_mesh(x+wx, y+wy, 0, 0.04, 0.06, 'y', 8,
                '#212121', 1.0, 'Wheel'))
    elif eq_key == 'eq_ultrasound':
        # Wheeled cart with monitor + probe
        fig.add_trace(box_mesh_global(x-0.22, y-0.18, 0,
            0.44, 0.36, 1.0, '#546e7a', 1.0, 'Ultrasound Cart'))
        # Big monitor
        fig.add_trace(box_mesh_global(x-0.18, y-0.14, 1.0,
            0.36, 0.04, 0.3, '#212121', 1.0, 'Monitor'))
        fig.add_trace(box_mesh_global(x-0.16, y-0.16, 1.03,
            0.32, 0.02, 0.24, '#4fc3f7', 1.0, 'Screen'))
        # Probe holder
        fig.add_trace(cylinder_mesh(x+0.15, y+0.1, 1.05, 0.03, 0.15, 'z', 8,
            '#212121', 1.0, 'Probe'))
    elif eq_key == 'eq_crashcart':
        # Red multi-drawer cart
        fig.add_trace(box_mesh_global(x-0.2, y-0.15, 0,
            0.4, 0.3, 1.0, '#d32f2f', 1.0, 'Crash Cart'))
        # 4 drawer lines
        for dz in [0.2, 0.4, 0.6, 0.8]:
            fig.add_trace(box_mesh_global(x-0.18, y-0.16, dz,
                0.36, 0.01, 0.02, '#fff', 1.0, 'Drawer Line'))
        # Top tray
        fig.add_trace(box_mesh_global(x-0.2, y-0.15, 1.0,
            0.4, 0.3, 0.04, '#fff', 1.0, 'Tray'))
    elif eq_key == 'eq_surgical_light':
        # Ceiling-mounted: cylinder hanging from ceiling
        fig.add_trace(cylinder_mesh(x, y, wall_height - 0.6,
            0.05, 0.6, 'z', 8, '#cfd8dc', 1.0, 'Arm'))
        # Light dome
        fig.add_trace(cylinder_mesh(x, y, wall_height - 0.8,
            0.35, 0.15, 'z', 16, '#fff59d', 0.9, 'Surgical Light'))
        # Inner glow
        fig.add_trace(cylinder_mesh(x, y, wall_height - 0.75,
            0.25, 0.08, 'z', 14, '#fff', 1.0, 'Light Glow'))
    elif eq_key == 'eq_warmer':
        # Orange radiant warmer (looks like an OR cradle)
        fig.add_trace(box_mesh_global(x-0.25, y-0.15, 0,
            0.5, 0.3, 0.9, '#ff9800', 1.0, 'Warmer'))
        # Overhead heater arm
        fig.add_trace(cylinder_mesh(x, y, 0.9, 0.03, 0.5, 'z', 8,
            '#37474f', 1.0, 'Arm'))
        fig.add_trace(box_mesh_global(x-0.15, y-0.05, 1.35,
            0.3, 0.1, 0.05, '#ff5722', 1.0, 'Heater'))
    elif eq_key == 'eq_infusion':
        # Small box on IV-style pole
        fig.add_trace(cylinder_mesh(x, y, 0, 0.025, 1.4, 'z', 6,
            '#9e9e9e', 1.0, 'Pole'))
        fig.add_trace(box_mesh_global(x-0.1, y-0.06, 1.0,
            0.2, 0.12, 0.18, '#42a5f5', 1.0, 'Infusion Pump'))
        fig.add_trace(box_mesh_global(x-0.08, y-0.07, 1.04,
            0.16, 0.01, 0.1, '#fff', 1.0, 'Display'))
    elif eq_key == 'eq_endoscopy':
        # Tall dark cart with stacked screens
        fig.add_trace(box_mesh_global(x-0.2, y-0.25, 0,
            0.4, 0.5, 1.6, '#37474f', 1.0, 'Endoscopy Tower'))
        for dz in [1.0, 1.25]:
            fig.add_trace(box_mesh_global(x-0.16, y-0.27, dz,
                0.32, 0.02, 0.2, '#0277bd', 1.0, 'Endo Screen'))
    elif eq_key == 'eq_ecg':
        # Cream-colored box on small wheeled cart
        fig.add_trace(box_mesh_global(x-0.18, y-0.12, 0.6,
            0.36, 0.24, 0.18, '#fff8e1', 1.0, 'ECG'))
        # Small cart legs
        for lx, ly in [(-0.14, -0.08), (0.12, -0.08),
                        (-0.14, 0.08), (0.12, 0.08)]:
            fig.add_trace(cylinder_mesh(x+lx, y+ly, 0, 0.015, 0.6, 'z', 6,
                '#9e9e9e', 1.0, 'Leg'))
        # Display
        fig.add_trace(box_mesh_global(x-0.14, y-0.13, 0.62,
            0.28, 0.01, 0.1, '#43a047', 1.0, 'ECG Display'))
    elif eq_key == 'eq_suction':
        # Yellow canister
        fig.add_trace(cylinder_mesh(x, y, 0, 0.08, 0.4, 'z', 12,
            '#fdd835', 1.0, 'Suction Canister'))
        # Hose
        fig.add_trace(cylinder_mesh(x, y, 0.4, 0.02, 0.3, 'z', 6,
            '#37474f', 1.0, 'Hose'))
    elif eq_key == 'eq_xray':
        # C-arm style: base + arm + emitter
        fig.add_trace(box_mesh_global(x-0.25, y-0.15, 0,
            0.5, 0.3, 0.7, '#90a4ae', 1.0, 'X-Ray Base'))
        # Vertical arm
        fig.add_trace(cylinder_mesh(x-0.15, y, 0.7, 0.04, 0.8, 'z', 8,
            '#607d8b', 1.0, 'Arm'))
        # Emitter at top
        fig.add_trace(box_mesh_global(x-0.2, y-0.1, 1.5,
            0.4, 0.2, 0.15, '#37474f', 1.0, 'X-Ray Emitter'))
    elif eq_key == 'eq_music':
        # Pink speaker on shelf
        fig.add_trace(cylinder_mesh(x, y, 1.4, 0.1, 0.15, 'z', 14,
            '#f48fb1', 1.0, 'Music Speaker'))
    elif eq_key == 'eq_bairhugger':
        # Light blue box (forced-air warmer unit)
        fig.add_trace(box_mesh_global(x-0.15, y-0.1, 0.4,
            0.3, 0.2, 0.25, '#80deea', 1.0, 'Bair Hugger'))
        # Hose
        fig.add_trace(cylinder_mesh(x+0.1, y, 0.55, 0.04, 0.4, 'z', 6,
            '#26c6da', 1.0, 'Warming Hose'))
    elif eq_key == 'eq_comm_board':
        # White wall board
        fig.add_trace(box_mesh_global(x-0.4, y, 1.4,
            0.8, 0.04, 0.5, '#fff', 1.0, 'Comm Board'))
        # Frame
        fig.add_trace(box_mesh_global(x-0.42, y-0.01, 1.38,
            0.84, 0.01, 0.54, '#37474f', 1.0, 'Frame'))
    elif eq_key == 'eq_fire_ext':
        # Wall-mounted fire extinguisher (red cylinder + neck + bracket)
        # Floor stand bracket
        fig.add_trace(box_mesh_global(x-0.08, y-0.08, 0,
            0.16, 0.16, 0.05, '#37474f', 1.0, 'FE Base'))
        # Main red body cylinder
        fig.add_trace(cylinder_mesh(x, y, 0.05, 0.09, 0.45, 'z', 14,
            '#c62828', 1.0, 'Fire Extinguisher'))
        # White label band
        fig.add_trace(cylinder_mesh(x, y, 0.18, 0.092, 0.08, 'z', 14,
            '#ffffff', 1.0, 'FE Label'))
        # Black neck/valve
        fig.add_trace(cylinder_mesh(x, y, 0.50, 0.05, 0.10, 'z', 10,
            '#212121', 1.0, 'FE Valve'))
        # Trigger handle (small box on top)
        fig.add_trace(box_mesh_global(x-0.03, y-0.02, 0.58,
            0.06, 0.04, 0.04, '#37474f', 1.0, 'FE Handle'))
        # Wall-mounted "FIRE" sign above (small red panel)
        fig.add_trace(box_mesh_global(x-0.15, y-0.02, 0.85,
            0.30, 0.02, 0.20, '#c62828', 0.9, 'FE Sign'))


# ════════════════════════════════════════════════════════════════
# 🔥 FIRE PROTECTION SYSTEMS (Batch SE-1)
# Ceiling + wall-mounted safety devices per NFPA 72 / 13 / 101.
# Counts auto-scale with floor area; placements follow code spacing.
# ════════════════════════════════════════════════════════════════
def get_fire_protection_layout(width, height, em_exit_positions=None):
    """Return positions for all 4 fire-protection device types.
    Returns dict with keys: smoke_detectors, sprinklers, pull_stations,
    emergency_lights. Each is a list of (x, y) tuples.
    em_exit_positions: list of (side, pos) for emergency exit doors so
    pull stations + emergency lights are placed near them."""
    import math

    # ─── 1. SMOKE DETECTORS — NFPA 72: ~84 m² coverage per detector ───
    # Grid with ~7m spacing (since 7×7 ≈ 49m² actual radius, 84m² typical)
    sd_spacing = 7.0
    n_x_sd = max(2, math.ceil(width / sd_spacing))
    n_y_sd = max(2, math.ceil(height / sd_spacing))
    sd_step_x = width / (n_x_sd + 1)
    sd_step_y = height / (n_y_sd + 1)
    smoke_detectors = []
    for ix in range(1, n_x_sd + 1):
        for iy in range(1, n_y_sd + 1):
            smoke_detectors.append((ix * sd_step_x, iy * sd_step_y))

    # ─── 2. SPRINKLERS — NFPA 13: 3.7m × 3.7m max spacing ───
    sp_spacing = 3.5  # slightly tighter than max for safety
    n_x_sp = max(2, math.ceil(width / sp_spacing))
    n_y_sp = max(2, math.ceil(height / sp_spacing))
    sp_step_x = width / (n_x_sp + 1)
    sp_step_y = height / (n_y_sp + 1)
    sprinklers = []
    for ix in range(1, n_x_sp + 1):
        for iy in range(1, n_y_sp + 1):
            sprinklers.append((ix * sp_step_x, iy * sp_step_y))

    # ─── 3. PULL STATIONS — near each exit (main + emergency) ───
    # Main door: bottom-center, station 0.6m to the right of door frame
    pull_stations = [(width/2 + 1.0, 0.4)]
    # Add near emergency exits
    if em_exit_positions:
        for side, pos in em_exit_positions:
            if side == 'left':
                pull_stations.append((0.5, pos - 0.6))
            elif side == 'right':
                pull_stations.append((width - 0.5, pos - 0.6))
            elif side == 'back':
                pull_stations.append((pos + 0.6, height - 0.4))

    # ─── 4. EMERGENCY LIGHTING — above every door + intermediate ───
    emergency_lights = []
    # Above main door
    emergency_lights.append((width/2, 0.3))
    # Above each emergency exit
    if em_exit_positions:
        for side, pos in em_exit_positions:
            if side == 'left':
                emergency_lights.append((0.3, pos))
            elif side == 'right':
                emergency_lights.append((width - 0.3, pos))
            elif side == 'back':
                emergency_lights.append((pos, height - 0.3))
    # Add intermediate egress-path lighting along walls (~6m intervals)
    el_spacing = 6.0
    if width > el_spacing:
        n_el_x = max(1, math.floor(width / el_spacing))
        for i in range(1, n_el_x + 1):
            emergency_lights.append((i * width / (n_el_x + 1), 0.3))
            emergency_lights.append((i * width / (n_el_x + 1), height - 0.3))

    return {
        'smoke_detectors': smoke_detectors,
        'sprinklers': sprinklers,
        'pull_stations': pull_stations,
        'emergency_lights': emergency_lights,
    }


def draw_fire_protection_2d(ax, layout, options):
    """Draw the 4 fire-protection systems on a 2D blueprint.
    options is a dict with bool flags for each system."""
    # ── 1. Smoke Detectors (ceiling, drawn as overlay) ──
    if options.get('smoke', False):
        for x, y in layout['smoke_detectors']:
            # Outer ring (white)
            ax.add_patch(_MCircle((x, y), 0.18, facecolor='#ffffff',
                edgecolor='#37474f', linewidth=0.8,
                alpha=0.85, zorder=11))
            # Inner red LED
            ax.add_patch(_MCircle((x, y), 0.06, facecolor='#e53935',
                edgecolor='#b71c1c', linewidth=0.5, zorder=12))
            # Label
            ax.text(x, y - 0.32, "SD", color='#37474f', fontsize=4,
                ha='center', weight='bold', zorder=12,
                bbox=dict(boxstyle="round,pad=0.1",
                          facecolor='white', edgecolor='none', alpha=0.7))

    # ── 2. Sprinklers (smaller, denser grid) ──
    if options.get('sprinklers', False):
        for x, y in layout['sprinklers']:
            # Sprinkler head: blue cross pattern + circle
            ax.add_patch(_MCircle((x, y), 0.12, facecolor='#bbdefb',
                edgecolor='#1976d2', linewidth=0.6,
                alpha=0.75, zorder=11))
            # X pattern inside (spray indicator)
            ax.plot([x-0.07, x+0.07], [y-0.07, y+0.07],
                color='#0d47a1', linewidth=0.6, zorder=12)
            ax.plot([x-0.07, x+0.07], [y+0.07, y-0.07],
                color='#0d47a1', linewidth=0.6, zorder=12)

    # ── 3. Pull Stations (wall-mounted, red squares) ──
    if options.get('pull_stations', False):
        for x, y in layout['pull_stations']:
            ax.add_patch(_MRect((x - 0.12, y - 0.12), 0.24, 0.24,
                facecolor='#e53935', edgecolor='#b71c1c',
                linewidth=1.2, zorder=11))
            # White triangle/arrow icon
            ax.text(x, y, "!", color='white', fontsize=8,
                ha='center', va='center', weight='bold', zorder=12)
            ax.text(x, y - 0.30, "PULL", color='#b71c1c', fontsize=4,
                ha='center', weight='bold', zorder=12)

    # ── 4. Emergency Lighting (green-bordered fixtures) ──
    if options.get('emergency_lights', False):
        for x, y in layout['emergency_lights']:
            # Green-bordered rectangle = emergency light fixture
            ax.add_patch(_MRect((x - 0.15, y - 0.05), 0.30, 0.10,
                facecolor='#c8e6c9', edgecolor='#2e7d32',
                linewidth=1.2, zorder=11))
            # Light bulb glow
            ax.add_patch(_MCircle((x, y), 0.04, facecolor='#fff59d',
                edgecolor='#f57f17', linewidth=0.4, zorder=12))


def draw_fire_protection_3d(fig, layout, options, wall_height=3.0):
    """Draw the 4 fire-protection systems in 3D."""
    # ── 1. Smoke Detectors (white half-spheres at ceiling) ──
    if options.get('smoke', False):
        for x, y in layout['smoke_detectors']:
            # White disc base mounted to ceiling
            fig.add_trace(cylinder_mesh(x, y, wall_height - 0.05,
                0.15, 0.05, 'z', 14, '#ffffff', 1.0, 'Smoke Detector'))
            # Red LED bottom
            fig.add_trace(cylinder_mesh(x, y, wall_height - 0.08,
                0.04, 0.03, 'z', 8, '#e53935', 1.0, 'SD LED'))

    # ── 2. Sprinklers (chrome heads at ceiling, pointing down) ──
    if options.get('sprinklers', False):
        for x, y in layout['sprinklers']:
            # Chrome cylinder mount
            fig.add_trace(cylinder_mesh(x, y, wall_height - 0.10,
                0.04, 0.10, 'z', 8, '#90a4ae', 1.0, 'Sprinkler'))
            # Deflector plate (sprinkler head)
            fig.add_trace(cylinder_mesh(x, y, wall_height - 0.15,
                0.08, 0.02, 'z', 12, '#cfd8dc', 1.0, 'Sprinkler Head'))

    # ── 3. Pull Stations (red boxes on walls, ~1.2m height) ──
    if options.get('pull_stations', False):
        for x, y in layout['pull_stations']:
            fig.add_trace(box_mesh_global(x - 0.10, y - 0.06, 1.2,
                0.20, 0.08, 0.25, '#e53935', 1.0, 'Pull Station'))
            # White label area
            fig.add_trace(box_mesh_global(x - 0.08, y - 0.07, 1.32,
                0.16, 0.01, 0.10, '#ffffff', 1.0, 'PULL Label'))

    # ── 4. Emergency Lighting (green-bordered wall fixtures at ~2.4m) ──
    if options.get('emergency_lights', False):
        for x, y in layout['emergency_lights']:
            # Light fixture body
            fig.add_trace(box_mesh_global(x - 0.18, y - 0.06, 2.4,
                0.36, 0.08, 0.18, '#c8e6c9', 1.0, 'Emergency Light'))
            # Green frame edge
            fig.add_trace(box_mesh_global(x - 0.19, y - 0.07, 2.39,
                0.38, 0.01, 0.20, '#2e7d32', 1.0, 'EL Frame'))
            # Bulb glow
            fig.add_trace(cylinder_mesh(x, y - 0.05, 2.45,
                0.05, 0.08, 'z', 8, '#fff59d', 0.9, 'EL Bulb'))


# ════════════════════════════════════════════════════════════════
# 🚨 MEDICAL EMERGENCY DEVICES (Batch SE-2)
# AED Stations + Eye Wash + Emergency Shower + Call Buttons
# Sources: AHA guidelines, ANSI Z358.1, OSHA 29 CFR 1910.151
# ════════════════════════════════════════════════════════════════

# AED Stations — applicable to nearly all clinical + public areas
_AED_APPLICABLE = {
    'Operating Room (OR)', 'Intensive Care Unit (ICU)',
    'Emergency Room (ER)', 'NICU (Neonatal ICU)',
    'General Patient Ward', 'Pediatric Ward',
    'Maternity & Delivery', 'Oncology', 'Dialysis Unit',
    'Cafeteria & Kitchen', 'Outpatient Clinic',
    'Reception & Waiting', 'Physical Therapy',
    'Dental Clinic', 'Pharmacy', 'Psychiatric Ward',
    'Radiology Department', 'MRI Room', 'CT Scan Room',
    'Laboratory', 'Blood Bank', 'Sterilization (CSSD)',
    'Administrative Offices',
}

# Eye Wash + Emergency Shower — ONLY facilities handling chemicals/biohazards
# Per ANSI Z358.1: required within 10 seconds reach
_CHEMICAL_FACILITIES = {
    'Laboratory', 'Pharmacy', 'Sterilization (CSSD)',
    'Blood Bank', 'Dialysis Unit', 'Radiology Department',
    'Oncology', 'Dental Clinic',
}

# Patient Call Buttons — ONLY in facilities with patient beds
_BED_FACILITIES = {
    'Operating Room (OR)', 'Intensive Care Unit (ICU)',
    'Emergency Room (ER)', 'NICU (Neonatal ICU)',
    'General Patient Ward', 'Pediatric Ward',
    'Maternity & Delivery', 'Oncology', 'Dialysis Unit',
    'Psychiatric Ward',
}


def get_medical_emergency_layout(width, height, n_beds=0, em_exit_positions=None):
    """Return positions for medical emergency devices.
    Returns dict with: aed, eye_wash, emergency_shower, call_buttons."""
    # ─── AED Stations — wall-mounted, 1-2 per room ───
    # Placement: prominently on a wall near the main entrance
    # For large rooms (> 80 m²): 2 AEDs
    aed_positions = []
    if width * height > 80:
        aed_positions.append((width * 0.20, 0.4))  # near entrance, left side
        aed_positions.append((width * 0.80, 0.4))  # near entrance, right side
    else:
        aed_positions.append((width * 0.85, 0.4))  # single, right of entrance

    # ─── Eye Wash + Emergency Shower — paired near sinks/work areas ───
    # Best placement: on a wall away from main traffic, near work surfaces
    eye_wash_positions = [(width * 0.05, height * 0.5)]  # left wall, middle
    emergency_shower_positions = [(width * 0.05, height * 0.65)]  # just above eye wash

    # ─── Patient Call Buttons — 1 per bed ───
    # Position alongside where beds would be drawn (perimeter walls in wards)
    call_button_positions = []
    if n_beds > 0:
        # Mirror the bed layout used in layout_ward_2d (top and bottom walls)
        max_per_wall = max(1, int((width - 1.5) / 1.5))
        top_beds = min((n_beds + 1) // 2, max_per_wall)
        bottom_beds = min(n_beds - top_beds, max_per_wall)
        if top_beds > 0:
            x_step = (width - 1) / top_beds
            for i in range(top_beds):
                bx = 0.5 + x_step * (i + 0.5)
                # Button just below the bed headboard area
                call_button_positions.append((bx + 0.55, height - 1.7))
        if bottom_beds > 0:
            x_step = (width - 1) / bottom_beds
            for i in range(bottom_beds):
                bx = 0.5 + x_step * (i + 0.5)
                # Button just above the bed headboard area
                call_button_positions.append((bx + 0.55, 1.7))

    return {
        'aed': aed_positions,
        'eye_wash': eye_wash_positions,
        'emergency_shower': emergency_shower_positions,
        'call_buttons': call_button_positions,
    }


def draw_medical_emergency_2d(ax, layout, options):
    """Draw 4 medical-emergency device types on the 2D blueprint."""
    # ── 1. AED Stations (red wall cabinets) ──
    if options.get('aed', False):
        for x, y in layout['aed']:
            # Cabinet body
            ax.add_patch(_MFBox((x - 0.18, y - 0.13), 0.36, 0.26,
                boxstyle="round,pad=0.02", facecolor='#e53935',
                edgecolor='#b71c1c', linewidth=1.2, zorder=11))
            # White heart + bolt icon area
            ax.add_patch(_MRect((x - 0.13, y - 0.08), 0.26, 0.16,
                facecolor='white', edgecolor='#b71c1c',
                linewidth=0.5, zorder=12))
            # Lightning bolt symbol
            ax.text(x, y, "⚡♥", fontsize=8, ha='center', va='center',
                color='#b71c1c', weight='bold', zorder=13)
            # Label below
            ax.text(x, y - 0.28, "AED", color='#b71c1c', fontsize=5,
                ha='center', weight='bold', zorder=12,
                bbox=dict(boxstyle="round,pad=0.15",
                          facecolor='white', edgecolor='#b71c1c',
                          linewidth=0.8))

    # ── 2. Eye Wash Stations (green wall units with eye icons) ──
    if options.get('eye_wash', False):
        for x, y in layout['eye_wash']:
            # Green backing plate
            ax.add_patch(_MFBox((x - 0.18, y - 0.12), 0.36, 0.24,
                boxstyle="round,pad=0.02", facecolor='#2e7d32',
                edgecolor='#1b5e20', linewidth=1.2, zorder=11))
            # Two white "eye" icons (paired wash nozzles)
            ax.add_patch(_MCircle((x - 0.08, y), 0.06,
                facecolor='white', edgecolor='#1b5e20',
                linewidth=0.6, zorder=12))
            ax.add_patch(_MCircle((x + 0.08, y), 0.06,
                facecolor='white', edgecolor='#1b5e20',
                linewidth=0.6, zorder=12))
            # Pupil dots
            ax.add_patch(_MCircle((x - 0.08, y), 0.025,
                facecolor='#1b5e20', zorder=13))
            ax.add_patch(_MCircle((x + 0.08, y), 0.025,
                facecolor='#1b5e20', zorder=13))
            # Label
            ax.text(x, y - 0.25, "EYE WASH", color='#1b5e20',
                fontsize=4, ha='center', weight='bold', zorder=12)

    # ── 3. Emergency Shower (green overhead head + drain) ──
    if options.get('emergency_shower', False):
        for x, y in layout['emergency_shower']:
            # Outer ring (showerhead from above)
            ax.add_patch(_MCircle((x, y), 0.20, facecolor='#a5d6a7',
                edgecolor='#1b5e20', linewidth=1.4, zorder=11))
            # Inner cross-hatch (spray pattern)
            for ang in [0, np.pi/3, 2*np.pi/3]:
                ax.plot([x + 0.15*np.cos(ang), x - 0.15*np.cos(ang)],
                        [y + 0.15*np.sin(ang), y - 0.15*np.sin(ang)],
                        color='#1b5e20', linewidth=0.8, zorder=12)
            # Center pull-chain dot
            ax.add_patch(_MCircle((x, y), 0.04, facecolor='#1b5e20',
                zorder=13))
            # Label
            ax.text(x, y - 0.32, "SAFETY\nSHOWER", color='#1b5e20',
                fontsize=4, ha='center', va='center',
                weight='bold', zorder=12,
                bbox=dict(boxstyle="round,pad=0.1",
                          facecolor='white', edgecolor='#2e7d32',
                          linewidth=0.6))

    # ── 4. Patient Call Buttons (small red dots on walls beside beds) ──
    if options.get('call_buttons', False):
        for x, y in layout['call_buttons']:
            # Square mounting plate
            ax.add_patch(_MRect((x - 0.06, y - 0.06), 0.12, 0.12,
                facecolor='#ffebee', edgecolor='#c62828',
                linewidth=0.8, zorder=11))
            # Red button
            ax.add_patch(_MCircle((x, y), 0.04, facecolor='#e53935',
                edgecolor='#b71c1c', linewidth=0.6, zorder=12))
            # Tiny label
            ax.text(x + 0.10, y, "📞", fontsize=4, ha='left', va='center',
                color='#c62828', weight='bold', zorder=12)


def draw_medical_emergency_3d(fig, layout, options, wall_height=3.0):
    """Draw 4 medical-emergency device types in 3D."""
    # ── 1. AED Stations (red wall cabinets at ~1.5m height) ──
    if options.get('aed', False):
        for x, y in layout['aed']:
            # Cabinet body mounted on wall (back of room facing)
            fig.add_trace(box_mesh_global(x - 0.18, y - 0.06, 1.3,
                0.36, 0.10, 0.32, '#e53935', 1.0, 'AED Cabinet'))
            # White face plate
            fig.add_trace(box_mesh_global(x - 0.14, y - 0.07, 1.34,
                0.28, 0.01, 0.24, '#ffffff', 1.0, 'AED Face'))
            # Red lightning bolt area (small panel)
            fig.add_trace(box_mesh_global(x - 0.06, y - 0.075, 1.42,
                0.12, 0.005, 0.08, '#e53935', 1.0, 'AED Symbol'))

    # ── 2. Eye Wash Stations (green wall units with two spouts) ──
    if options.get('eye_wash', False):
        for x, y in layout['eye_wash']:
            # Green backing plate
            fig.add_trace(box_mesh_global(x - 0.18, y - 0.06, 1.0,
                0.36, 0.10, 0.28, '#2e7d32', 1.0, 'Eye Wash Unit'))
            # Bowl/basin
            fig.add_trace(box_mesh_global(x - 0.16, y - 0.18, 0.95,
                0.32, 0.12, 0.06, '#cfd8dc', 1.0, 'Eye Wash Bowl'))
            # Two nozzles
            fig.add_trace(cylinder_mesh(x - 0.08, y - 0.10, 1.05,
                0.025, 0.06, 'z', 8, '#90a4ae', 1.0, 'EW Nozzle L'))
            fig.add_trace(cylinder_mesh(x + 0.08, y - 0.10, 1.05,
                0.025, 0.06, 'z', 8, '#90a4ae', 1.0, 'EW Nozzle R'))

    # ── 3. Emergency Shower (overhead pipe + showerhead) ──
    if options.get('emergency_shower', False):
        for x, y in layout['emergency_shower']:
            # Vertical pipe from ceiling
            fig.add_trace(cylinder_mesh(x, y, 1.8, 0.04,
                wall_height - 1.8, 'z', 8,
                '#2e7d32', 1.0, 'Shower Pipe'))
            # Showerhead disc
            fig.add_trace(cylinder_mesh(x, y, 1.85, 0.20, 0.04,
                'z', 16, '#43a047', 1.0, 'Shower Head'))
            # Pull chain (thin line from showerhead going down)
            fig.add_trace(cylinder_mesh(x + 0.22, y, 0.8, 0.005, 1.05,
                'z', 6, '#37474f', 1.0, 'Pull Chain'))
            # Pull triangle handle at bottom
            fig.add_trace(box_mesh_global(x + 0.20, y - 0.025, 0.75,
                0.05, 0.05, 0.08, '#2e7d32', 1.0, 'Pull Handle'))

    # ── 4. Patient Call Buttons (small red buttons on walls behind beds) ──
    if options.get('call_buttons', False):
        for x, y in layout['call_buttons']:
            # Square base plate on wall
            fig.add_trace(box_mesh_global(x - 0.08, y - 0.04, 1.0,
                0.16, 0.04, 0.16, '#ffebee', 1.0, 'Call Plate'))
            # Red button protrusion
            fig.add_trace(cylinder_mesh(x, y - 0.04, 1.07, 0.04, 0.025,
                'y', 8, '#e53935', 1.0, 'Call Button'))


# ════════════════════════════════════════════════════════════════
# 🛤️ WAYFINDING & EGRESS (Batch SE-3)
# Egress Path arrows + Wayfinding Signs + Refuge Areas + Stairs
# Sources: NFPA 101 (egress), ADA 4.3, ISO 7010 (signage)
# ════════════════════════════════════════════════════════════════

# Stairs apply only to multi-floor facilities. In single-room mode,
# the "stair access point" represents the connection to upper floors.
_STAIRS_APPLICABLE = {
    'Operating Room (OR)', 'Intensive Care Unit (ICU)',
    'Emergency Room (ER)', 'NICU (Neonatal ICU)',
    'General Patient Ward', 'Pediatric Ward',
    'Maternity & Delivery', 'Oncology', 'Dialysis Unit',
    'Cafeteria & Kitchen', 'Outpatient Clinic',
    'Reception & Waiting', 'Physical Therapy',
    'Pharmacy', 'Laboratory', 'Blood Bank',
    'Sterilization (CSSD)', 'Administrative Offices',
    'Psychiatric Ward', 'Dental Clinic',
    # Excluded: MRI Room, CT Scan Room (usually ground-floor only)
}


def get_wayfinding_layout(width, height, em_exit_positions=None,
                            include_main_door=True):
    """Compute placements for wayfinding & egress elements.
    Returns dict: egress_paths, wayfinding_signs, refuge_areas, stairs."""
    em_exit_positions = em_exit_positions or []

    # ─── 1. EGRESS PATHS — center of room → each exit door ───
    # Each path is a list of (x, y) waypoints from interior toward exit.
    egress_paths = []
    center = (width / 2, height / 2)

    # Main door egress path (always at bottom-center)
    if include_main_door:
        egress_paths.append({
            'points': [center, (width / 2, height * 0.25), (width / 2, 0.4)],
            'label': 'MAIN EXIT',
            'side': 'main',
        })

    # Emergency exit paths
    for side, pos in em_exit_positions:
        if side == 'left':
            egress_paths.append({
                'points': [center, (width * 0.30, pos), (0.5, pos)],
                'label': 'EMERGENCY EXIT', 'side': 'left'})
        elif side == 'right':
            egress_paths.append({
                'points': [center, (width * 0.70, pos), (width - 0.5, pos)],
                'label': 'EMERGENCY EXIT', 'side': 'right'})
        elif side == 'back':
            egress_paths.append({
                'points': [center, (pos, height * 0.70), (pos, height - 0.5)],
                'label': 'EMERGENCY EXIT', 'side': 'back'})

    # ─── 2. WAYFINDING SIGNS — directional signs on walls near doors ───
    # ISO 7010 standard: green safety signs at decision points
    wayfinding_signs = []
    # Main door wayfinding (mounted on side walls a few meters before door)
    wayfinding_signs.append({
        'x': width / 2 - 2.0, 'y': 1.5,
        'arrow': '↓', 'text': 'MAIN EXIT',
    })
    # Signs near each emergency exit
    for side, pos in em_exit_positions:
        if side == 'left':
            wayfinding_signs.append({
                'x': 1.5, 'y': pos + 0.7,
                'arrow': '←', 'text': 'EXIT'})
        elif side == 'right':
            wayfinding_signs.append({
                'x': width - 1.5, 'y': pos + 0.7,
                'arrow': '→', 'text': 'EXIT'})
        elif side == 'back':
            wayfinding_signs.append({
                'x': pos + 0.7, 'y': height - 1.5,
                'arrow': '↑', 'text': 'EXIT'})

    # ─── 3. REFUGE AREAS — protected spaces for non-ambulatory persons ───
    # NFPA 101 + ADA: 1 refuge area near each emergency exit (where
    # accommodating wheelchair-bound or mobility-impaired patients waits
    # for assisted evacuation). Typically 0.76m × 1.22m clear space.
    refuge_areas = []
    refuge_w, refuge_h = 1.2, 1.2  # m
    for side, pos in em_exit_positions:
        if side == 'left':
            refuge_areas.append({
                'x': 0.5, 'y': pos - 1.5,
                'w': refuge_w, 'h': refuge_h})
        elif side == 'right':
            refuge_areas.append({
                'x': width - 0.5 - refuge_w, 'y': pos - 1.5,
                'w': refuge_w, 'h': refuge_h})
        elif side == 'back':
            refuge_areas.append({
                'x': pos - refuge_w/2, 'y': height - 1.5 - refuge_h,
                'w': refuge_w, 'h': refuge_h})

    # ─── 4. EMERGENCY STAIRS — stair access symbol ───
    # In a single-room view this represents the connection to the
    # building's emergency stairwell (typically at a corner of the dept).
    stairs = []
    if width * height >= 30:  # only show for reasonably-sized rooms
        # Place in a corner away from main door
        stairs.append({
            'x': width - 1.6, 'y': height - 1.6,
            'w': 1.4, 'h': 1.4,
            'n_steps': 5,
        })

    return {
        'egress_paths': egress_paths,
        'wayfinding_signs': wayfinding_signs,
        'refuge_areas': refuge_areas,
        'stairs': stairs,
    }


def draw_wayfinding_2d(ax, layout, options):
    """Draw the 4 wayfinding/egress element types on the 2D blueprint."""
    # ── 1. Egress Paths (green dashed arrows from center → exits) ──
    if options.get('egress_path', False):
        for path in layout['egress_paths']:
            pts = path['points']
            # Draw dashed segments
            for i in range(len(pts) - 1):
                ax.plot([pts[i][0], pts[i+1][0]], [pts[i][1], pts[i+1][1]],
                    color='#2e7d32', linewidth=2.4,
                    linestyle=(0, (6, 3)),  # dashed pattern
                    alpha=0.85, zorder=8)
            # Arrowhead at the last segment, pointing toward exit
            x1, y1 = pts[-2]
            x2, y2 = pts[-1]
            ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle='->', color='#1b5e20',
                                 lw=2.5, mutation_scale=20),
                zorder=9)
            # Small "→ EXIT" label at midpoint of the path
            mid_idx = len(pts) // 2
            mx, my = pts[mid_idx]
            ax.text(mx + 0.3, my, f"→ {path['label']}",
                fontsize=5, color='#1b5e20', weight='bold', zorder=9,
                bbox=dict(boxstyle="round,pad=0.15",
                          facecolor='#e8f5e9', edgecolor='#2e7d32',
                          linewidth=0.6, alpha=0.9))

    # ── 2. Wayfinding Signs (green rectangles with arrow + EXIT text) ──
    if options.get('wayfinding', False):
        for sign in layout['wayfinding_signs']:
            sx, sy = sign['x'], sign['y']
            # Sign body
            ax.add_patch(_MFBox((sx - 0.20, sy - 0.10), 0.40, 0.20,
                boxstyle="round,pad=0.02", facecolor='#2e7d32',
                edgecolor='#1b5e20', linewidth=1.0, zorder=11))
            # White interior
            ax.add_patch(_MRect((sx - 0.17, sy - 0.07), 0.34, 0.14,
                facecolor='white', zorder=12))
            # Arrow + label
            ax.text(sx, sy, f"{sign['arrow']} {sign['text']}",
                fontsize=5, color='#1b5e20', ha='center', va='center',
                weight='bold', zorder=13)

    # ── 3. Refuge Areas (blue hatched rectangles with wheelchair icon) ──
    if options.get('refuge', False):
        for ref in layout['refuge_areas']:
            rx, ry, rw, rh = ref['x'], ref['y'], ref['w'], ref['h']
            # Filled blue zone
            ax.add_patch(_MRect((rx, ry), rw, rh,
                facecolor='#bbdefb', edgecolor='#1565c0',
                linewidth=1.5, alpha=0.65, zorder=7,
                hatch='///'))
            # Wheelchair symbol
            ax.text(rx + rw/2, ry + rh/2, "♿",
                fontsize=16, ha='center', va='center', zorder=8)
            # Label
            ax.text(rx + rw/2, ry + 0.15, "REFUGE",
                fontsize=5, color='#0d47a1', ha='center',
                weight='bold', zorder=8)

    # ── 4. Emergency Stairs (stair tread pattern) ──
    if options.get('stairs', False):
        for st in layout['stairs']:
            sx, sy, sw, sh = st['x'], st['y'], st['w'], st['h']
            n_steps = st['n_steps']
            # Background frame
            ax.add_patch(_MRect((sx, sy), sw, sh,
                facecolor='#eeeeee', edgecolor='#37474f',
                linewidth=1.5, zorder=7))
            # Step lines
            step_h = sh / n_steps
            for i in range(n_steps):
                ax.add_patch(_MRect((sx, sy + i * step_h), sw, step_h,
                    facecolor='#cfd8dc' if i % 2 == 0 else '#b0bec5',
                    edgecolor='#546e7a', linewidth=0.6, zorder=8))
                # Direction arrow on top
                if i == n_steps - 1:
                    ax.text(sx + sw/2, sy + sh + 0.15,
                        "↑ UP", fontsize=5, ha='center',
                        color='#37474f', weight='bold', zorder=9)
            # Label
            ax.text(sx + sw/2, sy - 0.20, "🪜 STAIRS",
                fontsize=5, ha='center', color='#37474f',
                weight='bold', zorder=9,
                bbox=dict(boxstyle="round,pad=0.15",
                          facecolor='white', edgecolor='#37474f',
                          linewidth=0.6))


def draw_wayfinding_3d(fig, layout, options, wall_height=3.0):
    """Draw the 4 wayfinding/egress element types in 3D."""
    # ── 1. Egress Paths (green strip on the floor along the route) ──
    if options.get('egress_path', False):
        for path in layout['egress_paths']:
            pts = path['points']
            for i in range(len(pts) - 1):
                x1, y1 = pts[i]
                x2, y2 = pts[i+1]
                # Compute segment midpoint + dimensions
                cx = (x1 + x2) / 2
                cy = (y1 + y2) / 2
                length = ((x2 - x1)**2 + (y2 - y1)**2) ** 0.5
                # Horizontal strip
                if abs(x2 - x1) > abs(y2 - y1):
                    fig.add_trace(box_mesh_global(
                        cx - length/2, cy - 0.15, 0.005,
                        length, 0.30, 0.02,
                        '#43a047', 0.7, 'Egress Path'))
                else:
                    fig.add_trace(box_mesh_global(
                        cx - 0.15, cy - length/2, 0.005,
                        0.30, length, 0.02,
                        '#43a047', 0.7, 'Egress Path'))

    # ── 2. Wayfinding Signs (hanging signs near doors at 2.4m) ──
    if options.get('wayfinding', False):
        for sign in layout['wayfinding_signs']:
            sx, sy = sign['x'], sign['y']
            # Sign panel hanging from ceiling
            fig.add_trace(box_mesh_global(sx - 0.30, sy - 0.04, 2.4,
                0.60, 0.08, 0.30, '#2e7d32', 1.0, 'Wayfinding Sign'))
            # White face
            fig.add_trace(box_mesh_global(sx - 0.26, sy - 0.045, 2.43,
                0.52, 0.005, 0.24, '#ffffff', 1.0, 'Sign Face'))
            # Hanging rod
            fig.add_trace(cylinder_mesh(sx, sy, 2.70, 0.015,
                wall_height - 2.70, 'z', 6,
                '#90a4ae', 1.0, 'Sign Rod'))

    # ── 3. Refuge Areas (blue floor zone with low railings) ──
    if options.get('refuge', False):
        for ref in layout['refuge_areas']:
            rx, ry, rw, rh = ref['x'], ref['y'], ref['w'], ref['h']
            # Floor zone (slightly raised)
            fig.add_trace(box_mesh_global(rx, ry, 0.005,
                rw, rh, 0.03, '#64b5f6', 0.65, 'Refuge Floor'))
            # Border outline (raised slightly more)
            fig.add_trace(box_mesh_global(rx - 0.02, ry - 0.02, 0.005,
                rw + 0.04, 0.02, 0.05, '#1565c0', 1.0, 'Refuge Edge'))
            fig.add_trace(box_mesh_global(rx - 0.02, ry + rh, 0.005,
                rw + 0.04, 0.02, 0.05, '#1565c0', 1.0, 'Refuge Edge'))
            # Wall-mounted sign indicating refuge area
            fig.add_trace(box_mesh_global(rx + rw/2 - 0.20, ry + rh/2 - 0.04,
                1.6, 0.40, 0.08, 0.30, '#1565c0', 1.0, 'Refuge Sign'))

    # ── 4. Emergency Stairs (stepped geometry going up) ──
    if options.get('stairs', False):
        for st in layout['stairs']:
            sx, sy, sw, sh = st['x'], st['y'], st['w'], st['h']
            n_steps = st['n_steps']
            step_h = sh / n_steps
            step_rise = 0.18  # standard stair rise per step
            for i in range(n_steps):
                step_z = i * step_rise
                fig.add_trace(box_mesh_global(
                    sx, sy + i * step_h, step_z,
                    sw, sh - i * step_h, step_rise,
                    '#cfd8dc' if i % 2 == 0 else '#b0bec5',
                    1.0, f'Step {i+1}'))
            # Side railings
            rail_z_top = n_steps * step_rise + 0.9
            fig.add_trace(box_mesh_global(sx - 0.05, sy, 0,
                0.05, sh, rail_z_top, '#37474f', 1.0, 'Stair Rail L'))
            fig.add_trace(box_mesh_global(sx + sw, sy, 0,
                0.05, sh, rail_z_top, '#37474f', 1.0, 'Stair Rail R'))


# ════════════════════════════════════════════════════════════════
# ☣️ HAZARD MARKERS (Batch SE-4)
# Radiation Shielding + Isolation Indicators + Hazmat Storage
# Sources: NFPA 704 (hazard diamond), CDC/OSHA biohazard,
#          IEC 60417 radiation symbol
# ════════════════════════════════════════════════════════════════

# Radiation Shielding — facilities with X-ray/radiation sources
_RADIATION_FACILITIES = {
    'Radiology Department', 'MRI Room', 'CT Scan Room',
    'Oncology', 'Dental Clinic',
}

# Isolation Room — facilities handling infectious patients
_ISOLATION_FACILITIES = {
    'Intensive Care Unit (ICU)', 'Emergency Room (ER)',
    'NICU (Neonatal ICU)', 'General Patient Ward',
    'Pediatric Ward', 'Oncology',
}

# Hazmat Storage — facilities with hazardous chemicals/biohazards
_HAZMAT_FACILITIES = {
    'Pharmacy', 'Laboratory', 'Sterilization (CSSD)',
    'Blood Bank', 'Oncology', 'Dialysis Unit',
    'Radiology Department', 'Dental Clinic',
}


# ════════════════════════════════════════════════════════════════
# 🌱 SUSTAINABILITY MATERIALS DATABASE (v22)
# Eco-friendly construction materials for hospital healthcare design.
# Each entry has: description, best-use locations in hospital, eco
# benefits, indicative cost (USD/m²), and the standard/reference.
# Categories: FLOORING / WALL / CEILING / INSULATION / PAINT
# Sources: USGBC LEED v4 Healthcare, ASHRAE 189.3, ASTM E2129, Cradle-to-Cradle
# ════════════════════════════════════════════════════════════════
MATERIALS_DB = {
    # ───────────── FLOORING ─────────────
    'flooring': {
        'Antimicrobial Vinyl (Seamless)': {
            'icon': '🧪',
            'description': 'Heat-welded PVC sheet with embedded silver-ion '
                           'biocide. Seamless joints prevent bacterial '
                           'colonization in cracks. Industry standard for '
                           'sterile healthcare environments.',
            'best_for': ['Operating Room (OR)', 'Intensive Care Unit (ICU)',
                         'NICU (Neonatal ICU)', 'Sterilization (CSSD)',
                         'Blood Bank'],
            'use_case': 'Sterile zones, surgical suites, isolation areas — '
                        'anywhere requiring frequent disinfection.',
            'eco_benefits': '15-yr lifespan, low-VOC adhesives (FloorScore '
                            'certified), recyclable at end of life.',
            'cost_per_m2_usd': 45,
            'reference': 'FGI 2018 §2.1-7.2.4.1 + ASTM F1700',
        },
        'Eco-Linoleum (Natural)': {
            'icon': '🌿',
            'description': 'Linseed oil + jute backing + cork dust + wood '
                           'flour — 100% bio-based. Naturally antimicrobial '
                           'due to ongoing linseed oxidation.',
            'best_for': ['General Patient Ward', 'Pediatric Ward',
                         'Outpatient Clinic', 'Administrative Offices',
                         'Physical Therapy', 'Reception & Waiting'],
            'use_case': 'Patient rooms, corridors, public areas — wherever '
                        'sustainability is prioritized over sterile-grade.',
            'eco_benefits': 'Carbon-negative manufacture, biodegradable, '
                            'no toxic off-gassing, 30+ yr lifespan.',
            'cost_per_m2_usd': 38,
            'reference': 'LEED v4 EQ 2 (Low-Emitting) + Cradle-to-Cradle Silver',
        },
        'High-Traffic Terrazzo': {
            'icon': '💎',
            'description': 'Cement + recycled glass/marble chips, ground '
                           'and polished. Hardest hospital floor — 100+ '
                           'year lifespan if maintained.',
            'best_for': ['Reception & Waiting', 'Cafeteria & Kitchen',
                         'Emergency Room (ER)', 'Outpatient Clinic',
                         'Radiology Department'],
            'use_case': 'Lobbies, dining, ER triage — high-traffic public '
                        'areas needing maximum durability and elegance.',
            'eco_benefits': 'Uses 60-70% recycled aggregate, no replacement '
                            'cycle, low embodied carbon over its lifespan.',
            'cost_per_m2_usd': 85,
            'reference': 'NTMA Standard + LEED v4 MR 1 (Building Reuse)',
        },
        'Rubber (Recycled)': {
            'icon': '⚫',
            'description': 'Sheet rubber from recycled tires + virgin '
                           'rubber binder. Shock-absorbing, slip-resistant, '
                           'naturally quiet underfoot.',
            'best_for': ['Pediatric Ward', 'Physical Therapy',
                         'Psychiatric Ward', 'NICU (Neonatal ICU)',
                         'Maternity & Delivery'],
            'use_case': 'Areas where falls are a major risk — pediatric, '
                        'geriatric, physical therapy, and behavioral '
                        'health where impact absorption is critical.',
            'eco_benefits': 'Diverts tires from landfill (~30 tires per '
                            '100 m²), recyclable again at end of life.',
            'cost_per_m2_usd': 52,
            'reference': 'ASTM F1859 (Static-Coefficient) + LEED v4 MR 3',
        },
        'Bamboo Engineered': {
            'icon': '🎋',
            'description': 'Strand-woven bamboo with low-VOC adhesive. '
                           'Harder than oak but renewable in 5 years (vs '
                           '60+ for hardwood). Warm aesthetic.',
            'best_for': ['Administrative Offices', 'Outpatient Clinic',
                         'Reception & Waiting', 'Physical Therapy',
                         'Cafeteria & Kitchen'],
            'use_case': 'Public-facing non-clinical areas where wood '
                        'warmth is desired. Not for wet/sterile zones.',
            'eco_benefits': 'Rapidly renewable (5-yr harvest cycle), '
                            'sequesters carbon, FSC-certified options.',
            'cost_per_m2_usd': 62,
            'reference': 'LEED v4 MR 2 (Rapidly Renewable) + FSC certification',
        },
        'Polished Concrete': {
            'icon': '🪨',
            'description': 'Existing concrete slab densified and polished '
                           'in place. No additional flooring material '
                           'needed — uses what\'s already there.',
            'best_for': ['Cafeteria & Kitchen', 'Sterilization (CSSD)',
                         'Pharmacy', 'Laboratory', 'Administrative Offices'],
            'use_case': 'Service/back-of-house areas, lab support, kitchens '
                        '— where exposed concrete is acceptable visually.',
            'eco_benefits': 'Zero added material, lowest embodied carbon, '
                            'eliminates entire flooring supply chain.',
            'cost_per_m2_usd': 25,
            'reference': 'ACI 310.1R + LEED v4 MR 1 (Building Reuse)',
        },
    },

    # ───────────── WALL MATERIAL ─────────────
    'wall': {
        'Antimicrobial Gypsum': {
            'icon': '🧱',
            'description': 'Gypsum board impregnated with copper-ion '
                           'biocide. Kills surface bacteria/viruses on '
                           'contact. Standard for healthcare partitions.',
            'best_for': ['Operating Room (OR)', 'Intensive Care Unit (ICU)',
                         'NICU (Neonatal ICU)', 'Emergency Room (ER)'],
            'use_case': 'Sterile and semi-sterile zones — primary partition '
                        'material in modern hospital construction.',
            'eco_benefits': 'Recycled paper facing (95%), low-VOC, '
                            'recyclable as construction waste.',
            'cost_per_m2_usd': 28,
            'reference': 'ASTM C1396 + EPA Antimicrobial Pesticide Reg.',
        },
        'Eco-Plaster (Lime-Hemp)': {
            'icon': '🌾',
            'description': 'Lime + hemp hurd + natural pigments. Breathable, '
                           'humidity-regulating, naturally antibacterial '
                           '(high pH). Sequesters CO₂ as it cures.',
            'best_for': ['General Patient Ward', 'Pediatric Ward',
                         'Maternity & Delivery', 'Psychiatric Ward',
                         'Outpatient Clinic'],
            'use_case': 'Patient bedrooms, calming therapeutic environments. '
                        'Reduces humidity-related infections naturally.',
            'eco_benefits': 'Carbon-negative (-30 kg CO₂/m³ over lifetime), '
                            'breathable, mold-resistant.',
            'cost_per_m2_usd': 42,
            'reference': 'Cradle-to-Cradle Gold + LEED v4 EQ 2.4',
        },
        'Glass-Reinforced Panel': {
            'icon': '🪟',
            'description': 'Fiberglass-reinforced plastic panel for wet/sterile '
                           'areas. Seamless, washable, impact-resistant.',
            'best_for': ['Sterilization (CSSD)', 'Laboratory', 'Pharmacy',
                         'Operating Room (OR)', 'Blood Bank'],
            'use_case': 'Wet labs, scrub-down areas, surgical wash zones — '
                        'where walls need pressure-washing daily.',
            'eco_benefits': '40-yr lifespan, no repainting cycle, recyclable '
                            'as composite material.',
            'cost_per_m2_usd': 65,
            'reference': 'ASTM E84 + FDA 21 CFR 174',
        },
        'Reclaimed Wood Panels': {
            'icon': '🪵',
            'description': 'FSC-certified reclaimed wood from demolished '
                           'buildings, treated with natural sealer. '
                           'Biophilic warmth.',
            'best_for': ['Reception & Waiting', 'Administrative Offices',
                         'Cafeteria & Kitchen', 'Outpatient Clinic'],
            'use_case': 'Public lobbies, accent walls, waiting rooms — '
                        'evidence-based design shows reduced patient '
                        'anxiety with biophilic materials.',
            'eco_benefits': 'Diverts demolition waste, captured carbon, '
                            'connects occupants to nature (biophilia).',
            'cost_per_m2_usd': 88,
            'reference': 'FSC + LEED v4 MR 2 (Salvaged Materials)',
        },
    },

    # ───────────── CEILING MATERIAL ─────────────
    'ceiling': {
        'Acoustic Mineral Fiber': {
            'icon': '🔇',
            'description': 'Suspended ceiling tiles made from rock wool '
                           '+ recycled cellulose. NRC > 0.85 — absorbs '
                           '85%+ of sound energy.',
            'best_for': ['General Patient Ward', 'NICU (Neonatal ICU)',
                         'Intensive Care Unit (ICU)', 'Reception & Waiting',
                         'Outpatient Clinic'],
            'use_case': 'Patient rest areas — HCAHPS survey scores improve '
                        'measurably when hospital noise drops below 45dB.',
            'eco_benefits': '75%+ recycled content, Cradle-to-Cradle Silver, '
                            'low-VOC.',
            'cost_per_m2_usd': 35,
            'reference': 'ASTM C423 (NRC) + FGI 2018 §1.2-5.6',
        },
        'Stretch Fabric Ceiling': {
            'icon': '🎨',
            'description': 'PVC-free polyester membrane stretched on '
                           'aluminum perimeter. Fully washable, mold-proof, '
                           'available in printable artwork.',
            'best_for': ['Pediatric Ward', 'Maternity & Delivery',
                         'Oncology', 'Reception & Waiting'],
            'use_case': 'Pediatric areas (printable sky/clouds reduce '
                        'anxiety in children) + spaces needing washable '
                        'continuous surfaces.',
            'eco_benefits': '25-yr warranty, no repainting, 100% PVC-free.',
            'cost_per_m2_usd': 75,
            'reference': 'CE EN 13501 (Fire) + REACH compliance',
        },
        'Exposed Cross-Laminated Timber': {
            'icon': '🌲',
            'description': 'CLT panels left exposed as ceiling/structure. '
                           'Sequesters carbon, structural + finish in one. '
                           'FSC-certified spruce/fir.',
            'best_for': ['Reception & Waiting', 'Cafeteria & Kitchen',
                         'Administrative Offices', 'Physical Therapy'],
            'use_case': 'Public areas where biophilic wood ceiling improves '
                        'occupant experience. Not for sterile zones.',
            'eco_benefits': 'Carbon-negative (1 m³ stores ~1 t CO₂), '
                            'fast assembly, structural + finish dual-use.',
            'cost_per_m2_usd': 110,
            'reference': 'APA PRG 320 + LEED v4 MR 6',
        },
        'Gypsum w/ LED Light Cove': {
            'icon': '💡',
            'description': 'Plain gypsum ceiling with integrated LED cove '
                           'lighting around perimeter. Indirect glare-free '
                           'illumination.',
            'best_for': ['Operating Room (OR)', 'Intensive Care Unit (ICU)',
                         'Radiology Department', 'CT Scan Room', 'MRI Room'],
            'use_case': 'Procedure rooms — indirect lighting reduces glare '
                        'on screens and surgical fields.',
            'eco_benefits': 'LED uses ~75% less energy than fluorescents; '
                            '50,000-hr lifespan (no relamping).',
            'cost_per_m2_usd': 48,
            'reference': 'Energy Star LED + IES RP-29 Healthcare Lighting',
        },
    },

    # ───────────── INSULATION ─────────────
    'insulation': {
        'Sheep Wool Batt': {
            'icon': '🐑',
            'description': 'Natural sheep wool batts. Self-regulates '
                           'humidity (absorbs 30% weight without losing '
                           'R-value). Naturally fire-resistant.',
            'best_for': ['General Patient Ward', 'NICU (Neonatal ICU)',
                         'Maternity & Delivery', 'Administrative Offices'],
            'use_case': 'Patient areas where indoor air quality is '
                        'paramount. Self-purifies formaldehyde from air.',
            'eco_benefits': 'Renewable annually (sheep shearing), absorbs '
                            'VOCs lifelong, biodegradable, low embodied '
                            'energy (~14 MJ/kg vs 60 for fiberglass).',
            'cost_per_m2_usd': 22,
            'reference': 'ASTM C518 (R-value) + Cradle-to-Cradle Gold',
        },
        'Recycled Denim/Cotton': {
            'icon': '👕',
            'description': 'Post-consumer denim shredded and treated with '
                           'borate (fire/pest retardant). Same R-value as '
                           'fiberglass without itch or formaldehyde.',
            'best_for': ['Pediatric Ward', 'Outpatient Clinic',
                         'Administrative Offices', 'Reception & Waiting'],
            'use_case': 'General partition insulation in non-critical '
                        'areas. Safer for occupants/installers (no glass '
                        'fibers).',
            'eco_benefits': '85% recycled content, fully recyclable, '
                            'diverts textile waste from landfill.',
            'cost_per_m2_usd': 18,
            'reference': 'ASTM C739 + LEED v4 MR 3 (Recycled Content)',
        },
        'Cork Board': {
            'icon': '🪵',
            'description': 'Compressed cork granules bonded with natural '
                           'resins. Excellent acoustic + thermal performance '
                           'in one material.',
            'best_for': ['NICU (Neonatal ICU)', 'Psychiatric Ward',
                         'Operating Room (OR)', 'Intensive Care Unit (ICU)'],
            'use_case': 'Areas needing combined sound + thermal isolation '
                        '— OR (sound + temp control), NICU (quiet sleep).',
            'eco_benefits': 'Harvested without harming the tree (9-yr '
                            'cycle), carbon-negative, biodegradable.',
            'cost_per_m2_usd': 35,
            'reference': 'ASTM E90 (STC) + LEED v4 MR 2',
        },
        'Aerogel Blanket (Premium)': {
            'icon': '✨',
            'description': 'Silica aerogel in flexible fiber blanket. '
                           'Highest R-value per inch of any insulation '
                           '(R-10/inch). Used by NASA.',
            'best_for': ['CT Scan Room', 'MRI Room', 'Operating Room (OR)',
                         'Sterilization (CSSD)'],
            'use_case': 'Tight spaces requiring max thermal performance — '
                        'narrow walls in equipment-dense rooms.',
            'eco_benefits': 'Tiny material volume = low embodied carbon; '
                            'enables thinner walls + larger usable space.',
            'cost_per_m2_usd': 95,
            'reference': 'ASTM C177 + ICC ES Report',
        },
    },

    # ───────────── PAINT ─────────────
    'paint': {
        'Zero-VOC Antimicrobial': {
            'icon': '🎨',
            'description': 'Water-based latex paint with embedded silver '
                           'or copper biocide. 0 g/L VOC. Continuously '
                           'kills bacteria/viruses on the surface.',
            'best_for': ['Operating Room (OR)', 'Intensive Care Unit (ICU)',
                         'NICU (Neonatal ICU)', 'Sterilization (CSSD)',
                         'Emergency Room (ER)', 'Blood Bank'],
            'use_case': 'Standard finish for all sterile + semi-sterile '
                        'walls. Reduces hospital-acquired infections (HAIs) '
                        'measurably (~20-30% on touch surfaces).',
            'eco_benefits': 'Zero VOC off-gassing, GREENGUARD Gold, no '
                            'air quality remediation needed after coating.',
            'cost_per_m2_usd': 8,
            'reference': 'EPA Method 24 + GREENGUARD Gold + EPA Reg.',
        },
        'Mineral Silicate (KEIM-style)': {
            'icon': '⛰️',
            'description': 'Potassium silicate mineral binder + natural '
                           'pigments. Chemically bonds to substrate '
                           '(petrifies). 50+ year lifespan.',
            'best_for': ['Reception & Waiting', 'Cafeteria & Kitchen',
                         'Administrative Offices', 'Outpatient Clinic'],
            'use_case': 'Long-lifespan finishes in non-clinical areas — '
                        'never needs repainting, extremely UV-stable.',
            'eco_benefits': 'No petrochemicals, mineral-based, allows wall '
                            'breathing (prevents mold), 5x lifespan of '
                            'standard latex.',
            'cost_per_m2_usd': 14,
            'reference': 'DIN 18363 + Cradle-to-Cradle Gold',
        },
        'Air-Purifying Photocatalytic': {
            'icon': '☀️',
            'description': 'Latex paint with TiO₂ nanoparticles. UV/visible '
                           'light activates surface → breaks down formaldehyde, '
                           'NOx, and odors into harmless compounds.',
            'best_for': ['Reception & Waiting', 'Cafeteria & Kitchen',
                         'Pediatric Ward', 'Outpatient Clinic',
                         'Administrative Offices'],
            'use_case': 'High-traffic public areas where ambient air quality '
                        'concerns exist. Actively cleans the air, not just '
                        'passive.',
            'eco_benefits': 'Reduces indoor formaldehyde by ~80%, NOx by '
                            '~50% under daylight, lasts as long as the paint.',
            'cost_per_m2_usd': 12,
            'reference': 'ISO 22197-1 (Air Purification) + LEED EQ Innovation',
        },
        'Lime Wash (Traditional)': {
            'icon': '🏛️',
            'description': 'Slaked lime + water + natural pigments. Has '
                           'naturally high pH which kills bacteria and '
                           'mold. Used since Roman times.',
            'best_for': ['Psychiatric Ward', 'Maternity & Delivery',
                         'General Patient Ward', 'Pediatric Ward'],
            'use_case': 'Calming environments where natural matte finish '
                        'is preferred. Excellent for moisture-rich rooms.',
            'eco_benefits': 'Reabsorbs CO₂ as it cures (carbonation '
                            'process makes it carbon-neutral), no '
                            'petrochemicals.',
            'cost_per_m2_usd': 10,
            'reference': 'EN 459-1 (Building Lime) + Cradle-to-Cradle',
        },
    },
}


def render_material_info_card(category, material_name):
    """Render an info card for the selected material in the sidebar.
    Uses st.sidebar.info / caption / metric pattern."""
    if category not in MATERIALS_DB or material_name not in MATERIALS_DB[category]:
        return
    m = MATERIALS_DB[category][material_name]

    with st.sidebar.expander(f"{m['icon']} About **{material_name}**",
                              expanded=False):
        st.markdown(f"**📋 What it is**")
        st.caption(m['description'])

        st.markdown(f"**🏥 Best for in hospital**")
        st.caption(m['use_case'])
        # Show specific departments
        depts_pretty = ", ".join(f"`{d}`" for d in m['best_for'][:4])
        if len(m['best_for']) > 4:
            depts_pretty += f", +{len(m['best_for']) - 4} more"
        st.caption(f"💡 Recommended: {depts_pretty}")

        st.markdown(f"**🌱 Eco benefits**")
        st.caption(m['eco_benefits'])

        # Cost + reference inline
        col_c, col_r = st.columns(2)
        with col_c:
            st.metric("💰 Cost", f"${m['cost_per_m2_usd']}/m²")
        with col_r:
            st.caption(f"📖 **Ref:**\n{m['reference']}")


def get_hazard_layout(width, height):
    """Return positions for hazard marker signs."""
    # Radiation markers: 1 sign per wall section + lead-lined wall indicators
    radiation_signs = [
        # Sign near main door + signs on perimeter walls
        (width * 0.15, 0.5),       # near main entrance
        (width * 0.85, 0.5),       # other side of main entrance
    ]
    # Lead-lined wall strips along all 4 walls (visible only when toggled)
    radiation_walls = {
        'thickness': 0.18,  # visual wall thickness for shielding indication
    }

    # Isolation room indicator: red diagonal border + biohazard sign
    # at a corner indicating an isolation pod within the room
    isolation_pods = [
        # A typical isolation room pod in one corner
        {'x': 0.4, 'y': height * 0.5, 'w': 2.5, 'h': 2.0},
    ]

    # Hazmat storage: 1 cabinet marker per room (typically on wall)
    hazmat_storage = [
        # Hazmat cabinet position (wall-mounted, with NFPA 704 diamond)
        (width * 0.92, height * 0.92),  # back-right corner
    ]

    return {
        'radiation_signs': radiation_signs,
        'radiation_walls': radiation_walls,
        'isolation_pods': isolation_pods,
        'hazmat_storage': hazmat_storage,
    }


def draw_hazards_2d(ax, layout, options, width, height):
    """Draw the 3 hazard marker types on the 2D blueprint."""
    # ── 1. Radiation Shielding (yellow trefoil + lead-lined walls) ──
    if options.get('radiation', False):
        # Yellow shielded-wall stripes along the perimeter (subtle)
        wall_t = layout['radiation_walls']['thickness']
        # Bottom inner stripe (yellow shielding indicator inside walls)
        for x, y, w, h in [
            (0, 0, width, wall_t),                       # bottom
            (0, height - wall_t, width, wall_t),         # top
            (0, 0, wall_t, height),                      # left
            (width - wall_t, 0, wall_t, height),         # right
        ]:
            ax.add_patch(_MRect((x, y), w, h,
                facecolor='#fff59d', edgecolor='#f57f17',
                linewidth=0.6, alpha=0.6, zorder=2,
                hatch='\\\\'))

        # Trefoil signs on the wall
        for sx, sy in layout['radiation_signs']:
            # Yellow triangular sign
            tri = np.array([[sx, sy + 0.18], [sx - 0.18, sy - 0.10],
                            [sx + 0.18, sy - 0.10]])
            ax.add_patch(Polygon(tri, closed=True,
                facecolor='#fdd835', edgecolor='#f57f17',
                linewidth=1.5, zorder=11))
            # Trefoil emoji (radiation hazard symbol)
            ax.text(sx, sy - 0.02, "☢", fontsize=12, ha='center',
                va='center', color='#bf360c', weight='bold', zorder=12)
            ax.text(sx, sy - 0.32, "RADIATION", color='#e65100',
                fontsize=4, ha='center', weight='bold', zorder=12,
                bbox=dict(boxstyle="round,pad=0.1",
                          facecolor='white', edgecolor='#f57f17',
                          linewidth=0.6))

    # ── 2. Isolation Pod (red diagonal border + biohazard) ──
    if options.get('isolation', False):
        for pod in layout['isolation_pods']:
            px, py, pw, ph = pod['x'], pod['y'], pod['w'], pod['h']
            # Striped red/yellow border (isolation hazard pattern)
            ax.add_patch(_MRect((px, py), pw, ph,
                facecolor='#ffebee', edgecolor='#c62828',
                linewidth=2.5, zorder=7,
                hatch='xx', alpha=0.55))
            # Biohazard symbol
            ax.text(px + pw/2, py + ph/2 + 0.2, "☣",
                fontsize=22, ha='center', va='center',
                color='#b71c1c', weight='bold', zorder=8)
            # Label
            ax.text(px + pw/2, py + ph/2 - 0.4,
                "ISOLATION\nNEG. PRESSURE",
                fontsize=5, color='#b71c1c', ha='center', va='center',
                weight='bold', zorder=8,
                bbox=dict(boxstyle="round,pad=0.15",
                          facecolor='white', edgecolor='#c62828',
                          linewidth=0.8))
            # Corner negative-pressure arrow indicator
            ax.annotate('', xy=(px + 0.2, py + 0.2),
                xytext=(px - 0.2, py - 0.2),
                arrowprops=dict(arrowstyle='->', color='#c62828',
                                 lw=1.2), zorder=9)

    # ── 3. Hazmat Storage (NFPA 704 diamond on wall) ──
    if options.get('hazmat', False):
        for hx, hy in layout['hazmat_storage']:
            # NFPA 704 diamond (rotated square)
            d = 0.22
            diamond_pts = np.array([
                [hx, hy + d],       # top   (red — flammability)
                [hx + d, hy],       # right (yellow — reactivity)
                [hx, hy - d],       # bot   (white — special)
                [hx - d, hy],       # left  (blue — health)
            ])
            # Background white diamond outline
            ax.add_patch(Polygon(diamond_pts, closed=True,
                facecolor='white', edgecolor='black',
                linewidth=1.8, zorder=11))
            # 4 colored quadrants (simplified NFPA 704)
            tri_top = np.array([[hx, hy + d], [hx + d*0.5, hy + d*0.5],
                                [hx - d*0.5, hy + d*0.5]])
            tri_right = np.array([[hx + d, hy], [hx + d*0.5, hy + d*0.5],
                                   [hx + d*0.5, hy - d*0.5]])
            tri_bot = np.array([[hx, hy - d], [hx - d*0.5, hy - d*0.5],
                                 [hx + d*0.5, hy - d*0.5]])
            tri_left = np.array([[hx - d, hy], [hx - d*0.5, hy + d*0.5],
                                  [hx - d*0.5, hy - d*0.5]])
            ax.add_patch(Polygon(tri_top, closed=True,
                facecolor='#e53935', edgecolor='black',
                linewidth=0.5, zorder=12))     # red — fire
            ax.add_patch(Polygon(tri_right, closed=True,
                facecolor='#fdd835', edgecolor='black',
                linewidth=0.5, zorder=12))     # yellow — reactivity
            ax.add_patch(Polygon(tri_bot, closed=True,
                facecolor='white', edgecolor='black',
                linewidth=0.5, zorder=12))     # white — special
            ax.add_patch(Polygon(tri_left, closed=True,
                facecolor='#1565c0', edgecolor='black',
                linewidth=0.5, zorder=12))     # blue — health
            # Numbers (typical hospital hazmat: health 3, fire 2, reactivity 1)
            ax.text(hx - d*0.55, hy, "3", fontsize=5,
                color='white', ha='center', va='center',
                weight='bold', zorder=13)
            ax.text(hx, hy + d*0.55, "2", fontsize=5,
                color='white', ha='center', va='center',
                weight='bold', zorder=13)
            ax.text(hx + d*0.55, hy, "1", fontsize=5,
                color='black', ha='center', va='center',
                weight='bold', zorder=13)
            # Label
            ax.text(hx, hy - 0.40, "HAZMAT", color='#37474f',
                fontsize=4, ha='center', weight='bold', zorder=12,
                bbox=dict(boxstyle="round,pad=0.1",
                          facecolor='white', edgecolor='black',
                          linewidth=0.6))


def draw_hazards_3d(fig, layout, options, width, height, wall_height=3.0):
    """Draw the 3 hazard marker types in 3D."""
    # ── 1. Radiation: yellow shielding strips on inner walls + signs ──
    if options.get('radiation', False):
        # Yellow shielding strips on bottom of all 4 walls
        shield_h = 0.4  # height of shielding indicator
        shield_z = 1.0   # mounted at ~1m height
        # Front (bottom) wall - skip due to main door, just stripes either side
        fig.add_trace(box_mesh_global(0, 0, shield_z,
            width * 0.30, 0.05, shield_h, '#fdd835', 0.8,
            'Radiation Shield'))
        fig.add_trace(box_mesh_global(width * 0.70, 0, shield_z,
            width * 0.30, 0.05, shield_h, '#fdd835', 0.8,
            'Radiation Shield'))
        # Back wall
        fig.add_trace(box_mesh_global(0, height - 0.05, shield_z,
            width, 0.05, shield_h, '#fdd835', 0.8, 'Radiation Shield'))
        # Side walls
        fig.add_trace(box_mesh_global(0, 0, shield_z,
            0.05, height, shield_h, '#fdd835', 0.8, 'Radiation Shield'))
        fig.add_trace(box_mesh_global(width - 0.05, 0, shield_z,
            0.05, height, shield_h, '#fdd835', 0.8, 'Radiation Shield'))

        # Trefoil signs (yellow panels) on the wall
        for sx, sy in layout['radiation_signs']:
            fig.add_trace(box_mesh_global(sx - 0.20, sy - 0.04, 1.6,
                0.40, 0.06, 0.40, '#fdd835', 1.0, 'Radiation Sign'))
            # Black trefoil mark
            fig.add_trace(box_mesh_global(sx - 0.10, sy - 0.045, 1.70,
                0.20, 0.005, 0.20, '#bf360c', 1.0, 'Trefoil'))

    # ── 2. Isolation Pod: red curtain/frame + biohazard sign ──
    if options.get('isolation', False):
        for pod in layout['isolation_pods']:
            px, py, pw, ph = pod['x'], pod['y'], pod['w'], pod['h']
            # Red-tinted curtain/wall frame
            curtain_h = 2.4
            # Front curtain
            fig.add_trace(box_mesh_global(px, py, 0,
                pw, 0.04, curtain_h, '#ef9a9a', 0.5,
                'Isolation Curtain'))
            # Back curtain
            fig.add_trace(box_mesh_global(px, py + ph - 0.04, 0,
                pw, 0.04, curtain_h, '#ef9a9a', 0.5,
                'Isolation Curtain'))
            # Sides
            fig.add_trace(box_mesh_global(px, py, 0,
                0.04, ph, curtain_h, '#ef9a9a', 0.5,
                'Isolation Curtain'))
            fig.add_trace(box_mesh_global(px + pw - 0.04, py, 0,
                0.04, ph, curtain_h, '#ef9a9a', 0.5,
                'Isolation Curtain'))
            # Biohazard sign on the curtain
            fig.add_trace(box_mesh_global(px + pw/2 - 0.20,
                py - 0.005, 1.8, 0.40, 0.01, 0.35,
                '#ffffff', 1.0, 'Biohazard Sign'))
            # Red triangle on sign
            fig.add_trace(box_mesh_global(px + pw/2 - 0.12,
                py - 0.01, 1.85, 0.24, 0.005, 0.25,
                '#c62828', 1.0, 'Biohazard Symbol'))

    # ── 3. Hazmat Storage: NFPA 704 diamond on wall ──
    if options.get('hazmat', False):
        for hx, hy in layout['hazmat_storage']:
            # Background white square (rotated to look like diamond)
            fig.add_trace(box_mesh_global(hx - 0.22, hy + 0.04, 1.4,
                0.44, 0.04, 0.44, '#ffffff', 1.0, 'Hazmat Sign Bg'))
            # 4 colored quadrants (simplified — small boxes)
            # Top (red - fire)
            fig.add_trace(box_mesh_global(hx - 0.10, hy + 0.03, 1.62,
                0.20, 0.01, 0.18, '#e53935', 1.0, 'NFPA Red'))
            # Bottom (white - special)
            fig.add_trace(box_mesh_global(hx - 0.10, hy + 0.03, 1.42,
                0.20, 0.01, 0.18, '#fafafa', 1.0, 'NFPA White'))
            # Left (blue - health)
            fig.add_trace(box_mesh_global(hx - 0.22, hy + 0.025, 1.50,
                0.10, 0.01, 0.20, '#1565c0', 1.0, 'NFPA Blue'))
            # Right (yellow - reactivity)
            fig.add_trace(box_mesh_global(hx + 0.12, hy + 0.025, 1.50,
                0.10, 0.01, 0.20, '#fdd835', 1.0, 'NFPA Yellow'))


# ════════════════════════════════════════════════════════════════
# 🎬 FIRE EVACUATION SIMULATION (the secret feature!)
# Animates occupants evacuating toward nearest exits.
# Uses NFPA SFPE walking speed (1.2 m/s) + RSET/ASET methodology.
# Source: NFPA 101 + SFPE Handbook of Fire Protection Engineering
# ════════════════════════════════════════════════════════════════
def simulate_fire_evacuation(width, height, fire_origin,
                              exit_positions, n_occupants=20,
                              n_frames=30, seed=42):
    """Run a fire evacuation simulation and return GIF bytes + metrics.
    fire_origin: (x, y) where fire starts
    exit_positions: list of (x, y) coords of all available exits
    Returns: (gif_bytes, metrics_dict)"""
    import matplotlib.pyplot as plt
    from matplotlib.patches import Rectangle, Circle, FancyBboxPatch
    rng = np.random.RandomState(seed)

    # NFPA SFPE constants
    walking_speed = 1.2          # m/s (typical adult)
    fire_growth_rate = 0.4       # m/s (t-squared growth, simplified)
    frame_duration_sec = 1.0     # each frame = 1 second of real time

    # Initialize occupant positions — scattered across the room
    occupants = []
    for i in range(n_occupants):
        ox = rng.uniform(width * 0.15, width * 0.85)
        oy = rng.uniform(height * 0.15, height * 0.85)
        occupants.append({'x': ox, 'y': oy, 'escaped': False,
                          'casualty': False, 'target_exit': None})

    # For each occupant, identify the NEAREST safe exit
    # (exits that are not on the same side as the fire)
    def nearest_exit(ox, oy):
        if not exit_positions:
            return None
        # Use Manhattan + small bonus for exits AWAY from fire
        best = None
        best_score = float('inf')
        fx, fy = fire_origin
        for ex, ey in exit_positions:
            dist = ((ex - ox)**2 + (ey - oy)**2) ** 0.5
            # Penalize exits close to fire (push occupants away from danger)
            fire_dist = ((ex - fx)**2 + (ey - fy)**2) ** 0.5
            score = dist - 0.3 * fire_dist  # prefer farther-from-fire exits
            if score < best_score:
                best_score = score
                best = (ex, ey)
        return best

    for o in occupants:
        o['target_exit'] = nearest_exit(o['x'], o['y'])

    # Generate frames
    from PIL import Image
    frames = []
    fire_radius = 0.3  # starting radius
    last_time = 0

    for frame_i in range(n_frames):
        t_sec = frame_i * frame_duration_sec
        fire_radius_now = 0.3 + fire_growth_rate * t_sec

        # Update each occupant's position
        for o in occupants:
            if o['escaped'] or o['casualty']:
                continue
            # Check if engulfed by fire
            fdist = ((o['x'] - fire_origin[0])**2 +
                     (o['y'] - fire_origin[1])**2) ** 0.5
            if fdist < fire_radius_now:
                o['casualty'] = True
                continue
            # Move toward target exit
            if o['target_exit']:
                ex, ey = o['target_exit']
                dx = ex - o['x']
                dy = ey - o['y']
                dist = (dx**2 + dy**2) ** 0.5
                step = walking_speed * frame_duration_sec
                if dist <= step + 0.3:  # arrived at exit (within 0.3m)
                    o['escaped'] = True
                else:
                    o['x'] += step * dx / dist
                    o['y'] += step * dy / dist

        # Build a frame plot
        fig, ax = plt.subplots(figsize=(9, 6.5))
        ax.set_xlim(-1, width + 1)
        ax.set_ylim(-1, height + 1)
        ax.set_facecolor('#f5f5dc')

        # Walls
        wt = 0.15
        for x, y, w, h in [(-wt, -wt, width + 2*wt, wt),
                            (-wt, height, width + 2*wt, wt),
                            (-wt, 0, wt, height),
                            (width, 0, wt, height)]:
            ax.add_patch(Rectangle((x, y), w, h,
                facecolor='#2c2c2c', zorder=3))

        # Exits (green)
        for ex, ey in exit_positions:
            ax.add_patch(Circle((ex, ey), 0.4,
                facecolor='#2e7d32', edgecolor='#1b5e20',
                linewidth=2, alpha=0.85, zorder=5))
            ax.text(ex, ey, "🚪", fontsize=14, ha='center',
                va='center', zorder=6)

        # Fire spread (glowing red circle, growing)
        # Outer glow
        ax.add_patch(Circle(fire_origin, fire_radius_now * 1.5,
            facecolor='#ffeb3b', edgecolor='none',
            alpha=0.30, zorder=4))
        # Mid glow
        ax.add_patch(Circle(fire_origin, fire_radius_now * 1.2,
            facecolor='#ff9800', edgecolor='none',
            alpha=0.55, zorder=4))
        # Core fire
        ax.add_patch(Circle(fire_origin, fire_radius_now,
            facecolor='#d32f2f', edgecolor='#b71c1c',
            linewidth=1, alpha=0.95, zorder=5))
        ax.text(fire_origin[0], fire_origin[1], "🔥",
            fontsize=18, ha='center', va='center', zorder=6)

        # Occupants
        n_escaped = 0
        n_casualty = 0
        for o in occupants:
            if o['escaped']:
                n_escaped += 1
                continue
            if o['casualty']:
                n_casualty += 1
                # X mark at last position
                ax.plot(o['x'], o['y'], marker='x',
                    color='#b71c1c', markersize=12, zorder=8,
                    markeredgewidth=2.5)
                continue
            # Active occupant — orange dot
            ax.add_patch(Circle((o['x'], o['y']), 0.18,
                facecolor='#fb8c00', edgecolor='#e65100',
                linewidth=1, zorder=8))
            # Trajectory line toward exit
            if o['target_exit']:
                ex, ey = o['target_exit']
                ax.plot([o['x'], ex], [o['y'], ey],
                    color='#fb8c00', alpha=0.35, linewidth=0.8,
                    linestyle=':', zorder=6)

        # Time + stats counter at top
        n_active = n_occupants - n_escaped - n_casualty
        ax.text(width / 2, height + 0.6,
            f"⏱️  T = {t_sec:.0f} sec     "
            f"✅ Escaped: {n_escaped}/{n_occupants}     "
            f"⚠️ Casualties: {n_casualty}     "
            f"👥 Active: {n_active}",
            fontsize=11, ha='center', weight='bold',
            color='#1b5e20', zorder=20,
            bbox=dict(boxstyle="round,pad=0.4",
                      facecolor='white', edgecolor='#1b5e20',
                      linewidth=1.5))

        ax.text(width / 2, -0.6,
            "🔥 Fire Evacuation Simulation — NFPA 101 / SFPE methodology",
            fontsize=8, ha='center', style='italic', color='#37474f')
        ax.set_aspect('equal')
        ax.axis('off')

        # Save to in-memory image
        buf = BytesIO()
        plt.savefig(buf, format='PNG', dpi=85,
                     bbox_inches='tight', facecolor='#f5f5dc')
        buf.seek(0)
        frames.append(Image.open(buf).copy())
        plt.close(fig)

        last_time = t_sec

        # Stop if everyone escaped or became casualty
        if n_escaped + n_casualty >= n_occupants:
            # Hold the final frame for 2 sec (2 extra frames)
            for _ in range(2):
                frames.append(frames[-1])
            break

    # Compile to GIF
    gif_buf = BytesIO()
    frames[0].save(gif_buf, format='GIF', append_images=frames[1:],
                    save_all=True, duration=600, loop=0,
                    optimize=True)
    gif_buf.seek(0)

    # Metrics
    n_escaped_final = sum(1 for o in occupants if o['escaped'])
    n_casualty_final = sum(1 for o in occupants if o['casualty'])
    evacuation_time = last_time
    # NFPA recommended max: 3 min (180s) for hospitals
    nfpa_compliant = (evacuation_time <= 180 and
                      n_casualty_final == 0)

    return gif_buf.getvalue(), {
        'evacuation_time_sec': evacuation_time,
        'n_escaped': n_escaped_final,
        'n_casualty': n_casualty_final,
        'n_total': n_occupants,
        'nfpa_compliant': nfpa_compliant,
        'n_frames': len(frames),
    }



# ────────────────────────────────────────────────────────────
# Smart placement: finds reasonable position for each equipment
# avoiding bed positions and other key zones
# ────────────────────────────────────────────────────────────
def get_equipment_positions(width, height, active_equipment):
    """Return [(eq_key, x, y), ...] for active equipment items.
    Positions avoid the center (where beds typically are) and prefer
    perimeter zones with adequate spacing."""
    if not active_equipment:
        return []
    # Candidate positions around the perimeter (avoiding door at bottom-center)
    candidates = [
        (width * 0.10, height * 0.30),  # left lower
        (width * 0.10, height * 0.55),  # left middle
        (width * 0.10, height * 0.80),  # left upper
        (width * 0.90, height * 0.30),  # right lower
        (width * 0.90, height * 0.55),  # right middle
        (width * 0.90, height * 0.80),  # right upper
        (width * 0.30, height * 0.92),  # top wall, left
        (width * 0.70, height * 0.92),  # top wall, right
        (width * 0.30, height * 0.08),  # bottom wall, left
        (width * 0.70, height * 0.08),  # bottom wall, right
        (width * 0.20, height * 0.20),  # corner
        (width * 0.80, height * 0.20),  # corner
        (width * 0.20, height * 0.80),  # corner
        (width * 0.80, height * 0.80),  # corner
    ]
    positions = []
    for i, eq_key in enumerate(active_equipment):
        if i < len(candidates):
            x, y = candidates[i]
            positions.append((eq_key, x, y))
    return positions


# ────────────────────────────────────────────────────────────
# Smart human position based on facility type (not collision-prone)
# ────────────────────────────────────────────────────────────
HUMAN_POSITIONS_REL = {
    'Operating Room (OR)':           (0.88, 0.50),   # right side, observing
    'Intensive Care Unit (ICU)':     (0.50, 0.50),   # at nursing station
    'Emergency Room (ER)':           (0.50, 0.50),   # at nursing station
    'NICU (Neonatal ICU)':           (0.50, 0.50),
    'General Patient Ward':          (0.50, 0.50),
    'Pediatric Ward':                (0.50, 0.50),
    'Radiology Department':          (0.88, 0.70),
    'MRI Room':                      (0.88, 0.70),   # control room side
    'CT Scan Room':                  (0.88, 0.70),
    'Laboratory':                    (0.20, 0.30),
    'Blood Bank':                    (0.50, 0.30),
    'Dialysis Unit':                 (0.20, 0.50),
    'Maternity & Delivery':          (0.88, 0.50),
    'Oncology':                      (0.88, 0.50),
    'Psychiatric Ward':              (0.88, 0.50),
    'Physical Therapy':              (0.20, 0.30),
    'Dental Clinic':                 (0.50, 0.30),
    'Outpatient Clinic':             (0.20, 0.50),
    'Pharmacy':                      (0.50, 0.30),   # patient side, in front
    'Sterilization (CSSD)':          (0.20, 0.50),
    'Cafeteria & Kitchen':           (0.50, 0.30),   # in dining area
    'Laundry':                       (0.20, 0.50),
    'Administrative Offices':        (0.20, 0.50),
    'Reception & Waiting':           (0.15, 0.30),   # behind desk
}


def get_smart_human_position(facility_type, width, height):
    """Get a smart (x, y) position for the human figure based on facility."""
    rel = HUMAN_POSITIONS_REL.get(facility_type, (0.88, 0.50))
    return rel[0] * width, rel[1] * height


# ────────────────────────────────────────────────────────────
# Layout dispatchers — central registry
# ────────────────────────────────────────────────────────────
CUSTOM_LAYOUTS_2D = {
    "Cafeteria & Kitchen": lambda ax, w, h, color, n, opts: layout_cafeteria_2d(ax, w, h, color, n_seats=n),
    "Pharmacy": lambda ax, w, h, color, n, opts: layout_pharmacy_2d(ax, w, h, color, n_counters=n),
    "Reception & Waiting": lambda ax, w, h, color, n, opts: layout_reception_2d(ax, w, h, color, n_seats=n),
    "General Patient Ward": lambda ax, w, h, color, n, opts: layout_ward_2d(ax, w, h, color, n, label_prefix="Bed", glass_partitions=opts.get('glass', False)),
    "Intensive Care Unit (ICU)": lambda ax, w, h, color, n, opts: layout_ward_2d(ax, w, h, color, n, label_prefix="ICU", glass_partitions=opts.get('glass', False)),
    "Emergency Room (ER)": lambda ax, w, h, color, n, opts: layout_ward_2d(ax, w, h, color, n, label_prefix="ER Bay", glass_partitions=opts.get('glass', False)),
    "Pediatric Ward": lambda ax, w, h, color, n, opts: layout_ward_2d(ax, w, h, color, n, label_prefix="Peds", glass_partitions=opts.get('glass', False)),
}

CUSTOM_LAYOUTS_3D = {
    "Cafeteria & Kitchen": lambda fig, w, h, color, wh, n, opts: layout_cafeteria_3d(fig, w, h, color, wh, n_seats=n),
    "Pharmacy": lambda fig, w, h, color, wh, n, opts: layout_pharmacy_3d(fig, w, h, color, wh, n_counters=n),
    "Reception & Waiting": lambda fig, w, h, color, wh, n, opts: layout_reception_3d(fig, w, h, color, wh, n_seats=n),
    "General Patient Ward": lambda fig, w, h, color, wh, n, opts: layout_ward_3d(fig, w, h, color, wh, n, glass_partitions=opts.get('glass', False)),
    "Intensive Care Unit (ICU)": lambda fig, w, h, color, wh, n, opts: layout_ward_3d(fig, w, h, color, wh, n, glass_partitions=opts.get('glass', False)),
    "Emergency Room (ER)": lambda fig, w, h, color, wh, n, opts: layout_ward_3d(fig, w, h, color, wh, n, glass_partitions=opts.get('glass', False)),
    "Pediatric Ward": lambda fig, w, h, color, wh, n, opts: layout_ward_3d(fig, w, h, color, wh, n, glass_partitions=opts.get('glass', False)),
}


# ============================================================
# PAGE CONFIG (must be first Streamlit command)
# ============================================================
st.set_page_config(
    page_title="Viridis - Green Hospital Planner",
    layout="wide",
    page_icon="🍃",
    initial_sidebar_state="collapsed"  # Collapsed during wizard
)

# ============================================================
# DARK ECO CSS — Professional SaaS-grade styling
# ============================================================
DARK_ECO_CSS = """
<style>
    /* ============ Base palette ============ */
    :root {
        --eco-bg-dark: #0d1f17;
        --eco-bg-mid: #142e22;
        --eco-bg-light: #1a3d2e;
        --eco-accent: #4ade80;
        --eco-accent-glow: #22c55e;
        --eco-accent-dark: #16a34a;
        --eco-text: #e8f5e9;
        --eco-text-dim: #94a3b8;
        --eco-card-bg: rgba(20, 46, 34, 0.6);
        --eco-card-border: rgba(74, 222, 128, 0.2);
        --eco-card-hover: rgba(74, 222, 128, 0.5);
    }

    /* ============ Main app background (wizard pages only) ============ */
    .wizard-active .stApp {
        background:
            radial-gradient(circle at 20% 30%, rgba(34, 197, 94, 0.08) 0%, transparent 50%),
            radial-gradient(circle at 80% 70%, rgba(74, 222, 128, 0.05) 0%, transparent 50%),
            linear-gradient(135deg, #0d1f17 0%, #142e22 50%, #0d1f17 100%);
        color: var(--eco-text);
        min-height: 100vh;
    }

    /* Hide default Streamlit padding on wizard pages */
    .wizard-active .block-container {
        padding-top: 1rem !important;
        max-width: 1100px;
    }

    /* ============ Animated leaf logo ============ */
    .viridis-logo {
        display: inline-block;
        font-size: 4.5rem;
        animation: leaf-sway 3s ease-in-out infinite;
        filter: drop-shadow(0 0 20px rgba(74, 222, 128, 0.5));
    }
    @keyframes leaf-sway {
        0%, 100% { transform: rotate(-5deg) scale(1); }
        50% { transform: rotate(5deg) scale(1.05); }
    }

    /* ============ Hero title ============ */
    .hero-title {
        font-size: 4rem;
        font-weight: 800;
        text-align: center;
        background: linear-gradient(135deg, #4ade80 0%, #22c55e 50%, #16a34a 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        margin: 0.5rem 0;
        letter-spacing: -2px;
        animation: fade-up 0.8s ease-out;
    }
    .hero-subtitle {
        font-size: 1.3rem;
        text-align: center;
        color: var(--eco-text-dim);
        margin-bottom: 2rem;
        font-weight: 300;
        animation: fade-up 1s ease-out;
    }
    @keyframes fade-up {
        from { opacity: 0; transform: translateY(20px); }
        to { opacity: 1; transform: translateY(0); }
    }

    /* ============ Feature cards (welcome page) ============ */
    .feature-card {
        background: var(--eco-card-bg);
        border: 1px solid var(--eco-card-border);
        border-radius: 16px;
        padding: 1.8rem 1.5rem;
        text-align: center;
        backdrop-filter: blur(10px);
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
        height: 100%;
        animation: fade-up 1.2s ease-out;
    }
    .feature-card:hover {
        border-color: var(--eco-card-hover);
        transform: translateY(-4px);
        box-shadow: 0 10px 40px rgba(74, 222, 128, 0.15);
    }
    .feature-icon {
        font-size: 2.5rem;
        margin-bottom: 0.8rem;
        display: block;
    }
    .feature-title {
        color: var(--eco-accent);
        font-size: 1.15rem;
        font-weight: 700;
        margin-bottom: 0.5rem;
    }
    .feature-desc {
        color: var(--eco-text-dim);
        font-size: 0.92rem;
        line-height: 1.5;
    }

    /* ============ Wizard progress bar ============ */
    .progress-container {
        max-width: 700px;
        margin: 0 auto 2rem auto;
    }
    .progress-track {
        display: flex;
        align-items: center;
        justify-content: space-between;
        margin-bottom: 0.5rem;
    }
    .progress-step {
        width: 36px;
        height: 36px;
        border-radius: 50%;
        background: var(--eco-bg-light);
        border: 2px solid rgba(74, 222, 128, 0.3);
        display: flex;
        align-items: center;
        justify-content: center;
        color: var(--eco-text-dim);
        font-weight: 700;
        font-size: 0.95rem;
        transition: all 0.4s ease;
        flex-shrink: 0;
        z-index: 2;
        position: relative;
    }
    .progress-step.active {
        background: var(--eco-accent);
        border-color: var(--eco-accent-glow);
        color: var(--eco-bg-dark);
        box-shadow: 0 0 20px rgba(74, 222, 128, 0.6);
        transform: scale(1.1);
    }
    .progress-step.complete {
        background: var(--eco-accent-dark);
        border-color: var(--eco-accent-dark);
        color: white;
    }
    .progress-line {
        flex: 1;
        height: 3px;
        background: var(--eco-bg-light);
        margin: 0 -2px;
        position: relative;
        overflow: hidden;
    }
    .progress-line.complete {
        background: var(--eco-accent-dark);
    }
    .progress-labels {
        display: flex;
        justify-content: space-between;
        font-size: 0.78rem;
        color: var(--eco-text-dim);
        margin-top: 0.4rem;
    }
    .progress-labels .label-active {
        color: var(--eco-accent);
        font-weight: 600;
    }

    /* ============ Step header ============ */
    .step-header {
        text-align: center;
        margin-bottom: 2rem;
        animation: fade-up 0.5s ease-out;
    }
    .step-number {
        display: inline-block;
        background: rgba(74, 222, 128, 0.15);
        color: var(--eco-accent);
        padding: 0.3rem 1rem;
        border-radius: 20px;
        font-size: 0.8rem;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 1px;
        margin-bottom: 0.8rem;
    }
    .step-title {
        font-size: 2.2rem;
        color: var(--eco-text);
        font-weight: 700;
        margin: 0.5rem 0;
    }
    .step-subtitle {
        color: var(--eco-text-dim);
        font-size: 1.05rem;
        max-width: 600px;
        margin: 0 auto;
    }

    /* ============ Mode selection cards (Step 1) ============ */
    .mode-card {
        background: var(--eco-card-bg);
        border: 2px solid var(--eco-card-border);
        border-radius: 20px;
        padding: 2.5rem 2rem;
        text-align: center;
        cursor: pointer;
        transition: all 0.3s ease;
        height: 100%;
    }
    .mode-card:hover {
        border-color: var(--eco-accent);
        transform: translateY(-6px) scale(1.02);
        box-shadow: 0 20px 60px rgba(74, 222, 128, 0.2);
    }
    .mode-card.selected {
        border-color: var(--eco-accent);
        background: rgba(74, 222, 128, 0.1);
        box-shadow: 0 0 30px rgba(74, 222, 128, 0.3);
    }
    .mode-icon {
        font-size: 4rem;
        display: block;
        margin-bottom: 1rem;
        filter: drop-shadow(0 0 15px rgba(74, 222, 128, 0.3));
    }
    .mode-title {
        font-size: 1.5rem;
        font-weight: 700;
        color: var(--eco-accent);
        margin-bottom: 0.5rem;
    }
    .mode-desc {
        color: var(--eco-text-dim);
        font-size: 0.95rem;
        line-height: 1.6;
    }
    .mode-features {
        list-style: none;
        padding: 0;
        margin-top: 1.2rem;
        text-align: left;
    }
    .mode-features li {
        color: var(--eco-text);
        font-size: 0.88rem;
        padding: 0.3rem 0;
        padding-left: 1.5rem;
        position: relative;
    }
    .mode-features li::before {
        content: '✓';
        color: var(--eco-accent);
        font-weight: 700;
        position: absolute;
        left: 0;
    }

    /* ============ Primary CTA button overrides ============ */
    .wizard-active .stButton > button {
        background: linear-gradient(135deg, var(--eco-accent-dark), var(--eco-accent));
        color: white;
        border: none;
        padding: 0.8rem 2rem;
        font-size: 1rem;
        font-weight: 700;
        border-radius: 12px;
        transition: all 0.3s ease;
        box-shadow: 0 4px 20px rgba(74, 222, 128, 0.3);
        letter-spacing: 0.3px;
    }
    .wizard-active .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 30px rgba(74, 222, 128, 0.5);
        background: linear-gradient(135deg, var(--eco-accent), var(--eco-accent-glow));
    }
    .wizard-active .stButton > button:active { transform: translateY(0); }

    /* Secondary buttons (Back) */
    .wizard-active .stButton button[kind="secondary"] {
        background: transparent;
        border: 1.5px solid var(--eco-card-border);
        color: var(--eco-text-dim);
        box-shadow: none;
    }
    .wizard-active .stButton button[kind="secondary"]:hover {
        border-color: var(--eco-accent);
        color: var(--eco-accent);
        background: rgba(74, 222, 128, 0.05);
    }

    /* ============ Form input styling on wizard ============ */
    .wizard-active .stSelectbox label,
    .wizard-active .stSlider label,
    .wizard-active .stRadio label,
    .wizard-active .stCheckbox label,
    .wizard-active .stNumberInput label {
        color: var(--eco-text) !important;
        font-weight: 600;
    }
    .wizard-active .stSelectbox > div > div,
    .wizard-active .stNumberInput > div > div > input {
        background: var(--eco-bg-light) !important;
        border: 1px solid var(--eco-card-border) !important;
        color: var(--eco-text) !important;
    }
    .wizard-active .stRadio > div { gap: 0.5rem; }
    .wizard-active div[data-baseweb="radio"] {
        background: var(--eco-card-bg);
        padding: 0.6rem 1rem;
        border-radius: 10px;
        border: 1px solid var(--eco-card-border);
        transition: all 0.2s ease;
    }
    .wizard-active div[data-baseweb="radio"]:hover {
        border-color: var(--eco-accent);
    }

    /* Wizard form sections (cards inside steps) */
    .wizard-section {
        background: var(--eco-card-bg);
        border: 1px solid var(--eco-card-border);
        border-radius: 14px;
        padding: 1.5rem;
        margin-bottom: 1rem;
        backdrop-filter: blur(10px);
        animation: fade-up 0.4s ease-out;
    }
    .section-label {
        color: var(--eco-accent);
        font-size: 0.85rem;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 1px;
        margin-bottom: 1rem;
        display: block;
    }

    /* Summary box at end */
    .summary-box {
        background: linear-gradient(135deg, rgba(74, 222, 128, 0.08), rgba(34, 197, 94, 0.05));
        border: 1px solid rgba(74, 222, 128, 0.3);
        border-radius: 16px;
        padding: 1.5rem;
        margin: 1rem 0;
    }
    .summary-item {
        display: flex;
        justify-content: space-between;
        padding: 0.5rem 0;
        border-bottom: 1px solid rgba(74, 222, 128, 0.1);
        color: var(--eco-text);
    }
    .summary-item:last-child { border-bottom: none; }
    .summary-key {
        color: var(--eco-text-dim);
        font-size: 0.92rem;
    }
    .summary-val {
        color: var(--eco-accent);
        font-weight: 600;
        font-size: 0.92rem;
    }

    /* Floating particles (subtle background ambience) */
    .particles {
        position: fixed;
        top: 0; left: 0;
        width: 100vw; height: 100vh;
        pointer-events: none;
        z-index: 0;
        overflow: hidden;
    }
    .particle {
        position: absolute;
        background: rgba(74, 222, 128, 0.4);
        border-radius: 50%;
        animation: float-up 12s linear infinite;
    }
    @keyframes float-up {
        0% { transform: translateY(100vh) translateX(0); opacity: 0; }
        10% { opacity: 0.6; }
        90% { opacity: 0.3; }
        100% { transform: translateY(-10vh) translateX(50px); opacity: 0; }
    }

    /* "Go Viridis" mega-button (Step 5) */
    .wizard-active .go-viridis-btn .stButton > button {
        background: linear-gradient(135deg, #16a34a, #22c55e, #4ade80);
        font-size: 1.4rem;
        padding: 1.2rem 3rem;
        border-radius: 16px;
        box-shadow: 0 10px 40px rgba(74, 222, 128, 0.5);
        animation: pulse-glow 2s ease-in-out infinite;
    }
    @keyframes pulse-glow {
        0%, 100% { box-shadow: 0 10px 40px rgba(74, 222, 128, 0.5); }
        50% { box-shadow: 0 10px 60px rgba(74, 222, 128, 0.8); }
    }

    /* Back-to-wizard pill (results page) */
    .back-pill {
        display: inline-block;
        background: rgba(74, 222, 128, 0.1);
        border: 1px solid rgba(74, 222, 128, 0.3);
        color: #16a34a;
        padding: 0.4rem 1rem;
        border-radius: 20px;
        font-size: 0.85rem;
        font-weight: 600;
        margin-bottom: 1rem;
    }
</style>
"""

st.markdown(DARK_ECO_CSS, unsafe_allow_html=True)

# ============================================================
# PAGE STATE INITIALIZATION
# ============================================================
if 'page' not in st.session_state:
    st.session_state.page = 'welcome'
if 'wizard_data' not in st.session_state:
    st.session_state.wizard_data = {}
if 'theme' not in st.session_state:
    st.session_state.theme = 'dark'  # 'dark' (Dark Eco) or 'light' (Medical White)
if 'lang' not in st.session_state:
    st.session_state.lang = 'en'  # 'en' or 'ar'
if 'reference' not in st.session_state:
    st.session_state.reference = 'fgi'  # 'fgi' or 'egyptian'
# v23: Empty-state on entry to Results — wait for user to click Generate.
# This flag flips True only when the Generate button is pressed; the
# Results page shows an educational dashboard until then, even if a
# valid `single_config` already exists from the wizard.
if 'design_generated' not in st.session_state:
    st.session_state.design_generated = False
if 'hospital_design_generated' not in st.session_state:
    st.session_state.hospital_design_generated = False

# Apply the active reference's overrides to MEDICAL_STANDARDS
# apply_reference_overrides()


# ════════════════════════════════════════════════════════════════
# BILINGUAL TRANSLATION SYSTEM
# Use t('key') to get the translated string in the active language.
# ════════════════════════════════════════════════════════════════
TRANSLATIONS = {
    # Welcome page
    'welcome_title': {'en': 'Welcome to Viridis', 'ar': 'مرحباً بك في Viridis'},
    'welcome_subtitle': {'en': 'Intelligent Green Medical Space Planner',
                          'ar': 'مخطط المساحات الطبية الذكي والصديق للبيئة'},
    'welcome_desc': {
        'en': 'Design eco-friendly hospitals automatically — from single rooms to full medical campuses. Get instant FGI-compliant blueprints, 3D renders, sustainability scores, and carbon savings reports.',
        'ar': 'صمم مستشفيات صديقة للبيئة تلقائياً — من الغرف المنفردة إلى المجمعات الطبية الكاملة. احصل على مخططات مطابقة لمعايير FGI، عروض ثلاثية الأبعاد، درجات استدامة، وتقارير توفير الكربون.'
    },
    'feature_depts': {'en': '24 Department Types', 'ar': '24 نوع قسم'},
    'feature_depts_desc': {
        'en': 'From Operating Rooms and NICU to Pharmacies and Cafeterias — covering every facility in a modern hospital.',
        'ar': 'من غرف العمليات وحضانات الأطفال إلى الصيدليات والمطاعم — تغطي كل منشأة في المستشفى الحديث.'
    },
    'feature_fgi': {'en': 'FGI Compliant Design', 'ar': 'تصميم متوافق مع FGI'},
    'feature_fgi_desc': {
        'en': 'Automated validation against Facility Guidelines Institute standards: minimum areas, ACH, pressure, and lighting.',
        'ar': 'تحقق تلقائي من معايير معهد إرشادات المنشآت: الحد الأدنى للمساحات، معدل تغيير الهواء، الضغط، والإضاءة.'
    },
    'feature_savings': {'en': 'Carbon & Cost Savings', 'ar': 'توفير الكربون والتكاليف'},
    'feature_savings_desc': {
        'en': 'Quantify environmental impact in real numbers — CO₂ tons, trees equivalent, energy MWh saved, dollars per year.',
        'ar': 'قياس التأثير البيئي بأرقام حقيقية — أطنان CO₂، ما يعادله من أشجار، ميجاوات/ساعة موفّرة، دولارات سنوياً.'
    },
    'start_designing': {'en': '🚀 Start Designing', 'ar': '🚀 ابدأ التصميم'},
    'footer_credit': {'en': 'Automated Eco-Friendly Hospital Design',
                       'ar': 'تصميم مستشفى أوتوماتيكي صديق للبيئة'},

    # Wizard steps
    'step_of': {'en': 'Step {} of {}', 'ar': 'الخطوة {} من {}'},
    'final_step': {'en': 'Step {} of {} · Final Step', 'ar': 'الخطوة {} من {} · الخطوة الأخيرة'},
    'back': {'en': '← Back', 'ar': '→ رجوع'},
    'next': {'en': 'Next →', 'ar': 'التالي ←'},
    'back_to_welcome': {'en': '← Back to Welcome', 'ar': '→ رجوع للترحيب'},
    'start_over': {'en': '🔄 Start Over', 'ar': '🔄 ابدأ من جديد'},
    'go_viridis': {'en': '🚀  Go Viridis!  🍃', 'ar': '🚀  انطلق مع Viridis!  🍃'},

    # Step 1
    'step1_title': {'en': 'Choose Your Design Scope', 'ar': 'اختر نطاق التصميم'},
    'step1_subtitle': {
        'en': 'Are you designing a single medical room/department in detail, or planning an entire hospital campus?',
        'ar': 'هل تريد تصميم غرفة/قسم طبي واحد بالتفصيل، أم تخطيط مجمع مستشفى كامل؟'
    },
    'mode_single': {'en': 'Single Department', 'ar': 'قسم واحد'},
    'mode_single_desc': {
        'en': 'Detailed 2D/3D blueprint of one medical space (OR, ICU, MRI, etc.)',
        'ar': 'مخطط تفصيلي ثنائي/ثلاثي الأبعاد لمساحة طبية واحدة (غرفة عمليات، عناية مركزة، أشعة، إلخ)'
    },
    'mode_hospital': {'en': 'Full Hospital', 'ar': 'مستشفى كامل'},
    'mode_hospital_desc': {
        'en': 'Complete hospital site plan with all departments, gardens & solar farm',
        'ar': 'مخطط موقع مستشفى كامل مع جميع الأقسام، الحدائق، ومزرعة الطاقة الشمسية'
    },
    'select_single': {'en': 'Select Single Department', 'ar': 'اختر قسم واحد'},
    'select_hospital': {'en': 'Select Full Hospital', 'ar': 'اختر مستشفى كامل'},

    # Step labels (progress)
    'lbl_mode': {'en': 'Mode', 'ar': 'النمط'},
    'lbl_facility': {'en': 'Facility', 'ar': 'المنشأة'},
    'lbl_spatial': {'en': 'Spatial', 'ar': 'المكاني'},
    'lbl_equipment': {'en': 'Equipment', 'ar': 'المعدات'},
    'lbl_sustainability': {'en': 'Sustainability', 'ar': 'الاستدامة'},

    # Sidebar
    'viridis_mode': {'en': '🏥 Viridis Mode', 'ar': '🏥 نمط Viridis'},
    'choose_scope': {'en': 'Choose Design Scope:', 'ar': 'اختر نطاق التصميم:'},
    'theme': {'en': '🎨 Theme', 'ar': '🎨 المظهر'},
    'theme_dark': {'en': '🌙 Dark', 'ar': '🌙 داكن'},
    'theme_light': {'en': '☀️ Light', 'ar': '☀️ فاتح'},
    'language': {'en': '🌐 Language', 'ar': '🌐 اللغة'},

    # Main title (results page)
    'app_title': {'en': '🍃 Viridis: Automated Eco-Friendly Hospital Design',
                   'ar': '🍃 Viridis: تصميم مستشفى أوتوماتيكي صديق للبيئة'},

    # Common UI strings
    'generate_design': {'en': '🚀 Generate Design', 'ar': '🚀 توليد التصميم'},
    'generate_hospital': {'en': '🚀 Generate Hospital Design',
                          'ar': '🚀 توليد تصميم المستشفى'},
    'back_to_wizard': {'en': '← Back to Wizard', 'ar': '→ رجوع للمعالج'},
    'lux_heatmap_title': {'en': '💡 Lux Levels Heatmap',
                          'ar': '💡 خريطة حرارية لمستويات الإضاءة'},
    'show_lux_heatmap': {'en': '💡 Show Lux Heatmap', 'ar': '💡 إظهار خريطة الإضاءة'},
    'pdf_export': {'en': '📑 Export PDF Report', 'ar': '📑 تصدير تقرير PDF'},
    'pdf_download': {'en': '⬇️ Download PDF Report', 'ar': '⬇️ تحميل تقرير PDF'},
}


def t(key, *args):
    """Translate a key. v24: forced English only — Arabic removed."""
    entry = TRANSLATIONS.get(key, {})
    txt = entry.get('en') or key
    if args:
        try: return txt.format(*args)
        except: return txt
    return txt


def go_to(page):
    """Navigate to another page."""
    st.session_state.page = page
    st.rerun()


# ============================================================
# MEDICAL STANDARDS DATABASE (FGI / WHO / ISO)
# Each department: minimum standards + base power consumption
# ============================================================
MEDICAL_STANDARDS = {
    # ---------- CRITICAL CARE ----------
    "Operating Room (OR)": {
        "category": "Critical Care", "icon": "🔪", "color": "#546e7a",
        "min_area_per_unit": 36, "min_ceiling": 3.0, "min_clearance": 1.2,
        "pressure": "+2.5 Pa", "ach": 25, "lux": 1000,
        "unit_name": "OR Table", "max_units": 1, "base_power": 1500,
        "water_usage": 80
    },
    "Intensive Care Unit (ICU)": {
        "category": "Critical Care", "icon": "💓", "color": "#1976d2",
        "min_area_per_unit": 18, "min_ceiling": 2.7, "min_clearance": 1.2,
        "pressure": "Neutral", "ach": 6, "lux": 300,
        "unit_name": "ICU Bed", "max_units": 8, "base_power": 1000,
        "water_usage": 60
    },
    "Emergency Room (ER)": {
        "category": "Critical Care", "icon": "🚑", "color": "#f57c00",
        "min_area_per_unit": 12, "min_ceiling": 2.7, "min_clearance": 1.2,
        "pressure": "Neutral", "ach": 12, "lux": 500,
        "unit_name": "ER Bed", "max_units": 10, "base_power": 950,
        "water_usage": 50
    },
    "NICU (Neonatal ICU)": {
        "category": "Critical Care", "icon": "👶", "color": "#ec407a",
        "min_area_per_unit": 11, "min_ceiling": 2.7, "min_clearance": 1.5,
        "pressure": "Positive", "ach": 6, "lux": 200,
        "unit_name": "Incubator", "max_units": 12, "base_power": 800,
        "water_usage": 30
    },

    # ---------- DIAGNOSTIC ----------
    "Radiology Department": {
        "category": "Diagnostic", "icon": "📡", "color": "#b71c1c",
        "min_area_per_unit": 25, "min_ceiling": 2.7, "min_clearance": 1.5,
        "pressure": "Negative", "ach": 6, "lux": 500,
        "unit_name": "X-Ray Unit", "max_units": 2, "base_power": 1800,
        "water_usage": 20
    },
    "MRI Room": {
        "category": "Diagnostic", "icon": "🧲", "color": "#4a148c",
        "min_area_per_unit": 40, "min_ceiling": 3.0, "min_clearance": 2.0,
        "pressure": "Neutral (Faraday Cage)", "ach": 8, "lux": 300,
        "unit_name": "MRI Scanner", "max_units": 1, "base_power": 3500,
        "water_usage": 30
    },
    "CT Scan Room": {
        "category": "Diagnostic", "icon": "🌀", "color": "#6a1b9a",
        "min_area_per_unit": 30, "min_ceiling": 2.7, "min_clearance": 1.5,
        "pressure": "Negative", "ach": 6, "lux": 400,
        "unit_name": "CT Scanner", "max_units": 1, "base_power": 2500,
        "water_usage": 20
    },
    "Laboratory": {
        "category": "Diagnostic", "icon": "🧪", "color": "#00695c",
        "min_area_per_unit": 8, "min_ceiling": 2.7, "min_clearance": 1.0,
        "pressure": "Negative", "ach": 10, "lux": 750,
        "unit_name": "Workbench", "max_units": 8, "base_power": 1200,
        "water_usage": 100
    },
    "Blood Bank": {
        "category": "Diagnostic", "icon": "🩸", "color": "#c62828",
        "min_area_per_unit": 6, "min_ceiling": 2.7, "min_clearance": 1.0,
        "pressure": "Positive", "ach": 8, "lux": 500,
        "unit_name": "Storage Refrigerator", "max_units": 6, "base_power": 1400,
        "water_usage": 10
    },

    # ---------- TREATMENT & SPECIALIZED ----------
    "Dialysis Unit": {
        "category": "Treatment", "icon": "💧", "color": "#0277bd",
        "min_area_per_unit": 9, "min_ceiling": 2.7, "min_clearance": 1.2,
        "pressure": "Neutral", "ach": 6, "lux": 300,
        "unit_name": "Dialysis Station", "max_units": 12, "base_power": 1100,
        "water_usage": 500
    },
    "Maternity & Delivery": {
        "category": "Specialized", "icon": "🤰", "color": "#f06292",
        "min_area_per_unit": 25, "min_ceiling": 2.7, "min_clearance": 1.5,
        "pressure": "Positive", "ach": 15, "lux": 500,
        "unit_name": "Delivery Bed", "max_units": 4, "base_power": 1100,
        "water_usage": 70
    },
    "Pediatric Ward": {
        "category": "Specialized", "icon": "🧸", "color": "#ffb300",
        "min_area_per_unit": 12, "min_ceiling": 2.7, "min_clearance": 1.2,
        "pressure": "Neutral", "ach": 6, "lux": 300,
        "unit_name": "Pediatric Bed", "max_units": 10, "base_power": 700,
        "water_usage": 40
    },
    "Oncology": {
        "category": "Specialized", "icon": "🎗️", "color": "#7b1fa2",
        "min_area_per_unit": 10, "min_ceiling": 2.7, "min_clearance": 1.2,
        "pressure": "Negative (Chemo)", "ach": 10, "lux": 400,
        "unit_name": "Chemo Chair", "max_units": 8, "base_power": 900,
        "water_usage": 40
    },
    "Psychiatric Ward": {
        "category": "Specialized", "icon": "🧠", "color": "#43a047",
        "min_area_per_unit": 14, "min_ceiling": 2.7, "min_clearance": 1.2,
        "pressure": "Neutral", "ach": 4, "lux": 500,
        "unit_name": "Safe Bed", "max_units": 8, "base_power": 600,
        "water_usage": 35
    },
    "Physical Therapy": {
        "category": "Treatment", "icon": "🦾", "color": "#26a69a",
        "min_area_per_unit": 10, "min_ceiling": 3.0, "min_clearance": 1.5,
        "pressure": "Neutral", "ach": 4, "lux": 400,
        "unit_name": "Therapy Station", "max_units": 8, "base_power": 500,
        "water_usage": 25
    },
    "Dental Clinic": {
        "category": "Treatment", "icon": "🦷", "color": "#00897b",
        "min_area_per_unit": 9, "min_ceiling": 2.7, "min_clearance": 1.0,
        "pressure": "Neutral", "ach": 8, "lux": 1500,
        "unit_name": "Dental Chair", "max_units": 4, "base_power": 700,
        "water_usage": 60
    },

    # ---------- GENERAL CARE ----------
    "General Patient Ward": {
        "category": "General Care", "icon": "🛏️", "color": "#5c6bc0",
        "min_area_per_unit": 10, "min_ceiling": 2.7, "min_clearance": 1.2,
        "pressure": "Neutral", "ach": 4, "lux": 300,
        "unit_name": "Patient Bed", "max_units": 12, "base_power": 500,
        "water_usage": 40
    },
    "Outpatient Clinic": {
        "category": "General Care", "icon": "🩺", "color": "#3949ab",
        "min_area_per_unit": 12, "min_ceiling": 2.7, "min_clearance": 1.0,
        "pressure": "Neutral", "ach": 4, "lux": 500,
        "unit_name": "Exam Table", "max_units": 6, "base_power": 450,
        "water_usage": 25
    },

    # ---------- SERVICE & SUPPORT ----------
    "Pharmacy": {
        "category": "Service", "icon": "💊", "color": "#388e3c",
        "min_area_per_unit": 4, "min_ceiling": 2.7, "min_clearance": 1.0,
        "pressure": "Positive", "ach": 4, "lux": 500,
        "unit_name": "Dispensing Counter", "max_units": 4, "base_power": 600,
        "water_usage": 15
    },
    "Sterilization (CSSD)": {
        "category": "Service", "icon": "♨️", "color": "#e64a19",
        "min_area_per_unit": 10, "min_ceiling": 2.7, "min_clearance": 1.2,
        "pressure": "Negative", "ach": 10, "lux": 500,
        "unit_name": "Autoclave", "max_units": 4, "base_power": 1800,
        "water_usage": 200
    },
    "Cafeteria & Kitchen": {
        "category": "Service", "icon": "🍽️", "color": "#fbc02d",
        "min_area_per_unit": 3, "min_ceiling": 3.0, "min_clearance": 0.8,
        "pressure": "Negative (Kitchen)", "ach": 15, "lux": 500,
        "unit_name": "Dining Seat", "max_units": 40, "base_power": 1300,
        "water_usage": 300
    },
    "Laundry": {
        "category": "Service", "icon": "🧺", "color": "#0288d1",
        "min_area_per_unit": 8, "min_ceiling": 3.0, "min_clearance": 1.2,
        "pressure": "Negative", "ach": 10, "lux": 400,
        "unit_name": "Washing Machine", "max_units": 6, "base_power": 2000,
        "water_usage": 800
    },
    "Administrative Offices": {
        "category": "Service", "icon": "💼", "color": "#757575",
        "min_area_per_unit": 6, "min_ceiling": 2.7, "min_clearance": 1.0,
        "pressure": "Neutral", "ach": 2, "lux": 500,
        "unit_name": "Office Desk", "max_units": 10, "base_power": 300,
        "water_usage": 10
    },
    "Reception & Waiting": {
        "category": "Service", "icon": "🪑", "color": "#8d6e63",
        "min_area_per_unit": 1.5, "min_ceiling": 3.0, "min_clearance": 1.0,
        "pressure": "Neutral", "ach": 4, "lux": 300,
        "unit_name": "Waiting Seat", "max_units": 30, "base_power": 350,
        "water_usage": 15
    },
}

# Category-based color scheme for Full Hospital site plan
CATEGORY_COLORS = {
    "Critical Care": "#e53935",
    "Diagnostic": "#5e35b1",
    "Treatment": "#00897b",
    "Specialized": "#d81b60",
    "General Care": "#3949ab",
    "Service": "#f57c00",
}

# ════════════════════════════════════════════════════════════════
# SCIENTIFIC REFERENCES & STANDARDS
# Two reference frameworks: US (FGI + ASHRAE) and Egyptian (MoH)
# Each provides citations, energy prices, EUI baselines, and any
# department-level overrides from the FGI baseline.
# ════════════════════════════════════════════════════════════════
SCIENTIFIC_REFERENCES = {
    'fgi': {
        'name': 'FGI 2018 + ASHRAE 170/90.1 (US)',
        'short': '🇺🇸 FGI / ASHRAE (US)',
        'flag': '🇺🇸',
        'currency': 'USD',
        'currency_symbol': '$',
        'electricity_price': 0.16,    # USD per kWh (EIA US avg 2024)
        'water_price': 4.50,           # USD per m³
        'eui_baseline': 380,           # kWh/m²/yr — average US hospital
        'eui_best': 180,               # kWh/m²/yr — LEED Platinum target
        'co2_per_kwh': 0.40,          # kg CO2 per kWh — US grid average
        # Department-specific overrides from baseline (none for FGI - it IS the baseline)
        'dept_overrides': {},
        'citations': {
            'min_area': 'FGI 2018 Table 2.1-1 (Healthcare Facility Guidelines)',
            'ach': 'ASHRAE 170-2017 Table 7-1 (Ventilation of Health Care Facilities)',
            'pressure': 'ASHRAE 170-2017 Table 7-1',
            'lux': 'IES RP-29-2017 (Illumination for Healthcare Facilities)',
            'eui': 'Energy Star Portfolio Manager — Hospitals (2023 benchmarks)',
            'co2': 'EPA eGRID 2022 — US national average emissions factor',
            'electricity_price': 'EIA Average US Commercial Electricity Price 2024',
            'water': 'AWWA Water and Wastewater Rate Survey 2023',
        },
    },
    'egyptian': {
        'name': 'Egyptian Hospital Code 2009 + MoH Guidelines',
        'short': '🇪🇬 Egyptian MoH',
        'flag': '🇪🇬',
        'currency': 'EGP',
        'currency_symbol': 'ج.م',
        'electricity_price': 2.50,    # EGP/kWh — commercial tariff 2024
        'water_price': 12.0,           # EGP/m³ — commercial
        'eui_baseline': 220,           # kWh/m²/yr — typical Egyptian hospital
        'eui_best': 130,               # kWh/m²/yr — green Egyptian target
        'co2_per_kwh': 0.55,           # kg CO2/kWh — Egypt grid (higher than US)
        # Egyptian standards are generally slightly more compact than FGI
        'dept_overrides': {
            'Operating Room (OR)': {
                'min_area_per_unit': 30, 'ach': 18, 'lux': 1000,
            },
            'Intensive Care Unit (ICU)': {
                'min_area_per_unit': 14, 'ach': 6, 'lux': 300,
            },
            'Emergency Room (ER)': {
                'min_area_per_unit': 9, 'ach': 10, 'lux': 500,
            },
            'NICU (Neonatal ICU)': {
                'min_area_per_unit': 9, 'ach': 6, 'lux': 200,
            },
            'General Patient Ward': {
                'min_area_per_unit': 8, 'ach': 4, 'lux': 300,
            },
            'Radiology Department': {
                'min_area_per_unit': 20, 'ach': 6, 'lux': 500,
            },
            'Pharmacy': {
                'min_area_per_unit': 3.5, 'ach': 4, 'lux': 500,
            },
        },
        'citations': {
            'min_area': 'الكود المصري لتصميم المستشفيات (2009)',
            'ach': 'وزارة الصحة المصرية — دليل تصميم المنشآت الصحية',
            'pressure': 'وزارة الصحة المصرية',
            'lux': 'الكود المصري للإضاءة الداخلية',
            'eui': 'الجهاز المركزي للتعبئة العامة والإحصاء — مؤشر استهلاك الطاقة',
            'co2': 'بيانات الشبكة الكهربائية المصرية 2023',
            'electricity_price': 'وزارة الكهرباء المصرية — تعريفة 2024 (2.50 ج.م/kWh)',
            'water': 'تعريفة شركة مياه الشرب التجارية 2024',
        },
    }
}

# Per-department EUI multipliers (energy-intensity vs hospital baseline).
# Sources: Energy Star Healthcare benchmarks + ASHRAE 90.1 typical loads.
DEPT_EUI_MULTIPLIERS = {
    'Operating Room (OR)':           2.0,   # surgical lights + HVAC + equipment
    'Intensive Care Unit (ICU)':     1.5,   # 24/7 monitoring + ventilators
    'Emergency Room (ER)':           1.3,
    'NICU (Neonatal ICU)':           1.4,
    'Radiology Department':          1.6,
    'MRI Room':                      3.5,   # superconducting magnet + cooling
    'CT Scan Room':                  2.5,   # high voltage X-ray
    'Laboratory':                    1.4,   # fume hoods + refrigeration
    'Blood Bank':                    1.3,   # 24/7 refrigeration
    'Dialysis Unit':                 1.6,   # pumps + water heating
    'Maternity & Delivery':          1.4,
    'Pediatric Ward':                1.0,
    'Oncology':                      1.3,
    'Psychiatric Ward':              0.9,
    'Physical Therapy':              1.0,
    'Dental Clinic':                 1.1,
    'General Patient Ward':          1.0,   # baseline
    'Outpatient Clinic':             0.8,
    'Pharmacy':                      0.9,
    'Sterilization (CSSD)':          2.2,   # autoclaves + steam
    'Cafeteria & Kitchen':           1.8,   # cooking equipment
    'Laundry':                       2.5,   # washers + dryers (biggest water user)
    'Administrative Offices':        0.4,   # low intensity
    'Reception & Waiting':           0.5,
}

# ════════════════════════════════════════════════════════════════
# CONSTRUCTION & EQUIPMENT COST DATABASE
# Per-department construction cost (per m²) + eco-feature costs
# Sources: RSMeans Healthcare Construction Cost Data 2024 (US)
#          + Egyptian Engineers Syndicate 2024 construction estimates
# ════════════════════════════════════════════════════════════════
CONSTRUCTION_COSTS = {
    'fgi': {  # US dollars per m²
        # Highly specialized rooms (HEPA, special HVAC, lead lining)
        'Operating Room (OR)':           4500,
        'NICU (Neonatal ICU)':           4000,
        'Sterilization (CSSD)':          3800,
        'MRI Room':                      6500,   # incl. magnetic shielding
        'CT Scan Room':                  4800,   # incl. lead shielding
        # Critical care
        'Intensive Care Unit (ICU)':     3500,
        'Emergency Room (ER)':           3000,
        # Diagnostic
        'Radiology Department':          3200,
        'Laboratory':                    3000,
        'Blood Bank':                    3500,
        # Treatment
        'Dialysis Unit':                 2800,
        'Maternity & Delivery':          3200,
        'Oncology':                      3500,
        # General clinical
        'General Patient Ward':          1800,
        'Pediatric Ward':                2200,
        'Psychiatric Ward':              2400,
        'Outpatient Clinic':             1800,
        'Dental Clinic':                 2200,
        'Physical Therapy':              1800,
        'Pharmacy':                      2400,
        # Service
        'Cafeteria & Kitchen':           2200,
        'Laundry':                       1600,
        'Administrative Offices':        1200,
        'Reception & Waiting':           1500,
    },
    'egyptian': {  # EGP per m² (typical Egyptian healthcare construction)
        'Operating Room (OR)':           45000,
        'NICU (Neonatal ICU)':           40000,
        'Sterilization (CSSD)':          38000,
        'MRI Room':                      65000,
        'CT Scan Room':                  48000,
        'Intensive Care Unit (ICU)':     35000,
        'Emergency Room (ER)':           30000,
        'Radiology Department':          32000,
        'Laboratory':                    28000,
        'Blood Bank':                    35000,
        'Dialysis Unit':                 28000,
        'Maternity & Delivery':          32000,
        'Oncology':                      35000,
        'General Patient Ward':          18000,
        'Pediatric Ward':                22000,
        'Psychiatric Ward':              22000,
        'Outpatient Clinic':             18000,
        'Dental Clinic':                 22000,
        'Physical Therapy':              18000,
        'Pharmacy':                      24000,
        'Cafeteria & Kitchen':           22000,
        'Laundry':                       15000,
        'Administrative Offices':        12000,
        'Reception & Waiting':           15000,
    },
}

# Cost of eco/sustainability features (in local currency, calculated by area)
ECO_FEATURE_COSTS = {
    'fgi': {  # USD
        'h_solar':       {'unit': 'per m² roof', 'cost': 380},   # PV ~ $380/m²
        'h_green_roof':  {'unit': 'per m²', 'cost': 200},
        'h_garden':      {'unit': 'per m² garden', 'cost': 120},
        'h_rainwater':   {'unit': 'flat', 'cost': 25000},
        'h_greywater':   {'unit': 'flat', 'cost': 45000},
        'h_ev_chargers': {'unit': 'per charger', 'cost': 3500},
        'h_smart_grid':  {'unit': 'flat', 'cost': 80000},
        'h_garage':      {'unit': 'per car space', 'cost': 12000},
        'laminar_ceiling':{'unit': 'per OR', 'cost': 75000},
        'partition_type':{'unit': 'flat', 'cost': 8000},  # glass walls
    },
    'egyptian': {  # EGP
        'h_solar':       {'unit': 'per m² roof', 'cost': 3500},
        'h_green_roof':  {'unit': 'per m²', 'cost': 1800},
        'h_garden':      {'unit': 'per m² garden', 'cost': 950},
        'h_rainwater':   {'unit': 'flat', 'cost': 200000},
        'h_greywater':   {'unit': 'flat', 'cost': 380000},
        'h_ev_chargers': {'unit': 'per charger', 'cost': 28000},
        'h_smart_grid':  {'unit': 'flat', 'cost': 650000},
        'h_garage':      {'unit': 'per car space', 'cost': 95000},
        'laminar_ceiling':{'unit': 'per OR', 'cost': 600000},
        'partition_type':{'unit': 'flat', 'cost': 65000},
    },
}

# Equipment cost (one-time purchase, per piece) — applied to active equipment
EQUIPMENT_COSTS_USD = {
    'eq_defib':       6000,
    'eq_ventilator':  25000,
    'eq_ultrasound':  35000,
    'eq_crashcart':   3500,
    'eq_surgical_light': 18000,
    'eq_warmer':      8000,
    'eq_infusion':    2500,
    'eq_endoscopy':   45000,
    'eq_ecg':         4000,
    'eq_suction':     1800,
    'eq_xray':        65000,
    'eq_music':       800,
    'eq_bairhugger':  3200,
    'eq_comm_board':  600,
    # Base equipment
    'add_monitor':    8500,
    'add_anesthesia': 35000,
    'add_iv_stand':   200,    # per stand
    'add_med_cart':   2500,
}


def calculate_construction_cost(facility_type, total_area, rooms_count):
    """Calculate total construction cost in the active reference's currency."""
    ref_key = st.session_state.get('reference', 'fgi')
    costs_table = CONSTRUCTION_COSTS.get(ref_key, CONSTRUCTION_COSTS['fgi'])
    cost_per_m2 = costs_table.get(facility_type, 2000)
    return cost_per_m2 * total_area * rooms_count


def calculate_equipment_cost(active_equipment, base_equipment,
                              units_per_room, rooms_count):
    """Sum all equipment costs (in USD — converted to local currency below)."""
    ref_key = st.session_state.get('reference', 'fgi')
    # Convert USD→EGP using ratio of electricity prices as approximation
    if ref_key == 'egyptian':
        usd_to_local = 2.50 / 0.16  # ~15.6× — close to actual FX rate
    else:
        usd_to_local = 1.0

    total = 0
    # Active extended equipment (per room)
    for eq_key in active_equipment:
        total += EQUIPMENT_COSTS_USD.get(eq_key, 0) * rooms_count
    # Base equipment
    for eq_key, enabled in base_equipment.items():
        if enabled:
            if eq_key == 'add_iv_stand':
                total += EQUIPMENT_COSTS_USD[eq_key] * units_per_room * rooms_count
            else:
                total += EQUIPMENT_COSTS_USD.get(eq_key, 0) * rooms_count
    return total * usd_to_local


def calculate_eco_features_cost(features_dict, total_area=0, n_chargers=4,
                                  n_or_rooms=0, garage_spaces=24):
    """Calculate cost of all eco features enabled."""
    ref_key = st.session_state.get('reference', 'fgi')
    costs = ECO_FEATURE_COSTS.get(ref_key, ECO_FEATURE_COSTS['fgi'])
    total = 0
    for feature, enabled in features_dict.items():
        if not enabled:
            continue
        cfg = costs.get(feature)
        if cfg is None:
            continue
        unit, c = cfg['unit'], cfg['cost']
        if unit == 'flat':
            total += c
        elif unit == 'per m² roof' or unit == 'per m²':
            total += c * total_area
        elif unit == 'per m² garden':
            total += c * total_area * 0.15  # ~15% of site is garden
        elif unit == 'per charger':
            total += c * n_chargers
        elif unit == 'per OR':
            total += c * n_or_rooms
        elif unit == 'per car space':
            total += c * garage_spaces
    return total


# ════════════════════════════════════════════════════════════════
# CLIMATE DATABASE — built-in (no internet dependency for defense)
# Sources: NASA POWER, ASHRAE Climate Data 2021, World Bank Climate
# Each city: climate zone, avg temp, peak sun hours (solar irradiance),
# annual rainfall, relative humidity
# ════════════════════════════════════════════════════════════════
CLIMATE_DATABASE = {
    # ───────────── EGYPTIAN CITIES 🇪🇬 ─────────────
    'Cairo':           {'zone': 'Hot', 'temp': 22, 'sun_hours': 6.4, 'rainfall': 25, 'humidity': 56, 'country': '🇪🇬'},
    'Alexandria':      {'zone': 'Moderate', 'temp': 20, 'sun_hours': 6.0, 'rainfall': 200, 'humidity': 69, 'country': '🇪🇬'},
    'Giza':            {'zone': 'Hot', 'temp': 22, 'sun_hours': 6.4, 'rainfall': 23, 'humidity': 55, 'country': '🇪🇬'},
    'Aswan':           {'zone': 'Hot', 'temp': 27, 'sun_hours': 7.2, 'rainfall': 1, 'humidity': 33, 'country': '🇪🇬'},
    'Luxor':           {'zone': 'Hot', 'temp': 26, 'sun_hours': 7.0, 'rainfall': 1, 'humidity': 36, 'country': '🇪🇬'},
    'Hurghada':        {'zone': 'Hot', 'temp': 24, 'sun_hours': 7.1, 'rainfall': 4, 'humidity': 48, 'country': '🇪🇬'},
    'Sharm El-Sheikh': {'zone': 'Hot', 'temp': 25, 'sun_hours': 7.3, 'rainfall': 5, 'humidity': 45, 'country': '🇪🇬'},
    'Mansoura':        {'zone': 'Moderate', 'temp': 21, 'sun_hours': 6.1, 'rainfall': 75, 'humidity': 64, 'country': '🇪🇬'},
    'Tanta':           {'zone': 'Moderate', 'temp': 21, 'sun_hours': 6.2, 'rainfall': 50, 'humidity': 62, 'country': '🇪🇬'},
    'Port Said':       {'zone': 'Moderate', 'temp': 21, 'sun_hours': 6.3, 'rainfall': 80, 'humidity': 70, 'country': '🇪🇬'},
    'Asyut':           {'zone': 'Hot', 'temp': 23, 'sun_hours': 6.8, 'rainfall': 5, 'humidity': 45, 'country': '🇪🇬'},
    'Suez':            {'zone': 'Hot', 'temp': 23, 'sun_hours': 6.7, 'rainfall': 20, 'humidity': 55, 'country': '🇪🇬'},
    # ───────────── GULF / MENA 🌍 ─────────────
    'Dubai':           {'zone': 'Hot', 'temp': 28, 'sun_hours': 6.9, 'rainfall': 100, 'humidity': 60, 'country': '🇦🇪'},
    'Abu Dhabi':       {'zone': 'Hot', 'temp': 28, 'sun_hours': 6.8, 'rainfall': 50, 'humidity': 58, 'country': '🇦🇪'},
    'Riyadh':          {'zone': 'Hot', 'temp': 26, 'sun_hours': 7.0, 'rainfall': 100, 'humidity': 32, 'country': '🇸🇦'},
    'Jeddah':          {'zone': 'Hot', 'temp': 29, 'sun_hours': 6.7, 'rainfall': 55, 'humidity': 60, 'country': '🇸🇦'},
    'Doha':            {'zone': 'Hot', 'temp': 28, 'sun_hours': 6.8, 'rainfall': 75, 'humidity': 58, 'country': '🇶🇦'},
    'Kuwait City':     {'zone': 'Hot', 'temp': 26, 'sun_hours': 6.9, 'rainfall': 110, 'humidity': 45, 'country': '🇰🇼'},
    'Amman':           {'zone': 'Moderate', 'temp': 18, 'sun_hours': 6.3, 'rainfall': 270, 'humidity': 52, 'country': '🇯🇴'},
    'Beirut':          {'zone': 'Moderate', 'temp': 21, 'sun_hours': 6.0, 'rainfall': 825, 'humidity': 70, 'country': '🇱🇧'},
    # ───────────── INTERNATIONAL 🌐 ─────────────
    'New York':        {'zone': 'Cold', 'temp': 13, 'sun_hours': 4.5, 'rainfall': 1200, 'humidity': 63, 'country': '🇺🇸'},
    'London':          {'zone': 'Cold', 'temp': 11, 'sun_hours': 2.8, 'rainfall': 600, 'humidity': 75, 'country': '🇬🇧'},
    'Los Angeles':     {'zone': 'Moderate', 'temp': 19, 'sun_hours': 5.6, 'rainfall': 380, 'humidity': 65, 'country': '🇺🇸'},
    'Berlin':          {'zone': 'Cold', 'temp': 10, 'sun_hours': 3.0, 'rainfall': 570, 'humidity': 72, 'country': '🇩🇪'},
    'Singapore':       {'zone': 'Hot', 'temp': 27, 'sun_hours': 4.5, 'rainfall': 2340, 'humidity': 84, 'country': '🇸🇬'},
    'Mumbai':          {'zone': 'Hot', 'temp': 27, 'sun_hours': 5.8, 'rainfall': 2200, 'humidity': 75, 'country': '🇮🇳'},
}


def detect_climate(city):
    """Return climate data dict for a city, default to Cairo if unknown."""
    return CLIMATE_DATABASE.get(city, CLIMATE_DATABASE['Cairo'])


def estimate_solar_potential(sun_hours, roof_area_m2, panel_efficiency=0.20):
    """Estimate annual solar generation (kWh/year).
    Formula: Area × efficiency × sun_hours × 365 × performance_ratio.
    Source: NREL PVWatts methodology (PR ~ 0.75)."""
    performance_ratio = 0.75
    # 1 m² of panel receives ~1 kW/m² peak; usable = area × eff
    daily_kwh = roof_area_m2 * panel_efficiency * sun_hours * performance_ratio
    return daily_kwh * 365


def estimate_rainwater_harvest(rainfall_mm, roof_area_m2, runoff_coeff=0.8):
    """Estimate annual rainwater harvest potential (liters/year).
    Formula: Rainfall(mm) × Area(m²) × runoff coefficient.
    1 mm rain on 1 m² = 1 liter."""
    return rainfall_mm * roof_area_m2 * runoff_coeff


# ════════════════════════════════════════════════════════════════
# LEED / ESTIDAMA / MOSTADAM CERTIFICATION PREDICTOR
# Maps design choices → green building certification credits
# Sources: LEED v4 Healthcare, Estidama Pearl Rating, Mostadam (KSA)
# ════════════════════════════════════════════════════════════════
LEED_CREDIT_MAP = {
    # credit_key: (category, max_points, description)
    'solar_pv':         ('Energy & Atmosphere', 10, 'On-site Renewable Energy'),
    'energy_star':      ('Energy & Atmosphere', 8, 'Optimize Energy Performance'),
    'smart_grid':       ('Energy & Atmosphere', 5, 'Advanced Energy Metering'),
    'natural_light':    ('Indoor Env. Quality', 6, 'Daylight & Quality Views'),
    'healing_garden':   ('Indoor Env. Quality', 4, 'Quality Views / Biophilia'),
    'rainwater':        ('Water Efficiency', 6, 'Rainwater Harvesting'),
    'greywater':        ('Water Efficiency', 5, 'Wastewater Reuse'),
    'eco_flooring':     ('Materials & Resources', 4, 'Low-Emitting Materials'),
    'waste_mgmt':       ('Materials & Resources', 5, 'Construction Waste Mgmt'),
    'green_roof':       ('Sustainable Sites', 5, 'Heat Island Reduction'),
    'ev_chargers':      ('Location & Transport', 4, 'Green Vehicles / EV'),
    'garden_site':      ('Sustainable Sites', 4, 'Open Space / Site Dev.'),
}


def predict_certification(choices):
    """Predict LEED level + Estidama Pearls + Mostadam from design choices.
    `choices` is a dict of boolean flags for each credit_key.
    Returns dict with scores and certification levels."""
    earned = 0
    max_possible = 0
    by_category = {}
    credits_detail = []

    for credit_key, (category, max_pts, desc) in LEED_CREDIT_MAP.items():
        max_possible += max_pts
        got = choices.get(credit_key, False)
        pts = max_pts if got else 0
        earned += pts
        by_category.setdefault(category, {'earned': 0, 'max': 0})
        by_category[category]['earned'] += pts
        by_category[category]['max'] += max_pts
        credits_detail.append({
            'credit': desc, 'category': category,
            'earned': pts, 'max': max_pts, 'achieved': got
        })

    # Scale to LEED 110-point system (our map is ~66 pts, scale up)
    leed_score = int((earned / max_possible) * 100) if max_possible else 0

    # LEED levels
    if leed_score >= 80:
        leed_level, leed_color = 'Platinum 🏆', '#e5e4e2'
    elif leed_score >= 60:
        leed_level, leed_color = 'Gold 🥇', '#ffd700'
    elif leed_score >= 50:
        leed_level, leed_color = 'Silver 🥈', '#c0c0c0'
    elif leed_score >= 40:
        leed_level, leed_color = 'Certified ✅', '#cd7f32'
    else:
        leed_level, leed_color = 'Not Certified ❌', '#e57373'

    # Estidama Pearl Rating (UAE) — 1 to 5 Pearls
    pearl_pct = earned / max_possible if max_possible else 0
    if pearl_pct >= 0.85: pearls = 5
    elif pearl_pct >= 0.70: pearls = 4
    elif pearl_pct >= 0.55: pearls = 3
    elif pearl_pct >= 0.40: pearls = 2
    else: pearls = 1

    # Mostadam (KSA) — similar tiering
    if pearl_pct >= 0.80: mostadam = 'Platinum'
    elif pearl_pct >= 0.65: mostadam = 'Gold'
    elif pearl_pct >= 0.50: mostadam = 'Silver'
    elif pearl_pct >= 0.35: mostadam = 'Green'
    else: mostadam = 'Not Rated'

    return {
        'leed_score': leed_score,
        'leed_level': leed_level,
        'leed_color': leed_color,
        'earned_points': earned,
        'max_points': max_possible,
        'pearls': pearls,
        'mostadam': mostadam,
        'by_category': by_category,
        'credits_detail': credits_detail,
    }


# ════════════════════════════════════════════════════════════════
# HVAC LOAD CALCULATOR
# ASHRAE-based cooling/heating load + airflow requirements
# Sources: ASHRAE Fundamentals 2021, ASHRAE 170 (ACH rates)
# ════════════════════════════════════════════════════════════════
def calculate_hvac_load(area_m2, ceiling_h, n_occupants, equipment_w,
                         lighting_w, climate, ach_required):
    """Cooling load estimate (ASHRAE simplified method).
    Returns load in kW, tons, BTU/h + required airflow."""
    # Envelope + solar gain by climate (W/m²)
    envelope_factor = {'Hot': 150, 'Moderate': 100, 'Cold': 70}.get(climate, 100)

    Q_envelope = area_m2 * envelope_factor       # building envelope + solar
    Q_people = n_occupants * 150                  # 100W sensible + 50W latent
    Q_lighting = lighting_w                        # lighting heat gain
    Q_equipment = equipment_w                      # medical equipment heat

    total_W = Q_envelope + Q_people + Q_lighting + Q_equipment
    total_kW = total_W / 1000
    tons = total_kW / 3.517                        # 1 ton = 3.517 kW
    btu_h = total_W * 3.412

    # Required ventilation airflow from ACH (ASHRAE 170)
    volume = area_m2 * ceiling_h                    # m³
    cmh = ach_required * volume                     # m³/h
    cfm = cmh * 0.5886                              # convert to CFM

    return {
        'total_kW': total_kW,
        'tons': tons,
        'btu_h': btu_h,
        'Q_envelope': Q_envelope,
        'Q_people': Q_people,
        'Q_lighting': Q_lighting,
        'Q_equipment': Q_equipment,
        'volume': volume,
        'cmh': cmh,
        'cfm': cfm,
        'ach': ach_required,
    }


# ════════════════════════════════════════════════════════════════
# FIRE SAFETY & EGRESS ANALYSIS
# NFPA 101 Life Safety Code + IBC egress requirements
# ════════════════════════════════════════════════════════════════
# Occupancy load factors (m² per person) — NFPA 101 Table 7.3.1.2
OCCUPANCY_FACTORS = {
    'Operating Room (OR)':           22.3,   # healthcare staff-only
    'Intensive Care Unit (ICU)':     22.3,
    'Emergency Room (ER)':           11.1,
    'NICU (Neonatal ICU)':           22.3,
    'General Patient Ward':          11.1,
    'Pediatric Ward':                11.1,
    'Reception & Waiting':           1.4,    # dense waiting area
    'Cafeteria & Kitchen':           1.4,    # assembly/dining
    'Outpatient Clinic':             9.3,
    'Administrative Offices':        9.3,
    'Pharmacy':                      9.3,
    'Laboratory':                    9.3,
    'Radiology Department':          11.1,
}


def calculate_fire_safety(area_m2, facility_type, stated_occupants=None):
    """NFPA 101 egress analysis. Returns exits, widths, compartments."""
    import math
    # Occupant load
    occ_factor = OCCUPANCY_FACTORS.get(facility_type, 11.1)
    calc_occupants = math.ceil(area_m2 / occ_factor)
    occupants = stated_occupants if stated_occupants else calc_occupants

    # Required number of exits (NFPA 101)
    if occupants > 1000:
        req_exits = 4
    elif occupants > 500:
        req_exits = 3
    elif occupants > 50:
        req_exits = 2
    else:
        req_exits = 1
    # Healthcare occupancies require minimum 2 exits
    req_exits = max(2, req_exits)

    # Egress width (NFPA 101: 5mm/person sprinklered, level egress)
    egress_factor_mm = 5  # mm per person (sprinklered)
    total_egress_mm = occupants * egress_factor_mm
    total_egress_m = total_egress_mm / 1000
    # Per-exit width
    per_exit_m = max(0.81, total_egress_m / req_exits)  # min 0.81m clear

    # Smoke compartments (healthcare: max 2,323 m² each, NFPA 101)
    n_smoke_compartments = max(1, math.ceil(area_m2 / 2323))

    # Max travel distance (sprinklered healthcare: 61m)
    max_travel_distance = 61

    # Fire extinguisher count (1 per 1000 sqft = ~93 m²)
    n_extinguishers = max(1, math.ceil(area_m2 / 93))

    return {
        'occupants': occupants,
        'calc_occupants': calc_occupants,
        'occ_factor': occ_factor,
        'req_exits': req_exits,
        'total_egress_m': total_egress_m,
        'per_exit_m': per_exit_m,
        'n_smoke_compartments': n_smoke_compartments,
        'max_travel_distance': max_travel_distance,
        'n_extinguishers': n_extinguishers,
    }


def build_water_sankey(total_water_lpm, has_rainwater, has_greywater,
                        rainfall_potential_lpm=0):
    """Build a Plotly Sankey diagram for hospital water balance.
    Returns a plotly Figure."""
    import plotly.graph_objects as go

    # Node labels
    labels = [
        'City Water',          # 0
        'Rainwater',           # 1
        'Total Supply',        # 2
        'Sanitary (40%)',      # 3
        'HVAC Cooling (25%)',  # 4
        'Process/Lab (20%)',   # 5
        'Irrigation (15%)',    # 6
        'Greywater Treatment', # 7
        'Recycled Water',      # 8
        'Sewage Out',          # 9
    ]
    node_colors = [
        '#1976d2', '#4fc3f7', '#0277bd', '#ef5350', '#42a5f5',
        '#7e57c2', '#66bb6a', '#26a69a', '#26c6da', '#8d6e63'
    ]

    # Flow values
    rain = rainfall_potential_lpm if has_rainwater else 0
    city = max(0, total_water_lpm - rain)

    sanitary = total_water_lpm * 0.40
    hvac = total_water_lpm * 0.25
    process = total_water_lpm * 0.20
    irrigation = total_water_lpm * 0.15

    sources, targets, values, link_colors = [], [], [], []

    def add_link(s, t, v, c):
        sources.append(s); targets.append(t); values.append(v)
        link_colors.append(c)

    # Supply → Total
    add_link(0, 2, city, 'rgba(25,118,210,0.4)')
    if rain > 0:
        add_link(1, 2, rain, 'rgba(79,195,247,0.4)')
    # Total → Uses
    add_link(2, 3, sanitary, 'rgba(239,83,80,0.3)')
    add_link(2, 4, hvac, 'rgba(66,165,245,0.3)')
    add_link(2, 5, process, 'rgba(126,87,194,0.3)')
    add_link(2, 6, irrigation, 'rgba(102,187,106,0.3)')
    # Greywater loop
    if has_greywater:
        recycled = (sanitary + hvac) * 0.45  # 45% recoverable
        add_link(3, 7, sanitary * 0.5, 'rgba(38,166,154,0.3)')
        add_link(4, 7, hvac * 0.4, 'rgba(38,166,154,0.3)')
        add_link(7, 8, recycled, 'rgba(38,198,218,0.4)')
        add_link(8, 6, recycled * 0.6, 'rgba(102,187,106,0.4)')  # to irrigation
        add_link(8, 4, recycled * 0.4, 'rgba(66,165,245,0.4)')   # back to HVAC
        add_link(3, 9, sanitary * 0.5, 'rgba(141,110,99,0.3)')   # rest to sewage
    else:
        add_link(3, 9, sanitary, 'rgba(141,110,99,0.3)')

    fig = go.Figure(data=[go.Sankey(
        node=dict(pad=15, thickness=20,
                  line=dict(color='black', width=0.5),
                  label=labels, color=node_colors),
        link=dict(source=sources, target=targets,
                  value=values, color=link_colors)
    )])
    fig.update_layout(
        title_text="💧 Hospital Water Balance (liters/month)",
        font_size=12, height=500,
        paper_bgcolor='rgba(0,0,0,0)'
    )
    return fig


def build_power_sld(has_solar, has_battery, total_load_kw,
                     has_critical=True):
    """Build a Power Single-Line Diagram (electrical schematic).
    Returns a matplotlib Figure."""
    fig, ax = plt.subplots(figsize=(11, 8))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis('off')

    def box(x, y, w, h, label, color='#e3f2fd', edge='#1565c0', fs=8):
        rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.05",
            facecolor=color, edgecolor=edge, linewidth=1.5, zorder=3)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2, label, ha='center', va='center',
            fontsize=fs, weight='bold', zorder=4, color='#0d47a1')

    def wire(x1, y1, x2, y2, color='#37474f', lw=2, style='-'):
        ax.plot([x1, x2], [y1, y2], color=color, linewidth=lw,
            linestyle=style, zorder=2)

    # Grid source (top)
    box(4, 9, 2, 0.8, "⚡ UTILITY GRID\n(National)", '#fff3e0', '#e65100', 9)

    # Main breaker
    wire(5, 9, 5, 8.5)
    box(4.3, 8, 1.4, 0.5, "Main Breaker", '#ffe0b2', '#e65100', 7)

    # Main Switchboard / Distribution
    wire(5, 8, 5, 7.3)
    box(3.5, 6.5, 3, 0.8, "MAIN SWITCHBOARD\n(MDB)", '#e3f2fd', '#1565c0', 9)

    # Solar PV (left, feeding into MDB)
    if has_solar:
        box(0.3, 8, 2, 0.8, "☀️ SOLAR PV\nArray", '#fff9c4', '#f9a825', 8)
        wire(2.3, 8.4, 3, 8.4, '#f9a825', 2)
        # Inverter
        box(2.5, 7.4, 1.2, 0.5, "Inverter", '#fff59d', '#f9a825', 7)
        wire(1.3, 8, 1.3, 7.65); wire(1.3, 7.65, 2.5, 7.65)
        wire(3.7, 7.65, 4, 7.65); wire(4, 7.65, 4, 7.3)

    # Battery storage
    if has_battery:
        box(0.3, 6.3, 2, 0.8, "🔋 BATTERY\nStorage", '#e8f5e9', '#2e7d32', 8)
        wire(2.3, 6.7, 3.5, 6.7, '#2e7d32', 2, '--')

    # Backup generator (right)
    box(7.7, 6.5, 2, 0.8, "🔌 BACKUP\nGenerator", '#fbe9e7', '#d84315', 8)
    wire(7.7, 6.9, 6.5, 6.9, '#d84315', 2, '--')

    # Distribution panels (bottom tier)
    wire(5, 6.5, 5, 5.8)
    wire(2.5, 5.8, 7.5, 5.8)  # horizontal bus

    panels = [
        (1.2, "Panel A\nLighting", '#e1f5fe'),
        (3.4, "Panel B\nHVAC", '#e1f5fe'),
        (5.6, "Panel C\nEquipment", '#e1f5fe'),
    ]
    for px, plabel, pcol in panels:
        wire(px + 0.6, 5.8, px + 0.6, 5.3)
        box(px, 4.5, 1.3, 0.8, plabel, pcol, '#0277bd', 7)

    # Critical loads with UPS (highlighted)
    if has_critical:
        wire(8.1, 5.8, 8.1, 5.3)
        box(7.4, 4.5, 1.5, 0.8, "🔴 UPS\nCritical Loads\n(OR/ICU)",
            '#ffebee', '#c62828', 7)
        ax.text(8.15, 4.2, "Uninterruptible", ha='center',
            fontsize=6, style='italic', color='#c62828')

    # Load labels at bottom
    for px, plabel, pcol in panels:
        wire(px + 0.6, 4.5, px + 0.6, 4.0)
        ax.plot(px + 0.6, 3.9, 'v', color='#37474f', markersize=8)

    # Title + total load
    ax.text(5, 9.7, "⚡ ELECTRICAL SINGLE-LINE DIAGRAM (SLD)",
        ha='center', fontsize=13, weight='bold', color='#0d47a1')
    ax.text(5, 0.5, f"Estimated Connected Load: {total_load_kw:.0f} kW  |  "
        f"NFPA 99 Healthcare Electrical System",
        ha='center', fontsize=9, style='italic', color='#546e7a')

    # Legend
    ax.plot([0.3, 0.8], [1.5, 1.5], color='#37474f', lw=2)
    ax.text(0.9, 1.5, "Normal Power", fontsize=7, va='center')
    ax.plot([0.3, 0.8], [1.1, 1.1], color='#37474f', lw=2, linestyle='--')
    ax.text(0.9, 1.1, "Backup/Emergency Power", fontsize=7, va='center')

    plt.tight_layout()
    return fig


# ════════════════════════════════════════════════════════════════
# BILL OF QUANTITIES (BoQ) GENERATOR
# Itemized list of construction, equipment, and eco features
# with quantities, unit rates, and line totals. Exportable.
# ════════════════════════════════════════════════════════════════
def build_boq_rows(mode, config):
    """Build Bill of Quantities line items.
    mode='single' or 'hospital'. Returns list of dicts (rows)."""
    ref = get_active_reference()
    ref_key = st.session_state.get('reference', 'fgi')
    curr = ref['currency_symbol']
    rows = []
    item_no = 1

    def add_row(category, desc, unit, qty, rate):
        nonlocal item_no
        rows.append({
            'No.': item_no,
            'Category': category,
            'Description': desc,
            'Unit': unit,
            'Qty': round(qty, 1),
            f'Rate ({curr})': round(rate, 2),
            f'Amount ({curr})': round(qty * rate, 2),
        })
        item_no += 1

    if mode == 'single':
        facility = config['facility_type']
        area = config['total_area']
        rooms = config['rooms_count']
        units = config['units_per_room']
        active_eq = config.get('active_equipment', [])

        # ─── Construction ───
        cost_m2 = CONSTRUCTION_COSTS.get(ref_key, {}).get(facility, 2000)
        add_row('Construction', f'{facility} — built area', 'm²',
                area * rooms, cost_m2)
        # Finishes
        floor = config.get('flooring', '')
        floor_rate = {'Antimicrobial Vinyl (Seamless)': 45,
                      'Eco-Linoleum (Natural)': 38,
                      'High-Traffic Terrazzo': 85}.get(floor, 45)
        if ref_key == 'egyptian': floor_rate *= 15.6
        add_row('Finishes', f'Flooring: {floor}', 'm²', area * rooms, floor_rate)

        # ─── HVAC (from cooling load) ───
        ach = config['std'].get('ach', 6)
        hvac_rate = 1200 if ref_key == 'fgi' else 18000  # per ton installed
        # rough tonnage
        est_tons = (area * rooms * 0.15)
        add_row('MEP — HVAC', f'HVAC system ({ach} ACH, medical grade)',
                'ton', est_tons, hvac_rate)

        # ─── Equipment ───
        usd_to_local = (2.50/0.16) if ref_key == 'egyptian' else 1.0
        base_eq = {
            'add_monitor': ('Vital Signs Monitor', config.get('add_monitor')),
            'add_anesthesia': ('Anesthesia Machine', config.get('add_anesthesia')),
            'add_iv_stand': ('IV Stand', config.get('add_iv_stand')),
            'add_med_cart': ('Medication Cart', config.get('add_med_cart')),
        }
        for k, (nm, enabled) in base_eq.items():
            if enabled:
                qty = units * rooms if k == 'add_iv_stand' else rooms
                add_row('Medical Equipment', nm, 'unit', qty,
                        EQUIPMENT_COSTS_USD.get(k, 0) * usd_to_local)
        for eq_key in active_eq:
            eq = next((e for e in EXTENDED_EQUIPMENT if e['key'] == eq_key), None)
            if eq:
                add_row('Medical Equipment', eq['label_en'], 'unit', rooms,
                        EQUIPMENT_COSTS_USD.get(eq_key, 0) * usd_to_local)

        # ─── Special features ───
        if config.get('laminar_ceiling'):
            lam_rate = ECO_FEATURE_COSTS[ref_key]['laminar_ceiling']['cost']
            add_row('Specialized', 'Laminar Flow Ceiling System', 'room',
                    rooms, lam_rate)
        if config.get('partition_type') == 'Antibacterial Glass Walls':
            add_row('Specialized', 'Antibacterial Glass Partitions', 'set',
                    rooms, ECO_FEATURE_COSTS[ref_key]['partition_type']['cost'])

    else:  # hospital mode
        depts = config['selected_depts']
        scale = config.get('scale', 1.0)
        # ─── Construction per department ───
        for dname, dinfo in depts.items():
            est_units = max(1, int(dinfo['max_units'] * 0.6))
            d_area = est_units * dinfo['min_area_per_unit'] * 1.4 * scale
            cost_m2 = CONSTRUCTION_COSTS.get(ref_key, {}).get(dname, 2000)
            add_row('Construction', f'{dname}', 'm²', d_area, cost_m2)

        # ─── Site-wide eco features ───
        total_area = sum(
            max(1, int(d['max_units']*0.6)) * d['min_area_per_unit'] * 1.4 * scale
            for d in depts.values())
        eco_items = [
            ('h_solar', 'Rooftop Solar PV', 'm²', total_area*0.4),
            ('h_green_roof', 'Green Roof', 'm²', total_area*0.3),
            ('h_garden', 'Healing Garden', 'm²', total_area*0.15),
            ('h_rainwater', 'Rainwater Harvesting System', 'system', 1),
            ('h_greywater', 'Greywater Recycling Plant', 'system', 1),
            ('h_smart_grid', 'Smart Grid + Battery', 'system', 1),
        ]
        for fkey, nm, unit, qty in eco_items:
            if config.get(fkey):
                cfg = ECO_FEATURE_COSTS[ref_key].get(fkey, {'cost': 0})
                add_row('Sustainability', nm, unit, qty, cfg['cost'])

    return rows


def build_functional_program_md(mode, config):
    """Build a Functional Program / Room Data Sheets document in Markdown.
    Returns markdown string (always works as fallback)."""
    ref = get_active_reference()
    from datetime import datetime as _dt
    lines = []
    lines.append("# FUNCTIONAL PROGRAM DOCUMENT")
    lines.append("## Viridis Eco-Hospital Design")
    lines.append(f"\n**Generated:** {_dt.now().strftime('%B %d, %Y')}  ")
    lines.append(f"**Standards Reference:** {ref['name']}  ")
    lines.append(f"**Currency:** {ref['currency']}\n")
    lines.append("---\n")

    def room_data_sheet(facility, std, area, units, rooms):
        s = []
        s.append(f"### {std.get('icon','🏥')} {facility}\n")
        s.append("| Parameter | Specification | Source |")
        s.append("|-----------|---------------|--------|")
        s.append(f"| Category | {std.get('category','—')} | — |")
        s.append(f"| Total Area | {area * rooms:.0f} m² ({rooms} room(s)) | Design |")
        s.append(f"| Units per Room | {units} × {std.get('unit_name','unit')} | Design |")
        s.append(f"| Min Area/Unit | {std.get('min_area_per_unit','—')} m² | {ref['citations']['min_area']} |")
        s.append(f"| Ceiling Height | {std.get('min_ceiling','—')} m | {ref['citations']['min_area']} |")
        s.append(f"| Air Changes/Hour | {std.get('ach','—')} ACH | {ref['citations']['ach']} |")
        s.append(f"| Pressure | {std.get('pressure','—')} | {ref['citations']['pressure']} |")
        s.append(f"| Illumination | {std.get('lux','—')} lux | {ref['citations']['lux']} |")
        s.append(f"| Power Load | {std.get('base_power','—')} W | EUI-based |")
        s.append(f"| Water Usage | {std.get('water_usage','—')} L/unit/day | Design |")
        s.append("")
        return "\n".join(s)

    if mode == 'single':
        facility = config['facility_type']
        std = config['std']
        lines.append("## 1. Project Overview\n")
        lines.append(f"This document defines the functional and technical "
                     f"requirements for a **{facility}** department designed "
                     f"to {ref['short']} standards.\n")
        lines.append("## 2. Room Data Sheet\n")
        lines.append(room_data_sheet(facility, std, config['total_area'],
                                     config['units_per_room'],
                                     config['rooms_count']))
        # Equipment list
        lines.append("## 3. Equipment Schedule\n")
        active_eq = config.get('active_equipment', [])
        if active_eq or config.get('add_monitor'):
            lines.append("| Equipment | Qty | Power (W) |")
            lines.append("|-----------|-----|-----------|")
            if config.get('add_monitor'):
                lines.append(f"| Vital Signs Monitor | {config['rooms_count']} | 200 |")
            if config.get('add_anesthesia'):
                lines.append(f"| Anesthesia Machine | {config['rooms_count']} | 500 |")
            for eq_key in active_eq:
                eq = next((e for e in EXTENDED_EQUIPMENT if e['key']==eq_key), None)
                if eq:
                    lines.append(f"| {eq['label_en']} | {config['rooms_count']} "
                                 f"| {eq['power_w']} |")
            lines.append("")
        else:
            lines.append("*No specialized equipment specified.*\n")
        # Adjacency
        lines.append("## 4. Adjacency Requirements\n")
        adj = {
            'Operating Room (OR)': 'Adjacent to: Sterilization (CSSD), ICU, '
                'Recovery. Direct access to sterile corridor.',
            'Emergency Room (ER)': 'Adjacent to: Radiology, Ambulance entry, '
                'ICU. Ground-floor with direct external access.',
            'Intensive Care Unit (ICU)': 'Adjacent to: OR, Emergency, '
                'Radiology. Central nursing station required.',
            'Radiology Department': 'Adjacent to: ER, ICU. Lead-lined walls. '
                'Ground or basement level preferred.',
        }
        lines.append(adj.get(facility,
            'Standard adjacency to circulation corridors and support services.'))
        lines.append("")

    else:  # hospital
        depts = config['selected_depts']
        lines.append("## 1. Project Overview\n")
        lines.append(f"Functional program for a **{config.get('hospital_scale','hospital')}** "
                     f"comprising **{len(depts)} departments**, designed to "
                     f"{ref['short']} standards with sustainability features.\n")
        lines.append("## 2. Department Schedule\n")
        lines.append("| # | Department | Category | Units | Area (m²) | ACH | Pressure |")
        lines.append("|---|-----------|----------|-------|-----------|-----|----------|")
        for i, (dname, dinfo) in enumerate(depts.items(), 1):
            est_units = max(1, int(dinfo['max_units']*0.6))
            d_area = est_units * dinfo['min_area_per_unit'] * 1.4 * config.get('scale',1)
            lines.append(f"| {i} | {dname} | {dinfo.get('category','—')} | "
                         f"{est_units} | {d_area:.0f} | {dinfo.get('ach','—')} | "
                         f"{dinfo.get('pressure','—')} |")
        lines.append("")
        lines.append("## 3. Room Data Sheets\n")
        for dname, dinfo in depts.items():
            est_units = max(1, int(dinfo['max_units']*0.6))
            d_area = est_units * dinfo['min_area_per_unit'] * 1.4 * config.get('scale',1)
            lines.append(room_data_sheet(dname, dinfo, d_area, est_units, 1))

    lines.append("\n---\n*Generated by Viridis — Automated Eco-Friendly "
                 "Hospital Design*")
    return "\n".join(lines)


# ════════════════════════════════════════════════════════════════
# 🔍 ANOMALY DETECTION
# Rule-based checks for unrealistic or non-compliant design choices
# ════════════════════════════════════════════════════════════════
def detect_anomalies(mode, config):
    """Scan a design for issues. Returns list of (severity, message) tuples.
    severity ∈ {'error', 'warning', 'info'}."""
    ref = get_active_reference()
    anomalies = []

    if mode == 'single':
        facility = config['facility_type']
        std = config['std']
        area = config['total_area']
        units = config['units_per_room']
        area_per_unit = area / max(1, units)

        # 1. Area below minimum
        min_req = std.get('min_area_per_unit', 0)
        if area_per_unit < min_req:
            anomalies.append(('error',
                f"🚨 **Area too small:** {area_per_unit:.0f} m²/unit is below "
                f"the {min_req} m² minimum for {facility} "
                f"({ref['short']}). Increase area or reduce units."))
        elif area_per_unit < min_req * 1.15:
            anomalies.append(('warning',
                f"⚠️ **Tight spacing:** {area_per_unit:.0f} m²/unit is close to "
                f"the {min_req} m² minimum. Consider more space for comfort."))

        # 2. Sterile zone without laminar
        if facility in ['Operating Room (OR)'] and not config.get('laminar_ceiling'):
            anomalies.append(('warning',
                "⚠️ **Missing laminar flow:** Operating Rooms typically require "
                "ultra-clean laminar airflow ceilings (ISO 5 / Class 100) over "
                "the surgical field."))

        # 3. Sterile zone without scrub sink
        if facility in ['Operating Room (OR)', 'NICU (Neonatal ICU)'] and \
           config.get('scrub_sink') == 'None':
            anomalies.append(('warning',
                f"⚠️ **No scrub station:** {facility} requires a surgical scrub "
                f"sink for hand antisepsis before procedures."))

        # 4. Critical care without monitoring
        if facility in ['Intensive Care Unit (ICU)', 'Operating Room (OR)',
                        'Emergency Room (ER)'] and not config.get('add_monitor'):
            anomalies.append(('warning',
                f"⚠️ **No patient monitoring:** {facility} should have vital "
                f"signs monitors for continuous patient observation."))

        # 5. Natural light recommendation
        if facility in ['Psychiatric Ward', 'Pediatric Ward',
                        'General Patient Ward', 'Oncology'] and \
           config.get('lighting_type') == 'Artificial Only':
            anomalies.append(('info',
                f"💡 **Daylight opportunity:** Patient recovery areas benefit "
                f"from natural light (reduces length of stay per evidence-based "
                f"design). Consider Natural or Mixed lighting."))

        # 6. Energy efficiency
        if config.get('equip_efficiency') == 'Legacy Systems (Standard)':
            anomalies.append(('info',
                "💡 **Efficiency upgrade:** Switching to Energy-Star certified "
                "equipment could cut energy use by ~12% and improve LEED score."))

        # 7. No renewable for high-energy facility
        if facility in ['MRI Room', 'CT Scan Room', 'Operating Room (OR)'] and \
           config.get('power_source') == 'National Grid Only':
            anomalies.append(('info',
                f"💡 **Renewable opportunity:** {facility} is energy-intensive. "
                f"Rooftop solar could offset significant operating costs."))

    else:  # hospital mode
        depts = config['selected_depts']
        n_depts = len(depts)

        # 1. Critical adjacencies missing
        has_or = any('Operating Room' in d for d in depts)
        has_cssd = any('Sterilization' in d for d in depts)
        has_er = any('Emergency' in d for d in depts)
        has_icu = any('Intensive Care' in d for d in depts)
        has_radiology = any('Radiology' in d for d in depts)

        if has_or and not has_cssd:
            anomalies.append(('warning',
                "⚠️ **Missing CSSD:** You have Operating Rooms but no "
                "Sterilization (CSSD) department. ORs require sterile "
                "instrument supply — add CSSD adjacent to surgery."))
        if has_er and not has_radiology:
            anomalies.append(('warning',
                "⚠️ **Missing Radiology:** Emergency departments need immediate "
                "imaging access. Add Radiology adjacent to the ER."))
        if has_er and not has_icu:
            anomalies.append(('info',
                "💡 **Consider ICU:** Emergency departments typically need an "
                "Intensive Care Unit for critical patient transfers."))
        if has_or and not has_icu:
            anomalies.append(('warning',
                "⚠️ **No ICU for surgery:** Operating Rooms should have an ICU "
                "for post-operative critical care."))

        # 2. Solar sizing
        if not config.get('h_solar'):
            anomalies.append(('info',
                "💡 **No solar:** A hospital this size has substantial roof "
                "area — rooftop solar PV is highly recommended for an "
                "eco-friendly design."))

        # 3. Water recycling
        if not config.get('h_greywater') and not config.get('h_rainwater'):
            anomalies.append(('info',
                "💡 **No water recycling:** Hospitals are major water consumers. "
                "Greywater recycling + rainwater harvesting cut water costs "
                "significantly."))

        # 4. Very small hospital with many departments
        if n_depts >= 15:
            anomalies.append(('info',
                f"💡 **Large program:** {n_depts} departments selected. Ensure "
                f"your site area and budget accommodate this scale."))

    return anomalies


# ════════════════════════════════════════════════════════════════
# ⚙️ OPTIMIZATION MODE
# Recommends best configuration for a chosen objective
# ════════════════════════════════════════════════════════════════
def optimize_design(facility_type, std, area, rooms, objective):
    """Recommend optimal choices for an objective.
    objective ∈ {'cost', 'co2', 'green', 'leed'}.
    Returns dict with recommendations + rationale."""
    ref = get_active_reference()
    recs = {}
    rationale = []

    if objective == 'cost':
        # Minimize total cost: cheapest viable options
        recs['power_source'] = 'National Grid Only'
        recs['lighting_type'] = 'Mixed'  # daylight cuts running cost cheaply
        recs['flooring'] = 'Eco-Linoleum (Natural)'  # cheapest eco option
        recs['equip_efficiency'] = 'Energy-Star / Eco-Mode Certified'
        rationale.append("Grid-only power avoids high upfront solar capital cost.")
        rationale.append("Mixed lighting uses free daylight to cut electricity bills.")
        rationale.append("Eco-Linoleum is the lowest-cost sustainable flooring.")
        rationale.append("Energy-Star equipment pays back quickly via lower bills.")

    elif objective == 'co2':
        # Minimize carbon: maximize renewables + efficiency
        recs['power_source'] = 'Hybrid (Solar + Battery Storage)'
        recs['lighting_type'] = 'Natural'
        recs['flooring'] = 'Eco-Linoleum (Natural)'
        recs['equip_efficiency'] = 'Energy-Star / Eco-Mode Certified'
        rationale.append("Solar + battery storage maximizes clean energy use, "
                         f"cutting grid CO₂ ({ref['co2_per_kwh']} kg/kWh).")
        rationale.append("Natural lighting eliminates daytime lighting emissions.")
        rationale.append("Energy-Star equipment minimizes total energy demand.")

    elif objective == 'green':
        # Maximize green score
        recs['power_source'] = 'Hybrid (Solar + Battery Storage)'
        recs['lighting_type'] = 'Mixed'
        recs['flooring'] = 'Eco-Linoleum (Natural)'
        recs['equip_efficiency'] = 'Energy-Star / Eco-Mode Certified'
        recs['waste_mgmt'] = 'Advanced Autoclaving & Shredding'
        recs['control_panel'] = 'Smart Touchscreen Console'
        rationale.append("Hybrid power gives the maximum green-score bonus (+20).")
        rationale.append("Advanced waste management adds sustainability points.")
        rationale.append("Smart controls optimize energy use dynamically.")

    elif objective == 'leed':
        # Maximize LEED credits
        recs['power_source'] = 'Hybrid (Solar + Battery Storage)'
        recs['lighting_type'] = 'Natural'
        recs['flooring'] = 'Eco-Linoleum (Natural)'
        recs['equip_efficiency'] = 'Energy-Star / Eco-Mode Certified'
        recs['waste_mgmt'] = 'Advanced Autoclaving & Shredding'
        rationale.append("Solar PV earns Energy & Atmosphere renewable credits.")
        rationale.append("Natural light earns Indoor Environmental Quality credits.")
        rationale.append("Low-emitting flooring earns Materials & Resources credits.")
        rationale.append("Construction waste management earns additional credits.")

    return {'recommendations': recs, 'rationale': rationale, 'objective': objective}


# ════════════════════════════════════════════════════════════════
# 💬 AI DESIGN ASSISTANT (rule-based knowledge engine)
# Pattern-matches questions → answers from standards knowledge base
# Can be upgraded to LLM API later (hook provided)
# ════════════════════════════════════════════════════════════════
def chatbot_answer(question, context=None):
    """Rule-based Q&A about hospital design standards.
    `context` is an optional dict with current design info.
    Returns a markdown answer string."""
    q = question.lower().strip()
    ref = get_active_reference()
    ctx = context or {}

    # Knowledge base: (keywords list, answer function)
    def kb_match(keywords):
        return all(k in q for k in keywords) if isinstance(keywords, list) \
            else keywords in q

    # ── ACH / ventilation ──
    if any(w in q for w in ['ach', 'air change', 'ventilation', 'air exchange']):
        if 'or' in q or 'operating' in q or 'surgery' in q:
            return ("**Operating Rooms require 20 ACH** (air changes per hour) "
                    "per ASHRAE 170-2017 Table 7-1. This high rate dilutes "
                    "airborne contaminants and maintains the sterile field. "
                    "At least 4 of those 20 changes must be outdoor air, and "
                    "ORs are kept at **positive pressure** to push air outward "
                    "and prevent contaminated air from entering.")
        return ("**Air changes per hour (ACH)** is how many times a room's full "
                "air volume is replaced hourly. Per ASHRAE 170: OR needs 20, "
                "ICU/ER need 6, patient rooms need 6, and isolation rooms need "
                "12. Higher ACH = better infection control but more energy.")

    # ── Pressure ──
    if 'pressure' in q and any(w in q for w in ['positive', 'negative', 'room', 'why']):
        return ("**Room pressurization** controls airflow direction. "
                "**Positive pressure** (OR, sterile storage) pushes air OUT to "
                "keep contaminants away from clean areas. **Negative pressure** "
                "(isolation, lab, soiled utility) pulls air IN so airborne "
                "pathogens can't escape. Per ASHRAE 170 + CDC guidelines.")

    # ── Lux / lighting ──
    if any(w in q for w in ['lux', 'lighting level', 'illuminance', 'how bright']):
        return ("**Illumination requirements** (IES RP-29-2017): Operating "
                "Rooms need 1000+ lux at the surgical task (up to 100,000 lux "
                "in the surgical light itself!), exam rooms 500 lux, patient "
                "rooms 100-300 lux, corridors 100-200 lux. Higher lux helps "
                "clinical accuracy but increases energy — daylighting helps.")

    # ── Area / size ──
    if any(w in q for w in ['minimum area', 'how big', 'room size', 'square met', 'area']):
        fac = ctx.get('facility_type')
        if fac and fac in MEDICAL_STANDARDS:
            std = MEDICAL_STANDARDS[fac]
            return (f"**{fac}** requires a minimum of "
                    f"**{std['min_area_per_unit']} m² per {std['unit_name']}** "
                    f"and a ceiling height of {std['min_ceiling']} m "
                    f"({ref['short']}, {ref['citations']['min_area']}). "
                    f"Add circulation space (~40%) for the actual room.")
        return ("**Minimum room areas** come from FGI 2018 Table 2.1-1 (US) or "
                "the Egyptian Hospital Code 2009. Examples (FGI): OR 37 m², "
                "ICU bed 18 m², ER bay 11 m², patient room 9.3 m². Egyptian "
                "standards are slightly more compact.")

    # ── Solar / renewable ──
    if any(w in q for w in ['solar', 'pv', 'renewable', 'photovolta']):
        base = ("**Solar PV** for hospitals: rooftop panels typically offset "
                "15-30% of energy use. Output depends on location — e.g., "
                "Aswan (7.2 peak sun hours) generates ~40% more than London "
                "(2.8h). ")
        if ctx.get('sel_city'):
            cd = detect_climate(ctx['sel_city'])
            base += (f"For **{ctx['sel_city']}** ({cd['sun_hours']}h sun/day), "
                     f"solar is {'highly recommended' if cd['sun_hours']>=6 else 'moderately effective'}.")
        return base

    # ── LEED / certification ──
    if any(w in q for w in ['leed', 'certif', 'green building', 'estidama', 'pearl', 'mostadam']):
        return ("**Green building certifications** rate sustainability: "
                "**LEED** (US, 110-point scale): Certified (40), Silver (50), "
                "Gold (60), Platinum (80+). **Estidama** (UAE): 1-5 Pearls. "
                "**Mostadam** (KSA): Green→Platinum. Earn credits via solar, "
                "water recycling, daylighting, low-emitting materials, and "
                "waste management. Check the Certification Forecast section!")

    # ── Cost / budget ──
    if any(w in q for w in ['cost', 'budget', 'price', 'expensive', 'how much', 'afford']):
        return (f"**Construction costs** (per m²): ORs are most expensive "
                f"(~{ref['currency_symbol']}{CONSTRUCTION_COSTS[st.session_state.get('reference','fgi')].get('Operating Room (OR)', 4500):,}/m²) "
                f"due to HVAC and sterile requirements, while wards are "
                f"cheaper (~{ref['currency_symbol']}{CONSTRUCTION_COSTS[st.session_state.get('reference','fgi')].get('General Patient Ward', 1800):,}/m²). "
                f"Use the **Budget Manager** to track your total and the "
                f"**Bill of Quantities** for itemized costs.")

    # ── Energy / EUI ──
    if any(w in q for w in ['energy', 'eui', 'consumption', 'kwh', 'power use']):
        return (f"**Energy Use Intensity (EUI)** measures kWh/m²/year. Baseline "
                f"hospital: {ref['eui_baseline']} kWh/m²/yr; green target: "
                f"{ref['eui_best']} kWh/m²/yr. Different departments vary — MRI "
                f"uses 3.5× the baseline (magnet cooling), while offices use "
                f"0.4×. We calculate energy as EUI × area × department factor.")

    # ── HVAC ──
    if any(w in q for w in ['hvac', 'cooling', 'air conditioning', 'chiller', 'ac']):
        return ("**HVAC sizing**: cooling load = envelope gain + people heat + "
                "lighting + equipment. We compute it in kW and tons (1 ton = "
                "3.517 kW). Hospitals need central chiller plants with N+1 "
                "redundancy (a backup chiller) so OR/ICU cooling never fails. "
                "See the HVAC tab in Engineering Analysis.")

    # ── Fire safety ──
    if any(w in q for w in ['fire', 'egress', 'exit', 'evacuat', 'nfpa', 'smoke']):
        return ("**Fire safety** (NFPA 101): healthcare needs ≥2 exits, max "
                "61 m travel distance (sprinklered), smoke compartments every "
                "≤2,323 m², and 'areas of refuge' for non-ambulatory patients. "
                "Egress width = 5 mm per occupant. See the Fire Safety tab.")

    # ── Greywater / water ──
    if any(w in q for w in ['water', 'greywater', 'rainwater', 'recycl']):
        return ("**Water systems**: hospitals use water for sanitary (40%), "
                "HVAC cooling (25%), processes/sterilization (20%), and "
                "irrigation (15%). **Greywater recycling** recovers ~30% of "
                "demand; **rainwater harvesting** captures roof runoff. See "
                "the Water Balance diagram for the full flow.")

    # ── Greeting ──
    if any(w in q for w in ['hello', 'hi', 'hey', 'salam', 'مرحبا', 'اهلا']):
        return ("👋 Hello! I'm the Viridis design assistant. Ask me about "
                "hospital design standards — ventilation (ACH), room sizes, "
                "lighting, solar, costs, LEED certification, HVAC, fire safety, "
                "and more. For example: *\"Why does an OR need 20 ACH?\"*")

    # ── Fallback ──
    return ("🤔 I'm not sure about that specific question, but I can help with:\n"
            "- **Ventilation** (ACH, pressure) — *\"Why 20 ACH for OR?\"*\n"
            "- **Room sizes** — *\"Minimum area for ICU?\"*\n"
            "- **Lighting** (lux levels)\n"
            "- **Solar & energy** (EUI, renewable potential)\n"
            "- **Costs & budget**\n"
            "- **LEED certification**\n"
            "- **HVAC & fire safety**\n\n"
            "Try rephrasing with one of these topics!")


def boq_to_excel_bytes(rows, summary_total, curr):
    """Convert BoQ rows to Excel bytes. Falls back to None if unavailable."""
    try:
        import pandas as pd
        df = pd.DataFrame(rows)
        buf = BytesIO()
        try:
            with pd.ExcelWriter(buf, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='BoQ',
                            startrow=2)
                ws = writer.sheets['BoQ']
                ws['A1'] = 'VIRIDIS — Bill of Quantities'
                # Add total row
                last = len(df) + 4
                ws[f'F{last}'] = 'GRAND TOTAL'
                ws[f'G{last}'] = round(summary_total, 2)
            buf.seek(0)
            return buf.getvalue()
        except Exception:
            return None
    except Exception:
        return None


def md_to_docx_bytes(md_text, title="Functional Program"):
    """Convert markdown-ish text to a .docx. Falls back to None if unavailable."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        doc = Document()
        for line in md_text.split('\n'):
            line = line.rstrip()
            if not line:
                continue
            if line.startswith('# '):
                h = doc.add_heading(line[2:], level=0)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=1)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=2)
            elif line.startswith('|'):
                # Skip table rendering for simplicity — add as monospace
                if '---' not in line:
                    p = doc.add_paragraph(line.replace('|', '  '))
                    p.runs[0].font.size = Pt(8) if p.runs else None
            elif line.startswith('---'):
                doc.add_paragraph('_' * 40)
            else:
                doc.add_paragraph(line.replace('**', '').replace('*', ''))
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
    except Exception:
        return None


def apply_reference_overrides():
    """Apply currently-selected reference's overrides to MEDICAL_STANDARDS.
    Mutates the global dict in place (re-runs every Streamlit script run)."""
    ref = SCIENTIFIC_REFERENCES.get(
        st.session_state.get('reference', 'fgi'), SCIENTIFIC_REFERENCES['fgi'])
    for dept_name, overrides in ref.get('dept_overrides', {}).items():
        if dept_name in MEDICAL_STANDARDS:
            for k, v in overrides.items():
                MEDICAL_STANDARDS[dept_name][k] = v


def get_active_reference():
    """Return the active reference dict (FGI or Egyptian)."""
    return SCIENTIFIC_REFERENCES.get(
        st.session_state.get('reference', 'fgi'), SCIENTIFIC_REFERENCES['fgi'])


# ════════════════════════════════════════════════════════════════
# ════════════════════════════════════════════════════════════════
#                    WIZARD PAGES
# ════════════════════════════════════════════════════════════════
# ════════════════════════════════════════════════════════════════

def render_progress_bar(current_step, total_steps=2):
    """Visual stepper showing wizard progress."""
    step_labels = ["Mode", "Facility"]
    html = '<div class="progress-container"><div class="progress-track">'
    for i in range(1, total_steps + 1):
        if i < current_step:
            cls = "complete"; icon = "✓"
        elif i == current_step:
            cls = "active"; icon = str(i)
        else:
            cls = ""; icon = str(i)
        html += f'<div class="progress-step {cls}">{icon}</div>'
        if i < total_steps:
            line_cls = "complete" if i < current_step else ""
            html += f'<div class="progress-line {line_cls}"></div>'
    html += '</div><div class="progress-labels">'
    for i, lbl in enumerate(step_labels, 1):
        cls = "label-active" if i == current_step else ""
        html += f'<span class="{cls}">{lbl}</span>'
    html += '</div></div>'
    st.markdown(html, unsafe_allow_html=True)

def render_floating_particles():
    """Subtle background particles for ambience."""
    particles_html = '<div class="particles">'
    import random
    random.seed(42)
    for _ in range(15):
        size = random.randint(3, 8)
        left = random.randint(0, 100)
        delay = random.randint(0, 12)
        duration = random.randint(8, 16)
        particles_html += (
            f'<div class="particle" style="width:{size}px;height:{size}px;'
            f'left:{left}%;animation-delay:{delay}s;'
            f'animation-duration:{duration}s;"></div>'
        )
    particles_html += '</div>'
    st.markdown(particles_html, unsafe_allow_html=True)

# Mark body so CSS scoping applies
def _wizard_body_class():
    """Inject wizard-only CSS that styles the dark eco theme.
    Since we can't add JS classes to body in Streamlit, we inject
    UNSCOPED CSS only when on wizard pages — it's removed when on
    the results page (which re-runs without this call)."""
    st.markdown(
        '<style>'
        # Main app dark gradient background
        '.stApp {'
        '  background: radial-gradient(circle at 20% 30%, rgba(34, 197, 94, 0.08) 0%, transparent 50%),'
        '              radial-gradient(circle at 80% 70%, rgba(74, 222, 128, 0.05) 0%, transparent 50%),'
        '              linear-gradient(135deg, #0d1f17 0%, #142e22 50%, #0d1f17 100%) !important;'
        '}'
        '.stApp .block-container {'
        '  padding-top: 1.5rem !important;'
        '  max-width: 1100px !important;'
        '}'
        # Text colors
        'h1, h2, h3, h4, h5 { color: #e8f5e9 !important; }'
        '.stApp p, .stApp label, .stApp span:not(.feature-icon):not(.mode-icon),'
        '.stApp div[class*="markdown"] { color: #cbd5e1; }'
        # Form input dark styling
        'div[data-baseweb="select"] > div {'
        '  background: #1a3d2e !important;'
        '  border-color: rgba(74,222,128,0.2) !important;'
        '  color: #e8f5e9 !important;'
        '}'
        'div[data-baseweb="select"] svg { color: #4ade80 !important; }'
        'div[data-baseweb="select"] input { color: #e8f5e9 !important; }'
        'div[data-baseweb="popover"] li { background: #1a3d2e !important; color: #e8f5e9 !important; }'
        'div[data-baseweb="popover"] li:hover { background: #4ade80 !important; color: #0d1f17 !important; }'
        'input[type="number"], textarea {'
        '  background: #1a3d2e !important;'
        '  color: #e8f5e9 !important;'
        '  border: 1px solid rgba(74,222,128,0.2) !important;'
        '}'
        # Slider
        '.stSlider [data-baseweb="slider"] > div > div > div { background: #4ade80 !important; }'
        '.stSlider [role="slider"] { background: #22c55e !important; border-color: #4ade80 !important; }'
        # Buttons (primary green)
        '.stButton > button {'
        '  background: linear-gradient(135deg, #16a34a, #22c55e) !important;'
        '  color: white !important;'
        '  border: none !important;'
        '  padding: 0.75rem 2rem !important;'
        '  font-weight: 700 !important;'
        '  border-radius: 12px !important;'
        '  box-shadow: 0 4px 20px rgba(74, 222, 128, 0.3) !important;'
        '  transition: all 0.3s ease !important;'
        '}'
        '.stButton > button:hover {'
        '  transform: translateY(-2px) !important;'
        '  box-shadow: 0 8px 30px rgba(74, 222, 128, 0.5) !important;'
        '  background: linear-gradient(135deg, #22c55e, #4ade80) !important;'
        '}'
        '.stButton > button[kind="secondary"] {'
        '  background: transparent !important;'
        '  border: 1.5px solid rgba(74, 222, 128, 0.3) !important;'
        '  color: #94a3b8 !important;'
        '  box-shadow: none !important;'
        '}'
        '.stButton > button[kind="secondary"]:hover {'
        '  border-color: #4ade80 !important;'
        '  color: #4ade80 !important;'
        '  background: rgba(74, 222, 128, 0.05) !important;'
        '}'
        # Checkbox & Radio
        '.stCheckbox label, .stRadio label { color: #e8f5e9 !important; }'
        '.stCheckbox label p, .stRadio label p { color: #e8f5e9 !important; }'
        'div[data-baseweb="checkbox"] > div:first-child {'
        '  background: #1a3d2e !important;'
        '  border: 2px solid rgba(74, 222, 128, 0.4) !important;'
        '}'
        'div[data-baseweb="checkbox"][aria-checked="true"] > div:first-child {'
        '  background: #4ade80 !important;'
        '  border-color: #4ade80 !important;'
        '}'
        # Expanders
        '.streamlit-expanderHeader, [data-testid="stExpander"] summary {'
        '  background: rgba(74, 222, 128, 0.08) !important;'
        '  color: #e8f5e9 !important;'
        '  border-radius: 8px !important;'
        '}'
        # Warning/Info text colors
        '.stAlert > div { color: #e8f5e9 !important; }'
        # Hide hamburger/footer on wizard
        '#MainMenu, footer, header[data-testid="stHeader"] { visibility: hidden; }'
        # Caption
        '.stApp small, [data-testid="stCaptionContainer"] { color: #94a3b8 !important; }'
        '</style>',
        unsafe_allow_html=True
    )


# ════════════════════════════════════════════════════════════════
# WELCOME PAGE
# ════════════════════════════════════════════════════════════════
def render_welcome():
    _wizard_body_class()
    render_floating_particles()

    # (v24: language switcher removed — site is English only)

    # Hero section
    st.markdown(
        f'<div style="text-align: center; padding-top: 1rem;">'
        f'<div class="viridis-logo">🍃</div>'
        f'<h1 class="hero-title">{t("welcome_title")}</h1>'
        f'<p class="hero-subtitle">{t("welcome_subtitle")}</p>'
        f'<p style="text-align: center; color: #94a3b8; max-width: 650px; '
        f'margin: 0 auto 2rem auto; line-height: 1.6; font-size: 1rem;">'
        f'{t("welcome_desc")}'
        f'</p></div>',
        unsafe_allow_html=True
    )

    # Feature cards
    c1, c2, c3 = st.columns(3, gap="medium")
    cards = [
        ("🏥", t("feature_depts"), t("feature_depts_desc")),
        ("📐", t("feature_fgi"), t("feature_fgi_desc")),
        ("🌞", t("feature_savings"), t("feature_savings_desc")),
    ]
    for col, (icon, title, desc) in zip([c1, c2, c3], cards):
        with col:
            st.markdown(
                f'<div class="feature-card">'
                f'<span class="feature-icon">{icon}</span>'
                f'<div class="feature-title">{title}</div>'
                f'<div class="feature-desc">{desc}</div>'
                f'</div>',
                unsafe_allow_html=True
            )

    st.markdown('<div style="height: 3rem;"></div>', unsafe_allow_html=True)

    # CTA Button (centered)
    spc1, mid, spc2 = st.columns([1, 1, 1])
    with mid:
        if st.button(t("start_designing"), use_container_width=True,
                     key="welcome_start"):
            go_to('step1')

    # Footer credit
    st.markdown(
        f'<div style="text-align: center; margin-top: 3rem; color: #64748b; '
        f'font-size: 0.85rem;">'
        f'{t("footer_credit")}</div>',
        unsafe_allow_html=True
    )


# ════════════════════════════════════════════════════════════════
# STEP 1: CHOOSE MODE
# ════════════════════════════════════════════════════════════════
def render_step1_mode():
    _wizard_body_class()
    render_progress_bar(1)

    st.markdown(
        '<div class="step-header">'
        '<span class="step-number">Step 1 of 2</span>'
        '<h2 class="step-title">Choose Your Design Scope</h2>'
        '<p class="step-subtitle">Are you designing a single medical room/department '
        'in detail, or planning an entire hospital campus?</p>'
        '</div>',
        unsafe_allow_html=True
    )

    current = st.session_state.wizard_data.get('app_mode', '')

    c1, c2 = st.columns(2, gap="large")
    with c1:
        sel = "selected" if current == "🏠 Single Department" else ""
        st.markdown(
            f'<div class="mode-card {sel}">'
            f'<span class="mode-icon">🏠</span>'
            f'<div class="mode-title">Single Department</div>'
            f'<div class="mode-desc">Detailed 2D/3D blueprint of one '
            f'medical space (OR, ICU, MRI, etc.)</div>'
            f'<ul class="mode-features">'
            f'<li>Realistic 3D equipment models</li>'
            f'<li>FGI compliance check</li>'
            f'<li>Sterile zone visualization</li>'
            f'<li>Energy & cost calculator</li>'
            f'</ul></div>',
            unsafe_allow_html=True
        )
        if st.button("Select Single Department",
                     use_container_width=True, key="m_single"):
            st.session_state.wizard_data['app_mode'] = "🏠 Single Department"
            go_to('step2')

    with c2:
        sel = "selected" if current == "🏥 Full Hospital Site Plan" else ""
        st.markdown(
            f'<div class="mode-card {sel}">'
            f'<span class="mode-icon">🏥</span>'
            f'<div class="mode-title">Full Hospital</div>'
            f'<div class="mode-desc">Complete hospital site plan with all '
            f'departments, gardens & solar farm</div>'
            f'<ul class="mode-features">'
            f'<li>Up to 24 departments on one site</li>'
            f'<li>Healing garden + walking paths</li>'
            f'<li>Bird\'s-eye 3D campus view</li>'
            f'<li>Annual carbon footprint report</li>'
            f'</ul></div>',
            unsafe_allow_html=True
        )
        if st.button("Select Full Hospital",
                     use_container_width=True, key="m_hospital"):
            st.session_state.wizard_data['app_mode'] = "🏥 Full Hospital Site Plan"
            go_to('step2')

    st.markdown('<div style="height: 2rem;"></div>', unsafe_allow_html=True)

    # Back button
    bspc1, bmid, bspc2 = st.columns([1, 1, 1])
    with bmid:
        if st.button("← Back to Welcome", type="secondary",
                     use_container_width=True, key="back_to_welcome"):
            go_to('welcome')


# ════════════════════════════════════════════════════════════════
# STEP 2: FACILITY SELECTION
# ════════════════════════════════════════════════════════════════
def render_step2_facility():
    _wizard_body_class()
    render_progress_bar(2)

    is_single = (st.session_state.wizard_data.get('app_mode')
                 == "🏠 Single Department")

    st.markdown(
        f'<div class="step-header">'
        f'<span class="step-number">Step 2 of 2 · Final Step</span>'
        f'<h2 class="step-title">'
        f'{"Select Department Type" if is_single else "Select Hospital Scale & Departments"}'
        f'</h2>'
        f'<p class="step-subtitle">'
        f'{"Choose the specific medical facility you want to design in detail." if is_single else "Pick the hospital size and which departments to include."}'
        f'</p>'
        f'</div>',
        unsafe_allow_html=True
    )

    if is_single:
        # Build category groups
        facility_groups = {}
        for name, info in MEDICAL_STANDARDS.items():
            cat = info["category"]
            facility_groups.setdefault(cat, []).append((name, info))

        st.markdown('<div class="wizard-section">', unsafe_allow_html=True)
        st.markdown('<span class="section-label">📁 Category</span>',
                    unsafe_allow_html=True)
        default_cat = st.session_state.wizard_data.get('category_choice',
                                                        list(facility_groups.keys())[0])
        if default_cat not in facility_groups:
            default_cat = list(facility_groups.keys())[0]
        category_choice = st.selectbox(
            "Department category",
            list(facility_groups.keys()),
            index=list(facility_groups.keys()).index(default_cat),
            label_visibility="collapsed",
            key="step2_cat"
        )
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="wizard-section">', unsafe_allow_html=True)
        st.markdown('<span class="section-label">🏥 Specific Facility</span>',
                    unsafe_allow_html=True)
        options = [f"{info['icon']} {name}" for name, info in facility_groups[category_choice]]
        default_facility = st.session_state.wizard_data.get('facility_label',
                                                             options[0])
        if default_facility not in options:
            default_facility = options[0]
        facility_label = st.selectbox(
            "Facility",
            options,
            index=options.index(default_facility),
            label_visibility="collapsed",
            key="step2_fac"
        )
        facility_type = facility_label.split(" ", 1)[1]
        st.markdown('</div>', unsafe_allow_html=True)

        # Show standards preview
        std = MEDICAL_STANDARDS[facility_type]
        st.markdown(
            f'<div class="summary-box">'
            f'<div class="summary-item"><span class="summary-key">Min area per unit:</span>'
            f'<span class="summary-val">≥ {std["min_area_per_unit"]} m²</span></div>'
            f'<div class="summary-item"><span class="summary-key">Air changes/hour (ACH):</span>'
            f'<span class="summary-val">{std["ach"]}</span></div>'
            f'<div class="summary-item"><span class="summary-key">Pressure:</span>'
            f'<span class="summary-val">{std["pressure"]}</span></div>'
            f'<div class="summary-item"><span class="summary-key">Illumination (lux):</span>'
            f'<span class="summary-val">{std["lux"]}</span></div>'
            f'<div class="summary-item"><span class="summary-key">Water usage:</span>'
            f'<span class="summary-val">~{std["water_usage"]} L/day per unit</span></div>'
            f'</div>',
            unsafe_allow_html=True
        )

        st.session_state.wizard_data['category_choice'] = category_choice
        st.session_state.wizard_data['facility_label'] = facility_label
        st.session_state.wizard_data['facility_type'] = facility_type

    else:
        # Full Hospital: scale + dept selection
        st.markdown('<div class="wizard-section">', unsafe_allow_html=True)
        st.markdown('<span class="section-label">📏 Hospital Scale</span>',
                    unsafe_allow_html=True)
        scale_options = ["Small Clinic (50 beds)", "Medium Hospital (150 beds)",
                         "Large Hospital (300 beds)", "Mega Medical City (600+ beds)"]
        default_scale = st.session_state.wizard_data.get('hospital_scale',
                                                          scale_options[1])
        hospital_scale = st.select_slider(
            "Hospital scale",
            options=scale_options,
            value=default_scale,
            label_visibility="collapsed",
            key="step2_scale"
        )
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="wizard-section">', unsafe_allow_html=True)
        st.markdown('<span class="section-label">🏥 Departments to Include</span>',
                    unsafe_allow_html=True)
        st.caption("Select departments by category (you can change later in sidebar):")

        current_depts = st.session_state.wizard_data.get('selected_dept_names', None)
        selected_dept_names = []

        for cat in ["Critical Care", "Diagnostic", "Treatment",
                    "Specialized", "General Care", "Service"]:
            with st.expander(f"📁 {cat}", expanded=(cat == "Critical Care")):
                for dept_name, dept_info in MEDICAL_STANDARDS.items():
                    if dept_info["category"] == cat:
                        if current_depts is None:
                            default = cat in ["Critical Care", "Service", "General Care"]
                        else:
                            default = dept_name in current_depts
                        if st.checkbox(f"{dept_info['icon']} {dept_name}",
                                        value=default,
                                        key=f"wiz_dept_{dept_name}"):
                            selected_dept_names.append(dept_name)
        st.markdown('</div>', unsafe_allow_html=True)

        st.session_state.wizard_data['hospital_scale'] = hospital_scale
        st.session_state.wizard_data['selected_dept_names'] = selected_dept_names

        if not selected_dept_names:
            st.warning("⚠️ Please select at least one department.")

    # Navigation
    st.markdown('<div style="height: 1.5rem;"></div>', unsafe_allow_html=True)
    nc1, nc2, nc3 = st.columns([1, 2, 1])
    with nc1:
        if st.button("← Back", type="secondary",
                     use_container_width=True, key="s2_back"):
            go_to('step1')
    with nc2:
        # Main launch button — commits all defaults + jumps straight to results.
        # Spatial config (area, climate, lighting), equipment, and
        # sustainability all auto-populate from sensible defaults in
        # commit_wizard_to_config(). The user fine-tunes everything from
        # the sidebar and tabs once on the results page.
        can_proceed = is_single or (not is_single and selected_dept_names)
        st.markdown('<div class="go-viridis-btn">', unsafe_allow_html=True)
        if st.button("🚀  Go Viridis!  🍃", use_container_width=True,
                     key="s2_go_viridis", disabled=not can_proceed):
            commit_wizard_to_config()
            go_to('results')
        st.markdown('</div>', unsafe_allow_html=True)
    with nc3:
        if st.button("🔄 Start Over", type="secondary",
                     use_container_width=True, key="s2_reset"):
            st.session_state.wizard_data = {}
            # Reset the empty-state flags so the next session shows the
            # educational dashboard again rather than stale results.
            st.session_state.design_generated = False
            st.session_state.hospital_design_generated = False
            # Also clear any prior snapshots so the sidebar widget defaults
            # reset properly.
            for k in ('single_config', 'hospital_config'):
                if k in st.session_state:
                    del st.session_state[k]
            go_to('welcome')


# ════════════════════════════════════════════════════════════════
# STEP 3: SPATIAL CONFIGURATION
# ════════════════════════════════════════════════════════════════
def render_step3_spatial():
    _wizard_body_class()
    render_progress_bar(3)

    is_single = (st.session_state.wizard_data.get('app_mode')
                 == "🏠 Single Department")

    st.markdown(
        f'<div class="step-header">'
        f'<span class="step-number">Step 3 of 3 · Final Step</span>'
        f'<h2 class="step-title">Spatial Configuration</h2>'
        f'<p class="step-subtitle">'
        f'{"Define the dimensions and capacity of the room." if is_single else "Set the climate and general operating conditions."}'
        f'</p>'
        f'</div>',
        unsafe_allow_html=True
    )

    if is_single:
        facility_type = st.session_state.wizard_data.get('facility_type',
                                                          'Operating Room (OR)')
        std = MEDICAL_STANDARDS[facility_type]

        c1, c2 = st.columns(2, gap="large")
        with c1:
            st.markdown('<div class="wizard-section">', unsafe_allow_html=True)
            st.markdown('<span class="section-label">📐 Room Dimensions</span>',
                        unsafe_allow_html=True)
            total_area = st.slider(
                "Total Room Area (m²)",
                min_value=20, max_value=200,
                value=st.session_state.wizard_data.get('total_area', 60),
                step=5, key="s3_area"
            )
            rooms_count = st.number_input(
                "Number of Rooms in Department",
                min_value=1, max_value=10,
                value=st.session_state.wizard_data.get('rooms_count', 2),
                key="s3_rooms"
            )
            units_per_room = st.number_input(
                f"Number of {std['unit_name']}s per Room",
                min_value=1, max_value=std['max_units'],
                value=min(st.session_state.wizard_data.get(
                    'units_per_room', 3), std['max_units']),
                key="s3_units"
            )
            st.markdown('</div>', unsafe_allow_html=True)

        with c2:
            st.markdown('<div class="wizard-section">', unsafe_allow_html=True)
            st.markdown('<span class="section-label">🌡️ Environmental</span>',
                        unsafe_allow_html=True)
            climate = st.radio(
                "Regional Climate Zone",
                ["Hot", "Cold", "Moderate"],
                index=["Hot", "Cold", "Moderate"].index(
                    st.session_state.wizard_data.get('climate', 'Moderate')),
                key="s3_climate"
            )
            lighting_type = st.radio(
                "Lighting Preference",
                ["Artificial Only", "Natural", "Mixed"],
                index=["Artificial Only", "Natural", "Mixed"].index(
                    st.session_state.wizard_data.get('lighting_type', 'Mixed')),
                key="s3_lighting"
            )
            st.markdown('</div>', unsafe_allow_html=True)

        # Validate area
        required = std['min_area_per_unit'] * units_per_room
        if total_area < required:
            st.warning(f"⚠️ Area too small: FGI standard requires ≥ {required} m² "
                       f"for {units_per_room} {std['unit_name']}s.")

        st.session_state.wizard_data.update({
            'total_area': total_area,
            'rooms_count': rooms_count,
            'units_per_room': units_per_room,
            'climate': climate,
            'lighting_type': lighting_type,
        })

    else:
        # Hospital mode: minimal step 3 (climate-level decisions)
        st.markdown('<div class="wizard-section">', unsafe_allow_html=True)
        st.markdown('<span class="section-label">🌍 Site & Climate</span>',
                    unsafe_allow_html=True)
        climate = st.radio(
            "Regional Climate Zone",
            ["Hot", "Cold", "Moderate"],
            index=["Hot", "Cold", "Moderate"].index(
                st.session_state.wizard_data.get('climate', 'Moderate')),
            horizontal=True,
            key="s3_climate_h"
        )
        outdoor_air = st.radio(
            "Outdoor Air Quality Status",
            ["Low Pollution (Rural/Suburban)", "High Pollution (Urban/Industrial)"],
            index=["Low Pollution (Rural/Suburban)",
                   "High Pollution (Urban/Industrial)"].index(
                st.session_state.wizard_data.get('outdoor_air',
                    'Low Pollution (Rural/Suburban)')),
            key="s3_air"
        )
        st.markdown('</div>', unsafe_allow_html=True)

        # Show summary so far
        n_depts = len(st.session_state.wizard_data.get('selected_dept_names', []))
        st.markdown(
            f'<div class="summary-box">'
            f'<div class="summary-item"><span class="summary-key">Hospital scale:</span>'
            f'<span class="summary-val">{st.session_state.wizard_data.get("hospital_scale", "—")}</span></div>'
            f'<div class="summary-item"><span class="summary-key">Departments selected:</span>'
            f'<span class="summary-val">{n_depts} of {len(MEDICAL_STANDARDS)}</span></div>'
            f'<div class="summary-item"><span class="summary-key">Climate:</span>'
            f'<span class="summary-val">{climate}</span></div>'
            f'</div>',
            unsafe_allow_html=True
        )

        st.session_state.wizard_data.update({
            'climate': climate,
            'outdoor_air': outdoor_air,
        })

    # Navigation
    st.markdown('<div style="height: 1.5rem;"></div>', unsafe_allow_html=True)
    nc1, nc2, nc3 = st.columns([1, 2, 1])
    with nc1:
        if st.button("← Back", type="secondary",
                     use_container_width=True, key="s3_back"):
            go_to('step2')
    with nc2:
        # Main launch button — commits all defaults + jumps straight to results.
        # Equipment and sustainability are auto-configured from sensible
        # defaults in commit_wizard_to_config(); user can fine-tune everything
        # from the sidebar and tabs once on the results page.
        st.markdown('<div class="go-viridis-btn">', unsafe_allow_html=True)
        if st.button("🚀  Go Viridis!  🍃", use_container_width=True,
                     key="s3_go_viridis"):
            commit_wizard_to_config()
            go_to('results')
        st.markdown('</div>', unsafe_allow_html=True)
    with nc3:
        if st.button("🔄 Start Over", type="secondary",
                     use_container_width=True, key="s3_reset"):
            st.session_state.wizard_data = {}
            go_to('welcome')


# ════════════════════════════════════════════════════════════════
# STEP 4: EQUIPMENT & COMPONENTS
# ════════════════════════════════════════════════════════════════
def render_step4_equipment():
    _wizard_body_class()
    render_progress_bar(4)

    is_single = (st.session_state.wizard_data.get('app_mode')
                 == "🏠 Single Department")

    st.markdown(
        f'<div class="step-header">'
        f'<span class="step-number">Step 4 of 5</span>'
        f'<h2 class="step-title">'
        f'{"Equipment & Components" if is_single else "Visualization & Layout"}'
        f'</h2>'
        f'<p class="step-subtitle">'
        f'{"Choose what medical equipment and infrastructure to include." if is_single else "Configure how your hospital is displayed."}'
        f'</p>'
        f'</div>',
        unsafe_allow_html=True
    )

    if is_single:
        facility_type = st.session_state.wizard_data.get('facility_type',
                                                          'Operating Room (OR)')
        std = MEDICAL_STANDARDS[facility_type]
        units_per_room = st.session_state.wizard_data.get('units_per_room', 1)

        c1, c2 = st.columns(2, gap="large")
        with c1:
            st.markdown('<div class="wizard-section">', unsafe_allow_html=True)
            st.markdown('<span class="section-label">🚪 Spatial Components</span>',
                        unsafe_allow_html=True)
            if std['max_units'] > 1:
                partition_type = st.selectbox(
                    "Unit Zoning Partitions",
                    ["Open Layout", "Antibacterial Glass Walls"],
                    index=["Open Layout", "Antibacterial Glass Walls"].index(
                        st.session_state.wizard_data.get('partition_type',
                                                          'Open Layout')),
                    key="s4_part"
                )
            else:
                partition_type = "Open Layout"
                st.caption("ℹ️ Single-unit room: partitions not applicable")

            scrub_sink = st.selectbox(
                "Scrub Sink Station",
                ["None", "Single Medical Sink", "Double Stainless Sink"],
                index=["None", "Single Medical Sink", "Double Stainless Sink"].index(
                    st.session_state.wizard_data.get('scrub_sink',
                                                      'Single Medical Sink')),
                key="s4_sink"
            )
            control_panel = st.radio(
                "Wall Control Interface",
                ["Analogue Gauges", "Smart Touchscreen Console"],
                index=["Analogue Gauges", "Smart Touchscreen Console"].index(
                    st.session_state.wizard_data.get('control_panel',
                                                      'Smart Touchscreen Console')),
                key="s4_panel"
            )
            laminar_ceiling = st.checkbox(
                "Deploy Ultra-Clean Laminar Flow Ceiling Grid",
                value=st.session_state.wizard_data.get(
                    'laminar_ceiling',
                    facility_type in ["Operating Room (OR)", "NICU (Neonatal ICU)"]),
                key="s4_laminar"
            )
            st.markdown('</div>', unsafe_allow_html=True)

        with c2:
            st.markdown('<div class="wizard-section">', unsafe_allow_html=True)
            st.markdown('<span class="section-label">🪑 Medical Equipment</span>',
                        unsafe_allow_html=True)
            add_monitor = st.checkbox(
                "Vital Signs Monitor",
                value=st.session_state.wizard_data.get('add_monitor',
                    facility_type in ["Operating Room (OR)",
                                       "Intensive Care Unit (ICU)",
                                       "Emergency Room (ER)",
                                       "NICU (Neonatal ICU)"]),
                key="s4_mon"
            )
            add_anesthesia = st.checkbox(
                "Anesthesia Machine",
                value=st.session_state.wizard_data.get('add_anesthesia',
                    facility_type == "Operating Room (OR)"),
                key="s4_anes"
            )
            add_iv_stand = st.checkbox(
                "IV Stands",
                value=st.session_state.wizard_data.get('add_iv_stand', True),
                key="s4_iv"
            )
            add_med_cart = st.checkbox(
                "Medication Cart",
                value=st.session_state.wizard_data.get('add_med_cart', True),
                key="s4_cart"
            )
            st.markdown('</div>', unsafe_allow_html=True)

        st.session_state.wizard_data.update({
            'partition_type': partition_type,
            'scrub_sink': scrub_sink,
            'control_panel': control_panel,
            'laminar_ceiling': laminar_ceiling,
            'add_monitor': add_monitor,
            'add_anesthesia': add_anesthesia,
            'add_iv_stand': add_iv_stand,
            'add_med_cart': add_med_cart,
        })

    else:
        # Hospital mode: visualization options
        st.markdown('<div class="wizard-section">', unsafe_allow_html=True)
        st.markdown('<span class="section-label">🎨 Display Options</span>',
                    unsafe_allow_html=True)
        c1, c2, c3 = st.columns(3)
        with c1:
            show_corridors = st.checkbox("Show Corridors", value=True,
                                          key="s4_corr")
        with c2:
            show_patient_flow = st.checkbox("Show Patient Flow", value=True,
                                             key="s4_flow")
        with c3:
            show_dept_labels = st.checkbox("Show Dept Labels", value=True,
                                            key="s4_lbl")
        st.markdown('</div>', unsafe_allow_html=True)

        st.session_state.wizard_data.update({
            'show_corridors': show_corridors,
            'show_patient_flow': show_patient_flow,
            'show_dept_labels': show_dept_labels,
        })

    # Navigation
    st.markdown('<div style="height: 1.5rem;"></div>', unsafe_allow_html=True)
    nc1, nc2, nc3 = st.columns([1, 1, 1])
    with nc1:
        if st.button("← Back", type="secondary",
                     use_container_width=True, key="s4_back"):
            go_to('step3')
    with nc3:
        if st.button("Next →", use_container_width=True, key="s4_next"):
            go_to('step5')


# ════════════════════════════════════════════════════════════════
# STEP 5: SUSTAINABILITY & POWER
# ════════════════════════════════════════════════════════════════
def render_step5_sustainability():
    _wizard_body_class()
    render_progress_bar(5)

    is_single = (st.session_state.wizard_data.get('app_mode')
                 == "🏠 Single Department")

    st.markdown(
        f'<div class="step-header">'
        f'<span class="step-number">Step 5 of 5 · Final Step</span>'
        f'<h2 class="step-title">Sustainability & Power</h2>'
        f'<p class="step-subtitle">Configure eco-friendly features. '
        f'These choices drive your Green Score, carbon savings, and ROI.</p>'
        f'</div>',
        unsafe_allow_html=True
    )

    if is_single:
        c1, c2 = st.columns(2, gap="large")
        with c1:
            st.markdown('<div class="wizard-section">', unsafe_allow_html=True)
            st.markdown('<span class="section-label">🏗️ Materials & Waste</span>',
                        unsafe_allow_html=True)
            flooring = st.selectbox(
                "Flooring Material",
                ["Antimicrobial Vinyl (Seamless)", "Eco-Linoleum (Natural)",
                 "High-Traffic Terrazzo"],
                index=["Antimicrobial Vinyl (Seamless)", "Eco-Linoleum (Natural)",
                       "High-Traffic Terrazzo"].index(
                    st.session_state.wizard_data.get('flooring',
                        'Eco-Linoleum (Natural)')),
                key="s5_floor"
            )
            waste_mgmt = st.selectbox(
                "Medical Waste System",
                ["Standard Centralized Disposal",
                 "On-site Color-Coded Segregation",
                 "Advanced Autoclaving & Shredding"],
                index=["Standard Centralized Disposal",
                       "On-site Color-Coded Segregation",
                       "Advanced Autoclaving & Shredding"].index(
                    st.session_state.wizard_data.get('waste_mgmt',
                        'Advanced Autoclaving & Shredding')),
                key="s5_waste"
            )
            equip_efficiency = st.radio(
                "Medical Equipment Class",
                ["Legacy Systems (Standard)",
                 "Energy-Star / Eco-Mode Certified"],
                index=["Legacy Systems (Standard)",
                       "Energy-Star / Eco-Mode Certified"].index(
                    st.session_state.wizard_data.get('equip_efficiency',
                        'Energy-Star / Eco-Mode Certified')),
                key="s5_equip"
            )
            st.markdown('</div>', unsafe_allow_html=True)

        with c2:
            st.markdown('<div class="wizard-section">', unsafe_allow_html=True)
            st.markdown('<span class="section-label">⚡ Power & Air Quality</span>',
                        unsafe_allow_html=True)
            outdoor_air = st.radio(
                "Outdoor Air Quality Status",
                ["Low Pollution (Rural/Suburban)",
                 "High Pollution (Urban/Industrial)"],
                index=["Low Pollution (Rural/Suburban)",
                       "High Pollution (Urban/Industrial)"].index(
                    st.session_state.wizard_data.get('outdoor_air',
                        'Low Pollution (Rural/Suburban)')),
                key="s5_air"
            )
            power_source = st.selectbox(
                "Primary Building Power Source",
                ["National Grid Only", "Grid + Rooftop Solar PV",
                 "Hybrid (Solar + Battery Storage)"],
                index=["National Grid Only", "Grid + Rooftop Solar PV",
                       "Hybrid (Solar + Battery Storage)"].index(
                    st.session_state.wizard_data.get('power_source',
                        'Grid + Rooftop Solar PV')),
                key="s5_power"
            )
            st.markdown('</div>', unsafe_allow_html=True)

        st.session_state.wizard_data.update({
            'flooring': flooring,
            'waste_mgmt': waste_mgmt,
            'equip_efficiency': equip_efficiency,
            'outdoor_air': outdoor_air,
            'power_source': power_source,
            # Defaults for missing
            'view_mode': '📊 Both Side-by-Side',
            'show_grid': True,
            'show_dimensions': True,
            'show_traffic_flow': True,
            'show_zones': True,
            'day_night': '☀️ Day',
        })

    else:
        # Hospital mode: full sustainability toggles
        c1, c2 = st.columns(2, gap="large")
        with c1:
            st.markdown('<div class="wizard-section">', unsafe_allow_html=True)
            st.markdown('<span class="section-label">🌞 Energy</span>',
                        unsafe_allow_html=True)
            h_solar = st.checkbox("🌞 Rooftop Solar PV Farm",
                value=st.session_state.wizard_data.get('h_solar', True),
                key="s5_solar")
            h_smart_grid = st.checkbox("⚡ Smart Grid + Battery Storage",
                value=st.session_state.wizard_data.get('h_smart_grid', False),
                key="s5_grid")
            h_ev_chargers = st.checkbox("🔌 EV Charging Stations",
                value=st.session_state.wizard_data.get('h_ev_chargers', True),
                key="s5_ev")
            st.markdown('</div>', unsafe_allow_html=True)

        with c2:
            st.markdown('<div class="wizard-section">', unsafe_allow_html=True)
            st.markdown('<span class="section-label">💧 Water & Green</span>',
                        unsafe_allow_html=True)
            h_garden = st.checkbox("🌳 Healing Garden (central courtyard)",
                value=st.session_state.wizard_data.get('h_garden', True),
                key="s5_gard")
            h_green_roof = st.checkbox("🌿 Green Roof",
                value=st.session_state.wizard_data.get('h_green_roof', True),
                key="s5_groof")
            h_rainwater = st.checkbox("💧 Rainwater Harvesting",
                value=st.session_state.wizard_data.get('h_rainwater', True),
                key="s5_rain")
            h_greywater = st.checkbox("♻️ Greywater Recycling",
                value=st.session_state.wizard_data.get('h_greywater', True),
                key="s5_grey")
            st.markdown('</div>', unsafe_allow_html=True)

        st.session_state.wizard_data.update({
            'h_solar': h_solar, 'h_garden': h_garden,
            'h_green_roof': h_green_roof, 'h_rainwater': h_rainwater,
            'h_ev_chargers': h_ev_chargers, 'h_greywater': h_greywater,
            'h_smart_grid': h_smart_grid,
        })

    # Final summary box
    st.markdown('<div style="height: 1rem;"></div>', unsafe_allow_html=True)
    wd = st.session_state.wizard_data
    if is_single:
        summary_html = (
            f'<div class="summary-box">'
            f'<div style="text-align:center; color:#4ade80; font-weight:700; '
            f'margin-bottom:1rem; font-size:1.1rem;">📋 Design Summary</div>'
            f'<div class="summary-item"><span class="summary-key">Mode:</span>'
            f'<span class="summary-val">Single Department</span></div>'
            f'<div class="summary-item"><span class="summary-key">Facility:</span>'
            f'<span class="summary-val">{wd.get("facility_type", "—")}</span></div>'
            f'<div class="summary-item"><span class="summary-key">Room area:</span>'
            f'<span class="summary-val">{wd.get("total_area", "—")} m²</span></div>'
            f'<div class="summary-item"><span class="summary-key">Power:</span>'
            f'<span class="summary-val">{wd.get("power_source", "—")}</span></div>'
            f'</div>'
        )
    else:
        n_depts = len(wd.get('selected_dept_names', []))
        eco_count = sum([wd.get('h_solar', False), wd.get('h_garden', False),
                         wd.get('h_green_roof', False), wd.get('h_rainwater', False),
                         wd.get('h_ev_chargers', False), wd.get('h_greywater', False),
                         wd.get('h_smart_grid', False)])
        summary_html = (
            f'<div class="summary-box">'
            f'<div style="text-align:center; color:#4ade80; font-weight:700; '
            f'margin-bottom:1rem; font-size:1.1rem;">📋 Hospital Summary</div>'
            f'<div class="summary-item"><span class="summary-key">Scale:</span>'
            f'<span class="summary-val">{wd.get("hospital_scale", "—")}</span></div>'
            f'<div class="summary-item"><span class="summary-key">Departments:</span>'
            f'<span class="summary-val">{n_depts} selected</span></div>'
            f'<div class="summary-item"><span class="summary-key">Eco features:</span>'
            f'<span class="summary-val">{eco_count} of 7 enabled</span></div>'
            f'<div class="summary-item"><span class="summary-key">Climate:</span>'
            f'<span class="summary-val">{wd.get("climate", "—")}</span></div>'
            f'</div>'
        )
    st.markdown(summary_html, unsafe_allow_html=True)

    # Mega CTA: Go Viridis
    st.markdown('<div style="height: 1.5rem;"></div>', unsafe_allow_html=True)
    bc1, bc2, bc3 = st.columns([1, 2, 1])
    with bc2:
        st.markdown('<div class="go-viridis-btn">', unsafe_allow_html=True)
        if st.button("🚀  Go Viridis!  🍃", use_container_width=True,
                     key="go_viridis"):
            commit_wizard_to_config()
            go_to('results')
        st.markdown('</div>', unsafe_allow_html=True)

    # Back / Reset
    st.markdown('<div style="height: 1rem;"></div>', unsafe_allow_html=True)
    nc1, nc2 = st.columns([1, 1])
    with nc1:
        if st.button("← Back", type="secondary",
                     use_container_width=True, key="s5_back"):
            go_to('step4')
    with nc2:
        if st.button("🔄 Start Over", type="secondary",
                     use_container_width=True, key="s5_reset"):
            st.session_state.wizard_data = {}
            go_to('welcome')


def commit_wizard_to_config():
    """Transfer wizard_data into the single_config/hospital_config that the
    main app reads from, so the results page renders without requiring
    another Generate click."""
    wd = st.session_state.wizard_data
    is_single = (wd.get('app_mode') == "🏠 Single Department")

    if is_single:
        facility_type = wd['facility_type']
        st.session_state.single_config = {
            'facility_type': facility_type,
            'std': MEDICAL_STANDARDS[facility_type],
            'total_area': wd.get('total_area', 60),
            'rooms_count': wd.get('rooms_count', 2),
            'units_per_room': wd.get('units_per_room', 1),
            'climate': wd.get('climate', 'Moderate'),
            'lighting_type': wd.get('lighting_type', 'Mixed'),
            'partition_type': wd.get('partition_type', 'Open Layout'),
            'scrub_sink': wd.get('scrub_sink', 'Single Medical Sink'),
            'control_panel': wd.get('control_panel', 'Smart Touchscreen Console'),
            'laminar_ceiling': wd.get('laminar_ceiling', False),
            'add_monitor': wd.get('add_monitor', True),
            'add_anesthesia': wd.get('add_anesthesia', False),
            'add_iv_stand': wd.get('add_iv_stand', True),
            'add_med_cart': wd.get('add_med_cart', True),
            'flooring': wd.get('flooring', 'Eco-Linoleum (Natural)'),
            'waste_mgmt': wd.get('waste_mgmt',
                                  'Advanced Autoclaving & Shredding'),
            'equip_efficiency': wd.get('equip_efficiency',
                                        'Energy-Star / Eco-Mode Certified'),
            'outdoor_air': wd.get('outdoor_air',
                                   'Low Pollution (Rural/Suburban)'),
            'power_source': wd.get('power_source', 'Grid + Rooftop Solar PV'),
            'view_mode': wd.get('view_mode', '📊 Both Side-by-Side'),
            'show_grid': wd.get('show_grid', True),
            'show_dimensions': wd.get('show_dimensions', True),
            'show_traffic_flow': wd.get('show_traffic_flow', True),
            'show_zones': wd.get('show_zones', True),
            'day_night': wd.get('day_night', '☀️ Day'),
        }
    else:
        # Build selected_depts from selected_dept_names
        sel_names = wd.get('selected_dept_names', [])
        scale_map = {"Small Clinic (50 beds)": 0.5,
                     "Medium Hospital (150 beds)": 1.0,
                     "Large Hospital (300 beds)": 1.8,
                     "Mega Medical City (600+ beds)": 3.0}
        h_scale = wd.get('hospital_scale', 'Medium Hospital (150 beds)')
        st.session_state.hospital_config = {
            'hospital_scale': h_scale,
            'scale': scale_map.get(h_scale, 1.0),
            'selected_depts': {n: MEDICAL_STANDARDS[n] for n in sel_names},
            'h_solar': wd.get('h_solar', True),
            'h_garden': wd.get('h_garden', True),
            'h_green_roof': wd.get('h_green_roof', True),
            'h_rainwater': wd.get('h_rainwater', True),
            'h_ev_chargers': wd.get('h_ev_chargers', True),
            'h_greywater': wd.get('h_greywater', True),
            'h_smart_grid': wd.get('h_smart_grid', False),
            'show_corridors': wd.get('show_corridors', True),
            'show_patient_flow': wd.get('show_patient_flow', True),
            'show_dept_labels': wd.get('show_dept_labels', True),
        }


# ════════════════════════════════════════════════════════════════
# WIZARD ROUTER — render wizard page if not on results
# ════════════════════════════════════════════════════════════════
if st.session_state.page == 'welcome':
    render_welcome()
    st.stop()
elif st.session_state.page == 'step1':
    render_step1_mode()
    st.stop()
elif st.session_state.page == 'step2':
    render_step2_facility()
    st.stop()
# Note: step3 (spatial), step4 (equipment), and step5 (sustainability) are
# intentionally removed from the wizard flow. The "🚀 Go Viridis!" button
# in step2 now commits sensible defaults for spatial, equipment, and
# sustainability — then jumps straight to the results page. Users can
# fine-tune everything from the sidebar and tabs on the results page.
# The step3/step4/step5 render functions remain defined but are not
# routed — kept for backward-compatibility / future re-enable.

# ════════════════════════════════════════════════════════════════
# Results page reached — show title + back-to-wizard pill
# ════════════════════════════════════════════════════════════════
back_col, _ = st.columns([1, 4])
with back_col:
    if st.button("← Back to Wizard", key="back_wizard"):
        go_to('welcome')

st.markdown(
    f'<h1 style="text-align: center; color: #1b5e20; margin-bottom: 0.5rem;">'
    f'{t("app_title")}'
    f'</h1>',
    unsafe_allow_html=True
)
st.markdown("---")

# ============================================================
# SIDEBAR — Standards Reference + Mode
# (Theme + Language removed in v24 — site is English/Dark only)
# ============================================================

# Scientific reference (FGI vs Egyptian) selector
ref_choice = st.sidebar.selectbox(
    "📚 Standards Reference",
    [SCIENTIFIC_REFERENCES['fgi']['short'],
     SCIENTIFIC_REFERENCES['egyptian']['short']],
    index=0 if st.session_state.reference == 'fgi' else 1,
    key='ref_select',
    help="Switch between US (FGI 2018 + ASHRAE) and Egyptian MoH standards. "
         "Changes minimum areas, ACH rates, energy benchmarks, "
         "electricity prices, and CO₂ emissions factor."
)
new_ref = 'fgi' if 'FGI' in ref_choice else 'egyptian'
if new_ref != st.session_state.reference:
    st.session_state.reference = new_ref
    # Clear cached configs since standards changed
    st.session_state.pop('single_config', None)
    st.session_state.pop('hospital_config', None)
    st.rerun()

# Show active citation panel (collapsed)
_active_ref = get_active_reference()
with st.sidebar.expander("📖 Active Citations", expanded=False):
    st.caption(f"**{_active_ref['name']}**")
    for key, citation in _active_ref['citations'].items():
        st.caption(f"• **{key}:** {citation}")

# Light-theme support removed in v24 — app is always dark eco theme.
if False:
    st.markdown(
        '<style>'
        '.stApp { background: #ffffff !important; color: #1b5e20 !important; }'
        '.stApp h1, .stApp h2, .stApp h3 { color: #1b5e20 !important; }'
        '.stApp p, .stApp label, .stApp .stMarkdown { color: #263238 !important; }'
        '.stApp [data-testid="stMetricLabel"] { color: #455a64 !important; }'
        '.stApp [data-testid="stMetricValue"] { color: #1b5e20 !important; }'
        'section[data-testid="stSidebar"] { background: #f1f8e9 !important; }'
        'section[data-testid="stSidebar"] * { color: #1b5e20 !important; }'
        '.stButton > button { background: linear-gradient(135deg, #16a34a, #22c55e) !important; color: white !important; }'
        '</style>',
        unsafe_allow_html=True
    )

# (v24: Arabic RTL CSS removed — site is English/LTR only)

st.sidebar.markdown("---")

# ============================================================
# SIDEBAR — MODE SELECTION (default from wizard)
# ============================================================
st.sidebar.header(t('viridis_mode'))
_mode_default = st.session_state.wizard_data.get(
    'app_mode', "🏠 Single Department")
_mode_options = ["🏠 Single Department", "🏥 Full Hospital Site Plan"]
app_mode = st.sidebar.radio(
    t('choose_scope'),
    _mode_options,
    index=_mode_options.index(_mode_default) if _mode_default in _mode_options else 0
)

# Show current generation status + reset option
_has_single = 'single_config' in st.session_state
_has_hospital = 'hospital_config' in st.session_state
if _has_single or _has_hospital:
    with st.sidebar.expander("📌 Active Design Status", expanded=False):
        if _has_single:
            st.success(f"✅ Single Dept: "
                       f"{st.session_state.single_config['facility_type']}")
        if _has_hospital:
            n_depts = len(st.session_state.hospital_config['selected_depts'])
            st.success(f"✅ Hospital: {n_depts} departments locked in")
        if st.button("🔄 Reset All Designs", use_container_width=True):
            st.session_state.pop('single_config', None)
            st.session_state.pop('hospital_config', None)
            st.rerun()

st.sidebar.markdown("---")

# ============================================================
# ════════════════════════════════════════════════════════════
# MODE 1: SINGLE DEPARTMENT (DETAILED DESIGN)
# ════════════════════════════════════════════════════════════
# ============================================================
if app_mode == "🏠 Single Department":

    st.sidebar.header("🛠️ Core Design Configuration")

    # Group facility selector by category
    facility_groups = {}
    for name, info in MEDICAL_STANDARDS.items():
        cat = info["category"]
        facility_groups.setdefault(cat, []).append(f"{info['icon']} {name}")

    # v23: pre-fill category + facility from the wizard's choice so the
    # sidebar reflects the user's wizard pick (rather than the first
    # category alphabetically). Only matters for the initial Results view.
    _wd_facility = st.session_state.wizard_data.get('facility_type', '')
    _default_cat_idx = 0
    _default_fac_idx = 0
    if _wd_facility and _wd_facility in MEDICAL_STANDARDS:
        _wf_cat = MEDICAL_STANDARDS[_wd_facility]["category"]
        _cats_list = list(facility_groups.keys())
        if _wf_cat in _cats_list:
            _default_cat_idx = _cats_list.index(_wf_cat)
            # And find the facility's index within that category list
            _wf_label = f"{MEDICAL_STANDARDS[_wd_facility]['icon']} {_wd_facility}"
            _facs_in_cat = facility_groups[_wf_cat]
            if _wf_label in _facs_in_cat:
                _default_fac_idx = _facs_in_cat.index(_wf_label)

    category_choice = st.sidebar.selectbox(
        "Department Category:",
        list(facility_groups.keys()),
        index=_default_cat_idx)
    # Facility list depends on the chosen category — recompute index safely
    _facs_list_now = facility_groups[category_choice]
    _safe_fac_idx = _default_fac_idx if _default_fac_idx < len(_facs_list_now) else 0
    facility_label = st.sidebar.selectbox(
        "Select Facility:",
        _facs_list_now,
        index=_safe_fac_idx)
    facility_type = facility_label.split(" ", 1)[1]  # strip icon
    std = MEDICAL_STANDARDS[facility_type]

    st.sidebar.caption(f"📋 Standard: ≥{std['min_area_per_unit']} m²/unit • "
                       f"ACH {std['ach']} • {std['pressure']}")

    total_area = st.sidebar.slider("Total Room Area (m²):",
                                    min_value=20, max_value=200, value=60, step=5)
    rooms_count = st.sidebar.number_input("Number of Rooms in Department:",
                                           min_value=1, max_value=10, value=2)
    units_per_room = st.sidebar.number_input(
        f"Number of {std['unit_name']}s per Room:",
        min_value=1, max_value=std['max_units'],
        value=min(3, std['max_units'])
    )

    # ─── Climate: auto-detect from city OR manual ───
    climate_mode = st.sidebar.radio("Climate Input:",
        ["🌍 Auto-Detect from City", "✋ Manual Zone"],
        help="Auto-detect uses a built-in climate database (NASA POWER + "
             "ASHRAE data) to set climate zone, solar potential, and "
             "rainfall for your city.")
    if climate_mode == "🌍 Auto-Detect from City":
        _city_list = sorted(CLIMATE_DATABASE.keys())
        # Default to Cairo
        sel_city = st.sidebar.selectbox("Select City:", _city_list,
            index=_city_list.index('Cairo'))
        _cdata = detect_climate(sel_city)
        climate = _cdata['zone']
        st.sidebar.caption(
            f"{_cdata['country']} **{sel_city}**: {climate} zone · "
            f"☀️ {_cdata['sun_hours']}h sun/day · "
            f"🌡️ {_cdata['temp']}°C avg · "
            f"🌧️ {_cdata['rainfall']}mm/yr · "
            f"💧 {_cdata['humidity']}% RH")
    else:
        sel_city = None
        climate = st.sidebar.radio("Regional Climate Zone:",
            ["Hot", "Cold", "Moderate"])
    lighting_type = st.sidebar.radio("Lighting Preference:", ["Artificial Only", "Natural", "Mixed"])

    # ═══════════════════════════════════════════════════════
    # SMART DISABLING — Context-aware option visibility
    # Only show options that make sense for the chosen facility
    # ═══════════════════════════════════════════════════════
    # Facility categories for relevance checks
    _SCRUB_FACILITIES = {
        'Operating Room (OR)', 'Intensive Care Unit (ICU)',
        'Emergency Room (ER)', 'NICU (Neonatal ICU)',
        'Maternity & Delivery', 'Dialysis Unit',
        'Laboratory', 'Sterilization (CSSD)',
        'Dental Clinic', 'Outpatient Clinic',
        'Radiology Department', 'Blood Bank',
    }
    _LAMINAR_FACILITIES = {
        'Operating Room (OR)', 'NICU (Neonatal ICU)',
        'Sterilization (CSSD)', 'Blood Bank', 'Oncology',
    }
    _SMART_PANEL_FACILITIES = {
        'Operating Room (OR)', 'Intensive Care Unit (ICU)',
        'Emergency Room (ER)', 'NICU (Neonatal ICU)',
        'Radiology Department', 'MRI Room', 'CT Scan Room',
        'Maternity & Delivery', 'Dialysis Unit',
        'Laboratory', 'Sterilization (CSSD)', 'Oncology',
    }
    _MONITOR_FACILITIES = {
        'Operating Room (OR)', 'Intensive Care Unit (ICU)',
        'Emergency Room (ER)', 'NICU (Neonatal ICU)',
        'Pediatric Ward', 'General Patient Ward',
        'Maternity & Delivery', 'Dialysis Unit', 'Oncology',
        'Outpatient Clinic', 'Dental Clinic',
    }
    _ANESTHESIA_FACILITIES = {
        'Operating Room (OR)', 'Maternity & Delivery',
    }
    _IV_FACILITIES = {
        'Operating Room (OR)', 'Intensive Care Unit (ICU)',
        'Emergency Room (ER)', 'NICU (Neonatal ICU)',
        'Pediatric Ward', 'General Patient Ward',
        'Maternity & Delivery', 'Dialysis Unit', 'Oncology',
        'Outpatient Clinic',
    }
    _MED_CART_FACILITIES = _IV_FACILITIES | {'Pharmacy'}
    _PARTITION_FACILITIES = {
        'General Patient Ward', 'Intensive Care Unit (ICU)',
        'Emergency Room (ER)', 'NICU (Neonatal ICU)',
        'Pediatric Ward', 'Dialysis Unit', 'Maternity & Delivery',
        'Operating Room (OR)', 'Oncology', 'Dental Clinic',
    }

    st.sidebar.markdown("---")
    st.sidebar.subheader("📐 Spatial & Visual Components")

    # Partition — only for multi-unit clinical rooms
    if std['max_units'] > 1 and facility_type in _PARTITION_FACILITIES:
        partition_type = st.sidebar.selectbox("Unit Zoning Partitions:",
            ["Open Layout", "Antibacterial Glass Walls"])
    else:
        partition_type = "Open Layout"
        st.sidebar.caption(
            f"ℹ️ *Partitions not applicable for {facility_type}*")

    # Scrub Sink — only for clinical/sterile facilities
    if facility_type in _SCRUB_FACILITIES:
        scrub_sink = st.sidebar.selectbox("Scrub Sink Station:",
            ["None", "Single Medical Sink", "Double Stainless Sink"],
            index=1 if facility_type in _LAMINAR_FACILITIES else 0)
    else:
        scrub_sink = "None"
        st.sidebar.caption(
            f"ℹ️ *Scrub sink not needed for {facility_type}*")

    # Smart Control Panel — only for technical facilities
    if facility_type in _SMART_PANEL_FACILITIES:
        control_panel = st.sidebar.radio("Wall Control Interface:",
            ["Analogue Gauges", "Smart Touchscreen Console"])
    else:
        control_panel = "Analogue Gauges"
        st.sidebar.caption(
            f"ℹ️ *Smart control panel not applicable for {facility_type}*")

    # Laminar Flow Ceiling — only for sterile zones
    if facility_type in _LAMINAR_FACILITIES:
        laminar_ceiling = st.sidebar.checkbox(
            "Deploy Ultra-Clean Laminar Flow Ceiling Grid",
            value=facility_type in ["Operating Room (OR)", "NICU (Neonatal ICU)"]
        )
    else:
        laminar_ceiling = False
        st.sidebar.caption(
            f"ℹ️ *Laminar flow ceiling reserved for sterile zones only*")

    st.sidebar.markdown("---")
    st.sidebar.subheader("🪑 Additional Equipment")
    st.sidebar.caption(
        f"All medical equipment available for **{facility_type}**. "
        f"Items not relevant to this department are hidden automatically.")

    # ─── Base 4 equipment (kept as direct variables for downstream code) ───
    if facility_type in _MONITOR_FACILITIES:
        add_monitor = st.sidebar.checkbox(
            "⚕️ Vital Signs Monitor",
            value=facility_type in ["Operating Room (OR)",
                                     "Intensive Care Unit (ICU)",
                                     "Emergency Room (ER)", "NICU (Neonatal ICU)"],
            key="base_monitor",
            help="Power: 200W · continuous patient monitoring")
    else:
        add_monitor = False

    if facility_type in _ANESTHESIA_FACILITIES:
        add_anesthesia = st.sidebar.checkbox(
            "💉 Anesthesia Machine",
            value=facility_type == "Operating Room (OR)",
            key="base_anesthesia",
            help="Power: 500W · for surgical anesthesia delivery")
    else:
        add_anesthesia = False

    if facility_type in _IV_FACILITIES:
        add_iv_stand = st.sidebar.checkbox(
            "💧 IV Stand", value=True, key="base_iv",
            help="Mobile pole for intravenous fluid bags")
    else:
        add_iv_stand = False

    if facility_type in _MED_CART_FACILITIES:
        add_med_cart = st.sidebar.checkbox(
            "🛒 Medication Cart", value=True, key="base_medcart",
            help="Mobile cart with drawers for dispensing meds")
    else:
        add_med_cart = False

    # ─── Extended/specialized equipment (flat under same section) ───
    # All 15 extended items filtered by facility applicability, including
    # the Fire Extinguisher (universal — applicable to every facility).
    _applicable_eq = [eq for eq in EXTENDED_EQUIPMENT
                      if _is_eq_applicable(facility_type, eq)]

    active_equipment = []
    for eq in _applicable_eq:
        default = _is_eq_default_on(facility_type, eq)
        checked = st.sidebar.checkbox(
            f"{eq['icon']} {eq['label_en']}",
            value=default, key=f"ext_{eq['key']}",
            help=f"Power: {eq['power_w']}W" if eq['power_w'] > 0
                 else "No power draw"
        )
        if checked:
            active_equipment.append(eq['key'])

    # ─── Total load summary across all selected equipment ───
    _base_loads = ((200 if add_monitor else 0) +
                   (500 if add_anesthesia else 0))
    _ext_loads = sum(e['power_w'] for e in EXTENDED_EQUIPMENT
                     if e['key'] in active_equipment)
    _total_eq_load = _base_loads + _ext_loads
    if _total_eq_load > 0:
        st.sidebar.caption(
            f"⚡ Total equipment load: **{_total_eq_load} W** "
            f"({_total_eq_load/1000:.2f} kW)")

    # If absolutely no equipment is applicable, show notice
    if (not _applicable_eq and not add_monitor and not add_anesthesia and
            not add_iv_stand and not add_med_cart):
        st.sidebar.caption(
            f"ℹ️ *No equipment is typically used in {facility_type}.*")

    # ════════════════════════════════════════════════════════════════
    # 🚨 SAFETY & EMERGENCY — emergency exits + fire protection systems
    # NFPA 101 (Life Safety) + NFPA 72 (alarm) + NFPA 13 (sprinklers)
    # ════════════════════════════════════════════════════════════════
    st.sidebar.markdown("---")
    st.sidebar.subheader("🚨 Safety & Emergency")

    # — Emergency Exits —
    show_emergency_exits = st.sidebar.checkbox(
        "🚪 Emergency Exits", value=True,
        key="show_em_exits",
        help="Mark emergency egress doors on side walls per NFPA 101. "
             "Count auto-derived from occupant load (Fire Safety tab).")

    # — Smoke Detectors —
    show_smoke_detectors = st.sidebar.checkbox(
        "💨 Smoke Detectors", value=True,
        key="show_smoke",
        help="NFPA 72: 1 detector per ~84 m². Ceiling-mounted. "
             "Auto-placed on a regular grid (~7m spacing).")

    # — Sprinkler System —
    show_sprinklers = st.sidebar.checkbox(
        "💧 Sprinkler System", value=True,
        key="show_sprinklers",
        help="NFPA 13: 3.7m × 3.7m max spacing. Ceiling-mounted heads.")

    # — Pull Stations —
    show_pull_stations = st.sidebar.checkbox(
        "🚨 Fire Alarm Pull Stations", value=True,
        key="show_pull",
        help="NFPA 72: manual fire alarm activation stations placed "
             "near each exit. Required on egress paths.")

    # — Emergency Lighting —
    show_emergency_lighting = st.sidebar.checkbox(
        "💡 Emergency Lighting", value=True,
        key="show_em_light",
        help="NFPA 101: battery-backup lights above exits and along "
             "egress paths. 90-minute minimum operation required.")

    # ─── Medical Emergency Devices (Batch SE-2) ───
    # AED Station — applicable in virtually any clinical/public area
    if facility_type in _AED_APPLICABLE:
        show_aed = st.sidebar.checkbox(
            "❤️‍🩹 AED Station", value=True,
            key="show_aed",
            help="Automated External Defibrillator cabinet. "
                 "AHA: place so any person is within ~90 sec reach. "
                 "Wall-mounted, prominently visible.")
    else:
        show_aed = False

    # Eye Wash + Emergency Shower — only for chemical-handling rooms
    if facility_type in _CHEMICAL_FACILITIES:
        show_eye_wash = st.sidebar.checkbox(
            "👁️ Eye Wash Station", value=True,
            key="show_eye_wash",
            help="ANSI Z358.1: required within 10 seconds reach where "
                 "chemicals or biohazards are handled. Paired with safety "
                 "shower for full decontamination.")
        show_emergency_shower = st.sidebar.checkbox(
            "🚿 Emergency Safety Shower", value=True,
            key="show_em_shower",
            help="ANSI Z358.1: overhead drench shower for chemical "
                 "exposure. Activates within 1 second, runs 15+ min.")
    else:
        show_eye_wash = False
        show_emergency_shower = False
        st.sidebar.caption(
            f"ℹ️ *Eye wash / safety shower not required for "
            f"{facility_type} (no chemical exposure risk).*")

    # Patient Call Buttons — only for facilities with patient beds
    if facility_type in _BED_FACILITIES:
        show_call_buttons = st.sidebar.checkbox(
            "📞 Patient Call Buttons", value=True,
            key="show_call_btn",
            help="Nurse-call buttons at every bedside. Required for "
                 "any inpatient area. Count auto-scales with bed count.")
    else:
        show_call_buttons = False

    # ─── Wayfinding & Egress (Batch SE-3) ───
    # Egress Path — universal (any room with doors needs egress)
    show_egress_path = st.sidebar.checkbox(
        "🛤️ Egress Path Arrows", value=True,
        key="show_egress",
        help="NFPA 101: visual egress routes from interior toward each "
             "exit. Green dashed lines with arrows show occupants the "
             "fastest path during evacuation.")

    # Wayfinding Signs — universal (ISO 7010 standard)
    show_wayfinding = st.sidebar.checkbox(
        "🪧 Wayfinding Signs", value=True,
        key="show_wayfinding",
        help="ISO 7010: green directional safety signs mounted at "
             "decision points and along egress paths. Required at "
             "every change of direction.")

    # Refuge Areas — universal (ADA + NFPA 101)
    show_refuge_areas = st.sidebar.checkbox(
        "♿ Refuge Areas (ADA)", value=True,
        key="show_refuge",
        help="NFPA 101 + ADA 4.3: protected waiting space near each "
             "emergency exit for non-ambulatory persons awaiting "
             "assisted evacuation. Required in multi-story buildings.")

    # Emergency Stairs Access — feature removed in v22 (single-room view
    # doesn't represent multi-story context well). The sustainability
    # of stairs is now better covered via the Hospital Mode's site plan.
    show_emergency_stairs = False  # kept as False for backward-compat

    # ─── Hazard Markers (Batch SE-4) ───
    # Radiation shielding markers — only for X-ray/radiation facilities
    if facility_type in _RADIATION_FACILITIES:
        show_radiation = st.sidebar.checkbox(
            "☢️ Radiation Shielding Markers", value=True,
            key="show_radiation",
            help="IEC 60417: trefoil hazard symbol + lead-lined wall "
                 "indicators for X-ray, MRI, CT, and radiation therapy rooms.")
    else:
        show_radiation = False

    # Isolation room indicator
    if facility_type in _ISOLATION_FACILITIES:
        show_isolation = st.sidebar.checkbox(
            "☣️ Isolation Pod Indicator", value=False,
            key="show_isolation",
            help="CDC + OSHA: biohazard symbol + negative-pressure marker "
                 "for isolation rooms (TB, COVID, immunocompromised). "
                 "Disabled by default — enable for infectious patient rooms.")
    else:
        show_isolation = False

    # Hazmat storage marker
    if facility_type in _HAZMAT_FACILITIES:
        show_hazmat = st.sidebar.checkbox(
            "⚠️ Hazmat Storage Marker", value=True,
            key="show_hazmat",
            help="NFPA 704 hazard diamond on storage cabinets. "
                 "Required where chemicals, biohazards, or radioactive "
                 "materials are stored.")
    else:
        show_hazmat = False

    # If no hazards apply, show a clean note
    if not (facility_type in _RADIATION_FACILITIES or
            facility_type in _ISOLATION_FACILITIES or
            facility_type in _HAZMAT_FACILITIES):
        st.sidebar.caption(
            f"ℹ️ *No hazard markers typically required for {facility_type}.*")

    # ─── 🎬 Fire Evacuation Simulation (the secret feature!) ───
    st.sidebar.markdown("")
    if st.sidebar.button("🎬  Simulate Fire Evacuation",
                           key="run_evac_sim",
                           use_container_width=True):
        st.session_state.show_evac_sim = True

    # Show count summary
    if any([show_emergency_exits, show_smoke_detectors, show_sprinklers,
            show_pull_stations, show_emergency_lighting,
            show_aed, show_eye_wash, show_emergency_shower,
            show_call_buttons, show_egress_path, show_wayfinding,
            show_refuge_areas, show_emergency_stairs,
            show_radiation, show_isolation, show_hazmat]):
        st.sidebar.caption(
            "ℹ️ All counts and placements are auto-computed per "
            "NFPA / AHA / ANSI / ADA / ISO / IEC codes. See the "
            "**Fire Safety** tab in Engineering Analysis for detailed "
            "compliance metrics.")

    st.sidebar.markdown("---")
    st.sidebar.subheader("🌱 Sustainability Variables")

    # ─── Eco Materials (v22: expanded from flooring-only to 5 categories) ───
    st.sidebar.markdown("**🏗️ Eco-Friendly Materials**")
    st.sidebar.caption(
        "Select a material in each category. The expander below each "
        "field shows description, hospital use-cases, eco benefits, "
        "cost, and the reference standard.")

    # 1. FLOORING
    flooring = st.sidebar.selectbox(
        "🦶 Flooring Material:",
        list(MATERIALS_DB['flooring'].keys()),
        index=1,  # default: Eco-Linoleum (Natural)
        key="mat_flooring")
    render_material_info_card('flooring', flooring)

    # 2. WALL MATERIAL
    wall_material = st.sidebar.selectbox(
        "🧱 Wall Material:",
        list(MATERIALS_DB['wall'].keys()),
        index=0,  # default: Antimicrobial Gypsum
        key="mat_wall")
    render_material_info_card('wall', wall_material)

    # 3. CEILING MATERIAL
    ceiling_material = st.sidebar.selectbox(
        "🏛️ Ceiling Material:",
        list(MATERIALS_DB['ceiling'].keys()),
        index=0,  # default: Acoustic Mineral Fiber
        key="mat_ceiling")
    render_material_info_card('ceiling', ceiling_material)

    # 4. INSULATION
    insulation_material = st.sidebar.selectbox(
        "🌡️ Wall/Roof Insulation:",
        list(MATERIALS_DB['insulation'].keys()),
        index=0,  # default: Sheep Wool Batt
        key="mat_insulation")
    render_material_info_card('insulation', insulation_material)

    # 5. PAINT
    paint_material = st.sidebar.selectbox(
        "🖌️ Interior Paint:",
        list(MATERIALS_DB['paint'].keys()),
        index=0,  # default: Zero-VOC Antimicrobial
        key="mat_paint")
    render_material_info_card('paint', paint_material)

    st.sidebar.markdown("---")
    st.sidebar.markdown("**♻️ Operational Sustainability**")
    waste_mgmt = st.sidebar.selectbox("Medical Waste System:",
        ["Standard Centralized Disposal", "On-site Color-Coded Segregation",
         "Advanced Autoclaving & Shredding"])
    equip_efficiency = st.sidebar.radio("Medical Equipment Class:",
        ["Legacy Systems (Standard)", "Energy-Star / Eco-Mode Certified"])
    outdoor_air = st.sidebar.radio("Outdoor Air Quality Status:",
        ["Low Pollution (Rural/Suburban)", "High Pollution (Urban/Industrial)"])
    power_source = st.sidebar.selectbox("Primary Building Power Source:",
        ["National Grid Only", "Grid + Rooftop Solar PV", "Hybrid (Solar + Battery Storage)"])

    st.sidebar.markdown("---")
    st.sidebar.subheader("🎨 Visualization Options")
    view_mode = st.sidebar.radio("Active View:",
        ["📐 2D Plan", "🌐 3D Render", "📊 Both Side-by-Side"])
    show_grid = st.sidebar.checkbox("Show Measurement Grid", value=True)
    show_dimensions = st.sidebar.checkbox("Show Dimensions", value=True)
    # Staff/Patient Traffic Flow removed in v22 — egress paths (in Safety
    # & Emergency) now show movement direction more meaningfully.
    show_traffic_flow = False
    show_zones = st.sidebar.checkbox("Show Sterile Zones", value=True)
    day_night = st.sidebar.radio("Lighting Mode (3D):", ["☀️ Day", "🌙 Night"])

    # ============================================================
    # 💰 BUDGET MANAGER
    # ============================================================
    st.sidebar.markdown("---")
    st.sidebar.subheader("💰 Budget Manager")
    _curr_sym_bud = get_active_reference()['currency_symbol']
    # Estimate a default budget based on facility + area
    _ref_key_bud = st.session_state.get('reference', 'fgi')
    _default_cost_per_m2 = CONSTRUCTION_COSTS.get(_ref_key_bud, {}).get(
        facility_type, 2000 if _ref_key_bud == 'fgi' else 20000)
    _default_budget = int(_default_cost_per_m2 * total_area * rooms_count * 1.5)

    enable_budget = st.sidebar.checkbox(
        "Enable Budget Tracking", value=False,
        help="Set a project budget and see if your design fits within it. "
             "The system will warn you when you exceed it.")

    if enable_budget:
        user_budget = st.sidebar.number_input(
            f"Your Total Budget ({_curr_sym_bud}):",
            min_value=0,
            value=_default_budget,
            step=10000 if _ref_key_bud == 'fgi' else 100000,
            format="%d",
            help="Total construction + equipment + sustainability budget. "
                 "Includes contingency reserve."
        )
        st.sidebar.caption(
            f"💡 Suggested for {facility_type}: "
            f"~{_curr_sym_bud}{_default_budget:,} "
            f"(based on standard rates)")
    else:
        user_budget = 0

    # ============================================================
    # GENERATE BUTTON — captures all inputs into session_state
    # ============================================================
    st.sidebar.markdown("---")
    # Contextual prompt above the button — guides first-time users
    if not st.session_state.get('design_generated', False):
        st.sidebar.markdown(
            "<div style='padding:10px; background:#fff3e0; "
            "border-left:4px solid #f57c00; border-radius:6px; "
            "margin-bottom:8px; font-size:0.85em; color:#e65100'>"
            "👇 <b>Click below</b> to generate your design — your "
            "configuration is ready!</div>",
            unsafe_allow_html=True)
    generate_btn = st.sidebar.button(
        "🚀 Generate Design", type="primary",
        use_container_width=True, key="gen_single"
    )

    if generate_btn:
        # Mark design as generated — unblocks the Results page render.
        st.session_state.design_generated = True
        # Snapshot every input into session state — the rest of the script
        # will read from this snapshot, so widget changes don't trigger
        # any expensive re-rendering until the user clicks again.
        st.session_state.single_config = {
            'facility_type': facility_type, 'std': std,
            'total_area': total_area, 'rooms_count': rooms_count,
            'units_per_room': units_per_room, 'climate': climate,
            'lighting_type': lighting_type, 'partition_type': partition_type,
            'scrub_sink': scrub_sink, 'control_panel': control_panel,
            'laminar_ceiling': laminar_ceiling, 'add_monitor': add_monitor,
            'add_anesthesia': add_anesthesia, 'add_iv_stand': add_iv_stand,
            'add_med_cart': add_med_cart, 'flooring': flooring,
            'wall_material': wall_material,
            'ceiling_material': ceiling_material,
            'insulation_material': insulation_material,
            'paint_material': paint_material,
            'waste_mgmt': waste_mgmt, 'equip_efficiency': equip_efficiency,
            'outdoor_air': outdoor_air, 'power_source': power_source,
            'view_mode': view_mode, 'show_grid': show_grid,
            'show_dimensions': show_dimensions,
            'show_traffic_flow': show_traffic_flow,
            'show_zones': show_zones, 'day_night': day_night,
            'active_equipment': active_equipment,
            'enable_budget': enable_budget,
            'user_budget': user_budget,
            'sel_city': sel_city,
            'show_emergency_exits': show_emergency_exits,
            'show_smoke_detectors': show_smoke_detectors,
            'show_sprinklers': show_sprinklers,
            'show_pull_stations': show_pull_stations,
            'show_emergency_lighting': show_emergency_lighting,
            'show_aed': show_aed,
            'show_eye_wash': show_eye_wash,
            'show_emergency_shower': show_emergency_shower,
            'show_call_buttons': show_call_buttons,
            'show_egress_path': show_egress_path,
            'show_wayfinding': show_wayfinding,
            'show_refuge_areas': show_refuge_areas,
            'show_emergency_stairs': show_emergency_stairs,
            'show_radiation': show_radiation,
            'show_isolation': show_isolation,
            'show_hazmat': show_hazmat,
        }

    # ════════════════════════════════════════════════════════════════
    # v23: Empty-state guard — wait for user to click Generate before
    # rendering any results. Even if single_config exists (from the
    # wizard's commit), nothing renders until the Generate button is
    # pressed. After first Generate, subsequent sidebar changes auto-
    # rerun via Streamlit's normal flow (live update).
    # ════════════════════════════════════════════════════════════════
    if not st.session_state.get('design_generated', False):
        # Educational dashboard — clear 2-step guide
        st.markdown(
            "<div style='text-align:center; padding:30px 10px 20px 10px'>"
            "<h1 style='color:#1b5e20; margin-bottom:8px'>"
            "🍃 Viridis Design Studio</h1>"
            "<p style='color:#388e3c; font-size:1.1em; margin-top:0'>"
            "Your facility configuration is ready in the sidebar — "
            "follow the 2 steps below to generate the full design report."
            "</p></div>",
            unsafe_allow_html=True)

        col1, col2 = st.columns(2)
        with col1:
            st.markdown("""
            <div style="border:2px solid #2e7d32; border-radius:14px;
                        padding:24px 26px; background:#e8f5e9;
                        min-height:420px; box-sizing:border-box;
                        box-shadow:0 2px 8px rgba(46,125,50,0.08)">
              <div style="font-size:2.6em; text-align:center;
                          line-height:1; margin-bottom:8px">👈</div>
              <h3 style="color:#1b5e20; text-align:center;
                         margin:6px 0 14px 0; font-size:1.35em">
                Step 1 — Review &amp; Adjust
              </h3>
              <p style="color:#2e7d32; font-size:0.95em;
                        margin:0 0 10px 0">
                Open the sidebar (◀ icon if collapsed) and review
                your configuration:
              </p>
              <ul style="color:#2e7d32; font-size:0.9em;
                         line-height:1.65; margin:0 0 12px 0;
                         padding-left:22px">
                <li>🏥 Facility &amp; spatial dimensions</li>
                <li>🪑 Medical equipment selection</li>
                <li>🚨 Safety &amp; emergency systems</li>
                <li>🌱 Sustainability materials (5 categories)</li>
                <li>🌍 Site location &amp; budget</li>
              </ul>
              <p style="color:#388e3c; font-size:0.82em;
                        font-style:italic; margin:0;
                        padding-top:6px;
                        border-top:1px dashed #a5d6a7">
                💡 Tip: hover over any setting for the relevant
                code/standard reference.
              </p>
            </div>
            """, unsafe_allow_html=True)
        with col2:
            st.markdown("""
            <div style="border:2px solid #1565c0; border-radius:14px;
                        padding:24px 26px; background:#e3f2fd;
                        min-height:420px; box-sizing:border-box;
                        box-shadow:0 2px 8px rgba(21,101,192,0.08)">
              <div style="font-size:2.6em; text-align:center;
                          line-height:1; margin-bottom:8px">🚀</div>
              <h3 style="color:#0d47a1; text-align:center;
                         margin:6px 0 14px 0; font-size:1.35em">
                Step 2 — Press Generate
              </h3>
              <p style="color:#1565c0; font-size:0.95em;
                        margin:0 0 10px 0">
                Scroll to the bottom of the sidebar and click
                <b>🚀 Generate Design</b>. You will get:
              </p>
              <ul style="color:#1565c0; font-size:0.9em;
                         line-height:1.65; margin:0 0 12px 0;
                         padding-left:22px">
                <li>📐 2D blueprint with measurements</li>
                <li>🌐 Interactive 3D rendering</li>
                <li>🔬 HVAC, fire, water, power analyses</li>
                <li>🏆 LEED / Estidama / Mostadam forecast</li>
                <li>📑 Bill of Quantities + Functional Program</li>
                <li>🤖 AI design checks &amp; chatbot</li>
              </ul>
              <p style="color:#0d47a1; font-size:0.82em;
                        font-style:italic; margin:0;
                        padding-top:6px;
                        border-top:1px dashed #90caf9">
                ⚡ After the first generation, sidebar changes
                will update results live.
              </p>
            </div>
            """, unsafe_allow_html=True)

        # Optional quick stats showing what's available in the design
        st.markdown(
            "<div style='text-align:center; margin-top:30px; "
            "padding:15px; background:#f1f8e9; border-radius:8px'>"
            "<span style='color:#558b2f; font-size:0.9em'>"
            "Currently configured: "
            f"<b>{facility_type}</b> · "
            f"{total_area:.0f} m² × {rooms_count} room(s) · "
            f"<b>{climate}</b> climate · "
            f"<b>{flooring}</b> flooring"
            "</span></div>",
            unsafe_allow_html=True)

        st.stop()

    # Restore variables from the locked-in config snapshot
    cfg = st.session_state.single_config
    facility_type = cfg['facility_type']
    std = cfg['std']
    total_area = cfg['total_area']
    rooms_count = cfg['rooms_count']
    units_per_room = cfg['units_per_room']
    climate = cfg['climate']
    lighting_type = cfg['lighting_type']
    partition_type = cfg['partition_type']
    scrub_sink = cfg['scrub_sink']
    control_panel = cfg['control_panel']
    laminar_ceiling = cfg['laminar_ceiling']
    add_monitor = cfg['add_monitor']
    add_anesthesia = cfg['add_anesthesia']
    add_iv_stand = cfg['add_iv_stand']
    add_med_cart = cfg['add_med_cart']
    active_equipment = cfg.get('active_equipment', [])
    enable_budget = cfg.get('enable_budget', False)
    user_budget = cfg.get('user_budget', 0)
    sel_city = cfg.get('sel_city', None)
    show_emergency_exits = cfg.get('show_emergency_exits', True)
    show_smoke_detectors = cfg.get('show_smoke_detectors', True)
    show_sprinklers = cfg.get('show_sprinklers', True)
    show_pull_stations = cfg.get('show_pull_stations', True)
    show_emergency_lighting = cfg.get('show_emergency_lighting', True)
    show_aed = cfg.get('show_aed', True)
    show_eye_wash = cfg.get('show_eye_wash', False)
    show_emergency_shower = cfg.get('show_emergency_shower', False)
    show_call_buttons = cfg.get('show_call_buttons', False)
    show_egress_path = cfg.get('show_egress_path', True)
    show_wayfinding = cfg.get('show_wayfinding', True)
    show_refuge_areas = cfg.get('show_refuge_areas', True)
    show_emergency_stairs = cfg.get('show_emergency_stairs', True)
    show_radiation = cfg.get('show_radiation', False)
    show_isolation = cfg.get('show_isolation', False)
    show_hazmat = cfg.get('show_hazmat', False)
    flooring = cfg['flooring']
    # New material categories (v22) — with safe defaults for old snapshots
    wall_material = cfg.get('wall_material', 'Antimicrobial Gypsum')
    ceiling_material = cfg.get('ceiling_material', 'Acoustic Mineral Fiber')
    insulation_material = cfg.get('insulation_material', 'Sheep Wool Batt')
    paint_material = cfg.get('paint_material', 'Zero-VOC Antimicrobial')
    waste_mgmt = cfg['waste_mgmt']
    equip_efficiency = cfg['equip_efficiency']
    outdoor_air = cfg['outdoor_air']
    power_source = cfg['power_source']
    view_mode = cfg['view_mode']
    show_grid = cfg['show_grid']
    show_dimensions = cfg['show_dimensions']
    show_traffic_flow = cfg['show_traffic_flow']
    show_zones = cfg['show_zones']
    day_night = cfg['day_night']

    # ============================================================
    # METRICS CALCULATION
    # ============================================================
    def calculate_metrics():
        """Energy/cost calculation based on EUI (Energy Use Intensity)
        sourced from FGI/ASHRAE OR Egyptian MoH benchmarks.

        Formula (per Energy Star Healthcare benchmarks):
            Annual Energy = EUI_dept × Total Area
            EUI_dept = EUI_baseline × department_multiplier
            Monthly Energy = Annual / 12
            Monthly Cost = Monthly Energy × electricity_price
        """
        ref = get_active_reference()
        green_score = 30
        if facility_type in ["Operating Room (OR)", "Radiology Department",
                              "MRI Room", "CT Scan Room"]:
            green_score += 10
        if lighting_type in ["Natural", "Mixed"]: green_score += 10
        if flooring == "Eco-Linoleum (Natural)": green_score += 10
        if waste_mgmt == "Advanced Autoclaving & Shredding": green_score += 10
        if equip_efficiency == "Energy-Star / Eco-Mode Certified": green_score += 10
        if control_panel == "Smart Touchscreen Console": green_score += 5
        if power_source == "Grid + Rooftop Solar PV": green_score += 15
        elif power_source == "Hybrid (Solar + Battery Storage)": green_score += 20
        green_score = min(green_score, 100)

        # ─── EUI-based energy calculation (REPLACES the old guesswork) ───
        # Source: Energy Star Portfolio Manager — Hospital benchmarks
        eui_base = ref['eui_baseline']  # kWh/m²/year
        dept_mult = DEPT_EUI_MULTIPLIERS.get(facility_type, 1.0)
        # Effective EUI for this department type
        eui_dept = eui_base * dept_mult  # kWh/m²/year

        # Total floor area (m²)
        total_floor_area = total_area * rooms_count

        # Annual energy (kWh/year) → Monthly (kWh/month)
        annual_normal = eui_dept * total_floor_area
        monthly_normal = annual_normal / 12

        # ─── Savings percentage (additive contributions, capped) ───
        # Source: ASHRAE 90.1 + LEED Healthcare scorecard
        savings_pct = 0.06  # baseline (better building envelope)
        if climate in ["Hot", "Cold"]:
            savings_pct += 0.08  # smart HVAC for extreme climate
        if lighting_type in ["Natural", "Mixed"]:
            savings_pct += 0.10  # daylight harvesting (IES recommended)
        if equip_efficiency == "Energy-Star / Eco-Mode Certified":
            savings_pct += 0.12  # Energy Star equipment typical savings
        if flooring == "Eco-Linoleum (Natural)":
            savings_pct += 0.02  # marginal (thermal mass)
        if waste_mgmt == "Advanced Autoclaving & Shredding":
            savings_pct += 0.03  # reduced transport emissions
        if power_source == "Grid + Rooftop Solar PV":
            savings_pct += 0.20  # solar offset (typical 15-25%)
        elif power_source == "Hybrid (Solar + Battery Storage)":
            savings_pct += 0.30  # solar + battery deeper offset
        savings_pct = min(savings_pct, 0.70)  # realistic cap

        # Apply savings
        monthly_green = monthly_normal * (1 - savings_pct)
        energy_saved = monthly_normal - monthly_green

        # ─── Cost calculation in local currency ───
        money_saved = energy_saved * ref['electricity_price']  # currency/month

        return (green_score, monthly_normal, monthly_green,
                energy_saved, savings_pct, money_saved)

    def check_compliance():
        required_area = std["min_area_per_unit"] * units_per_room
        ref = get_active_reference()
        cite = ref['citations'].get('min_area', '')
        issues = []
        if total_area < required_area:
            issues.append(f"⚠️ Area too small: need ≥ {required_area} m² for "
                          f"{units_per_room} {std['unit_name']}s "
                          f"({ref['short']})\n   📖 Source: {cite}")
        if total_area / units_per_room < std["min_area_per_unit"]:
            issues.append(f"⚠️ Per-unit area below minimum "
                          f"{std['min_area_per_unit']} m² ({ref['short']})")
        return issues

    green_score, normal_env, green_env, saved_env, saved_pct, saved_money = calculate_metrics()
    compliance_issues = check_compliance()
    _active_ref = get_active_reference()
    _curr = _active_ref['currency_symbol']
    _ref_short = _active_ref['short']

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("🌲 Green Score", f"{green_score}%",
              help=f"Composite sustainability score (0-100%). "
                   f"Source: LEED Healthcare scorecard adapted")
    c2.metric("📉 Energy Saved", f"{saved_env:.0f} kWh/mo",
              delta=f"-{saved_pct*100:.1f}%",
              help=f"Monthly energy reduction vs traditional design. "
                   f"Formula: EUI × Area × savings%. "
                   f"Source: {_active_ref['citations']['eui']}")
    c3.metric("💰 Money Saved", f"{_curr}{saved_money:,.0f}/mo",
              help=f"Energy cost savings at "
                   f"{_curr}{_active_ref['electricity_price']:.2f}/kWh. "
                   f"Source: {_active_ref['citations']['electricity_price']}")
    c4.metric(f"📋 {_ref_short} Compliance",
              "✅ Compliant" if not compliance_issues else f"⚠️ {len(compliance_issues)} issues",
              help=f"Validation against {_active_ref['name']}")

    if compliance_issues:
        with st.expander("⚠️ Compliance Warnings", expanded=False):
            for i in compliance_issues:
                st.warning(i)

    # ════════════════════════════════════════════════════════════════
    # 💰 BUDGET DASHBOARD (only if enabled by user)
    # ════════════════════════════════════════════════════════════════
    if enable_budget and user_budget > 0:
        st.markdown("---")
        st.subheader("💰 Budget Analysis")

        # Calculate component costs
        construction_cost = calculate_construction_cost(
            facility_type, total_area, rooms_count)
        base_eq_dict = {
            'add_monitor': add_monitor, 'add_anesthesia': add_anesthesia,
            'add_iv_stand': add_iv_stand, 'add_med_cart': add_med_cart,
        }
        equipment_cost = calculate_equipment_cost(
            active_equipment, base_eq_dict, units_per_room, rooms_count)
        # Eco features applicable to single-room: glass partitions + laminar
        eco_dict = {
            'partition_type': partition_type == "Antibacterial Glass Walls",
            'laminar_ceiling': laminar_ceiling,
        }
        eco_cost = calculate_eco_features_cost(
            eco_dict, total_area=total_area * rooms_count,
            n_or_rooms=rooms_count if facility_type == "Operating Room (OR)" else 0)
        # Solar/power features (for single mode, approximate)
        if power_source == "Grid + Rooftop Solar PV":
            eco_cost += ECO_FEATURE_COSTS[_ref_key_bud]['h_solar']['cost'] * total_area * 0.5
        elif power_source == "Hybrid (Solar + Battery Storage)":
            eco_cost += (ECO_FEATURE_COSTS[_ref_key_bud]['h_solar']['cost'] * total_area * 0.5
                         + ECO_FEATURE_COSTS[_ref_key_bud]['h_smart_grid']['cost'])

        total_cost = construction_cost + equipment_cost + eco_cost
        # Add contingency reserve (~10%)
        contingency = total_cost * 0.10
        grand_total = total_cost + contingency

        # ─── Cost breakdown metrics ───
        bc1, bc2, bc3, bc4 = st.columns(4)
        bc1.metric("🏗️ Construction",
                   f"{_curr}{construction_cost:,.0f}",
                   help=f"Source: RSMeans Healthcare Construction Data 2024 "
                        f"({_curr}{construction_cost/(total_area*rooms_count):,.0f}/m²)")
        bc2.metric("🏥 Equipment",
                   f"{_curr}{equipment_cost:,.0f}",
                   help="Sum of all selected medical equipment costs")
        bc3.metric("🌱 Eco Features",
                   f"{_curr}{eco_cost:,.0f}",
                   help="Sustainability + premium features (glass, laminar, solar)")
        bc4.metric("📊 Total (incl. 10% contingency)",
                   f"{_curr}{grand_total:,.0f}",
                   help="Construction + Equipment + Eco + Contingency reserve")

        # ─── Affordability check ───
        budget_pct = (grand_total / user_budget) * 100
        if grand_total <= user_budget:
            remaining = user_budget - grand_total
            st.success(
                f"✅ **Within Budget!** Your design costs "
                f"{_curr}{grand_total:,.0f} ({budget_pct:.0f}% of "
                f"{_curr}{user_budget:,}). "
                f"You have **{_curr}{remaining:,.0f}** remaining "
                f"({100-budget_pct:.0f}% buffer) — consider adding more "
                f"sustainability features or upgrades!"
            )
        else:
            overage = grand_total - user_budget
            st.error(
                f"❌ **Over Budget by {_curr}{overage:,.0f}** "
                f"({budget_pct:.0f}% of budget). "
                f"Consider:\n"
                f"- Reducing total area by {(overage/construction_cost*total_area*rooms_count):.0f}m²\n"
                f"- Removing non-essential equipment ({_curr}{equipment_cost*0.3:,.0f} possible savings)\n"
                f"- Choosing 'National Grid Only' to remove solar/storage cost"
            )

        # ─── ROI Calculator ───
        with st.expander("📈 ROI & Payback Analysis", expanded=False):
            # Monthly savings already calculated above
            annual_savings = saved_money * 12
            if annual_savings > 0:
                payback_years = grand_total / annual_savings
                roi_10yr = (annual_savings * 10 - grand_total) / grand_total * 100
                rc1, rc2, rc3 = st.columns(3)
                rc1.metric("Annual Savings",
                           f"{_curr}{annual_savings:,.0f}/year",
                           help="From energy efficiency vs traditional design")
                rc2.metric("Payback Period",
                           f"{payback_years:.1f} years" if payback_years < 50 else "N/A",
                           help="Time to recover total investment from energy savings")
                rc3.metric("10-Year ROI",
                           f"{roi_10yr:+.0f}%",
                           help="Return on investment over 10 years")
                st.caption(f"💡 Assuming stable electricity prices at "
                           f"{_curr}{_active_ref['electricity_price']:.2f}/kWh "
                           f"and {saved_pct*100:.0f}% energy savings rate.")
            else:
                st.info("Enable more sustainability features to calculate ROI.")

    # ════════════════════════════════════════════════════════════════
    # 🏆 GREEN BUILDING CERTIFICATION PREDICTOR (LEED / Estidama / Mostadam)
    # ════════════════════════════════════════════════════════════════
    st.markdown("---")
    st.subheader("🏆 Green Building Certification Forecast")

    # Build the choices dict from current design
    cert_choices = {
        'solar_pv': power_source in ["Grid + Rooftop Solar PV",
                                      "Hybrid (Solar + Battery Storage)"],
        'energy_star': equip_efficiency == "Energy-Star / Eco-Mode Certified",
        'smart_grid': power_source == "Hybrid (Solar + Battery Storage)",
        'natural_light': lighting_type in ["Natural", "Mixed"],
        'healing_garden': False,  # single dept — no garden
        'rainwater': False,
        'greywater': False,
        'eco_flooring': flooring == "Eco-Linoleum (Natural)",
        'waste_mgmt': waste_mgmt == "Advanced Autoclaving & Shredding",
        'green_roof': False,
        'ev_chargers': False,
        'garden_site': False,
    }
    cert = predict_certification(cert_choices)

    cert_c1, cert_c2, cert_c3 = st.columns(3)
    with cert_c1:
        st.markdown(f"### LEED v4")
        st.markdown(f"<div style='padding:10px;border-radius:8px;"
                    f"background:{cert['leed_color']};text-align:center;"
                    f"color:#1a1a1a;font-weight:bold;font-size:1.1em'>"
                    f"{cert['leed_level']}<br>{cert['leed_score']}/100 pts</div>",
                    unsafe_allow_html=True)
    with cert_c2:
        st.markdown(f"### Estidama 🇦🇪")
        st.markdown(f"<div style='padding:10px;border-radius:8px;"
                    f"background:#e8f5e9;text-align:center;color:#1b5e20;"
                    f"font-weight:bold;font-size:1.3em'>"
                    f"{'⭐' * cert['pearls']}<br>{cert['pearls']} Pearl"
                    f"{'s' if cert['pearls'] > 1 else ''}</div>",
                    unsafe_allow_html=True)
    with cert_c3:
        st.markdown(f"### Mostadam 🇸🇦")
        st.markdown(f"<div style='padding:10px;border-radius:8px;"
                    f"background:#e8f5e9;text-align:center;color:#1b5e20;"
                    f"font-weight:bold;font-size:1.1em'>"
                    f"{cert['mostadam']}<br>Rating</div>",
                    unsafe_allow_html=True)

    with st.expander("📋 Certification Credit Breakdown", expanded=False):
        st.caption(f"Earned **{cert['earned_points']}/{cert['max_points']}** "
                   f"sustainability credits. Here's how to improve:")
        for detail in cert['credits_detail']:
            icon = "✅" if detail['achieved'] else "⬜"
            st.markdown(f"{icon} **{detail['credit']}** "
                        f"({detail['category']}) — "
                        f"{detail['earned']}/{detail['max']} pts")
        # Suggestions for unearned credits
        missing = [d for d in cert['credits_detail'] if not d['achieved']]
        if missing:
            st.markdown("---")
            st.markdown("**💡 To improve your rating:**")
            for d in missing[:4]:
                st.caption(f"• Add **{d['credit']}** → +{d['max']} pts")

    # ════════════════════════════════════════════════════════════════
    # ☀️ SOLAR & RAINWATER POTENTIAL (if city auto-detected)
    # ════════════════════════════════════════════════════════════════
    if sel_city:
        st.markdown("---")
        st.subheader(f"☀️ Site Renewable Potential — {sel_city}")
        _cd = detect_climate(sel_city)
        roof_area = total_area * rooms_count
        solar_annual = estimate_solar_potential(_cd['sun_hours'], roof_area * 0.6)
        rainwater_annual = estimate_rainwater_harvest(_cd['rainfall'], roof_area)

        sp1, sp2, sp3 = st.columns(3)
        sp1.metric("☀️ Solar Potential",
                   f"{solar_annual/1000:,.1f} MWh/yr",
                   help=f"If 60% of roof ({roof_area*0.6:.0f}m²) had PV panels "
                        f"at {_cd['sun_hours']}h peak sun. Source: NREL PVWatts")
        sp2.metric("💧 Rainwater Harvest",
                   f"{rainwater_annual/1000:,.1f} m³/yr",
                   help=f"From {_cd['rainfall']}mm annual rainfall on "
                        f"{roof_area:.0f}m² roof (0.8 runoff coefficient)")
        # Solar offset of energy use
        solar_offset_pct = min(100, (solar_annual / (normal_env * 12)) * 100) if normal_env > 0 else 0
        sp3.metric("⚡ Potential Energy Offset",
                   f"{solar_offset_pct:.0f}%",
                   help="% of annual energy demand that rooftop solar could cover")
        if solar_offset_pct > 50:
            st.success(f"🌟 Excellent solar site! Rooftop PV could offset "
                       f"{solar_offset_pct:.0f}% of this facility's energy use.")
        elif _cd['sun_hours'] >= 6:
            st.info(f"☀️ {sel_city} has strong solar resource "
                    f"({_cd['sun_hours']}h/day) — solar PV highly recommended.")

    # ════════════════════════════════════════════════════════════════
    # 🏗️ ECO-MATERIALS SELECTED — summary cards
    # ════════════════════════════════════════════════════════════════
    st.markdown("---")
    st.subheader("🏗️ Eco-Materials Specified for This Design")

    _mats_used = [
        ('Flooring',   'flooring',   flooring,           '🦶'),
        ('Wall',       'wall',       wall_material,      '🧱'),
        ('Ceiling',    'ceiling',    ceiling_material,   '🏛️'),
        ('Insulation', 'insulation', insulation_material,'🌡️'),
        ('Paint',      'paint',      paint_material,     '🖌️'),
    ]
    # Render as a 5-column row of compact cards
    mat_cols = st.columns(5)
    _total_mat_cost = 0
    for (cat_label, cat_key, mat_name, cat_icon), col in zip(_mats_used,
                                                              mat_cols):
        m_info = MATERIALS_DB.get(cat_key, {}).get(mat_name, {})
        m_icon = m_info.get('icon', '📦')
        m_cost = m_info.get('cost_per_m2_usd', 0)
        _total_mat_cost += m_cost
        with col:
            st.markdown(
                f"<div style='border:1.5px solid #2e7d32; "
                f"border-radius:8px; padding:10px 8px; "
                f"background:#e8f5e9; text-align:center; "
                f"min-height:110px'>"
                f"<div style='font-size:1.3em'>{cat_icon} {cat_label}</div>"
                f"<div style='font-size:1.6em; margin:4px 0'>{m_icon}</div>"
                f"<div style='font-size:0.85em; font-weight:bold; "
                f"color:#1b5e20'>{mat_name}</div>"
                f"<div style='font-size:0.75em; color:#388e3c; "
                f"margin-top:4px'>${m_cost}/m²</div>"
                f"</div>",
                unsafe_allow_html=True)

    # Combined materials info expander
    with st.expander("📚 Full Material Specifications & References",
                      expanded=False):
        for cat_label, cat_key, mat_name, cat_icon in _mats_used:
            m_info = MATERIALS_DB.get(cat_key, {}).get(mat_name, {})
            if not m_info:
                continue
            st.markdown(f"### {cat_icon} {cat_label}: **{mat_name}**")
            mc1, mc2 = st.columns([2, 1])
            with mc1:
                st.markdown(f"**📋 Description:** {m_info['description']}")
                st.markdown(f"**🏥 Best for:** {m_info['use_case']}")
                st.markdown(f"**🌱 Eco benefits:** {m_info['eco_benefits']}")
            with mc2:
                st.metric("💰 Cost", f"${m_info['cost_per_m2_usd']}/m²")
                st.caption(f"📖 **Ref:** {m_info['reference']}")
                # Recommended departments
                rec_depts = ', '.join(m_info['best_for'][:3])
                if len(m_info['best_for']) > 3:
                    rec_depts += f' (+{len(m_info["best_for"])-3} more)'
                st.caption(f"🏢 **Departments:** {rec_depts}")
            st.markdown("---")

    # Combined materials cost summary
    mat_summary_cols = st.columns(3)
    mat_summary_cols[0].metric(
        "💰 Combined Material Cost",
        f"${_total_mat_cost}/m²",
        help="Sum of cost per m² across all 5 selected materials.")
    mat_summary_cols[1].metric(
        "🏗️ Total Project Material Cost (est.)",
        f"${_total_mat_cost * total_area * rooms_count:,.0f}",
        help=f"Approx. {total_area * rooms_count:.0f} m² × $/m² total.")
    mat_summary_cols[2].metric(
        "🌱 Sustainability Score",
        f"{sum(1 for c, k, m, i in _mats_used if 'Eco' in m or 'Recycled' in m or 'Bamboo' in m or 'Reclaimed' in m or 'Mineral' in m or 'Air-Purifying' in m or 'Cork' in m or 'Sheep' in m or 'Lime' in m)}/5 eco picks",
        help="Number of categories where you chose a renewable / "
             "recycled / bio-based material.")

    # ════════════════════════════════════════════════════════════════
    # 🔬 ENGINEERING ANALYSIS (HVAC + Fire Safety)
    # ════════════════════════════════════════════════════════════════
    st.markdown("---")
    st.subheader("🔬 Engineering Analysis")

    eng_tab1, eng_tab2 = st.tabs(["💨 HVAC Load", "🚨 Fire Safety & Egress"])

    with eng_tab1:
        # Calculate equipment heat load from active equipment
        _eq_heat = sum(e['power_w'] for e in EXTENDED_EQUIPMENT
                       if e['key'] in active_equipment)
        if add_monitor: _eq_heat += 200
        if add_anesthesia: _eq_heat += 500
        # Lighting load estimate
        _light_load = total_area * 15  # ~15 W/m² for LED hospital lighting
        # Occupancy estimate
        _n_occ = max(2, units_per_room + 2)  # patients + staff

        hvac = calculate_hvac_load(
            total_area, std['min_ceiling'], _n_occ,
            _eq_heat, _light_load, climate, std['ach'])

        st.markdown(f"**Cooling load analysis for {facility_type}** "
                    f"(per room, ASHRAE method)")
        h1, h2, h3 = st.columns(3)
        h1.metric("❄️ Cooling Load",
                  f"{hvac['total_kW']:.1f} kW",
                  help="Total sensible + latent cooling required")
        h2.metric("🔧 AC Capacity Needed",
                  f"{hvac['tons']:.1f} tons",
                  help="1 ton = 3.517 kW = 12,000 BTU/h. "
                       "Size your chiller/AC to this.")
        h3.metric("💨 Ventilation Airflow",
                  f"{hvac['cmh']:,.0f} m³/h",
                  help=f"At {hvac['ach']} ACH (ASHRAE 170). "
                       f"= {hvac['cfm']:,.0f} CFM")

        st.markdown("**Load breakdown:**")
        load_data = {
            'Component': ['🏠 Envelope + Solar', '👥 People',
                          '💡 Lighting', '🏥 Equipment'],
            'Heat Gain (W)': [hvac['Q_envelope'], hvac['Q_people'],
                              hvac['Q_lighting'], hvac['Q_equipment']],
        }
        st.bar_chart(load_data, x='Component', y='Heat Gain (W)')
        st.caption(f"📖 Method: ASHRAE Fundamentals 2021. Climate factor for "
                   f"**{climate}** = {({'Hot':150,'Moderate':100,'Cold':70}.get(climate,100))} W/m². "
                   f"ACH from ASHRAE 170 Table 7-1.")

    with eng_tab2:
        fire = calculate_fire_safety(total_area * rooms_count, facility_type)
        st.markdown(f"**NFPA 101 Life Safety Code analysis** "
                    f"(total area: {total_area*rooms_count:.0f} m²)")
        f1, f2, f3 = st.columns(3)
        f1.metric("👥 Occupant Load",
                  f"{fire['occupants']} people",
                  help=f"NFPA 101 factor: {fire['occ_factor']} m²/person")
        f2.metric("🚪 Required Exits",
                  f"{fire['req_exits']}",
                  help="Healthcare requires minimum 2 means of egress")
        f3.metric("📏 Egress Width",
                  f"{fire['per_exit_m']:.2f} m/exit",
                  help=f"Total {fire['total_egress_m']:.2f}m "
                       f"(5mm/person, sprinklered)")

        f4, f5, f6 = st.columns(3)
        f4.metric("🔥 Smoke Compartments",
                  f"{fire['n_smoke_compartments']}",
                  help="NFPA 101: max 2,323 m² per compartment in healthcare")
        f5.metric("🧯 Fire Extinguishers",
                  f"{fire['n_extinguishers']}",
                  help="1 per ~93 m² (NFPA 10)")
        f6.metric("🏃 Max Travel Distance",
                  f"{fire['max_travel_distance']} m",
                  help="Max distance to nearest exit (sprinklered healthcare)")

        # Compliance note
        if fire['req_exits'] >= 2:
            st.success(f"✅ Design provides adequate egress for "
                       f"{fire['occupants']} occupants with "
                       f"{fire['req_exits']} exits.")
        st.caption("📖 Method: NFPA 101 Life Safety Code + NFPA 99 Healthcare "
                   "Facilities Code. Assumes fully sprinklered building.")

    # ════════════════════════════════════════════════════════════════
    # 📑 PROJECT DOCUMENTS (Bill of Quantities + Functional Program)
    # ════════════════════════════════════════════════════════════════
    st.markdown("---")
    st.subheader("📑 Project Documents")

    doc_tab1, doc_tab2 = st.tabs(["📋 Bill of Quantities",
                                   "📄 Functional Program"])

    with doc_tab1:
        boq_config = {
            'facility_type': facility_type, 'std': std,
            'total_area': total_area, 'rooms_count': rooms_count,
            'units_per_room': units_per_room,
            'flooring': flooring, 'active_equipment': active_equipment,
            'add_monitor': add_monitor, 'add_anesthesia': add_anesthesia,
            'add_iv_stand': add_iv_stand, 'add_med_cart': add_med_cart,
            'laminar_ceiling': laminar_ceiling, 'partition_type': partition_type,
        }
        boq_rows = build_boq_rows('single', boq_config)
        boq_total = sum(r[f'Amount ({_curr})'] for r in boq_rows)
        # Add 10% contingency
        boq_contingency = boq_total * 0.10
        boq_grand = boq_total + boq_contingency

        st.markdown(f"**Itemized cost breakdown for {facility_type}**")
        try:
            import pandas as pd
            df_boq = pd.DataFrame(boq_rows)
            st.dataframe(df_boq, use_container_width=True, hide_index=True)
        except Exception:
            # Fallback: render as markdown table
            for r in boq_rows:
                st.text(f"{r['No.']}. {r['Description']} — "
                        f"{r['Qty']} {r['Unit']} × {_curr}{r[f'Rate ({_curr})']} "
                        f"= {_curr}{r[f'Amount ({_curr})']:,.0f}")

        bsum1, bsum2, bsum3 = st.columns(3)
        bsum1.metric("Subtotal", f"{_curr}{boq_total:,.0f}")
        bsum2.metric("Contingency (10%)", f"{_curr}{boq_contingency:,.0f}")
        bsum3.metric("GRAND TOTAL", f"{_curr}{boq_grand:,.0f}")

        # Export buttons
        exp1, exp2 = st.columns(2)
        with exp1:
            # CSV (always works)
            try:
                import pandas as pd
                csv_data = pd.DataFrame(boq_rows).to_csv(index=False)
            except Exception:
                # Manual CSV
                hdr = list(boq_rows[0].keys()) if boq_rows else []
                csv_data = ','.join(hdr) + '\n'
                for r in boq_rows:
                    csv_data += ','.join(str(r[h]) for h in hdr) + '\n'
            st.download_button("⬇️ Download CSV", csv_data,
                file_name=f"BoQ_{facility_type.replace(' ','_').replace('/','_')}.csv",
                mime="text/csv", use_container_width=True)
        with exp2:
            xlsx_bytes = boq_to_excel_bytes(boq_rows, boq_grand, _curr)
            if xlsx_bytes:
                st.download_button("⬇️ Download Excel", xlsx_bytes,
                    file_name=f"BoQ_{facility_type.replace(' ','_').replace('/','_')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True)
            else:
                st.caption("Excel export needs openpyxl — CSV works above.")

    with doc_tab2:
        fp_config = boq_config.copy()
        fp_md = build_functional_program_md('single', fp_config)
        st.markdown("**Functional Program Document** — room data sheets, "
                    "equipment schedule, and adjacency requirements.")
        with st.expander("👁️ Preview Document", expanded=False):
            st.markdown(fp_md)
        # Export buttons
        fpexp1, fpexp2 = st.columns(2)
        with fpexp1:
            st.download_button("⬇️ Download Markdown", fp_md,
                file_name=f"FunctionalProgram_{facility_type.replace(' ','_').replace('/','_')}.md",
                mime="text/markdown", use_container_width=True)
        with fpexp2:
            docx_bytes = md_to_docx_bytes(fp_md, "Functional Program")
            if docx_bytes:
                st.download_button("⬇️ Download Word", docx_bytes,
                    file_name=f"FunctionalProgram_{facility_type.replace(' ','_').replace('/','_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True)
            else:
                st.caption("Word export needs python-docx — Markdown works above.")

    # ════════════════════════════════════════════════════════════════
    # 🤖 AI SMART FEATURES (Anomaly Detection + Optimization + Chatbot)
    # ════════════════════════════════════════════════════════════════
    st.markdown("---")
    st.subheader("🤖 AI Smart Assistant")

    ai_tab1, ai_tab2, ai_tab3 = st.tabs([
        "🔍 Design Check", "⚙️ Optimize", "💬 Ask the Assistant"])

    # ─── Anomaly Detection ───
    with ai_tab1:
        anomaly_config = {
            'facility_type': facility_type, 'std': std,
            'total_area': total_area, 'units_per_room': units_per_room,
            'laminar_ceiling': laminar_ceiling, 'scrub_sink': scrub_sink,
            'add_monitor': add_monitor, 'lighting_type': lighting_type,
            'equip_efficiency': equip_efficiency, 'power_source': power_source,
        }
        anomalies = detect_anomalies('single', anomaly_config)
        if not anomalies:
            st.success("✅ **No issues detected!** Your design follows good "
                       "practice and meets the selected standards.")
        else:
            n_err = sum(1 for s, _ in anomalies if s == 'error')
            n_warn = sum(1 for s, _ in anomalies if s == 'warning')
            n_info = sum(1 for s, _ in anomalies if s == 'info')
            st.markdown(f"Found **{n_err} errors**, **{n_warn} warnings**, "
                        f"**{n_info} suggestions**:")
            for severity, msg in anomalies:
                if severity == 'error':
                    st.error(msg)
                elif severity == 'warning':
                    st.warning(msg)
                else:
                    st.info(msg)

    # ─── Optimization Mode ───
    with ai_tab2:
        st.markdown("**Let the system recommend the best configuration** for "
                    "your priority:")
        opt_objective = st.radio("Optimize for:",
            ["💰 Minimum Cost", "🌍 Minimum CO₂", "🌱 Maximum Green Score",
             "🏆 Maximum LEED"],
            horizontal=True, key="opt_obj_single")
        obj_map = {"💰 Minimum Cost": 'cost', "🌍 Minimum CO₂": 'co2',
                   "🌱 Maximum Green Score": 'green', "🏆 Maximum LEED": 'leed'}
        if st.button("⚙️ Generate Optimal Design", key="opt_btn_single"):
            opt = optimize_design(facility_type, std, total_area,
                                  rooms_count, obj_map[opt_objective])
            st.markdown(f"#### 🎯 Recommended configuration for "
                        f"**{opt_objective}**:")
            for param, value in opt['recommendations'].items():
                pretty = param.replace('_', ' ').title()
                st.markdown(f"- **{pretty}:** `{value}`")
            st.markdown("**Why these choices?**")
            for r in opt['rationale']:
                st.caption(f"• {r}")
            st.info("💡 Apply these settings in the sidebar, then regenerate "
                    "to see the optimized metrics.")

    # ─── AI Chatbot ───
    with ai_tab3:
        st.markdown("Ask me anything about hospital design standards! "
                    "*(ventilation, room sizes, lighting, solar, costs, "
                    "LEED, HVAC, fire safety...)*")
        # Init chat history
        if 'chat_history_single' not in st.session_state:
            st.session_state.chat_history_single = []
        # Quick suggestion buttons
        sugg_cols = st.columns(3)
        suggestions = ["Why does OR need 20 ACH?",
                       "Minimum area for this room?",
                       "How much can solar save?"]
        for i, sugg in enumerate(suggestions):
            if sugg_cols[i].button(sugg, key=f"sugg_s_{i}"):
                ans = chatbot_answer(sugg, anomaly_config)
                st.session_state.chat_history_single.append((sugg, ans))
        # Text input
        user_q = st.chat_input("Type your question...",
                                key="chat_input_single")
        if user_q:
            ans = chatbot_answer(user_q, anomaly_config)
            st.session_state.chat_history_single.append((user_q, ans))
        # Display history (most recent first)
        for q_text, a_text in reversed(st.session_state.chat_history_single[-6:]):
            with st.chat_message("user"):
                st.markdown(q_text)
            with st.chat_message("assistant"):
                st.markdown(a_text)

    st.markdown("---")

    # ============================================================
    # DIMENSIONS & UNIT POSITIONING
    # ============================================================
    def smart_dimensions(area, units):
        if units <= 1: ratio = 1.0
        elif units == 2: ratio = 1.4
        elif units <= 4: ratio = 1.5
        else: ratio = 1.7
        width = np.sqrt(area * ratio)
        height = area / width
        return width, height

    width, height = smart_dimensions(total_area, units_per_room)

    def calc_unit_positions(width, height, units):
        positions = []
        if units == 1:
            positions.append((width/2, height/2))
        elif units == 2:
            positions = [(width*0.3, height*0.55), (width*0.7, height*0.55)]
        elif units == 3:
            positions = [(width*0.22, height*0.55), (width*0.5, height*0.55),
                         (width*0.78, height*0.55)]
        elif units == 4:
            positions = [(width*0.25, height*0.7), (width*0.75, height*0.7),
                         (width*0.25, height*0.3), (width*0.75, height*0.3)]
        else:
            cols = int(np.ceil(units / 2))
            x_slots = np.linspace(width*0.15, width*0.85, cols)
            for x in x_slots:
                positions.append((x, height*0.7))
                if len(positions) < units:
                    positions.append((x, height*0.3))
        return positions

    unit_positions = calc_unit_positions(width, height, units_per_room)

    # ============================================================
    # 2D BLUEPRINT
    # ============================================================
    def draw_2d_blueprint():
        fig, ax = plt.subplots(figsize=(9, 7))
        ax.set_xlim(-1.5, width + 2.5)
        ax.set_ylim(-1.5, height + 2.5)
        ax.set_facecolor('#fafbf7')

        if show_grid:
            for x in np.arange(0, width+0.1, 1):
                ax.axvline(x, color='#d4dccc', linewidth=0.3, zorder=0)
            for y in np.arange(0, height+0.1, 1):
                ax.axhline(y, color='#d4dccc', linewidth=0.3, zorder=0)

        if show_zones:
            ax.add_patch(patches.Rectangle((width*0.1, height*0.2), width*0.8, height*0.7,
                facecolor='#c8e6c9', alpha=0.3, zorder=1))
            ax.add_patch(patches.Rectangle((0, 0), width, height*0.2,
                facecolor='#fff9c4', alpha=0.3, zorder=1))
            ax.text(width*0.5, height*0.55, "STERILE ZONE", fontsize=7, ha='center',
                    color='#2e7d32', alpha=0.5, weight='bold', zorder=2)
            ax.text(width*0.5, height*0.1, "TRANSITION ZONE", fontsize=6, ha='center',
                    color='#f57f17', alpha=0.5, weight='bold', zorder=2)

        # Walls
        wt = 0.15
        for x, y, w, h in [(-wt, -wt, width+2*wt, wt), (-wt, height, width+2*wt, wt),
                            (-wt, 0, wt, height), (width, 0, wt, height)]:
            ax.add_patch(patches.Rectangle((x, y), w, h, facecolor='#2c2c2c', zorder=3))

        # Door
        door_w = 1.2
        door_x = width/2 - door_w/2
        ax.add_patch(patches.Rectangle((door_x, -wt), door_w, wt,
            facecolor='#fafbf7', edgecolor='#8B4513', linewidth=2, zorder=4))
        ax.add_patch(patches.Arc((door_x, 0), 2.4, 2.4, theta1=0, theta2=90,
            color='#8B4513', linewidth=0.8, linestyle='--', zorder=4))
        ax.plot([door_x, door_x + door_w], [0, 0], color='#8B4513', linewidth=1.5, zorder=4)
        ax.text(width/2, -0.55, "Hermetic Sliding Door", color="#5d4037",
                fontsize=8, ha="center", weight="bold", zorder=4)

        # ═══════════════════════════════════════════════════════
        # 🚨 EMERGENCY EXITS (NFPA 101 Life Safety Code)
        # Count derived from occupant load — healthcare min 2 exits.
        # Placed on side walls (perpendicular to main entrance).
        # ═══════════════════════════════════════════════════════
        # exit_positions is also used by fire-protection placement
        # (pull stations, emergency lighting) below.
        exit_positions = []
        if show_emergency_exits:
            try:
                _fire = calculate_fire_safety(width * height, facility_type)
                _n_exits = _fire['req_exits']
            except Exception:
                _n_exits = 2  # safe default for healthcare
            # Main door already counts as 1 exit, so we add (n_exits - 1) more
            extra_exits = max(1, _n_exits - 1)
            em_door_w = 1.0
            em_color = '#2e7d32'      # safety green
            em_label_bg = '#1b5e20'

            # Decide placement: left wall first, then right wall, then back
            if extra_exits >= 1:
                # Left wall (vertical), at 2/3 height (away from front)
                exit_positions.append(('left', height * 0.65))
            if extra_exits >= 2:
                # Right wall, at 2/3 height
                exit_positions.append(('right', height * 0.65))
            if extra_exits >= 3:
                # Back wall, off-center
                exit_positions.append(('back', width * 0.25))

            for side, pos in exit_positions:
                if side == 'left':
                    # Door opening in the left wall
                    ax.add_patch(patches.Rectangle(
                        (-wt, pos - em_door_w/2), wt, em_door_w,
                        facecolor=em_color, edgecolor=em_label_bg,
                        linewidth=2, zorder=4))
                    # Arc showing swing direction (outward)
                    ax.add_patch(patches.Arc(
                        (0, pos - em_door_w/2), 2*em_door_w, 2*em_door_w,
                        theta1=90, theta2=180,
                        color=em_label_bg, linewidth=0.8,
                        linestyle='--', zorder=4))
                    # EXIT label outside
                    ax.text(-0.55, pos, "🚪 EXIT", color=em_label_bg,
                            fontsize=7, ha="center", va='center',
                            weight="bold", rotation=90, zorder=4,
                            bbox=dict(boxstyle="round,pad=0.2",
                                      facecolor='white',
                                      edgecolor=em_color, linewidth=1.2))
                elif side == 'right':
                    ax.add_patch(patches.Rectangle(
                        (width, pos - em_door_w/2), wt, em_door_w,
                        facecolor=em_color, edgecolor=em_label_bg,
                        linewidth=2, zorder=4))
                    ax.add_patch(patches.Arc(
                        (width, pos - em_door_w/2), 2*em_door_w, 2*em_door_w,
                        theta1=0, theta2=90,
                        color=em_label_bg, linewidth=0.8,
                        linestyle='--', zorder=4))
                    ax.text(width + 0.55, pos, "🚪 EXIT", color=em_label_bg,
                            fontsize=7, ha="center", va='center',
                            weight="bold", rotation=-90, zorder=4,
                            bbox=dict(boxstyle="round,pad=0.2",
                                      facecolor='white',
                                      edgecolor=em_color, linewidth=1.2))
                else:  # back
                    ax.add_patch(patches.Rectangle(
                        (pos - em_door_w/2, height), em_door_w, wt,
                        facecolor=em_color, edgecolor=em_label_bg,
                        linewidth=2, zorder=4))
                    ax.add_patch(patches.Arc(
                        (pos - em_door_w/2, height), 2*em_door_w, 2*em_door_w,
                        theta1=270, theta2=360,
                        color=em_label_bg, linewidth=0.8,
                        linestyle='--', zorder=4))
                    ax.text(pos, height + 0.55, "🚪 EXIT", color=em_label_bg,
                            fontsize=7, ha="center", va='center',
                            weight="bold", zorder=4,
                            bbox=dict(boxstyle="round,pad=0.2",
                                      facecolor='white',
                                      edgecolor=em_color, linewidth=1.2))

        # ═══════════════════════════════════════════════════════
        # 🔥 FIRE PROTECTION SYSTEMS (2D)
        # Smoke detectors + sprinklers (ceiling overlay) +
        # pull stations + emergency lighting (wall fixtures)
        # ═══════════════════════════════════════════════════════
        if (show_smoke_detectors or show_sprinklers or
                show_pull_stations or show_emergency_lighting):
            _fp_layout = get_fire_protection_layout(
                width, height, exit_positions)
            draw_fire_protection_2d(ax, _fp_layout, {
                'smoke': show_smoke_detectors,
                'sprinklers': show_sprinklers,
                'pull_stations': show_pull_stations,
                'emergency_lights': show_emergency_lighting,
            })

        # ═══════════════════════════════════════════════════════
        # 🚨 MEDICAL EMERGENCY DEVICES (2D)
        # AED + Eye Wash + Emergency Shower + Call Buttons
        # ═══════════════════════════════════════════════════════
        if (show_aed or show_eye_wash or show_emergency_shower or
                show_call_buttons):
            _med_layout = get_medical_emergency_layout(
                width, height,
                n_beds=units_per_room if show_call_buttons else 0,
                em_exit_positions=exit_positions)
            draw_medical_emergency_2d(ax, _med_layout, {
                'aed': show_aed,
                'eye_wash': show_eye_wash,
                'emergency_shower': show_emergency_shower,
                'call_buttons': show_call_buttons,
            })

        # ═══════════════════════════════════════════════════════
        # 🛤️ WAYFINDING & EGRESS (2D)
        # Egress Path arrows + Wayfinding Signs + Refuge Areas +
        # Emergency Stairs
        # ═══════════════════════════════════════════════════════
        if (show_egress_path or show_wayfinding or show_refuge_areas or
                show_emergency_stairs):
            _wf_layout = get_wayfinding_layout(
                width, height, exit_positions, include_main_door=True)
            draw_wayfinding_2d(ax, _wf_layout, {
                'egress_path': show_egress_path,
                'wayfinding': show_wayfinding,
                'refuge': show_refuge_areas,
                'stairs': show_emergency_stairs,
            })

        # ═══════════════════════════════════════════════════════
        # ☣️ HAZARD MARKERS (2D)
        # Radiation Shielding + Isolation Pod + Hazmat NFPA 704
        # ═══════════════════════════════════════════════════════
        if show_radiation or show_isolation or show_hazmat:
            _hz_layout = get_hazard_layout(width, height)
            draw_hazards_2d(ax, _hz_layout, {
                'radiation': show_radiation,
                'isolation': show_isolation,
                'hazmat': show_hazmat,
            }, width, height)

        unit_color = std["color"]

        # ═══════════════════════════════════════════════════════
        # CUSTOM LAYOUT DISPATCHER — use architectural layouts
        # for departments that need specialized arrangement
        # (cafeteria, wards, pharmacy, reception, etc.)
        # ═══════════════════════════════════════════════════════
        if facility_type in CUSTOM_LAYOUTS_2D:
            CUSTOM_LAYOUTS_2D[facility_type](
                ax, width, height, unit_color, units_per_room,
                {'glass': partition_type == "Antibacterial Glass Walls"})
            # Skip the generic unit loop — custom layout handled it
            _used_custom_2d = True
        else:
            _used_custom_2d = False

        # Draw units — DIFFERENT SHAPES per facility type!
        # (Only runs if no custom layout above)
        for idx, (bx, by) in enumerate(unit_positions if not _used_custom_2d else []):

            # ─── Clearance halo ───
            ax.add_patch(patches.Rectangle(
                (bx - 0.6 - std["min_clearance"]/2, by - 1.0 - std["min_clearance"]/2),
                1.2 + std["min_clearance"], 2.0 + std["min_clearance"],
                facecolor='none', edgecolor='#aed581',
                linewidth=0.6, linestyle=':', zorder=5))

            # ─── Unit shape varies by department ───
            if facility_type == "MRI Room":
                # Donut-shaped MRI scanner
                ax.add_patch(Circle((bx, by), 1.0, facecolor='#ede7f6',
                    edgecolor=unit_color, linewidth=2, zorder=6))
                ax.add_patch(Circle((bx, by), 0.4, facecolor='#2c2c2c', zorder=7))
                ax.add_patch(patches.Rectangle((bx-0.4, by-1.5), 0.8, 1.0,
                    facecolor='#eceff1', edgecolor=unit_color, linewidth=1, zorder=6))
                ax.text(bx, by + 1.3, f"MRI {idx+1}", color=unit_color,
                        fontsize=8, ha="center", weight="bold", zorder=8)

            elif facility_type == "CT Scan Room":
                # Ring scanner
                ax.add_patch(Circle((bx, by), 0.8, facecolor=unit_color,
                    edgecolor='black', linewidth=1.5, zorder=6))
                ax.add_patch(Circle((bx, by), 0.45, facecolor='#fafbf7', zorder=7))
                ax.add_patch(patches.Rectangle((bx-0.35, by-1.5), 0.7, 1.0,
                    facecolor='#eceff1', edgecolor=unit_color, linewidth=1, zorder=6))
                ax.text(bx, by + 1.1, f"CT {idx+1}", color=unit_color,
                        fontsize=8, ha="center", weight="bold", zorder=8)

            elif facility_type == "Dialysis Unit":
                # Reclining chair + dialysis machine
                ax.add_patch(FancyBboxPatch((bx-0.4, by-0.8), 0.8, 1.6,
                    boxstyle="round,pad=0.05", facecolor='#e1f5fe',
                    edgecolor=unit_color, linewidth=1.5, zorder=6))
                ax.add_patch(patches.Rectangle((bx+0.5, by-0.3), 0.35, 0.8,
                    facecolor=unit_color, edgecolor='black', linewidth=0.5, zorder=6))
                ax.text(bx+0.67, by+0.1, "💧", fontsize=8, ha='center', zorder=7)
                ax.text(bx, by + 1.0, f"Station {idx+1}", color=unit_color,
                        fontsize=7, ha="center", weight="bold", zorder=8)

            elif facility_type == "NICU (Neonatal ICU)":
                # Incubator (small transparent box)
                ax.add_patch(FancyBboxPatch((bx-0.4, by-0.5), 0.8, 1.0,
                    boxstyle="round,pad=0.05", facecolor='#fce4ec', alpha=0.6,
                    edgecolor=unit_color, linewidth=2, zorder=6))
                ax.text(bx, by, "👶", fontsize=12, ha='center', va='center', zorder=7)
                ax.text(bx, by + 0.7, f"Incub {idx+1}", color=unit_color,
                        fontsize=7, ha="center", weight="bold", zorder=8)

            elif facility_type == "Laboratory":
                # Workbench with equipment
                ax.add_patch(patches.Rectangle((bx-0.7, by-0.3), 1.4, 0.6,
                    facecolor='#b2dfdb', edgecolor=unit_color, linewidth=1.5, zorder=6))
                # Microscope + tubes
                ax.add_patch(Circle((bx-0.4, by+0.1), 0.1, facecolor='#37474f', zorder=7))
                ax.add_patch(patches.Rectangle((bx, by+0.05), 0.4, 0.15,
                    facecolor='#80cbc4', edgecolor='black', linewidth=0.3, zorder=7))
                ax.text(bx, by - 0.5, f"Bench {idx+1}", color=unit_color,
                        fontsize=7, ha="center", weight="bold", zorder=8)

            elif facility_type == "Pharmacy":
                # Dispensing counter with shelves behind
                ax.add_patch(patches.Rectangle((bx-0.8, by+0.3), 1.6, 0.3,
                    facecolor='#a5d6a7', edgecolor=unit_color, linewidth=1.5, zorder=6))
                for sh in range(3):
                    ax.add_patch(patches.Rectangle((bx-0.7, by-0.7+sh*0.25), 1.4, 0.15,
                        facecolor=unit_color, alpha=0.4, edgecolor='black',
                        linewidth=0.3, zorder=6))
                ax.text(bx, by + 0.8, f"💊 Pharm {idx+1}", color=unit_color,
                        fontsize=7, ha="center", weight="bold", zorder=8)

            elif facility_type == "Blood Bank":
                # Refrigerator unit
                ax.add_patch(FancyBboxPatch((bx-0.4, by-0.8), 0.8, 1.6,
                    boxstyle="round,pad=0.03", facecolor='#ffcdd2',
                    edgecolor=unit_color, linewidth=2, zorder=6))
                ax.text(bx, by, "🩸", fontsize=12, ha='center', va='center', zorder=7)
                ax.text(bx, by - 1.0, f"Fridge {idx+1}", color=unit_color,
                        fontsize=7, ha="center", weight="bold", zorder=8)

            elif facility_type == "Dental Clinic":
                # Dental chair (reclining)
                ax.add_patch(FancyBboxPatch((bx-0.35, by-0.9), 0.7, 1.8,
                    boxstyle="round,pad=0.1", facecolor='#b2dfdb',
                    edgecolor=unit_color, linewidth=1.5, zorder=6))
                # Light overhead
                ax.add_patch(Circle((bx+0.5, by-0.5), 0.12, facecolor='#fff59d',
                    edgecolor='black', linewidth=0.5, zorder=7))
                ax.text(bx, by + 1.1, f"🦷 Chair {idx+1}", color=unit_color,
                        fontsize=7, ha="center", weight="bold", zorder=8)

            elif facility_type == "Oncology":
                # Comfortable chemo recliner
                ax.add_patch(FancyBboxPatch((bx-0.5, by-0.7), 1.0, 1.4,
                    boxstyle="round,pad=0.15", facecolor='#e1bee7',
                    edgecolor=unit_color, linewidth=1.5, zorder=6))
                # IV drip
                ax.add_patch(Circle((bx-0.6, by-0.4), 0.07, facecolor='#9e9e9e', zorder=7))
                ax.text(bx, by + 0.9, f"🎗️ Chair {idx+1}", color=unit_color,
                        fontsize=7, ha="center", weight="bold", zorder=8)

            elif facility_type == "Maternity & Delivery":
                # Wider birthing bed + baby warmer
                ax.add_patch(FancyBboxPatch((bx-0.6, by-1.0), 1.2, 2.0,
                    boxstyle="round,pad=0.05", facecolor='#fce4ec',
                    edgecolor=unit_color, linewidth=1.5, zorder=6))
                # Baby warmer (small box next to bed)
                ax.add_patch(patches.Rectangle((bx+0.7, by-0.3), 0.4, 0.6,
                    facecolor='#fff59d', edgecolor=unit_color, linewidth=1, zorder=7))
                ax.text(bx+0.9, by, "👶", fontsize=7, ha='center', va='center', zorder=8)
                ax.text(bx, by + 1.2, f"Delivery {idx+1}", color=unit_color,
                        fontsize=7, ha="center", weight="bold", zorder=8)

            elif facility_type == "Physical Therapy":
                # Exercise mat + parallel bars
                ax.add_patch(patches.Rectangle((bx-0.8, by-1.0), 1.6, 2.0,
                    facecolor='#b2dfdb', edgecolor=unit_color, linewidth=1.5, zorder=6))
                # Bars
                ax.plot([bx-0.6, bx-0.6], [by-0.8, by+0.8],
                        color=unit_color, linewidth=2, zorder=7)
                ax.plot([bx+0.6, bx+0.6], [by-0.8, by+0.8],
                        color=unit_color, linewidth=2, zorder=7)
                ax.text(bx, by + 1.2, f"🦾 Station {idx+1}", color=unit_color,
                        fontsize=7, ha="center", weight="bold", zorder=8)

            elif facility_type == "Cafeteria & Kitchen":
                # Round table with chairs
                ax.add_patch(Circle((bx, by), 0.5, facecolor='#fff8e1',
                    edgecolor=unit_color, linewidth=1.5, zorder=6))
                for ang in range(0, 360, 90):
                    rad = np.radians(ang)
                    cx, cy = bx + 0.8*np.cos(rad), by + 0.8*np.sin(rad)
                    ax.add_patch(Circle((cx, cy), 0.15, facecolor=unit_color, zorder=6))
                ax.text(bx, by + 1.0, f"🍽️ T{idx+1}", color=unit_color,
                        fontsize=6, ha="center", weight="bold", zorder=8)

            elif facility_type == "Administrative Offices":
                # Desk + chair
                ax.add_patch(patches.Rectangle((bx-0.6, by-0.3), 1.2, 0.6,
                    facecolor='#bcaaa4', edgecolor=unit_color, linewidth=1.2, zorder=6))
                ax.add_patch(Circle((bx, by-0.7), 0.2, facecolor=unit_color, zorder=6))
                ax.text(bx, by + 0.5, f"💼 Desk {idx+1}", color=unit_color,
                        fontsize=6, ha="center", weight="bold", zorder=8)

            elif facility_type == "Reception & Waiting":
                # Row of chairs
                for cw in range(3):
                    ax.add_patch(patches.Rectangle((bx-0.6+cw*0.4, by-0.2), 0.3, 0.4,
                        facecolor=unit_color, alpha=0.6, edgecolor='black',
                        linewidth=0.3, zorder=6))
                ax.text(bx, by + 0.4, f"🪑 Row {idx+1}", color=unit_color,
                        fontsize=6, ha="center", weight="bold", zorder=8)

            elif facility_type == "Sterilization (CSSD)":
                # Autoclave (cylindrical)
                ax.add_patch(FancyBboxPatch((bx-0.5, by-0.7), 1.0, 1.4,
                    boxstyle="round,pad=0.1", facecolor='#ffccbc',
                    edgecolor=unit_color, linewidth=2, zorder=6))
                ax.add_patch(Circle((bx, by), 0.3, facecolor='#37474f', zorder=7))
                ax.text(bx, by, "♨️", fontsize=10, ha='center', va='center', zorder=8)
                ax.text(bx, by + 0.9, f"Auto {idx+1}", color=unit_color,
                        fontsize=7, ha="center", weight="bold", zorder=8)

            elif facility_type == "Laundry":
                # Washing machine
                ax.add_patch(patches.Rectangle((bx-0.45, by-0.5), 0.9, 1.0,
                    facecolor='#b3e5fc', edgecolor=unit_color, linewidth=1.5, zorder=6))
                ax.add_patch(Circle((bx, by), 0.3, facecolor='white',
                    edgecolor=unit_color, linewidth=1, zorder=7))
                ax.text(bx, by + 0.7, f"🧺 W{idx+1}", color=unit_color,
                        fontsize=7, ha="center", weight="bold", zorder=8)

            else:
                # ─── Default: Hospital bed (for wards, ICU, OR, ER, Pediatric, etc.) ───
                bed_w, bed_l = 0.9, 2.0
                ax.add_patch(FancyBboxPatch((bx - bed_w/2, by - bed_l/2 + 0.2),
                    bed_w, bed_l - 0.2, boxstyle="round,pad=0.02",
                    facecolor='#eceff1', edgecolor=unit_color, linewidth=1.5, zorder=6))
                ax.add_patch(patches.Rectangle((bx - bed_w/2, by - bed_l/2),
                    bed_w, 0.25, facecolor=unit_color, zorder=6))
                ax.add_patch(FancyBboxPatch((bx - bed_w/2 + 0.1, by - bed_l/2 + 0.3),
                    bed_w - 0.2, 0.4, boxstyle="round,pad=0.02",
                    facecolor='white', edgecolor='#bdbdbd', linewidth=0.5, zorder=7))
                for wx, wy in [(-bed_w/2+0.1, -bed_l/2+0.3), (bed_w/2-0.1, -bed_l/2+0.3),
                                (-bed_w/2+0.1, bed_l/2-0.1), (bed_w/2-0.1, bed_l/2-0.1)]:
                    ax.add_patch(Circle((bx+wx, by+wy), 0.08,
                        facecolor='#212121', zorder=7))
                ax.text(bx, by + bed_l/2 + 0.2, f"{std['unit_name']} {idx+1}",
                        color=unit_color, fontsize=7, ha="center", weight="bold", zorder=8)

            # HEPA overlay
            if laminar_ceiling:
                ax.add_patch(patches.Rectangle((bx - 0.9, by - 1.2), 1.8, 2.4,
                    linewidth=1.2, linestyle='--', edgecolor='#2196f3',
                    facecolor='#bbdefb', alpha=0.2, zorder=4))
                ax.text(bx, by - 1.35, "HEPA H14", color="#1565c0", fontsize=5,
                        ha="center", style='italic', zorder=4)

            # Vital monitor
            if add_monitor and facility_type not in ["Pharmacy", "Cafeteria & Kitchen",
                "Laundry", "Administrative Offices", "Reception & Waiting",
                "Sterilization (CSSD)", "Laboratory"]:
                ax.add_patch(patches.Rectangle((bx + 0.55, by - 1.0),
                    0.35, 0.35, facecolor='#1a237e',
                    edgecolor='black', linewidth=0.5, zorder=7))
                ax.text(bx + 0.72, by - 0.82, "M", color='lime',
                        fontsize=6, ha='center', va='center',
                        weight='bold', zorder=8)

            # IV stand
            if add_iv_stand and facility_type not in ["Pharmacy", "Cafeteria & Kitchen",
                "Laundry", "Administrative Offices", "Reception & Waiting",
                "Sterilization (CSSD)", "Laboratory", "Blood Bank"]:
                ax.add_patch(Circle((bx - 0.65, by - 0.9), 0.08,
                    facecolor='#9e9e9e', edgecolor='black',
                    linewidth=0.4, zorder=7))
                ax.add_patch(Circle((bx - 0.65, by - 0.9), 0.04,
                    facecolor='#e3f2fd', zorder=8))

            # Glass partitions
            if partition_type == "Antibacterial Glass Walls" and \
               idx % 2 == 0 and idx+1 < len(unit_positions):
                next_bx = unit_positions[idx+1][0]
                mid_x = (bx + next_bx) / 2
                ax.plot([mid_x, mid_x], [height*0.15, height*0.9],
                        color='#4fc3f7', linewidth=3, alpha=0.7, zorder=5)
                ax.text(mid_x, height*0.92, "Glass", color="#0277bd",
                        fontsize=5, ha="center", style='italic', zorder=5)

        # Anesthesia
        if add_anesthesia and facility_type == "Operating Room (OR)":
            ax.add_patch(FancyBboxPatch((width*0.05, height*0.4), 0.6, 0.5,
                boxstyle="round,pad=0.02", facecolor='#37474f',
                edgecolor='black', linewidth=0.8, zorder=6))
            ax.text(width*0.05 + 0.3, height*0.65, "Anesthesia\nMachine",
                    color='#37474f', fontsize=5, ha='center',
                    weight='bold', zorder=6)

        # Med cart
        if add_med_cart and facility_type not in ["Cafeteria & Kitchen",
            "Laundry", "Administrative Offices", "Reception & Waiting"]:
            ax.add_patch(FancyBboxPatch((width*0.85, height*0.4), 0.45, 0.7,
                boxstyle="round,pad=0.02", facecolor='#fff59d',
                edgecolor='#f57f17', linewidth=1, zorder=6))
            ax.text(width*0.85 + 0.22, height*0.78, "Med\nCart",
                    color='#e65100', fontsize=5, ha='center',
                    weight='bold', zorder=6)

        # Scrub sink
        if scrub_sink != "None":
            sink_w = 1.0 if scrub_sink == "Double Stainless Sink" else 0.5
            ax.add_patch(FancyBboxPatch((width*0.05, 0.2), sink_w, 0.5,
                boxstyle="round,pad=0.02", facecolor='#90a4ae',
                edgecolor='#37474f', linewidth=1, zorder=6))
            n_faucets = 2 if scrub_sink == "Double Stainless Sink" else 1
            for i in range(n_faucets):
                fx = width*0.05 + sink_w*(i+1)/(n_faucets+1)
                ax.add_patch(Circle((fx, 0.45), 0.06, facecolor='#cfd8dc',
                    edgecolor='black', linewidth=0.5, zorder=7))
            ax.text(width*0.05 + sink_w/2, 0.85, "Scrub Sink", color="#263238",
                    fontsize=6, ha="center", weight="bold", zorder=6)

        # Smart Touchscreen Console — mounted on the right wall, next to the
        # main entrance door (not on any door). Previous placement on the
        # left wall conflicted with the Emergency Exit there.
        if control_panel == "Smart Touchscreen Console":
            # Wall-mounted bezel hugging the right wall, just inside the room
            ax.add_patch(FancyBboxPatch((width - 0.18, 0.9), 0.18, 0.7,
                boxstyle="round,pad=0.02", facecolor='#1a237e',
                edgecolor='black', linewidth=0.8, zorder=6))
            # Touchscreen face (lighter, inset)
            ax.add_patch(patches.Rectangle((width - 0.15, 0.98), 0.12, 0.55,
                facecolor='#4fc3f7', edgecolor='#0277bd',
                linewidth=0.4, zorder=7))
            # Power-indicator LED
            ax.add_patch(patches.Circle((width - 0.09, 1.55), 0.02,
                facecolor='#76ff03', edgecolor='#33691e',
                linewidth=0.4, zorder=8))
            # Label (placed inside the room so it doesn't overlap walls)
            ax.text(width - 0.55, 1.25, "Smart\nConsole",
                    color="#1a237e", fontsize=5.5, va="center",
                    ha="center", weight="bold", zorder=6,
                    bbox=dict(boxstyle="round,pad=0.15",
                              facecolor='white', edgecolor='#1a237e',
                              linewidth=0.6, alpha=0.9))

        # ─── EXTENDED EQUIPMENT (specialized devices placed smartly) ───
        if active_equipment:
            eq_positions = get_equipment_positions(width, height, active_equipment)
            for eq_key, ex, ey in eq_positions:
                draw_equipment_2d(ax, ex, ey, eq_key, std.get("color", "#1976d2"))

        # Window
        if lighting_type in ["Natural", "Mixed"]:
            ax.add_patch(patches.Rectangle((width - 0.05, height/2 - 1.2), 0.15, 2.4,
                facecolor='#81d4fa', edgecolor='#0277bd', linewidth=1.5, zorder=6))
            ax.plot([width + 0.075, width + 0.075], [height/2 - 1.2, height/2 + 1.2],
                    color='#0277bd', linewidth=1, zorder=7)
            ax.plot([width - 0.05, width + 0.1], [height/2, height/2],
                    color='#0277bd', linewidth=1, zorder=7)
            ax.text(width + 0.25, height/2, "Window", color='#01579b',
                    fontsize=6, va='center', rotation=90, weight='bold', zorder=6)

        # Traffic flow
        if show_traffic_flow:
            staff_arrow = FancyArrowPatch((width/2 - 0.4, 0.2), (width*0.15, height*0.4),
                arrowstyle='->', color='#2e7d32', linewidth=2,
                alpha=0.6, zorder=9, connectionstyle="arc3,rad=0.2")
            ax.add_patch(staff_arrow)
            ax.text(width*0.15, height*0.42, "Staff", color='#2e7d32',
                    fontsize=6, weight='bold', zorder=9)
            patient_arrow = FancyArrowPatch((width/2 + 0.4, 0.2), (width*0.5, height*0.4),
                arrowstyle='->', color='#ef6c00', linewidth=2,
                alpha=0.6, zorder=9, connectionstyle="arc3,rad=-0.2")
            ax.add_patch(patient_arrow)
            ax.text(width*0.55, height*0.42, "Patient", color='#ef6c00',
                    fontsize=6, weight='bold', zorder=9)

        # Dimensions
        if show_dimensions:
            ax.annotate('', xy=(width, height + 0.6), xytext=(0, height + 0.6),
                arrowprops=dict(arrowstyle='<->', color='#424242', lw=1))
            ax.text(width/2, height + 0.8, f"{width:.2f} m", ha='center',
                    fontsize=8, color='#424242', weight='bold')
            ax.annotate('', xy=(width + 0.6, height), xytext=(width + 0.6, 0),
                arrowprops=dict(arrowstyle='<->', color='#424242', lw=1))
            ax.text(width + 0.8, height/2, f"{height:.2f} m", va='center',
                    rotation=90, fontsize=8, color='#424242', weight='bold')

        # Scale bar
        sb_x, sb_y = -1.2, -1.1
        ax.add_patch(patches.Rectangle((sb_x, sb_y), 1, 0.08,
            facecolor='black', zorder=10))
        ax.add_patch(patches.Rectangle((sb_x + 1, sb_y), 1, 0.08,
            facecolor='white', edgecolor='black', zorder=10))
        ax.text(sb_x, sb_y - 0.25, "0", fontsize=6, ha='center', zorder=10)
        ax.text(sb_x + 1, sb_y - 0.25, "1m", fontsize=6, ha='center', zorder=10)
        ax.text(sb_x + 2, sb_y - 0.25, "2m", fontsize=6, ha='center', zorder=10)

        # North arrow
        na_x, na_y = width + 1.8, height + 0.5
        ax.add_patch(Circle((na_x, na_y), 0.4, facecolor='white',
            edgecolor='black', linewidth=1, zorder=10))
        ax.annotate('', xy=(na_x, na_y + 0.3), xytext=(na_x, na_y - 0.2),
            arrowprops=dict(arrowstyle='->', color='red', lw=1.5))
        ax.text(na_x, na_y + 0.55, "N", fontsize=8, ha='center',
                weight='bold', color='red', zorder=10)

        # Title
        ax.text(0, height + 1.5, f"VIRIDIS BLUEPRINT — {facility_type}",
            fontsize=10, weight='bold', color='#1b5e20')
        ax.text(0, height + 1.2, f"Area: {total_area} m² | "
            f"{std['unit_name']}s: {units_per_room} | "
            f"ACH: {std['ach']} | Lux: {std['lux']}",
            fontsize=7, color='#424242', style='italic')

        # Legend
        legend_elements = [
            patches.Patch(facecolor='#c8e6c9', alpha=0.4, label='Sterile Zone'),
            patches.Patch(facecolor='#fff9c4', alpha=0.4, label='Transition Zone'),
            Line2D([0], [0], color='#2e7d32', lw=2, label='Staff Flow'),
            Line2D([0], [0], color='#ef6c00', lw=2, label='Patient Flow'),
            Line2D([0], [0], color='#aed581', lw=1, linestyle=':', label='Clearance'),
        ]
        ax.legend(handles=legend_elements, loc='upper right',
            bbox_to_anchor=(1.0, -0.02), ncol=3, fontsize=6, framealpha=0.9)

        ax.set_aspect('equal')
        ax.axis('off')
        plt.tight_layout()
        return fig

    # ============================================================
    # 3D RENDER (reused from previous version, simplified)
    # ============================================================
    def box_mesh(x0, y0, z0, dx, dy, dz, color, opacity=1.0, name=''):
        x = [x0, x0+dx, x0+dx, x0, x0, x0+dx, x0+dx, x0]
        y = [y0, y0, y0+dy, y0+dy, y0, y0, y0+dy, y0+dy]
        z = [z0, z0, z0, z0, z0+dz, z0+dz, z0+dz, z0+dz]
        i = [7,0,0,0,4,4,6,6,4,0,3,2]
        j = [3,4,1,2,5,6,5,2,0,1,6,3]
        k = [0,7,2,3,6,7,1,1,5,5,7,6]
        return go.Mesh3d(x=x, y=y, z=z, i=i, j=j, k=k, color=color,
            opacity=opacity, name=name, hovertext=name, hoverinfo='text', flatshading=True)

    def draw_3d_render():
        wall_height = max(std["min_ceiling"], 3.0)
        fig = go.Figure()
        is_night = (day_night == "🌙 Night")
        floor_color = '#37474f' if is_night else '#e8f5e9'
        wall_color = '#263238' if is_night else '#f5f5f5'

        # ── Floor ──
        fig.add_trace(go.Mesh3d(x=[0, width, width, 0], y=[0, 0, height, height],
            z=[0, 0, 0, 0], i=[0, 0], j=[1, 2], k=[2, 3],
            color=floor_color, opacity=0.95, name='Floor',
            hovertext=f'Flooring: {flooring}', hoverinfo='text'))

        # ── Walls (back + sides) ──
        fig.add_trace(box_mesh(0, height, 0, width, 0.1, wall_height,
            wall_color, 0.85, 'Back Wall'))
        fig.add_trace(box_mesh(-0.1, 0, 0, 0.1, height, wall_height,
            wall_color, 0.4, 'Left Wall'))
        fig.add_trace(box_mesh(width, 0, 0, 0.1, height, wall_height,
            wall_color, 0.4, 'Right Wall'))

        # ═══════════════════════════════════════════════════════
        # 🚨 EMERGENCY EXITS in 3D (NFPA 101 — green-marked doors)
        # ═══════════════════════════════════════════════════════
        # exits3d is also passed to fire-protection drawing for
        # pull-station / emergency-light placement near each door.
        exits3d = []
        if show_emergency_exits:
            try:
                _fire3d = calculate_fire_safety(width * height, facility_type)
                _n_exits3d = _fire3d['req_exits']
            except Exception:
                _n_exits3d = 2
            _extra3d = max(1, _n_exits3d - 1)
            door_w3d = 1.0
            door_h3d = 2.1   # standard egress door height
            em_green = '#2e7d32'
            em_dark = '#1b5e20'

            if _extra3d >= 1:
                exits3d.append(('left',  height * 0.65))
            if _extra3d >= 2:
                exits3d.append(('right', height * 0.65))
            if _extra3d >= 3:
                exits3d.append(('back',  width * 0.25))

            for side, pos in exits3d:
                if side == 'left':
                    # Door panel set into the left wall
                    fig.add_trace(box_mesh_global(
                        -0.105, pos - door_w3d/2, 0,
                        0.11, door_w3d, door_h3d,
                        em_green, 0.95, 'Emergency Exit Door'))
                    # Glowing EXIT sign above the door
                    fig.add_trace(box_mesh_global(
                        -0.11, pos - 0.3, door_h3d + 0.05,
                        0.02, 0.6, 0.3,
                        em_dark, 1.0, 'EXIT Sign'))
                    fig.add_trace(box_mesh_global(
                        -0.12, pos - 0.25, door_h3d + 0.10,
                        0.01, 0.5, 0.20,
                        '#ffffff', 1.0, 'EXIT Light'))
                elif side == 'right':
                    fig.add_trace(box_mesh_global(
                        width - 0.005, pos - door_w3d/2, 0,
                        0.11, door_w3d, door_h3d,
                        em_green, 0.95, 'Emergency Exit Door'))
                    fig.add_trace(box_mesh_global(
                        width + 0.09, pos - 0.3, door_h3d + 0.05,
                        0.02, 0.6, 0.3,
                        em_dark, 1.0, 'EXIT Sign'))
                    fig.add_trace(box_mesh_global(
                        width + 0.11, pos - 0.25, door_h3d + 0.10,
                        0.01, 0.5, 0.20,
                        '#ffffff', 1.0, 'EXIT Light'))
                else:  # back
                    fig.add_trace(box_mesh_global(
                        pos - door_w3d/2, height - 0.005, 0,
                        door_w3d, 0.11, door_h3d,
                        em_green, 0.95, 'Emergency Exit Door'))
                    fig.add_trace(box_mesh_global(
                        pos - 0.3, height + 0.09, door_h3d + 0.05,
                        0.6, 0.02, 0.3,
                        em_dark, 1.0, 'EXIT Sign'))
                    fig.add_trace(box_mesh_global(
                        pos - 0.25, height + 0.11, door_h3d + 0.10,
                        0.5, 0.01, 0.20,
                        '#ffffff', 1.0, 'EXIT Light'))

        # ═══════════════════════════════════════════════════════
        # 🔥 FIRE PROTECTION SYSTEMS (3D)
        # ═══════════════════════════════════════════════════════
        if (show_smoke_detectors or show_sprinklers or
                show_pull_stations or show_emergency_lighting):
            _fp_layout_3d = get_fire_protection_layout(
                width, height, exits3d)
            draw_fire_protection_3d(fig, _fp_layout_3d, {
                'smoke': show_smoke_detectors,
                'sprinklers': show_sprinklers,
                'pull_stations': show_pull_stations,
                'emergency_lights': show_emergency_lighting,
            }, wall_height=wall_height)

        # ═══════════════════════════════════════════════════════
        # 🚨 MEDICAL EMERGENCY DEVICES (3D)
        # ═══════════════════════════════════════════════════════
        if (show_aed or show_eye_wash or show_emergency_shower or
                show_call_buttons):
            _med_layout_3d = get_medical_emergency_layout(
                width, height,
                n_beds=units_per_room if show_call_buttons else 0,
                em_exit_positions=exits3d)
            draw_medical_emergency_3d(fig, _med_layout_3d, {
                'aed': show_aed,
                'eye_wash': show_eye_wash,
                'emergency_shower': show_emergency_shower,
                'call_buttons': show_call_buttons,
            }, wall_height=wall_height)

        # ═══════════════════════════════════════════════════════
        # 🛤️ WAYFINDING & EGRESS (3D)
        # ═══════════════════════════════════════════════════════
        if (show_egress_path or show_wayfinding or show_refuge_areas or
                show_emergency_stairs):
            _wf_layout_3d = get_wayfinding_layout(
                width, height, exits3d, include_main_door=True)
            draw_wayfinding_3d(fig, _wf_layout_3d, {
                'egress_path': show_egress_path,
                'wayfinding': show_wayfinding,
                'refuge': show_refuge_areas,
                'stairs': show_emergency_stairs,
            }, wall_height=wall_height)

        # ═══════════════════════════════════════════════════════
        # ☣️ HAZARD MARKERS (3D)
        # ═══════════════════════════════════════════════════════
        if show_radiation or show_isolation or show_hazmat:
            _hz_layout_3d = get_hazard_layout(width, height)
            draw_hazards_3d(fig, _hz_layout_3d, {
                'radiation': show_radiation,
                'isolation': show_isolation,
                'hazmat': show_hazmat,
            }, width, height, wall_height=wall_height)

        unit_color = std["color"]

        # ═══════════════════════════════════════════════════════
        # CUSTOM LAYOUT DISPATCHER — architecturally realistic
        # 3D layouts (wards with beds on walls, cafeteria with
        # service bar, pharmacy with shelves, etc.)
        # ═══════════════════════════════════════════════════════
        if facility_type in CUSTOM_LAYOUTS_3D:
            CUSTOM_LAYOUTS_3D[facility_type](
                fig, width, height, unit_color, wall_height, units_per_room,
                {'glass': partition_type == "Antibacterial Glass Walls"})
            _used_custom_3d = True
        else:
            _used_custom_3d = False

        # ──────────────────────────────────────────────────────
        # REALISTIC UNITS — different 3D model per department
        # (Only runs if no custom layout above)
        # ──────────────────────────────────────────────────────
        for idx, (bx, by) in enumerate(unit_positions if not _used_custom_3d else []):

            if facility_type == "MRI Room":
                for t in mri_scanner_3d(bx, by, 0, unit_color,
                                         f'MRI {idx+1}'):
                    fig.add_trace(t)

            elif facility_type == "CT Scan Room":
                for t in ct_scanner_3d(bx, by, 0, unit_color,
                                        f'CT {idx+1}'):
                    fig.add_trace(t)

            elif facility_type == "Dialysis Unit":
                for t in recliner_3d(bx, by, 0, '#0277bd',
                                      f'Dialysis Chair {idx+1}'):
                    fig.add_trace(t)
                # Dialysis machine beside the chair
                fig.add_trace(box_mesh_global(bx+0.5, by-0.3, 0,
                                               0.4, 0.5, 1.4, unit_color, 1.0,
                                               f'Dialysis Machine {idx+1}'))
                # Display screen on machine
                fig.add_trace(box_mesh_global(bx+0.52, by-0.32, 1.0,
                                               0.36, 0.02, 0.3,
                                               '#4fc3f7', 1.0, 'Display'))

            elif facility_type == "Oncology":
                for t in recliner_3d(bx, by, 0, '#7b1fa2',
                                      f'Chemo Chair {idx+1}'):
                    fig.add_trace(t)
                # IV stand for chemo
                fig.add_trace(cylinder_mesh(bx-0.6, by-0.4, 0, 0.03, 1.7,
                                             'z', 8, '#9e9e9e', 1.0, 'IV Pole'))
                fig.add_trace(box_mesh_global(bx-0.68, by-0.46, 1.45,
                                               0.16, 0.12, 0.2,
                                               '#e1f5fe', 0.7, 'IV Bag'))

            elif facility_type == "NICU (Neonatal ICU)":
                for t in incubator_3d(bx, by, 0, unit_color,
                                       f'Incubator {idx+1}'):
                    fig.add_trace(t)

            elif facility_type == "Laboratory":
                for t in lab_bench_3d(bx, by, 0, unit_color,
                                       f'Bench {idx+1}'):
                    fig.add_trace(t)

            elif facility_type == "Pharmacy":
                for t in pharmacy_counter_3d(bx, by, 0, unit_color,
                                              f'Pharm {idx+1}'):
                    fig.add_trace(t)

            elif facility_type == "Blood Bank":
                for t in blood_fridge_3d(bx, by, 0, unit_color,
                                          f'Blood Fridge {idx+1}'):
                    fig.add_trace(t)

            elif facility_type == "Sterilization (CSSD)":
                for t in autoclave_3d(bx, by, 0, unit_color,
                                       f'Autoclave {idx+1}'):
                    fig.add_trace(t)

            elif facility_type == "Laundry":
                for t in washer_3d(bx, by, 0, unit_color,
                                    f'Washer {idx+1}'):
                    fig.add_trace(t)

            elif facility_type == "Cafeteria & Kitchen":
                for t in dining_set_3d(bx, by, 0, unit_color,
                                        f'Table {idx+1}'):
                    fig.add_trace(t)

            elif facility_type == "Administrative Offices":
                for t in desk_3d(bx, by, 0, unit_color,
                                  f'Desk {idx+1}'):
                    fig.add_trace(t)

            elif facility_type == "Physical Therapy":
                for t in parallel_bars_3d(bx, by, 0, unit_color,
                                           f'PT Station {idx+1}'):
                    fig.add_trace(t)

            elif facility_type == "Maternity & Delivery":
                for t in maternity_bed_3d(bx, by, 0, unit_color,
                                           f'Delivery Bed {idx+1}'):
                    fig.add_trace(t)

            elif facility_type == "Reception & Waiting":
                # Row of 3 chairs side-by-side
                for cw in range(3):
                    for t in chair_3d(bx-0.6+cw*0.5, by, 0,
                                       0.4, 0.45, 0.45, 0.5, unit_color,
                                       f'Seat {idx*3+cw+1}'):
                        fig.add_trace(t)

            elif facility_type == "Pediatric Ward":
                # Smaller, brightly colored bed
                for t in hospital_bed_3d(bx, by, 0, '#ffb300',
                                          size=0.75, name=f'Pediatric Bed {idx+1}'):
                    fig.add_trace(t)

            elif facility_type == "Dental Clinic":
                # Reclining dental chair + overhead lamp + side cabinet
                for t in recliner_3d(bx, by, 0, '#00897b',
                                      f'Dental Chair {idx+1}'):
                    fig.add_trace(t)
                # Overhead lamp on articulating arm
                fig.add_trace(cylinder_mesh(bx+0.5, by, 1.5, 0.03, 0.6, 'z',
                                             8, '#bdbdbd', 1.0, 'Lamp Arm'))
                fig.add_trace(cylinder_mesh(bx+0.5, by, 2.1, 0.18, 0.08,
                                             'z', 16, '#fff59d', 1.0,
                                             'Dental Lamp'))

            else:
                # ── Default: standard hospital bed (OR/ICU/ER/Ward) ──
                for t in hospital_bed_3d(bx, by, 0, unit_color,
                                          size=1.0, name=f'{std["unit_name"]} {idx+1}'):
                    fig.add_trace(t)
                # Vital signs monitor on wall
                if add_monitor:
                    fig.add_trace(box_mesh_global(bx-0.2, by-1.05, 1.5,
                                                   0.4, 0.05, 0.3,
                                                   '#0d47a1', 1.0,
                                                   f'Monitor {idx+1}'))
                    fig.add_trace(box_mesh_global(bx-0.18, by-1.06, 1.55,
                                                   0.36, 0.01, 0.22,
                                                   '#00e676' if not is_night else '#69f0ae',
                                                   1.0, 'Screen'))
                # IV stand
                if add_iv_stand:
                    fig.add_trace(cylinder_mesh(bx-0.7, by-0.7, 0, 0.025, 1.6,
                                                 'z', 8, '#9e9e9e', 1.0, 'IV Pole'))
                    fig.add_trace(box_mesh_global(bx-0.78, by-0.76, 1.4,
                                                   0.16, 0.12, 0.2,
                                                   '#e1f5fe', 0.7, 'IV Bag'))

            # ── HEPA grid (mounted from ceiling) ──
            if laminar_ceiling:
                fig.add_trace(box_mesh_global(bx-0.9, by-1.1, wall_height-0.25,
                                               1.8, 2.2, 0.08, '#42a5f5', 0.55,
                                               f'HEPA H14 Grid {idx+1}'))

        # ── Window ──
        if lighting_type in ["Natural", "Mixed"]:
            fig.add_trace(box_mesh_global(width - 0.02, height/2 - 1.2, 1.0,
                                           0.04, 2.4, 1.5, '#81d4fa', 0.5, 'Window'))

        # ── Ceiling LED panels ──
        light_color = '#fff59d' if not is_night else '#fff176'
        for lx in np.linspace(width*0.25, width*0.75, 3):
            for ly in np.linspace(height*0.3, height*0.8, 2):
                fig.add_trace(box_mesh_global(lx-0.2, ly-0.1, wall_height-0.1,
                                               0.4, 0.2, 0.05, light_color,
                                               1.0 if not is_night else 0.5,
                                               'LED Panel'))

        # ── Anesthesia (OR) ──
        if add_anesthesia and facility_type == "Operating Room (OR)":
            fig.add_trace(box_mesh_global(width*0.05, height*0.45, 0,
                                           0.5, 0.4, 1.4, '#37474f', 1.0,
                                           'Anesthesia Machine'))
            fig.add_trace(box_mesh_global(width*0.05+0.05, height*0.45+0.05, 1.4,
                                           0.4, 0.05, 0.3, '#1a237e', 1.0,
                                           'Anesthesia Display'))

        # ── Scrub Sink ──
        if scrub_sink != "None":
            sink_w = 1.0 if scrub_sink == "Double Stainless Sink" else 0.5
            fig.add_trace(box_mesh_global(width*0.05, 0.2, 0,
                                           sink_w, 0.5, 0.9, '#90a4ae', 1.0,
                                           scrub_sink))
            n_faucets = 2 if scrub_sink == "Double Stainless Sink" else 1
            for i in range(n_faucets):
                fx = width*0.05 + sink_w*(i+1)/(n_faucets+1) - 0.03
                fig.add_trace(box_mesh_global(fx, 0.25, 0.9,
                                               0.06, 0.06, 0.3,
                                               '#cfd8dc', 1.0, 'Faucet'))

        # ── Glass Partitions in 3D (for non-custom-layout depts that
        # still use the center-of-room unit_positions arrangement) ──
        if (partition_type == "Antibacterial Glass Walls"
                and not _used_custom_3d and len(unit_positions) > 1):
            for idx in range(0, len(unit_positions) - 1, 2):
                bx1, by1 = unit_positions[idx]
                bx2, by2 = unit_positions[idx + 1]
                # Only if beds are roughly in same row
                if abs(by1 - by2) < 1.0:
                    mid_x = (bx1 + bx2) / 2
                    # Glass wall — full height, slim, transparent
                    fig.add_trace(box_mesh_global(
                        mid_x - 0.025, height*0.15, 0,
                        0.05, height*0.7, wall_height*0.85,
                        '#81d4fa', 0.3,
                        f'Glass Partition {idx//2 + 1}'))
                    # Metal frame on top
                    fig.add_trace(box_mesh_global(
                        mid_x - 0.075, height*0.15, wall_height*0.85 - 0.05,
                        0.15, height*0.7, 0.05,
                        '#cfd8dc', 1.0, 'Glass Frame'))

        # ── Smart console: mounted on the RIGHT wall next to main door ──
        # Previous placement on left wall (y=0.6h) overlapped with the
        # Emergency Exit door. Now placed on right wall just inside the
        # main entrance, mounted at 1.2m height (ADA-compliant reach).
        if control_panel == "Smart Touchscreen Console":
            # Console housing — wall-mounted, faces into the room
            fig.add_trace(box_mesh_global(width - 0.06, 0.9, 1.2,
                                           0.06, 0.6, 0.8, '#1a237e', 1.0,
                                           'Smart Control Console'))
            # Touchscreen face (lighter, slightly proud of the bezel)
            fig.add_trace(box_mesh_global(width - 0.07, 0.98, 1.30,
                                           0.005, 0.44, 0.6, '#4fc3f7', 1.0,
                                           'Touchscreen'))
            # Small status-LED strip at the top
            fig.add_trace(box_mesh_global(width - 0.065, 1.05, 1.92,
                                           0.005, 0.30, 0.03, '#76ff03', 1.0,
                                           'Status LED'))

        # ─── EXTENDED EQUIPMENT in 3D (placed smartly around the room) ───
        if active_equipment:
            eq_positions_3d = get_equipment_positions(width, height, active_equipment)
            for eq_key, ex, ey in eq_positions_3d:
                draw_equipment_3d(fig, ex, ey, eq_key, wall_height)

        # ── Human figure for scale (realistic: legs + torso + arms + head + scrubs) ──
        # Smart position: avoids overlap with beds/equipment based on facility type
        hx, hy = get_smart_human_position(facility_type, width, height)
        # ─ Legs (2 separate cylinders) ─
        leg_color = '#1976d2'  # scrub pants (medical blue)
        fig.add_trace(cylinder_mesh(hx - 0.08, hy, 0, 0.07, 0.85, 'z', 8,
                                     leg_color, 1.0, 'Leg L'))
        fig.add_trace(cylinder_mesh(hx + 0.08, hy, 0, 0.07, 0.85, 'z', 8,
                                     leg_color, 1.0, 'Leg R'))
        # ─ Shoes ─
        fig.add_trace(box_mesh_global(hx - 0.15, hy - 0.1, 0,
                                       0.14, 0.22, 0.06, '#212121', 1.0, 'Shoe L'))
        fig.add_trace(box_mesh_global(hx + 0.01, hy - 0.1, 0,
                                       0.14, 0.22, 0.06, '#212121', 1.0, 'Shoe R'))
        # ─ Torso (scrub top - lab coat white) ─
        torso_color = '#ffffff'
        fig.add_trace(cylinder_mesh(hx, hy, 0.85, 0.22, 0.45, 'z', 12,
                                     torso_color, 1.0, 'Lab Coat'))
        # ─ Scrub V-neck (color accent on chest) ─
        fig.add_trace(box_mesh_global(hx - 0.06, hy - 0.18, 1.0,
                                       0.12, 0.04, 0.2, '#0277bd', 1.0, 'V-Neck'))
        # ─ Arms (2 angled cylinders) ─
        arm_color = torso_color
        fig.add_trace(cylinder_mesh(hx - 0.25, hy, 0.95, 0.06, 0.4, 'z', 8,
                                     arm_color, 1.0, 'Arm L'))
        fig.add_trace(cylinder_mesh(hx + 0.25, hy, 0.95, 0.06, 0.4, 'z', 8,
                                     arm_color, 1.0, 'Arm R'))
        # ─ Hands (small spheres approximated as short cylinders) ─
        fig.add_trace(cylinder_mesh(hx - 0.25, hy, 0.55, 0.06, 0.08, 'z', 8,
                                     '#ffccbc', 1.0, 'Hand L'))
        fig.add_trace(cylinder_mesh(hx + 0.25, hy, 0.55, 0.06, 0.08, 'z', 8,
                                     '#ffccbc', 1.0, 'Hand R'))
        # ─ Neck ─
        fig.add_trace(cylinder_mesh(hx, hy, 1.3, 0.06, 0.08, 'z', 8,
                                     '#ffccbc', 1.0, 'Neck'))
        # ─ Head (slightly egg-shaped using cylinder) ─
        fig.add_trace(cylinder_mesh(hx, hy, 1.38, 0.13, 0.22, 'z', 14,
                                     '#ffccbc', 1.0, 'Staff (1.7m scale)'))
        # ─ Hair (darker cap on top) ─
        fig.add_trace(cylinder_mesh(hx, hy, 1.55, 0.135, 0.06, 'z', 14,
                                     '#3e2723', 1.0, 'Hair'))
        # ─ Stethoscope around neck (small torus) ─
        fig.add_trace(torus_mesh(hx, hy, 1.22, R=0.13, r=0.015,
                                  axis='z', n_major=16, n_minor=6,
                                  color='#212121', opacity=1.0,
                                  name='Stethoscope'))

        # ── Layout ──
        fig.update_layout(
            scene=dict(
                xaxis=dict(title='Width (m)', range=[-0.5, width+0.5],
                    backgroundcolor='#1a1a2e' if is_night else '#f5fdf7'),
                yaxis=dict(title='Length (m)', range=[-0.5, height+1],
                    backgroundcolor='#1a1a2e' if is_night else '#f5fdf7'),
                zaxis=dict(title='Height (m)', range=[0, wall_height+0.5],
                    backgroundcolor='#0d0d1e' if is_night else '#fafffe'),
                aspectmode='data',
                camera=dict(eye=dict(x=1.6, y=-1.8, z=1.2),
                    center=dict(x=0, y=0, z=-0.2)),
                bgcolor='#0a0a18' if is_night else '#ffffff'),
            margin=dict(l=0, r=0, b=0, t=30),
            height=580, showlegend=False,
            title=dict(text=f"3D Render — {facility_type} ({day_night})",
                font=dict(size=12, color='#1b5e20')))
        return fig

    # ============================================================
    # RENDER UI
    # ============================================================
    left, right = st.columns([1, 1.3])

    with left:
        st.subheader("📋 Viridis Advisory Report")
        rep = f"### 🏥 {std['icon']} {facility_type}\n"
        rep += f"- **Category:** {std['category']}\n"
        rep += f"- **Department Scale:** {rooms_count} rooms × {units_per_room} {std['unit_name']}s each.\n"
        rep += f"- **Per-Unit Area:** {total_area/units_per_room:.1f} m² "
        rep += f"(Standard min: {std['min_area_per_unit']} m²)\n"
        rep += f"- **Required ACH:** {std['ach']} | **Pressure:** {std['pressure']} | **Lux:** {std['lux']}\n"
        rep += f"- **Water Usage:** ~{std['water_usage']} L/day per unit\n"

        rep += f"\n### 🌬️ Department-Specific Notes\n"
        if facility_type == "Operating Room (OR)":
            rep += "- Mandatory **Positive Pressure** with **HEPA H14** filters.\n"
        elif facility_type == "MRI Room":
            rep += "- **Faraday Cage** lining required to block external EM interference.\n"
            rep += "- **Helium cooling** for superconducting magnet — sustainability concern.\n"
        elif facility_type == "Radiology Department":
            rep += "- **Lead lining** 2–3 mm on all walls.\n"
        elif facility_type == "Dialysis Unit":
            rep += "- **High water demand** (~500 L/day per station) → prime greywater recycling candidate.\n"
        elif facility_type == "Laboratory":
            rep += "- **Negative pressure** to contain biohazards; **fume hoods** required.\n"
        elif facility_type == "NICU (Neonatal ICU)":
            rep += "- **Dimmable, sound-controlled environment**; positive pressure for sterility.\n"
        elif facility_type == "Oncology":
            rep += "- **Chemo safety hood** with negative pressure required for compounding zone.\n"
        elif facility_type == "Pharmacy":
            rep += "- **Temperature/humidity logging** mandatory for drug storage.\n"
        elif facility_type == "Sterilization (CSSD)":
            rep += "- Major **steam & heat producer** — heat recovery loops can offset 25% energy.\n"
        elif facility_type == "Laundry":
            rep += "- **Highest water consumer** in hospital → ozone wash systems cut water by 60%.\n"
        elif facility_type == "Cafeteria & Kitchen":
            rep += "- **Composting + heat-recovery ventilation** turns this from cost center to eco win.\n"

        if laminar_ceiling:
            rep += "- ✅ **Laminar Flow** ceiling deployed above units.\n"
        if scrub_sink != "None":
            rep += f"- ✅ **{scrub_sink}** with hands-free sensor taps.\n"
        if partition_type == "Antibacterial Glass Walls":
            rep += "- ✅ **Antibacterial glass partitions** between units.\n"

        st.info(rep)

        # Energy chart
        fig_e, ax_e = plt.subplots(figsize=(5, 2.5))
        ax_e.bar(['Traditional', 'Viridis'], [normal_env, green_env],
            color=['#ef5350', '#66bb6a'], width=0.5,
            edgecolor='black', linewidth=0.8)
        ax_e.set_ylabel("Electricity (kWh/mo)", fontsize=9)
        ax_e.set_title("Energy Comparison", fontsize=10, weight='bold')
        for i, v in enumerate([normal_env, green_env]):
            ax_e.text(i, v + 50, f"{v:.0f}", ha='center',
                fontsize=8, weight='bold')
        ax_e.grid(axis='y', alpha=0.3)
        st.pyplot(fig_e)

    with right:
        if view_mode == "📐 2D Plan":
            st.subheader("📐 2D Floor Plan Blueprint")
            f2d = draw_2d_blueprint()
            st.pyplot(f2d)
            buf = BytesIO()
            f2d.savefig(buf, format='png', dpi=200, bbox_inches='tight')
            st.download_button("⬇️ Download Blueprint (PNG)", buf.getvalue(),
                file_name=f"viridis_{facility_type.replace(' ','_')}.png",
                mime="image/png")
        elif view_mode == "🌐 3D Render":
            st.subheader("🌐 3D Interactive Render")
            st.plotly_chart(draw_3d_render(), use_container_width=True)
        else:
            st.subheader("📐 2D Plan")
            st.pyplot(draw_2d_blueprint())
            st.subheader("🌐 3D Render")
            st.plotly_chart(draw_3d_render(), use_container_width=True)

    # ════════════════════════════════════════════════════════════════
    # 🎬 FIRE EVACUATION SIMULATION (sidebar-triggered modal section)
    # ════════════════════════════════════════════════════════════════
    if st.session_state.get('show_evac_sim', False):
        st.markdown("---")
        st.subheader("🎬 Fire Evacuation Simulation")
        st.caption("**NFPA 101 / SFPE methodology** — occupants move toward "
                   "nearest exits at 1.2 m/s (typical adult walking speed). "
                   "Fire spreads outward over time. Target: full evacuation "
                   "in under 180 seconds.")

        sim_col_l, sim_col_r = st.columns([1, 2])
        with sim_col_l:
            st.markdown("##### 🔥 Fire Origin")
            # Preset location selector
            preset = st.selectbox("Preset Location:", [
                "🎲 Random",
                "🔥 Kitchen / Cafeteria (electrical)",
                "🔥 Operating Room (gas leak)",
                "🔥 Storage Room",
                "✋ Custom Coordinates",
            ], key="evac_preset")

            # Compute room dimensions for fire-origin defaults
            _sim_width = total_area / np.sqrt(total_area / 1.5)
            _sim_height = total_area / _sim_width

            if preset == "🎲 Random":
                import random as _rand
                fire_x = _rand.uniform(_sim_width * 0.2, _sim_width * 0.8)
                fire_y = _rand.uniform(_sim_height * 0.2, _sim_height * 0.8)
            elif preset == "🔥 Kitchen / Cafeteria (electrical)":
                fire_x = _sim_width * 0.85  # back-right (kitchen prep area)
                fire_y = _sim_height * 0.85
            elif preset == "🔥 Operating Room (gas leak)":
                fire_x = _sim_width * 0.50  # center (surgical table area)
                fire_y = _sim_height * 0.50
            elif preset == "🔥 Storage Room":
                fire_x = _sim_width * 0.15  # back-left (storage corner)
                fire_y = _sim_height * 0.85
            else:  # Custom
                fire_x = st.number_input("Fire X (m):",
                    0.5, _sim_width - 0.5, _sim_width / 2,
                    step=0.5, key="evac_fx")
                fire_y = st.number_input("Fire Y (m):",
                    0.5, _sim_height - 0.5, _sim_height / 2,
                    step=0.5, key="evac_fy")

            # Number of occupants (NFPA-based default)
            try:
                _fire_safety = calculate_fire_safety(
                    total_area, facility_type)
                default_occupants = _fire_safety.get('occupants', 20)
            except Exception:
                default_occupants = 20

            n_occ = st.slider("👥 Number of Occupants:", 5, 60,
                min(default_occupants, 30), key="evac_occ",
                help="NFPA 101: defaults to the calculated occupant "
                     "load for this facility")

            run_btn = st.button("▶️  Run Simulation",
                key="run_sim_btn", use_container_width=True,
                type="primary")
            close_btn = st.button("✕  Close",
                key="close_sim_btn", use_container_width=True)
            if close_btn:
                st.session_state.show_evac_sim = False
                st.rerun()

        with sim_col_r:
            if run_btn:
                # Build exit positions for the simulation
                sim_exits = []
                # Main door at bottom-center
                sim_exits.append((_sim_width / 2, 0.3))
                # Emergency exits (if shown)
                if show_emergency_exits:
                    try:
                        _fs = calculate_fire_safety(
                            _sim_width * _sim_height, facility_type)
                        _ne = _fs['req_exits']
                    except Exception:
                        _ne = 2
                    extra = max(1, _ne - 1)
                    if extra >= 1:
                        sim_exits.append((0.3, _sim_height * 0.65))
                    if extra >= 2:
                        sim_exits.append((_sim_width - 0.3,
                                           _sim_height * 0.65))
                    if extra >= 3:
                        sim_exits.append((_sim_width * 0.25,
                                           _sim_height - 0.3))

                with st.spinner("🎬 Generating evacuation animation..."):
                    gif_bytes, metrics = simulate_fire_evacuation(
                        _sim_width, _sim_height,
                        (fire_x, fire_y), sim_exits,
                        n_occupants=n_occ, n_frames=40)

                st.image(gif_bytes, use_container_width=True)

                # Metrics dashboard
                st.markdown("#### 📊 Simulation Results")
                m1, m2, m3, m4 = st.columns(4)
                m1.metric("⏱️ Total Time",
                    f"{metrics['evacuation_time_sec']:.0f} s",
                    help="Time until all occupants escaped or became "
                         "casualties")
                m2.metric("✅ Escaped",
                    f"{metrics['n_escaped']}/{metrics['n_total']}")
                m3.metric("⚠️ Casualties",
                    f"{metrics['n_casualty']}")
                m4.metric("🏥 NFPA Compliance",
                    "✅ PASS" if metrics['nfpa_compliant']
                    else "❌ FAIL",
                    help="NFPA 101: target <180s + zero casualties")

                if metrics['nfpa_compliant']:
                    st.success(
                        f"🎉 **Evacuation Successful!** All "
                        f"{metrics['n_escaped']} occupants escaped in "
                        f"**{metrics['evacuation_time_sec']:.0f}s**, "
                        f"well within the NFPA 101 target of 180s.")
                else:
                    if metrics['n_casualty'] > 0:
                        st.error(
                            f"❌ **{metrics['n_casualty']} casualties** — "
                            f"fire spread engulfed occupants before they "
                            f"reached an exit. Consider: adding more "
                            f"emergency exits, sprinkler coverage, or "
                            f"earlier alarm activation.")
                    else:
                        st.warning(
                            f"⚠️ Evacuation took "
                            f"{metrics['evacuation_time_sec']:.0f}s, "
                            f"exceeding the NFPA 101 target. Consider "
                            f"widening egress paths or adding exits.")
            else:
                st.info("⬅️ Choose a fire origin and click **Run Simulation** "
                        "to generate the evacuation animation.")

    # Save config
    st.markdown("---")
    col_save, col_pdf, col_info = st.columns([1, 1, 2])
    with col_save:
        design_data = {
            "facility": facility_type, "area": total_area, "rooms": rooms_count,
            "units": units_per_room, "climate": climate, "lighting": lighting_type,
            "partition": partition_type, "sink": scrub_sink, "panel": control_panel,
            "laminar": laminar_ceiling, "flooring": flooring, "waste": waste_mgmt,
            "equip": equip_efficiency, "power": power_source,
            "green_score": green_score, "energy_saved": saved_env,
            "money_saved": saved_money
        }
        st.download_button("💾 JSON",
            json.dumps(design_data, indent=2),
            file_name="viridis_design.json", mime="application/json",
            use_container_width=True)
    with col_pdf:
        # Build PDF on demand
        pdf_config = {
            'facility_type': facility_type, 'total_area': total_area,
            'rooms_count': rooms_count, 'units_per_room': units_per_room,
            'climate': climate, 'lighting_type': lighting_type,
            'partition_type': partition_type, 'scrub_sink': scrub_sink,
            'control_panel': control_panel, 'laminar_ceiling': laminar_ceiling,
            'flooring': flooring, 'waste_mgmt': waste_mgmt,
            'equip_efficiency': equip_efficiency, 'power_source': power_source,
        }
        pdf_metrics = {
            'green_score': green_score,
            'energy_saved': saved_env,
            'money_saved': saved_money,
            'savings_pct': saved_pct,
            'normal_env': normal_env,
            'green_env': green_env,
        }
        try:
            # Build a fresh blueprint figure for the PDF (avoid mutating displayed one)
            pdf_blueprint = draw_2d_blueprint()
            pdf_bytes = generate_pdf_report('single', pdf_config, pdf_metrics,
                                             blueprint_fig=pdf_blueprint)
            plt.close(pdf_blueprint)
            st.download_button(t('pdf_download'), pdf_bytes,
                file_name=f"viridis_report_{facility_type.replace(' ', '_').replace('/', '_')}.pdf",
                mime="application/pdf", use_container_width=True)
        except Exception as e:
            st.caption(f"PDF unavailable: {e}")
    with col_info:
        st.caption("💡 Tip: Switch to **Full Hospital Mode** in the sidebar "
                   "to design the entire facility at once.")

    # ───────── Optional: Lux Heatmap (in sidebar toggle) ─────────
    show_lux = st.sidebar.checkbox(t('show_lux_heatmap'),
                                    value=False, key='sb_lux')
    if show_lux:
        st.markdown("---")
        st.subheader(t('lux_heatmap_title'))
        try:
            lux_fig, lux_data = generate_lux_heatmap(
                width, height, facility_type, lighting_type, std['lux'])
            st.pyplot(lux_fig)
            avg_lux = float(np.mean(lux_data))
            min_lux = float(np.min(lux_data))
            req = std['lux']
            cM1, cM2, cM3 = st.columns(3)
            cM1.metric("Avg Lux", f"{avg_lux:.0f}",
                       delta=f"{(avg_lux-req)/req*100:+.0f}% vs req")
            cM2.metric("Min Lux", f"{min_lux:.0f}",
                       delta="✅ Adequate" if min_lux >= req*0.5 else "⚠️ Low")
            cM3.metric("Required Lux", f"{req}")
            st.caption("🔴 Red dashed line = required lux level boundary. "
                       "Yellow stars = LED ceiling panels. "
                       "Blue thick line = window (if natural light enabled).")
        except Exception as e:
            st.error(f"Could not generate heatmap: {e}")

# ============================================================
# ════════════════════════════════════════════════════════════
# MODE 2: FULL HOSPITAL SITE PLAN
# ════════════════════════════════════════════════════════════
# ============================================================
else:
    st.sidebar.header("🏥 Hospital Configuration")

    # v23: pre-fill scale + departments from the wizard's choice
    _wd = st.session_state.wizard_data
    _wd_scale = _wd.get('hospital_scale', "Medium Hospital (150 beds)")
    _scale_opts = ["Small Clinic (50 beds)", "Medium Hospital (150 beds)",
                   "Large Hospital (300 beds)", "Mega Medical City (600+ beds)"]
    if _wd_scale not in _scale_opts:
        _wd_scale = "Medium Hospital (150 beds)"

    hospital_scale = st.sidebar.select_slider(
        "Hospital Scale:",
        options=_scale_opts,
        value=_wd_scale
    )

    # Hospital scale → multiplier for areas
    scale_map = {
        "Small Clinic (50 beds)": 0.5,
        "Medium Hospital (150 beds)": 1.0,
        "Large Hospital (300 beds)": 1.8,
        "Mega Medical City (600+ beds)": 3.0
    }
    scale = scale_map[hospital_scale]

    st.sidebar.markdown("---")
    st.sidebar.subheader("📋 Departments to Include")
    st.sidebar.caption("Toggle which departments are part of this hospital design")

    # v23: when wizard provided a dept list, use that as defaults; else
    # fall back to category-based defaults.
    _wd_dept_names = set(_wd.get('selected_dept_names', []))
    _use_wizard_depts = bool(_wd_dept_names)

    selected_depts = {}
    for cat in ["Critical Care", "Diagnostic", "Treatment",
                "Specialized", "General Care", "Service"]:
        with st.sidebar.expander(f"📁 {cat}", expanded=(cat == "Critical Care")):
            for dept_name, dept_info in MEDICAL_STANDARDS.items():
                if dept_info["category"] == cat:
                    if _use_wizard_depts:
                        default = dept_name in _wd_dept_names
                    else:
                        default = cat in ["Critical Care", "Service",
                                           "General Care"]
                    if st.checkbox(f"{dept_info['icon']} {dept_name}",
                                    value=default, key=f"dept_{dept_name}"):
                        selected_depts[dept_name] = dept_info

    st.sidebar.markdown("---")
    st.sidebar.subheader("🏗️ Hospital Architecture")
    # v24: All-in-One Building option removed. Hospitals are always rendered
    # as separate wings (multi-building campus). Variable kept for backward
    # compatibility with downstream code that checks `building_mode`.
    building_mode = "🏘️ Separate Wings (multiple buildings)"
    h_garage = st.sidebar.checkbox("🚗 Parking Garage (covered)", value=True)
    h_roads = st.sidebar.checkbox("🛣️ Road Network (driveways + paths)", value=True)
    h_outdoor_seating = st.sidebar.checkbox("🪑 Outdoor Seating (benches in garden + perimeter)", value=True)

    st.sidebar.markdown("---")
    st.sidebar.subheader("🌍 Site Location")
    h_climate_mode = st.sidebar.radio("Location Input:",
        ["🌍 Select City", "✋ Skip"],
        help="Selecting a city enables solar potential, rainwater harvest, "
             "and climate-aware recommendations for the whole hospital.")
    if h_climate_mode == "🌍 Select City":
        _h_city_list = sorted(CLIMATE_DATABASE.keys())
        h_sel_city = st.sidebar.selectbox("Hospital City:", _h_city_list,
            index=_h_city_list.index('Cairo'), key='h_city')
        _hcd = detect_climate(h_sel_city)
        st.sidebar.caption(
            f"{_hcd['country']} **{h_sel_city}**: {_hcd['zone']} · "
            f"☀️ {_hcd['sun_hours']}h sun · 🌧️ {_hcd['rainfall']}mm/yr")
    else:
        h_sel_city = None

    st.sidebar.markdown("---")
    st.sidebar.subheader("🌱 Hospital-Wide Sustainability")
    h_solar = st.sidebar.checkbox("🌞 Rooftop Solar PV Farm", value=True)
    h_garden = st.sidebar.checkbox("🌳 Healing Garden (central courtyard)", value=True)
    h_green_roof = st.sidebar.checkbox("🌿 Green Roof", value=True)
    h_rainwater = st.sidebar.checkbox("💧 Rainwater Harvesting", value=True)
    h_ev_chargers = st.sidebar.checkbox("🔌 EV Charging Stations", value=True)
    h_greywater = st.sidebar.checkbox("♻️ Greywater Recycling", value=True)
    h_smart_grid = st.sidebar.checkbox("⚡ Smart Grid + Battery Storage", value=False)

    st.sidebar.markdown("---")
    st.sidebar.subheader("🎨 Visualization")
    show_corridors = st.sidebar.checkbox("Show Corridors", value=True)
    show_patient_flow = st.sidebar.checkbox("Show Patient Flow", value=True)
    show_dept_labels = st.sidebar.checkbox("Show Department Labels", value=True)

    if not selected_depts:
        st.warning("⚠️ Please select at least one department from the sidebar.")
        st.stop()

    # ============================================================
    # 💰 HOSPITAL BUDGET MANAGER
    # ============================================================
    st.sidebar.markdown("---")
    st.sidebar.subheader("💰 Project Budget")
    _curr_h_bud = get_active_reference()['currency_symbol']
    _ref_key_h = st.session_state.get('reference', 'fgi')

    # Estimate default hospital budget
    _est_total_area_h = sum(
        max(1, int(d['max_units'] * 0.6)) * d['min_area_per_unit'] * 1.4 * scale
        for d in selected_depts.values()
    )
    _avg_cost_per_m2 = (2500 if _ref_key_h == 'fgi' else 25000)
    _default_h_budget = int(_est_total_area_h * _avg_cost_per_m2 * 1.4)

    h_enable_budget = st.sidebar.checkbox(
        "Enable Hospital Budget Tracking", value=False,
        help="Track total project cost (construction + equipment + sustainability) "
             "against your budget.")

    if h_enable_budget:
        h_user_budget = st.sidebar.number_input(
            f"Total Project Budget ({_curr_h_bud}):",
            min_value=0,
            value=_default_h_budget,
            step=100000 if _ref_key_h == 'fgi' else 1000000,
            format="%d",
            help="Total project budget for the entire hospital."
        )
        st.sidebar.caption(
            f"💡 Estimated cost: ~{_curr_h_bud}{_default_h_budget:,}")
    else:
        h_user_budget = 0

    # ============================================================
    # GENERATE BUTTON — captures hospital config into session_state
    # ============================================================
    st.sidebar.markdown("---")
    if not st.session_state.get('hospital_design_generated', False):
        st.sidebar.markdown(
            "<div style='padding:10px; background:#fff3e0; "
            "border-left:4px solid #f57c00; border-radius:6px; "
            "margin-bottom:8px; font-size:0.85em; color:#e65100'>"
            "👇 <b>Click below</b> to generate your hospital — your "
            "configuration is ready!</div>",
            unsafe_allow_html=True)
    generate_btn = st.sidebar.button(
        "🚀 Generate Hospital Design", type="primary",
        use_container_width=True, key="gen_hospital"
    )

    if generate_btn:
        # Mark hospital design as generated — unblocks the Results page.
        st.session_state.hospital_design_generated = True
        # Snapshot every input. The 3D hospital render is expensive — we
        # only want it to re-run when the user explicitly asks for it.
        st.session_state.hospital_config = {
            'hospital_scale': hospital_scale, 'scale': scale,
            'selected_depts': dict(selected_depts),  # snapshot of dict
            'h_solar': h_solar, 'h_garden': h_garden,
            'h_green_roof': h_green_roof, 'h_rainwater': h_rainwater,
            'h_ev_chargers': h_ev_chargers, 'h_greywater': h_greywater,
            'h_smart_grid': h_smart_grid,
            'building_mode': building_mode,
            'h_garage': h_garage, 'h_roads': h_roads,
            'h_outdoor_seating': h_outdoor_seating,
            'h_enable_budget': h_enable_budget,
            'h_user_budget': h_user_budget,
            'h_sel_city': h_sel_city,
            'show_corridors': show_corridors,
            'show_patient_flow': show_patient_flow,
            'show_dept_labels': show_dept_labels,
        }

    # ════════════════════════════════════════════════════════════════
    # v23: Empty-state guard for Hospital Mode (matches Single Mode)
    # Educational dashboard until first Generate click.
    # ════════════════════════════════════════════════════════════════
    if not st.session_state.get('hospital_design_generated', False):
        st.markdown(
            "<div style='text-align:center; padding:30px 10px 20px 10px'>"
            "<h1 style='color:#1b5e20; margin-bottom:8px'>"
            "🏥 Viridis Hospital Studio</h1>"
            "<p style='color:#388e3c; font-size:1.1em; margin-top:0'>"
            "Your hospital configuration is ready in the sidebar — "
            "follow the 2 steps below to generate the full site plan."
            "</p></div>",
            unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("""
            <div style="border:2px solid #2e7d32; border-radius:14px;
                        padding:24px 26px; background:#e8f5e9;
                        min-height:440px; box-sizing:border-box;
                        box-shadow:0 2px 8px rgba(46,125,50,0.08)">
              <div style="font-size:2.6em; text-align:center;
                          line-height:1; margin-bottom:8px">👈</div>
              <h3 style="color:#1b5e20; text-align:center;
                         margin:6px 0 14px 0; font-size:1.35em">
                Step 1 — Review &amp; Adjust
              </h3>
              <p style="color:#2e7d32; font-size:0.95em;
                        margin:0 0 10px 0">
                Open the sidebar and review your hospital:
              </p>
              <ul style="color:#2e7d32; font-size:0.9em;
                         line-height:1.65; margin:0 0 12px 0;
                         padding-left:22px">
                <li>🏢 Hospital scale &amp; departments</li>
                <li>🌳 Healing garden + green roof</li>
                <li>🌞 Solar PV farm &amp; smart grid</li>
                <li>💧 Rainwater + greywater systems</li>
                <li>🚗 EV chargers, garage, drop-off</li>
                <li>🌍 Site location &amp; budget</li>
              </ul>
              <p style="color:#388e3c; font-size:0.82em;
                        font-style:italic; margin:0;
                        padding-top:6px;
                        border-top:1px dashed #a5d6a7">
                💡 Tip: hover over any setting for the relevant
                code/standard reference.
              </p>
            </div>
            """, unsafe_allow_html=True)
        with col2:
            st.markdown("""
            <div style="border:2px solid #1565c0; border-radius:14px;
                        padding:24px 26px; background:#e3f2fd;
                        min-height:440px; box-sizing:border-box;
                        box-shadow:0 2px 8px rgba(21,101,192,0.08)">
              <div style="font-size:2.6em; text-align:center;
                          line-height:1; margin-bottom:8px">🚀</div>
              <h3 style="color:#0d47a1; text-align:center;
                         margin:6px 0 14px 0; font-size:1.35em">
                Step 2 — Press Generate
              </h3>
              <p style="color:#1565c0; font-size:0.95em;
                        margin:0 0 10px 0">
                Scroll to the bottom of the sidebar and click
                <b>🚀 Generate Hospital Design</b>:
              </p>
              <ul style="color:#1565c0; font-size:0.9em;
                         line-height:1.65; margin:0 0 12px 0;
                         padding-left:22px">
                <li>🗺️ 2D campus site plan with road network</li>
                <li>🌐 Interactive 3D bird's-eye render</li>
                <li>📊 Energy, water, CO₂ savings dashboard</li>
                <li>🏆 Green certification forecast</li>
                <li>📑 Functional program + bill of quantities</li>
                <li>🤖 AI optimization + design chatbot</li>
              </ul>
              <p style="color:#0d47a1; font-size:0.82em;
                        font-style:italic; margin:0;
                        padding-top:6px;
                        border-top:1px dashed #90caf9">
                ⚡ After the first generation, sidebar changes
                will update results live.
              </p>
            </div>
            """, unsafe_allow_html=True)
        # Quick stats
        n_d = len(selected_depts) if selected_depts else 0
        st.markdown(
            "<div style='text-align:center; margin-top:30px; "
            "padding:15px; background:#f1f8e9; border-radius:8px'>"
            "<span style='color:#558b2f; font-size:0.9em'>"
            "Currently configured: "
            f"<b>{hospital_scale}</b> · "
            f"<b>{n_d}</b> departments selected · "
            f"Site location: <b>{h_sel_city or 'not set'}</b>"
            "</span></div>",
            unsafe_allow_html=True)
        st.stop()

    # Restore variables from snapshot
    hcfg = st.session_state.hospital_config
    hospital_scale = hcfg['hospital_scale']
    scale = hcfg['scale']
    selected_depts = hcfg['selected_depts']
    h_solar = hcfg['h_solar']
    h_garden = hcfg['h_garden']
    h_green_roof = hcfg['h_green_roof']
    h_rainwater = hcfg['h_rainwater']
    h_ev_chargers = hcfg['h_ev_chargers']
    h_greywater = hcfg['h_greywater']
    h_smart_grid = hcfg['h_smart_grid']
    building_mode = hcfg.get('building_mode', "🏘️ Separate Wings (multiple buildings)")
    h_garage = hcfg.get('h_garage', True)
    h_roads = hcfg.get('h_roads', True)
    h_outdoor_seating = hcfg.get('h_outdoor_seating', True)
    h_enable_budget = hcfg.get('h_enable_budget', False)
    h_user_budget = hcfg.get('h_user_budget', 0)
    h_sel_city = hcfg.get('h_sel_city', None)
    show_corridors = hcfg['show_corridors']
    show_patient_flow = hcfg['show_patient_flow']
    show_dept_labels = hcfg['show_dept_labels']

    # ============================================================
    # HOSPITAL-WIDE CALCULATIONS (EUI-based, FGI/Egyptian-aware)
    # ============================================================
    def calculate_hospital_metrics():
        """Hospital-wide energy/water/CO2 using EUI methodology.
        Formula per Energy Star Healthcare:
            Annual kWh = Σ (EUI_dept × area_dept) for each department
        Where EUI_dept = EUI_baseline × dept_multiplier (table-based).
        """
        ref = get_active_reference()
        eui_base = ref['eui_baseline']  # kWh/m²/year
        elec_price = ref['electricity_price']
        co2_factor = ref['co2_per_kwh']
        water_price = ref['water_price']

        total_area = 0
        total_annual_energy_normal = 0  # kWh/year (annual baseline)
        total_water = 0  # L/month

        for dept_name, dept_info in selected_depts.items():
            est_units = max(1, int(dept_info["max_units"] * 0.6))
            dept_area = est_units * dept_info["min_area_per_unit"] * 1.4 * scale
            total_area += dept_area
            # Get EUI multiplier for this department type
            dept_mult = DEPT_EUI_MULTIPLIERS.get(dept_name, 1.0)
            eui_dept = eui_base * dept_mult  # kWh/m²/year
            # Annual energy for this department
            annual_dept = eui_dept * dept_area
            total_annual_energy_normal += annual_dept
            total_water += dept_info["water_usage"] * est_units * 30  # L/mo

        # Monthly energy (kWh/month)
        total_energy_normal = total_annual_energy_normal / 12

        # Savings from sustainability features (additive, capped)
        savings_pct = 0.08
        if h_solar: savings_pct += 0.22
        if h_green_roof: savings_pct += 0.06
        if h_garden: savings_pct += 0.03
        if h_smart_grid: savings_pct += 0.10
        if h_ev_chargers: savings_pct += 0.02
        savings_pct = min(savings_pct, 0.70)

        water_savings = 0.05
        if h_rainwater: water_savings += 0.20
        if h_greywater: water_savings += 0.25
        water_savings = min(water_savings, 0.55)

        total_energy_green = total_energy_normal * (1 - savings_pct)
        energy_saved = total_energy_normal - total_energy_green
        money_saved = energy_saved * elec_price  # in active currency
        water_saved = total_water * water_savings

        # Carbon footprint based on local grid emissions factor
        co2_saved = energy_saved * co2_factor  # kg CO2/month
        trees_equiv = co2_saved / 21  # 1 tree ~ 21 kg CO2/year

        return {
            "total_area": total_area,
            "energy_normal": total_energy_normal,
            "energy_green": total_energy_green,
            "energy_saved": energy_saved,
            "money_saved": money_saved,
            "water_normal": total_water,
            "water_saved": water_saved,
            "co2_saved": co2_saved,
            "trees_equiv": trees_equiv,
            "savings_pct": savings_pct,
            "currency_symbol": ref['currency_symbol'],
            "reference_name": ref['short'],
        }

    hm = calculate_hospital_metrics()

    # ============================================================
    # TOP HOSPITAL METRICS
    # ============================================================
    _ref_h = get_active_reference()
    _curr_h = hm['currency_symbol']
    st.subheader(f"📊 {hospital_scale} — Sustainability Dashboard "
                 f"({hm['reference_name']})")
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("🏢 Total Built Area", f"{hm['total_area']:,.0f} m²",
              help=f"Sum of all department areas (incl. circulation factor 1.4). "
                   f"Standards: {_ref_h['citations']['min_area']}")
    c2.metric("📉 Energy Saved", f"{hm['energy_saved']/1000:,.1f} MWh/mo",
              delta=f"-{hm['savings_pct']*100:.1f}%",
              help=f"EUI baseline: {_ref_h['eui_baseline']} kWh/m²/year. "
                   f"Source: {_ref_h['citations']['eui']}")
    c3.metric("💰 Money Saved", f"{_curr_h}{hm['money_saved']*12/1000:,.0f}K/year",
              help=f"Electricity: {_curr_h}{_ref_h['electricity_price']:.2f}/kWh. "
                   f"Source: {_ref_h['citations']['electricity_price']}")
    c4.metric("🌍 CO₂ Reduced", f"{hm['co2_saved']*12/1000:,.1f} tons/year",
              help=f"Grid emissions factor: {_ref_h['co2_per_kwh']} kg CO₂/kWh. "
                   f"Source: {_ref_h['citations']['co2']}")
    c5.metric("🌳 Trees Equivalent", f"{hm['trees_equiv']*12:,.0f} trees/year",
              help="Based on 21 kg CO₂ absorbed per tree per year (US Forest Service)")

    if hm['water_saved'] > 0:
        st.info(f"💧 **Water Saved:** {hm['water_saved']/1000:,.1f} m³/month "
                f"({hm['water_saved']*12/1000:,.0f} m³/year) — enough to fill "
                f"{int(hm['water_saved']*12/2500)} Olympic-sized swimming pools annually 🏊")

    # ════════════════════════════════════════════════════════════════
    # 💰 HOSPITAL BUDGET DASHBOARD
    # ════════════════════════════════════════════════════════════════
    if h_enable_budget and h_user_budget > 0:
        st.markdown("---")
        st.subheader("💰 Hospital Project Budget Analysis")

        # Construction cost = sum of all dept construction costs
        total_construction = 0
        for dname, dinfo in selected_depts.items():
            est_units = max(1, int(dinfo['max_units'] * 0.6))
            dept_area = est_units * dinfo['min_area_per_unit'] * 1.4 * scale
            total_construction += calculate_construction_cost(dname, dept_area, 1)

        # Eco features cost (calculated for whole site)
        h_eco_dict = {
            'h_solar': h_solar, 'h_green_roof': h_green_roof,
            'h_garden': h_garden, 'h_rainwater': h_rainwater,
            'h_greywater': h_greywater, 'h_ev_chargers': h_ev_chargers,
            'h_smart_grid': h_smart_grid, 'h_garage': h_garage,
        }
        n_or = sum(1 for k in selected_depts if 'Operating Room' in k)
        eco_cost_h = calculate_eco_features_cost(
            h_eco_dict, total_area=hm['total_area'],
            n_chargers=8, n_or_rooms=n_or, garage_spaces=30)

        # Equipment cost — estimate as 15% of construction
        equip_cost_h = total_construction * 0.15

        # Site work (parking, roads, landscape) ~5% of construction
        site_work = total_construction * 0.05

        total_h_cost = total_construction + eco_cost_h + equip_cost_h + site_work
        contingency_h = total_h_cost * 0.10
        grand_total_h = total_h_cost + contingency_h

        # Display metrics
        hb1, hb2, hb3, hb4, hb5 = st.columns(5)
        hb1.metric("🏗️ Construction",
                   f"{_curr_h}{total_construction/1e6:,.2f}M",
                   help=f"All departments × cost per m²")
        hb2.metric("🏥 Equipment",
                   f"{_curr_h}{equip_cost_h/1e6:,.2f}M",
                   help="Medical equipment (~15% of construction)")
        hb3.metric("🌱 Sustainability",
                   f"{_curr_h}{eco_cost_h/1e6:,.2f}M",
                   help="Solar, garden, water systems, garage, etc.")
        hb4.metric("🛣️ Site Work",
                   f"{_curr_h}{site_work/1e6:,.2f}M",
                   help="Roads, landscaping, paving (~5%)")
        hb5.metric("📊 GRAND TOTAL",
                   f"{_curr_h}{grand_total_h/1e6:,.2f}M",
                   help="Including 10% contingency reserve")

        # Affordability indicator
        h_budget_pct = (grand_total_h / h_user_budget) * 100
        if grand_total_h <= h_user_budget:
            h_remaining = h_user_budget - grand_total_h
            st.success(
                f"✅ **Within Budget!** Total: {_curr_h}{grand_total_h/1e6:,.2f}M "
                f"({h_budget_pct:.0f}% of {_curr_h}{h_user_budget/1e6:,.2f}M). "
                f"Remaining buffer: **{_curr_h}{h_remaining/1e6:,.2f}M** "
                f"({100-h_budget_pct:.0f}%)."
            )
        else:
            h_overage = grand_total_h - h_user_budget
            st.error(
                f"❌ **Over Budget by {_curr_h}{h_overage/1e6:,.2f}M** "
                f"({h_budget_pct:.0f}% of {_curr_h}{h_user_budget/1e6:,.2f}M). "
                f"\n\n**Cost-cutting suggestions:**\n"
                f"- Disable Smart Grid: save ~{_curr_h}{ECO_FEATURE_COSTS[_ref_key_h]['h_smart_grid']['cost']/1e6:.2f}M\n"
                f"- Disable Greywater Recycling: save ~{_curr_h}{ECO_FEATURE_COSTS[_ref_key_h]['h_greywater']['cost']/1e3:.0f}K\n"
                f"- Reduce hospital scale to 'Small Clinic': save ~30-50%\n"
                f"- Remove parking garage: save ~{_curr_h}{ECO_FEATURE_COSTS[_ref_key_h]['h_garage']['cost']*30/1e6:.2f}M"
            )

        # ROI breakdown
        with st.expander("📈 Hospital ROI Analysis", expanded=False):
            h_annual_savings = hm['money_saved'] * 12
            if h_annual_savings > 0:
                h_payback = grand_total_h / h_annual_savings
                h_roi_20yr = (h_annual_savings * 20 - grand_total_h) / grand_total_h * 100
                rch1, rch2, rch3 = st.columns(3)
                rch1.metric("Annual Energy Savings",
                            f"{_curr_h}{h_annual_savings/1e6:,.2f}M/year")
                rch2.metric("Energy Payback Period",
                            f"{h_payback:.1f} years" if h_payback < 100 else "N/A",
                            help="From electricity savings alone")
                rch3.metric("20-Year ROI",
                            f"{h_roi_20yr:+.0f}%")
                st.caption(f"💡 Note: This is energy-only ROI. Total ROI is "
                           f"much higher when including operational benefits, "
                           f"staff productivity, patient outcomes, and asset "
                           f"appreciation.")

    # ════════════════════════════════════════════════════════════════
    # 🏆 HOSPITAL CERTIFICATION FORECAST (LEED / Estidama / Mostadam)
    # ════════════════════════════════════════════════════════════════
    st.markdown("---")
    st.subheader("🏆 Green Building Certification Forecast")

    h_cert_choices = {
        'solar_pv': h_solar,
        'energy_star': True,  # hospital-grade equipment assumed efficient
        'smart_grid': h_smart_grid,
        'natural_light': True,  # hospitals designed with daylighting
        'healing_garden': h_garden,
        'rainwater': h_rainwater,
        'greywater': h_greywater,
        'eco_flooring': True,
        'waste_mgmt': True,
        'green_roof': h_green_roof,
        'ev_chargers': h_ev_chargers,
        'garden_site': h_garden,
    }
    h_cert = predict_certification(h_cert_choices)

    hcert1, hcert2, hcert3 = st.columns(3)
    with hcert1:
        st.markdown("### LEED v4 HC")
        st.markdown(f"<div style='padding:12px;border-radius:8px;"
                    f"background:{h_cert['leed_color']};text-align:center;"
                    f"color:#1a1a1a;font-weight:bold;font-size:1.2em'>"
                    f"{h_cert['leed_level']}<br>{h_cert['leed_score']}/100 pts</div>",
                    unsafe_allow_html=True)
    with hcert2:
        st.markdown("### Estidama 🇦🇪")
        st.markdown(f"<div style='padding:12px;border-radius:8px;"
                    f"background:#e8f5e9;text-align:center;color:#1b5e20;"
                    f"font-weight:bold;font-size:1.4em'>"
                    f"{'⭐' * h_cert['pearls']}<br>{h_cert['pearls']} Pearls</div>",
                    unsafe_allow_html=True)
    with hcert3:
        st.markdown("### Mostadam 🇸🇦")
        st.markdown(f"<div style='padding:12px;border-radius:8px;"
                    f"background:#e8f5e9;text-align:center;color:#1b5e20;"
                    f"font-weight:bold;font-size:1.2em'>"
                    f"{h_cert['mostadam']}<br>Rating</div>",
                    unsafe_allow_html=True)

    with st.expander("📋 Certification Credit Breakdown by Category",
                      expanded=False):
        for cat, vals in h_cert['by_category'].items():
            pct = (vals['earned'] / vals['max'] * 100) if vals['max'] else 0
            st.markdown(f"**{cat}**: {vals['earned']}/{vals['max']} pts "
                        f"({pct:.0f}%)")
            st.progress(pct / 100)
        missing = [d for d in h_cert['credits_detail'] if not d['achieved']]
        if missing:
            st.markdown("**💡 Credits not yet earned:**")
            for d in missing:
                st.caption(f"⬜ {d['credit']} (+{d['max']} pts) — "
                           f"{d['category']}")

    # ════════════════════════════════════════════════════════════════
    # ☀️ HOSPITAL RENEWABLE POTENTIAL (if city selected)
    # ════════════════════════════════════════════════════════════════
    if h_sel_city:
        st.markdown("---")
        st.subheader(f"☀️ Site Renewable Potential — {h_sel_city}")
        _hcd2 = detect_climate(h_sel_city)
        # Estimate roof area as ~40% of total built area (single-story footprint)
        h_roof_area = hm['total_area'] * 0.4
        h_solar_annual = estimate_solar_potential(_hcd2['sun_hours'],
                                                   h_roof_area * 0.7)
        h_rain_annual = estimate_rainwater_harvest(_hcd2['rainfall'], h_roof_area)

        hsp1, hsp2, hsp3 = st.columns(3)
        hsp1.metric("☀️ Solar Farm Potential",
                    f"{h_solar_annual/1e6:,.2f} GWh/yr",
                    help=f"70% of {h_roof_area:.0f}m² roof with PV at "
                         f"{_hcd2['sun_hours']}h peak sun")
        hsp2.metric("💧 Rainwater Harvest",
                    f"{h_rain_annual/1e6:,.2f} ML/yr",
                    help=f"{_hcd2['rainfall']}mm rainfall on {h_roof_area:.0f}m² roof")
        h_solar_offset = min(100, (h_solar_annual / (hm['energy_normal']*12)) * 100) if hm['energy_normal'] > 0 else 0
        hsp3.metric("⚡ Energy Self-Sufficiency",
                    f"{h_solar_offset:.0f}%",
                    help="% of hospital energy demand solar could supply")
        if h_solar_offset > 30:
            st.success(f"🌟 {h_sel_city}'s solar resource could make this "
                       f"hospital {h_solar_offset:.0f}% energy self-sufficient "
                       f"with rooftop PV alone!")

    # ════════════════════════════════════════════════════════════════
    # 🔬 HOSPITAL ENGINEERING ANALYSIS
    # HVAC + Fire Safety + Water Balance + Power SLD
    # ════════════════════════════════════════════════════════════════
    st.markdown("---")
    st.subheader("🔬 Engineering Systems Analysis")

    he_tab1, he_tab2, he_tab3, he_tab4 = st.tabs([
        "💨 HVAC", "🚨 Fire Safety", "💧 Water Balance", "⚡ Power SLD"])

    with he_tab1:
        # Hospital-wide HVAC: aggregate across departments
        total_cooling_kw = 0
        worst_ach = 6
        for dname, dinfo in selected_depts.items():
            est_units = max(1, int(dinfo['max_units'] * 0.6))
            d_area = est_units * dinfo['min_area_per_unit'] * 1.4 * scale
            d_hvac = calculate_hvac_load(
                d_area, dinfo.get('min_ceiling', 3.0),
                est_units + 2, dinfo['base_power'],
                d_area * 15, 'Hot', dinfo.get('ach', 6))
            total_cooling_kw += d_hvac['total_kW']
            worst_ach = max(worst_ach, dinfo.get('ach', 6))

        total_tons = total_cooling_kw / 3.517
        ht1, ht2, ht3 = st.columns(3)
        ht1.metric("❄️ Total Cooling Load",
                   f"{total_cooling_kw:,.0f} kW",
                   help="Sum across all departments")
        ht2.metric("🔧 Central Plant Size",
                   f"{total_tons:,.0f} tons",
                   help="Required chiller plant capacity")
        ht3.metric("💨 Peak ACH Required",
                   f"{worst_ach} ACH",
                   help="Highest ventilation rate (likely OR/Isolation)")
        st.caption("📖 ASHRAE 170 + 90.1. Central chilled-water plant "
                   "recommended for hospitals > 100 tons.")
        st.info("💡 **Design tip:** For this load, consider a central chiller "
                "plant with N+1 redundancy (one backup chiller) to maintain "
                "cooling during maintenance — critical for OR/ICU.")

    with he_tab2:
        total_h_area = hm['total_area']
        # Use ER occupancy factor as representative, or sum
        h_fire = calculate_fire_safety(total_h_area, 'Emergency Room (ER)')
        hf1, hf2, hf3 = st.columns(3)
        hf1.metric("👥 Total Occupant Load", f"{h_fire['occupants']:,} people")
        hf2.metric("🚪 Required Exits (min)", f"{h_fire['req_exits']}")
        hf3.metric("🔥 Smoke Compartments", f"{h_fire['n_smoke_compartments']}")
        hf4, hf5 = st.columns(2)
        hf4.metric("🧯 Fire Extinguishers", f"{h_fire['n_extinguishers']}")
        hf5.metric("📏 Total Egress Width", f"{h_fire['total_egress_m']:.1f} m")
        st.warning("⚠️ **Hospital fire safety requires:** fire-rated smoke "
                   "barriers between compartments, emergency generators for "
                   "egress lighting, areas of refuge for non-ambulatory "
                   "patients, and a fully addressable fire alarm system "
                   "(NFPA 72).")
        st.caption("📖 NFPA 101 Life Safety Code + NFPA 99.")

    with he_tab3:
        # Water balance Sankey
        h_water_lpm = hm['water_normal']  # liters/month
        rain_potential = 0
        if h_sel_city and h_rainwater:
            _wcd = detect_climate(h_sel_city)
            rain_potential = estimate_rainwater_harvest(
                _wcd['rainfall'], hm['total_area'] * 0.4) / 12  # monthly
        try:
            water_fig = build_water_sankey(
                h_water_lpm, h_rainwater, h_greywater, rain_potential)
            st.plotly_chart(water_fig, use_container_width=True)
        except Exception as e:
            st.error(f"Water diagram error: {e}")
        wc1, wc2 = st.columns(2)
        wc1.metric("💧 Total Water Demand",
                   f"{h_water_lpm/1000:,.0f} m³/mo")
        if h_greywater:
            wc2.metric("♻️ Greywater Recovered",
                       f"~{h_water_lpm*0.30/1000:,.0f} m³/mo",
                       help="~30% of demand recovered via recycling")
        st.caption("📖 Flows based on typical hospital water use distribution "
                   "(sanitary 40%, HVAC 25%, process 20%, irrigation 15%).")

    with he_tab4:
        # Power SLD
        total_load_kw = total_cooling_kw * 1.3 + hm['energy_normal'] / 720
        try:
            power_fig = build_power_sld(
                h_solar, h_smart_grid, total_load_kw, has_critical=True)
            st.pyplot(power_fig)
            plt.close(power_fig)
        except Exception as e:
            st.error(f"Power diagram error: {e}")
        st.caption("📖 NFPA 99 Healthcare Electrical Systems. Critical "
                   "branch (OR/ICU/Life-support) must have emergency power "
                   "within 10 seconds via automatic transfer switch (ATS).")

    # ════════════════════════════════════════════════════════════════
    # 📑 HOSPITAL PROJECT DOCUMENTS (BoQ + Functional Program)
    # ════════════════════════════════════════════════════════════════
    st.markdown("---")
    st.subheader("📑 Project Documents")

    hdoc1, hdoc2 = st.tabs(["📋 Bill of Quantities", "📄 Functional Program"])

    with hdoc1:
        h_boq_config = {
            'selected_depts': selected_depts, 'scale': scale,
            'hospital_scale': hospital_scale,
            'h_solar': h_solar, 'h_green_roof': h_green_roof,
            'h_garden': h_garden, 'h_rainwater': h_rainwater,
            'h_greywater': h_greywater, 'h_smart_grid': h_smart_grid,
        }
        h_boq_rows = build_boq_rows('hospital', h_boq_config)
        h_boq_total = sum(r[f'Amount ({_curr_h})'] for r in h_boq_rows)
        h_boq_cont = h_boq_total * 0.10
        h_boq_grand = h_boq_total + h_boq_cont

        st.markdown(f"**Itemized cost breakdown — {hospital_scale}**")
        try:
            import pandas as pd
            st.dataframe(pd.DataFrame(h_boq_rows),
                         use_container_width=True, hide_index=True)
        except Exception:
            for r in h_boq_rows:
                st.text(f"{r['No.']}. {r['Description']}: "
                        f"{_curr_h}{r[f'Amount ({_curr_h})']:,.0f}")

        hbs1, hbs2, hbs3 = st.columns(3)
        hbs1.metric("Subtotal", f"{_curr_h}{h_boq_total/1e6:,.2f}M")
        hbs2.metric("Contingency", f"{_curr_h}{h_boq_cont/1e6:,.2f}M")
        hbs3.metric("GRAND TOTAL", f"{_curr_h}{h_boq_grand/1e6:,.2f}M")

        hexp1, hexp2 = st.columns(2)
        with hexp1:
            try:
                import pandas as pd
                h_csv = pd.DataFrame(h_boq_rows).to_csv(index=False)
            except Exception:
                hdr = list(h_boq_rows[0].keys()) if h_boq_rows else []
                h_csv = ','.join(hdr) + '\n' + '\n'.join(
                    ','.join(str(r[h]) for h in hdr) for r in h_boq_rows)
            st.download_button("⬇️ Download CSV", h_csv,
                file_name="Hospital_BoQ.csv", mime="text/csv",
                use_container_width=True)
        with hexp2:
            h_xlsx = boq_to_excel_bytes(h_boq_rows, h_boq_grand, _curr_h)
            if h_xlsx:
                st.download_button("⬇️ Download Excel", h_xlsx,
                    file_name="Hospital_BoQ.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True)
            else:
                st.caption("Excel needs openpyxl — CSV works above.")

    with hdoc2:
        h_fp_md = build_functional_program_md('hospital', h_boq_config)
        st.markdown("**Hospital Functional Program** — department schedule + "
                    "room data sheets for all departments.")
        with st.expander("👁️ Preview Document", expanded=False):
            st.markdown(h_fp_md)
        hfp1, hfp2 = st.columns(2)
        with hfp1:
            st.download_button("⬇️ Download Markdown", h_fp_md,
                file_name="Hospital_FunctionalProgram.md",
                mime="text/markdown", use_container_width=True)
        with hfp2:
            h_docx = md_to_docx_bytes(h_fp_md, "Hospital Functional Program")
            if h_docx:
                st.download_button("⬇️ Download Word", h_docx,
                    file_name="Hospital_FunctionalProgram.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True)
            else:
                st.caption("Word needs python-docx — Markdown works above.")

    # ════════════════════════════════════════════════════════════════
    # 🤖 HOSPITAL AI SMART FEATURES
    # ════════════════════════════════════════════════════════════════
    st.markdown("---")
    st.subheader("🤖 AI Smart Assistant")

    hai1, hai2, hai3 = st.tabs([
        "🔍 Design Check", "⚙️ Optimize", "💬 Ask the Assistant"])

    with hai1:
        h_anomaly_config = {
            'selected_depts': selected_depts,
            'h_solar': h_solar, 'h_greywater': h_greywater,
            'h_rainwater': h_rainwater,
        }
        h_anomalies = detect_anomalies('hospital', h_anomaly_config)
        if not h_anomalies:
            st.success("✅ **No issues detected!** Your hospital program "
                       "follows good adjacency and sustainability practice.")
        else:
            n_w = sum(1 for s, _ in h_anomalies if s == 'warning')
            n_i = sum(1 for s, _ in h_anomalies if s == 'info')
            st.markdown(f"Found **{n_w} warnings**, **{n_i} suggestions**:")
            for severity, msg in h_anomalies:
                if severity == 'error':
                    st.error(msg)
                elif severity == 'warning':
                    st.warning(msg)
                else:
                    st.info(msg)

    with hai2:
        st.markdown("**Optimization recommendations** for your hospital:")
        h_opt_obj = st.radio("Optimize for:",
            ["💰 Minimum Cost", "🌍 Minimum CO₂", "🌱 Maximum Green Score",
             "🏆 Maximum LEED"],
            horizontal=True, key="opt_obj_hosp")
        h_obj_map = {"💰 Minimum Cost": 'cost', "🌍 Minimum CO₂": 'co2',
                     "🌱 Maximum Green Score": 'green', "🏆 Maximum LEED": 'leed'}
        if st.button("⚙️ Generate Recommendations", key="opt_btn_hosp"):
            sel_obj = h_obj_map[h_opt_obj]
            st.markdown(f"#### 🎯 For **{h_opt_obj}**:")
            if sel_obj == 'cost':
                st.markdown("- Disable **Smart Grid** (highest capital cost)\n"
                            "- Keep **Rainwater Harvesting** (cheap, high ROI)\n"
                            "- Defer **EV Chargers** to phase 2\n"
                            "- Consider downsizing the **Hospital Scale**")
            elif sel_obj == 'co2':
                st.markdown("- Enable **Solar PV Farm** (max clean energy)\n"
                            "- Enable **Smart Grid + Battery** (store solar)\n"
                            "- Enable **Green Roof** (insulation + carbon)\n"
                            "- Enable **all water recycling** (pump energy savings)")
            elif sel_obj == 'green':
                st.markdown("- Enable **ALL** sustainability features\n"
                            "- **Healing Garden** + **Green Roof** for biophilia\n"
                            "- **Greywater + Rainwater** for water credits\n"
                            "- **Smart Grid** for energy management")
            else:  # leed
                st.markdown("- **Solar PV** → Energy & Atmosphere credits\n"
                            "- **Healing Garden** → Indoor Env. Quality + Sites\n"
                            "- **Water recycling** → Water Efficiency credits\n"
                            "- **EV Chargers** → Location & Transportation\n"
                            "- **Green Roof** → Sustainable Sites (heat island)")
            st.info("💡 Toggle these features in the sidebar and regenerate.")

    with hai3:
        st.markdown("Ask me about hospital design, standards, or sustainability!")
        if 'chat_history_hosp' not in st.session_state:
            st.session_state.chat_history_hosp = []
        h_sugg_cols = st.columns(3)
        h_suggestions = ["What is LEED certification?",
                         "Why recycle greywater?",
                         "How much can solar save?"]
        for i, sugg in enumerate(h_suggestions):
            if h_sugg_cols[i].button(sugg, key=f"sugg_h_{i}"):
                ans = chatbot_answer(sugg, {'sel_city': h_sel_city})
                st.session_state.chat_history_hosp.append((sugg, ans))
        h_user_q = st.chat_input("Type your question...", key="chat_input_hosp")
        if h_user_q:
            ans = chatbot_answer(h_user_q, {'sel_city': h_sel_city})
            st.session_state.chat_history_hosp.append((h_user_q, ans))
        for q_text, a_text in reversed(st.session_state.chat_history_hosp[-6:]):
            with st.chat_message("user"):
                st.markdown(q_text)
            with st.chat_message("assistant"):
                st.markdown(a_text)

    st.markdown("---")

    # ============================================================
    # HOSPITAL SITE PLAN DRAWING
    # ============================================================
    def draw_site_plan():
        """Lay out all departments on a grid representing the hospital site.
        Uses a 'garden-first' grid that reserves the central 2x2 cells for
        the healing garden BEFORE placing any departments — guaranteeing no
        overlap. Site dimensions scale dynamically with department count."""
        fig, ax = plt.subplots(figsize=(14, 11))

        # ─── Grid sizing: depends on number of departments ───
        depts = list(selected_depts.items())
        n_depts = len(depts)
        garden_cells = 4 if h_garden else 0  # 2x2 block reserved for garden
        total_cells_needed = n_depts + garden_cells

        # Target aspect ratio for the site (4:3 roughly)
        target_aspect = 4.0 / 3.0
        cols = max(4, int(np.ceil(np.sqrt(total_cells_needed * target_aspect))))
        rows = max(3, int(np.ceil(total_cells_needed / cols)))
        # Re-expand if cells insufficient after rounding
        while cols * rows < total_cells_needed:
            cols += 1

        # Each cell is a hospital wing of ~12m × 10m, multiplied by scale
        cell_w = 12 * (scale ** 0.4)
        cell_h = 10 * (scale ** 0.4)
        margin = 5

        site_w = cols * cell_w + 2*margin
        site_h = rows * cell_h + 2*margin

        ax.set_xlim(-6, site_w + 6)
        ax.set_ylim(-6, site_h + 7)
        ax.set_facecolor('#e8f5e9')  # Lawn background

        # ─── Site boundary ───
        ax.add_patch(patches.Rectangle((0, 0), site_w, site_h,
            facecolor='#f1f8e9', edgecolor='#2e7d32',
            linewidth=2, linestyle='--', zorder=1))

        # ─── Green perimeter (trees) — spaced by site size ───
        n_trees_x = max(8, int(site_w / 5))
        n_trees_y = max(5, int(site_h / 5))
        for x in np.linspace(2, site_w - 2, n_trees_x):
            for y_pos in [1.5, site_h - 1.5]:
                ax.add_patch(Circle((x, y_pos), 0.8, facecolor='#2e7d32',
                    edgecolor='#1b5e20', linewidth=0.5, zorder=2))
        for y in np.linspace(4, site_h - 4, n_trees_y):
            for x_pos in [1.5, site_w - 1.5]:
                ax.add_patch(Circle((x_pos, y), 0.8, facecolor='#2e7d32',
                    edgecolor='#1b5e20', linewidth=0.5, zorder=2))

        # ─── Identify garden cells (center 2×2 block) BEFORE placing depts ───
        garden_cells_set = set()
        g_col_start, g_row_start = -1, -1
        if h_garden:
            g_col_start = (cols - 2) // 2
            g_row_start = (rows - 2) // 2
            for r in range(g_row_start, g_row_start + 2):
                for c in range(g_col_start, g_col_start + 2):
                    garden_cells_set.add((r, c))

        # ─── Build list of available (non-garden) cells in priority order ───
        # Cells closer to center get higher priority for Critical Care
        center_r, center_c = rows / 2, cols / 2
        all_cells = []
        for r in range(rows):
            for c in range(cols):
                if (r, c) not in garden_cells_set:
                    dist_from_center = np.sqrt((r-center_r)**2 + (c-center_c)**2)
                    all_cells.append((dist_from_center, r, c))
        all_cells.sort()  # Closest to center first
        available_cells = [(r, c) for _, r, c in all_cells]

        # ─── Sort departments by priority (Critical Care closest to garden) ───
        priority_order = {"Critical Care": 0, "Diagnostic": 1, "Treatment": 2,
                          "Specialized": 3, "General Care": 4, "Service": 5}
        depts_sorted = sorted(depts,
            key=lambda x: priority_order.get(x[1]["category"], 99))

        # ─── Place departments in cells ───
        positions = []
        for idx, (dname, dinfo) in enumerate(depts_sorted):
            if idx >= len(available_cells):
                break  # Safety check (shouldn't happen with correct sizing)
            r, c = available_cells[idx]
            # Cell coordinates (rows top-down → flip y)
            dx = margin + c * cell_w + cell_w * 0.08
            dy = margin + (rows - 1 - r) * cell_h + cell_h * 0.08
            dw = cell_w * 0.84
            dh = cell_h * 0.84

            color = CATEGORY_COLORS[dinfo["category"]]

            # Department block
            ax.add_patch(FancyBboxPatch((dx, dy), dw, dh,
                boxstyle="round,pad=0.15", facecolor=color,
                edgecolor='black', linewidth=1.5, alpha=0.88, zorder=5))
            # Inner highlight
            ax.add_patch(patches.Rectangle((dx+0.4, dy+0.4), dw-0.8, dh-0.8,
                facecolor='none', edgecolor='white', linewidth=0.6,
                alpha=0.4, zorder=6))

            # Labels
            if show_dept_labels:
                est_units = max(1, int(dinfo["max_units"] * 0.6))
                dept_area = est_units * dinfo["min_area_per_unit"] * 1.4 * scale
                short_name = dname.split('(')[0].strip()
                if len(short_name) > 18: short_name = short_name[:16] + "…"
                # Icon
                ax.text(dx + dw/2, dy + dh*0.72, dinfo['icon'],
                    fontsize=min(18, cell_w*1.5), ha='center', va='center', zorder=7)
                # Name
                ax.text(dx + dw/2, dy + dh*0.40, short_name,
                    fontsize=min(8, cell_w*0.8), ha='center', va='center',
                    weight='bold', color='white', zorder=7)
                # Area
                ax.text(dx + dw/2, dy + dh*0.18, f"{dept_area:.0f} m²",
                    fontsize=min(7, cell_w*0.7), ha='center', va='center',
                    color='white', style='italic', zorder=7)

            positions.append((dname, dinfo, dx + dw/2, dy + dh/2, dw, dh))

        # ─── HEALING GARDEN (occupies reserved 2×2 cells, no overlap risk) ───
        if h_garden:
            # Garden grass square fills the 2x2 cell area
            g_dx = margin + g_col_start * cell_w + cell_w * 0.1
            g_dy = margin + (rows - 2 - g_row_start) * cell_h + cell_h * 0.1
            g_dw = 2 * cell_w - cell_w * 0.2
            g_dh = 2 * cell_h - cell_h * 0.2
            g_cx = g_dx + g_dw / 2
            g_cy = g_dy + g_dh / 2

            # Soft grass square (rounded)
            ax.add_patch(FancyBboxPatch((g_dx, g_dy), g_dw, g_dh,
                boxstyle="round,pad=0.2", facecolor='#c8e6c9',
                edgecolor='#2e7d32', linewidth=2, alpha=0.95, zorder=4))

            # Garden inner area (slightly darker green ring)
            inner_pad = 0.5
            ax.add_patch(FancyBboxPatch(
                (g_dx + inner_pad, g_dy + inner_pad),
                g_dw - 2*inner_pad, g_dh - 2*inner_pad,
                boxstyle="round,pad=0.15", facecolor='#a5d6a7',
                edgecolor='#558b2f', linewidth=1, alpha=0.7, zorder=5))

            # Walking paths (cross shape through garden)
            ax.plot([g_dx, g_dx + g_dw], [g_cy, g_cy],
                color='#bcaaa4', linewidth=2.5, alpha=0.6, zorder=6)
            ax.plot([g_cx, g_cx], [g_dy, g_dy + g_dh],
                color='#bcaaa4', linewidth=2.5, alpha=0.6, zorder=6)

            # Trees inside garden (4 corners-ish)
            tree_radius = min(g_dw, g_dh) * 0.35
            for ang in np.linspace(np.pi/4, 2*np.pi + np.pi/4, 8, endpoint=False):
                tx = g_cx + tree_radius*np.cos(ang)
                ty = g_cy + tree_radius*np.sin(ang)
                # Tree shadow
                ax.add_patch(Circle((tx+0.1, ty-0.1), 0.7,
                    facecolor='#1b5e20', alpha=0.3, zorder=6))
                # Tree leaves
                ax.add_patch(Circle((tx, ty), 0.7,
                    facecolor='#2e7d32', edgecolor='#1b5e20',
                    linewidth=0.8, zorder=7))
                # Tree highlight
                ax.add_patch(Circle((tx-0.15, ty+0.15), 0.3,
                    facecolor='#43a047', alpha=0.6, zorder=8))

            # Central fountain (precisely centered, no overlap)
            fountain_r = min(g_dw, g_dh) * 0.12
            # Outer stone ring
            ax.add_patch(Circle((g_cx, g_cy), fountain_r * 1.4,
                facecolor='#9e9e9e', edgecolor='#424242',
                linewidth=1.5, zorder=7))
            # Water basin
            ax.add_patch(Circle((g_cx, g_cy), fountain_r,
                facecolor='#4fc3f7', edgecolor='#0277bd',
                linewidth=1.2, zorder=8))
            # Water ripples
            for ripple_r in [0.4, 0.7]:
                ax.add_patch(Circle((g_cx, g_cy), fountain_r * ripple_r,
                    facecolor='none', edgecolor='white',
                    linewidth=0.8, alpha=0.7, zorder=9))
            # Center water spout
            ax.add_patch(Circle((g_cx, g_cy), fountain_r * 0.15,
                facecolor='white', alpha=0.9, zorder=10))

            # Label OUTSIDE the garden (above, in margin) — won't overlap
            ax.text(g_cx, g_dy + g_dh + 0.5, "🌳 Healing Garden",
                fontsize=10, ha='center', weight='bold',
                color='#1b5e20', zorder=11,
                bbox=dict(boxstyle='round,pad=0.3',
                    facecolor='#f1f8e9', edgecolor='#2e7d32', alpha=0.9))

        # Set center reference for corridors and patient flow
        cx, cy = site_w / 2, site_h / 2

        # ─── Corridors (connect departments to nearest neighbors via center) ───
        if show_corridors:
            for (_, _, px, py, pw, ph) in positions:
                ax.plot([px, cx], [py, cy], color='#9e9e9e',
                    linewidth=1.5, alpha=0.4, linestyle='-', zorder=3)

        # ─── Patient Flow (ER → Radiology → OR → ICU → Ward) ───
        if show_patient_flow:
            flow_order = ["Emergency Room (ER)", "Radiology Department",
                          "Operating Room (OR)", "Intensive Care Unit (ICU)",
                          "General Patient Ward"]
            flow_positions = []
            for fname in flow_order:
                for (dname, _, px, py, _, _) in positions:
                    if dname == fname:
                        flow_positions.append((px, py, fname))
                        break

            for i in range(len(flow_positions) - 1):
                x1, y1, _ = flow_positions[i]
                x2, y2, _ = flow_positions[i+1]
                arrow = FancyArrowPatch((x1, y1), (x2, y2),
                    arrowstyle='->', color='#ff5722', linewidth=2.5,
                    alpha=0.7, zorder=9,
                    connectionstyle="arc3,rad=0.15", mutation_scale=20)
                ax.add_patch(arrow)

            if flow_positions:
                ax.text(flow_positions[0][0], flow_positions[0][1] + 2.5,
                    "🩹 Patient Flow", fontsize=8, ha='center', weight='bold',
                    color='#bf360c', bbox=dict(boxstyle='round,pad=0.3',
                    facecolor='#ffe0b2', edgecolor='#ff5722'), zorder=11)

        # ─── Solar panels on roof (icons in top strip) ───
        if h_solar:
            solar_y = site_h - 0.3
            for sx in np.linspace(margin + 2, site_w - margin - 2, 12):
                ax.add_patch(patches.Rectangle((sx - 0.6, solar_y), 1.2, 0.6,
                    facecolor='#1565c0', edgecolor='#0d47a1', linewidth=0.5,
                    hatch='//', zorder=8))
            ax.text(site_w/2, site_h + 1, "🌞 Rooftop Solar Farm",
                fontsize=8, ha='center', weight='bold', color='#0d47a1', zorder=10)

        # ─── Rainwater tank ───
        if h_rainwater:
            ax.add_patch(Circle((site_w - 3, 4), 1.5, facecolor='#4fc3f7',
                edgecolor='#01579b', linewidth=1.5, zorder=8))
            ax.text(site_w - 3, 4, "💧", fontsize=12, ha='center',
                va='center', zorder=9)
            ax.text(site_w - 3, 1.5, "Rainwater\nTank", fontsize=6,
                ha='center', weight='bold', color='#01579b', zorder=10)

        # ─── EV chargers ───
        if h_ev_chargers:
            for ex in np.linspace(3, 9, 4):
                ax.add_patch(patches.Rectangle((ex - 0.3, 1.5), 0.6, 1.2,
                    facecolor='#66bb6a', edgecolor='#1b5e20',
                    linewidth=0.8, zorder=8))
                ax.text(ex, 2.1, "⚡", fontsize=6, ha='center',
                    va='center', zorder=9)
            ax.text(6, 0.5, "EV Charging", fontsize=6,
                ha='center', weight='bold', color='#1b5e20', zorder=10)

        # ─── Main entrance (bottom center) ───
        ax.add_patch(patches.Rectangle((cx - 2, -0.5), 4, 0.5,
            facecolor='#5d4037', edgecolor='black', linewidth=1, zorder=8))
        ax.text(cx, -1.5, "🚪 Main Entrance", fontsize=9, ha='center',
            weight='bold', color='#3e2723', zorder=10)

        # ─── Title block ───
        ax.text(0, site_h + 3.5, f"VIRIDIS HOSPITAL SITE PLAN — {hospital_scale}",
            fontsize=14, weight='bold', color='#1b5e20')
        ax.text(0, site_h + 2.5,
            f"Departments: {len(selected_depts)} | Total Built Area: {hm['total_area']:,.0f} m² | "
            f"Annual Energy Savings: {hm['energy_saved']*12/1000:,.0f} MWh",
            fontsize=8, color='#424242', style='italic')

        # ─── Legend (category colors) ───
        legend_elements = [patches.Patch(facecolor=col, label=cat)
                          for cat, col in CATEGORY_COLORS.items()]
        if show_patient_flow:
            legend_elements.append(Line2D([0], [0], color='#ff5722',
                lw=2, label='Patient Flow'))
        if show_corridors:
            legend_elements.append(Line2D([0], [0], color='#9e9e9e',
                lw=1, label='Corridors'))
        ax.legend(handles=legend_elements, loc='upper right',
            bbox_to_anchor=(1.0, -0.02), ncol=4, fontsize=7, framealpha=0.95)

        # ─── North arrow ───
        na_x, na_y = site_w + 2, site_h - 2
        ax.add_patch(Circle((na_x, na_y), 1.2, facecolor='white',
            edgecolor='black', linewidth=1, zorder=10))
        ax.annotate('', xy=(na_x, na_y + 0.8), xytext=(na_x, na_y - 0.5),
            arrowprops=dict(arrowstyle='->', color='red', lw=2))
        ax.text(na_x, na_y + 1.6, "N", fontsize=10, ha='center',
            weight='bold', color='red', zorder=10)

        # ─── Scale bar ───
        sb_x, sb_y = 0, -3
        ax.add_patch(patches.Rectangle((sb_x, sb_y), 5, 0.3,
            facecolor='black', zorder=10))
        ax.add_patch(patches.Rectangle((sb_x + 5, sb_y), 5, 0.3,
            facecolor='white', edgecolor='black', zorder=10))
        ax.text(sb_x, sb_y - 1, "0", fontsize=7, ha='center', zorder=10)
        ax.text(sb_x + 5, sb_y - 1, "5m", fontsize=7, ha='center', zorder=10)
        ax.text(sb_x + 10, sb_y - 1, "10m", fontsize=7, ha='center', zorder=10)

        ax.set_aspect('equal')
        ax.axis('off')
        plt.tight_layout()
        return fig, positions

    # ============================================================
    # HOSPITAL 3D RENDER (Bird's-eye 3D view of whole site)
    # ============================================================
    def draw_hospital_3d():
        """OPTIMIZED 3D site plan. Uses MeshAccumulator to merge hundreds of
        small meshes into ~12 grouped meshes for smooth interactivity.
        Plotly with 500+ Mesh3d traces lags badly; with ~12 it's fluid."""
        # Re-compute the same grid as draw_site_plan so positions match
        depts = list(selected_depts.items())
        n_depts = len(depts)
        garden_cells = 4 if h_garden else 0
        total_cells_needed = n_depts + garden_cells
        target_aspect = 4.0 / 3.0
        cols = max(4, int(np.ceil(np.sqrt(total_cells_needed * target_aspect))))
        rows = max(3, int(np.ceil(total_cells_needed / cols)))
        while cols * rows < total_cells_needed:
            cols += 1
        cell_w = 12 * (scale ** 0.4)
        cell_h = 10 * (scale ** 0.4)
        margin = 5
        site_w = cols * cell_w + 2*margin
        site_h = rows * cell_h + 2*margin

        # Building heights by category (m)
        height_map = {
            "Critical Care": 9.0, "Diagnostic": 6.0, "Treatment": 6.0,
            "Specialized": 7.5, "General Care": 10.5, "Service": 4.5,
        }

        # Identify garden cells
        garden_cells_set = set()
        g_col_start, g_row_start = -1, -1
        if h_garden:
            g_col_start = (cols - 2) // 2
            g_row_start = (rows - 2) // 2
            for r in range(g_row_start, g_row_start + 2):
                for c in range(g_col_start, g_col_start + 2):
                    garden_cells_set.add((r, c))

        # Sort cells by distance from center
        center_r, center_c = rows / 2, cols / 2
        all_cells = []
        for r in range(rows):
            for c in range(cols):
                if (r, c) not in garden_cells_set:
                    dist = np.sqrt((r-center_r)**2 + (c-center_c)**2)
                    all_cells.append((dist, r, c))
        all_cells.sort()
        available_cells = [(r, c) for _, r, c in all_cells]

        # Sort departments
        priority_order = {"Critical Care": 0, "Diagnostic": 1, "Treatment": 2,
                          "Specialized": 3, "General Care": 4, "Service": 5}
        depts_sorted = sorted(depts,
            key=lambda x: priority_order.get(x[1]["category"], 99))

        # ═══════════════════════════════════════════════════════
        # ACCUMULATORS — group everything by color for performance
        # ═══════════════════════════════════════════════════════
        building_accs = {cat: MeshAccumulator(col, f'{cat} Buildings', 0.92)
                         for cat, col in CATEGORY_COLORS.items()}
        roof_acc = MeshAccumulator('#37474f', 'Roofs', 1.0)
        window_acc = MeshAccumulator('#4fc3f7', 'Windows', 0.75)
        solar_acc = MeshAccumulator('#1565c0', 'Solar Panels', 1.0)
        green_roof_acc = MeshAccumulator('#66bb6a', 'Green Roofs', 1.0)
        tree_trunk_acc = MeshAccumulator('#5d4037', 'Tree Trunks', 1.0)
        tree_leaves_acc = MeshAccumulator('#2e7d32', 'Tree Leaves', 1.0)

        # Hospital labels stored separately (text annotations)
        building_labels = []  # list of (x, y, z, name)

        # ═══════════════════════════════════════════════════════
        # Place department buildings (all into accumulators)
        # ═══════════════════════════════════════════════════════
        is_unified = "All-in-One" in building_mode

        if is_unified:
            # ─── UNIFIED BUILDING MODE: one large mass enclosing all depts ───
            # Calculate bounding box covering all available cells (excluding garden)
            min_r = min(r for (r, c) in available_cells[:len(depts_sorted)])
            max_r = max(r for (r, c) in available_cells[:len(depts_sorted)])
            min_c = min(c for (r, c) in available_cells[:len(depts_sorted)])
            max_c = max(c for (r, c) in available_cells[:len(depts_sorted)])

            big_x = margin + min_c * cell_w + cell_w * 0.05
            big_y = margin + (rows - 1 - max_r) * cell_h + cell_h * 0.05
            big_w = (max_c - min_c + 1) * cell_w - cell_w * 0.1
            big_h = (max_r - min_r + 1) * cell_h - cell_h * 0.1
            big_height = 12.0  # unified hospital tower height

            # Cut out garden hole if garden lies inside the unified building
            # For simplicity, draw the building as one big mass; garden is separate
            # The garden cells were already excluded from available_cells

            # Main unified mass — use the "General Care" color (hospital identity)
            building_accs["General Care"].add_box(
                big_x, big_y, 0, big_w, big_h, big_height)
            # Bigger overhanging roof
            roof_acc.add_box(big_x - 0.4, big_y - 0.4, big_height,
                             big_w + 0.8, big_h + 0.8, 0.25)
            # Solar panels covering most of the roof (it's a big single roof)
            if h_solar:
                panel_w, panel_d = 1.2, 1.5
                nx = max(1, int((big_w - 1.0) / panel_w))
                ny = max(1, int((big_h - 1.0) / panel_d))
                for i_p in range(nx):
                    for j_p in range(ny):
                        px = big_x + 0.5 + i_p * panel_w
                        py = big_y + 0.5 + j_p * panel_d
                        solar_acc.add_box(px, py, big_height + 0.28,
                                          panel_w * 0.85, panel_d * 0.85, 0.05)
            # A green roof patch in one corner
            if h_green_roof:
                green_roof_acc.add_box(
                    big_x + big_w * 0.05, big_y + big_h * 0.05, big_height + 0.26,
                    big_w * 0.35, big_h * 0.35, 0.1)

            # Floor-level windows on all 4 facades (one strip per floor per side)
            floors = max(1, int(big_height / 1.2))
            for f in range(floors):
                wz = 0.5 + f * 1.2
                # Front (south)
                window_acc.add_box(big_x + 0.5, big_y - 0.02, wz,
                                    big_w - 1.0, 0.04, 0.5)
                # Back (north)
                window_acc.add_box(big_x + 0.5, big_y + big_h - 0.02, wz,
                                    big_w - 1.0, 0.04, 0.5)
                # Left (west) — only if narrow enough
                if big_h > 6:
                    window_acc.add_box(big_x - 0.02, big_y + 0.5, wz,
                                        0.04, big_h - 1.0, 0.5)
                    window_acc.add_box(big_x + big_w - 0.02, big_y + 0.5, wz,
                                        0.04, big_h - 1.0, 0.5)

            # Department labels positioned at their cell center but at building height
            for idx, (dname, dinfo) in enumerate(depts_sorted):
                if idx >= len(available_cells):
                    break
                r, c = available_cells[idx]
                dx = margin + c * cell_w + cell_w * 0.1
                dy = margin + (rows - 1 - r) * cell_h + cell_h * 0.1
                dw = cell_w * 0.8
                dh = cell_h * 0.8
                # Floor-zone partitions visible from above? Add subtle interior
                # division markers — slight darker color strips on the roof
                roof_acc.add_box(dx, dy + dh - 0.04, big_height + 0.001,
                                  dw, 0.04, 0.05)
                # Labels float just above the unified roof level (clear visibility)
                building_labels.append(
                    (dx + dw/2, dy + dh/2, big_height + 1.0,
                     f"{dinfo['icon']} {dname.split('(')[0].strip()}"))

            # Single big entrance label
            building_labels.append(
                (big_x + big_w/2, big_y - 0.5, big_height + 2.5,
                 "🏥 VIRIDIS MEDICAL CENTER"))
        else:
            # ─── SEPARATE WINGS MODE: each dept is its own building ───
            for idx, (dname, dinfo) in enumerate(depts_sorted):
                if idx >= len(available_cells):
                    break
                r, c = available_cells[idx]
                dx = margin + c * cell_w + cell_w * 0.1
                dy = margin + (rows - 1 - r) * cell_h + cell_h * 0.1
                dw = cell_w * 0.8
                dh = cell_h * 0.8
                cat = dinfo["category"]
                building_h = height_map.get(cat, 6.0)

                # Main building mass
                building_accs[cat].add_box(dx, dy, 0, dw, dh, building_h)
                # Roof slab
                roof_acc.add_box(dx-0.2, dy-0.2, building_h, dw+0.4, dh+0.4, 0.15)

                # Solar panels on Service/Diagnostic/Treatment roofs
                if h_solar and cat in ["Service", "Diagnostic", "Treatment"]:
                    panel_w, panel_d = 1.0, 1.3
                    nx = max(1, int((dw-0.4) / panel_w))
                    ny = max(1, int((dh-0.4) / panel_d))
                    for i_p in range(nx):
                        for j_p in range(ny):
                            px = dx + 0.2 + i_p*panel_w
                            py = dy + 0.2 + j_p*panel_d
                            solar_acc.add_box(px, py, building_h+0.18,
                                              panel_w*0.85, panel_d*0.85, 0.04)
                # Green roofs on Wards/Specialized
                elif h_green_roof and cat in ["General Care", "Specialized"]:
                    green_roof_acc.add_box(dx+0.15, dy+0.15, building_h+0.16,
                                           dw-0.3, dh-0.3, 0.08)

                # Simplified windows — one strip per floor
                floors = max(1, int(building_h / 1.2))
                for f in range(floors):
                    wz = 0.5 + f * 1.2
                    window_acc.add_box(dx + 0.3, dy - 0.02,
                                       wz, dw - 0.6, 0.04, 0.5)

                # Label position
                building_labels.append(
                    (dx + dw/2, dy + dh/2, building_h + 0.8,
                     f"{dinfo['icon']} {dname.split('(')[0].strip()}"))

        # ═══════════════════════════════════════════════════════
        # Trees around the perimeter (into accumulators)
        # ═══════════════════════════════════════════════════════
        n_trees_x = max(6, int(site_w / 7))
        n_trees_y = max(4, int(site_h / 7))
        for x in np.linspace(2, site_w - 2, n_trees_x):
            for y_pos in [-1.5, site_h + 1.5]:
                trunk_h = 1.0
                tree_trunk_acc.add_cylinder(x, y_pos, 0, 0.12, trunk_h, 8)
                tree_leaves_acc.add_cone(x, y_pos, trunk_h, 0.6, 1.8, 10)
        for y in np.linspace(3, site_h - 3, n_trees_y):
            for x_pos in [-1.5, site_w + 1.5]:
                trunk_h = 1.0
                tree_trunk_acc.add_cylinder(x_pos, y, 0, 0.12, trunk_h, 8)
                tree_leaves_acc.add_cone(x_pos, y, trunk_h, 0.6, 1.8, 10)

        # ═══════════════════════════════════════════════════════
        # Healing Garden (also into accumulators where possible)
        # ═══════════════════════════════════════════════════════
        garden_grass_acc = MeshAccumulator('#7cb342', 'Garden Lawn', 1.0)
        garden_path_acc = MeshAccumulator('#d7ccc8', 'Garden Paths', 1.0)
        fountain_stone_acc = MeshAccumulator('#9e9e9e', 'Fountain Stone', 1.0)
        fountain_water_acc = MeshAccumulator('#4fc3f7', 'Fountain Water', 0.85)

        if h_garden:
            g_dx = margin + g_col_start * cell_w + cell_w * 0.15
            g_dy = margin + (rows - 2 - g_row_start) * cell_h + cell_h * 0.15
            g_dw = 2 * cell_w - cell_w * 0.3
            g_dh = 2 * cell_h - cell_h * 0.3
            g_cx = g_dx + g_dw/2
            g_cy = g_dy + g_dh/2

            # Garden grass (single box)
            garden_grass_acc.add_box(g_dx, g_dy, 0.05, g_dw, g_dh, 0.05)

            # Cross paths
            path_w = 0.6
            garden_path_acc.add_box(g_dx, g_cy-path_w/2, 0.11,
                                     g_dw, path_w, 0.04)
            garden_path_acc.add_box(g_cx-path_w/2, g_dy, 0.11,
                                     path_w, g_dh, 0.04)

            # Trees in garden (into the same accumulators as perimeter trees)
            tree_radius = min(g_dw, g_dh) * 0.32
            for ang in np.linspace(np.pi/4, 2*np.pi + np.pi/4, 8, endpoint=False):
                tx = g_cx + tree_radius*np.cos(ang)
                ty = g_cy + tree_radius*np.sin(ang)
                trunk_h = 1.1
                tree_trunk_acc.add_cylinder(tx, ty, 0.1, 0.12, trunk_h, 8)
                tree_leaves_acc.add_cone(tx, ty, 0.1 + trunk_h, 0.7, 2.0, 10)

            # Fountain: stone ring + water basin (static, in accumulators)
            fountain_r = min(g_dw, g_dh) * 0.10
            fountain_stone_acc.add_cylinder(g_cx, g_cy, 0.1,
                                             fountain_r*1.4, 0.4, 20)
            fountain_water_acc.add_cylinder(g_cx, g_cy, 0.35,
                                             fountain_r*1.1, 0.15, 20)
            # NOTE: Water spout is animated below as a separate trace,
            # NOT added to the static accumulator. Position stored for frames.
            _fountain_pos = (g_cx, g_cy, fountain_r)
            building_labels.append(
                (g_cx, g_cy, 3.0, "🌳 Healing Garden"))
        else:
            _fountain_pos = None

        # ═══════════════════════════════════════════════════════
        # Build the figure with ALL accumulated meshes as ONE trace each
        # ═══════════════════════════════════════════════════════
        fig = go.Figure()

        # Ground (lawn) - 2 simple Mesh3d
        fig.add_trace(go.Mesh3d(
            x=[-3, site_w+3, site_w+3, -3],
            y=[-3, -3, site_h+3, site_h+3],
            z=[0, 0, 0, 0],
            i=[0, 0], j=[1, 2], k=[2, 3],
            color='#aed581', opacity=1.0, name='Site Lawn',
            hovertext='Hospital Grounds', hoverinfo='text'))
        fig.add_trace(go.Mesh3d(
            x=[0, site_w, site_w, 0],
            y=[0, 0, site_h, site_h],
            z=[0.02, 0.02, 0.02, 0.02],
            i=[0, 0], j=[1, 2], k=[2, 3],
            color='#f1f8e9', opacity=1.0, name='Hospital Pavement',
            hovertext='Pavement', hoverinfo='text'))

        # Add all accumulators (only if they contain anything)
        for acc in [*building_accs.values(), roof_acc, window_acc,
                    solar_acc, green_roof_acc, tree_trunk_acc, tree_leaves_acc,
                    garden_grass_acc, garden_path_acc,
                    fountain_stone_acc, fountain_water_acc]:
            mesh = acc.to_mesh3d()
            if mesh is not None:
                fig.add_trace(mesh)

        # Entrance canopy (only 3 small meshes — keep separate for clarity)
        fig.add_trace(go.Mesh3d(
            x=[site_w/2-4, site_w/2+4, site_w/2+4, site_w/2-4],
            y=[-3, -3, 0, 0], z=[0.05]*4,
            i=[0, 0], j=[1, 2], k=[2, 3],
            color='#616161', opacity=1.0, name='Main Driveway',
            hovertext='Driveway', hoverinfo='text'))

        # Rainwater tank (single cylinder via accumulator)
        if h_rainwater:
            tank_acc = MeshAccumulator('#01579b', 'Rainwater Tank', 0.95)
            tank_acc.add_cylinder(site_w - 3, 3, 0, 1.2, 3.0, 16)
            fig.add_trace(tank_acc.to_mesh3d())
            building_labels.append((site_w - 3, 3, 3.5, "💧 Rainwater"))

        # EV chargers (combined into one accumulator)
        if h_ev_chargers:
            ev_acc = MeshAccumulator('#66bb6a', 'EV Chargers', 1.0)
            for ex_i in range(4):
                ex = 3 + ex_i * 1.5
                ev_acc.add_box(ex - 0.15, 1.5, 0, 0.3, 0.3, 1.5)
            fig.add_trace(ev_acc.to_mesh3d())

        # ═══════════════════════════════════════════════════════
        # PARKING GARAGE (covered multi-level structure on the side)
        # ═══════════════════════════════════════════════════════
        if h_garage:
            gar_w = min(12, site_w * 0.25)
            gar_d = min(8, site_h * 0.18)
            gar_x = site_w + 0.5  # right side of site
            gar_y = site_h * 0.3
            gar_floors = 3
            floor_h = 2.5
            gar_h = gar_floors * floor_h

            # Main garage structure (light gray concrete)
            fig.add_trace(box_mesh_global(gar_x, gar_y, 0,
                gar_w, gar_d, gar_h, '#bdbdbd', 0.85,
                'Parking Garage'))
            # Roof
            fig.add_trace(box_mesh_global(gar_x - 0.3, gar_y - 0.3, gar_h,
                gar_w + 0.6, gar_d + 0.6, 0.2,
                '#37474f', 1.0, 'Garage Roof'))
            # Floor dividers (visible openings between floors)
            for f_ in range(1, gar_floors):
                fig.add_trace(box_mesh_global(gar_x, gar_y - 0.05, f_ * floor_h - 0.05,
                    gar_w, 0.05, 0.1, '#424242', 1.0, f'Floor Slab {f_}'))
            # Open parking deck appearance — front facade with openings
            for f_ in range(gar_floors):
                # Strip openings (darker)
                fig.add_trace(box_mesh_global(gar_x + 0.5, gar_y - 0.02,
                    f_ * floor_h + 0.3, gar_w - 1, 0.04, floor_h - 0.6,
                    '#1a1a1a', 0.9, f'Opening F{f_+1}'))
            # Parked cars (small boxes inside on ground floor)
            car_colors = ['#e53935', '#1976d2', '#43a047', '#fb8c00',
                          '#ab47bc', '#00897b', '#ef6c00', '#5e35b1']
            cars_per_row = max(2, int(gar_w / 2.2))
            for cr in range(2):  # 2 rows
                for ci in range(cars_per_row):
                    cx = gar_x + 0.8 + ci * 2.0
                    cy = gar_y + 1.0 + cr * 3.0
                    c_color = car_colors[(cr * cars_per_row + ci) % len(car_colors)]
                    # Car body (low box)
                    fig.add_trace(box_mesh_global(cx, cy, 0.15,
                        1.6, 1.8, 0.5, c_color, 1.0, 'Car'))
                    # Car cabin (smaller box on top)
                    fig.add_trace(box_mesh_global(cx + 0.15, cy + 0.3, 0.65,
                        1.3, 1.2, 0.4, c_color, 1.0, 'Cabin'))
                    # Windshield
                    fig.add_trace(box_mesh_global(cx + 0.18, cy + 0.28, 0.7,
                        1.24, 0.04, 0.35, '#0d47a1', 0.7, 'Windshield'))
            building_labels.append(
                (gar_x + gar_w/2, gar_y + gar_d/2, gar_h + 1.2,
                 "🚗 Parking Garage"))

        # ═══════════════════════════════════════════════════════
        # ROAD NETWORK (driveway loop + paths + INTERNAL roads)
        # ═══════════════════════════════════════════════════════
        if h_roads:
            road_color = '#424242'
            line_color = '#fff59d'
            ambulance_color = '#b71c1c'  # red for ambulance route
            sidewalk_color = '#cfd8dc'   # light gray for pedestrian
            service_road_color = '#37474f'  # darker for service roads

            # ─── Main access road (entrance → garage) ───
            fig.add_trace(box_mesh_global(
                site_w + 0.2, -4, 0.03,
                3.0, site_h * 0.35, 0.04,
                road_color, 1.0, 'Access Road'))
            for ly in np.arange(-4, site_h * 0.35 - 4, 1.5):
                fig.add_trace(box_mesh_global(
                    site_w + 1.55, ly, 0.06,
                    0.1, 0.7, 0.02,
                    line_color, 1.0, 'Road Line'))

            # ─── Drop-off lane (front main entrance) ───
            fig.add_trace(box_mesh_global(
                site_w/2 - 6, -5, 0.03,
                12, 2, 0.04,
                road_color, 1.0, 'Drop-off Lane'))
            for lx in np.arange(site_w/2 - 5, site_w/2 + 5, 1.5):
                fig.add_trace(box_mesh_global(
                    lx, -4.1, 0.06,
                    0.8, 0.1, 0.02,
                    line_color, 1.0, 'Lane Marking'))

            # ─── Front sidewalk ───
            fig.add_trace(box_mesh_global(
                0, -1, 0.03,
                site_w, 1, 0.04,
                sidewalk_color, 1.0, 'Sidewalk Front'))

            # ═══════════════════════════════════════════════════════
            # NEW: INTERNAL HOSPITAL ROAD NETWORK
            # ═══════════════════════════════════════════════════════

            # ─── Perimeter loop road around the site ───
            # Left perimeter road (vertical, west side)
            fig.add_trace(box_mesh_global(
                -3.5, 2, 0.025,
                2.5, site_h - 4, 0.04,
                road_color, 1.0, 'West Perimeter Road'))
            # Dashed lines on west road
            for ly in np.arange(2, site_h - 2, 1.5):
                fig.add_trace(box_mesh_global(
                    -2.3, ly, 0.06,
                    0.1, 0.7, 0.02,
                    line_color, 1.0, 'West Line'))

            # Back perimeter road (horizontal, north side - service access)
            fig.add_trace(box_mesh_global(
                -1, site_h + 1, 0.025,
                site_w + 2, 2.5, 0.04,
                service_road_color, 1.0, 'North Service Road'))
            # Service road markings (white dashed for service vehicles)
            for lx in np.arange(0, site_w, 1.5):
                fig.add_trace(box_mesh_global(
                    lx, site_h + 2.15, 0.06,
                    0.7, 0.1, 0.02,
                    '#ffffff', 1.0, 'Service Line'))

            # ─── AMBULANCE ROUTE — direct to ER zone ───
            # Red-painted dedicated emergency lane
            # Bottom-left corner approach to ER
            fig.add_trace(box_mesh_global(
                -3.5, -3, 0.025,
                site_w * 0.35, 2.0, 0.05,
                ambulance_color, 0.85, 'Ambulance Route'))
            # White "AMBULANCE" stripes
            for lx in np.arange(-2.5, site_w * 0.35 - 4, 2.0):
                fig.add_trace(box_mesh_global(
                    lx, -2.5, 0.08,
                    1.2, 0.15, 0.02,
                    '#ffffff', 1.0, 'Ambulance Stripe'))
            # ER label marker
            building_labels.append(
                (site_w * 0.18, -2, 1.0, "🚑 AMBULANCE ENTRY"))

            # ─── Inter-building pedestrian paths (sidewalks) ───
            # Horizontal walkway through middle of site
            fig.add_trace(box_mesh_global(
                1, site_h * 0.45, 0.025,
                site_w - 2, 1.0, 0.03,
                sidewalk_color, 1.0, 'Mid Walkway H'))
            # Vertical walkway down the middle (front-to-back)
            fig.add_trace(box_mesh_global(
                site_w * 0.5 - 0.5, 1, 0.025,
                1.0, site_h - 2, 0.03,
                sidewalk_color, 1.0, 'Mid Walkway V'))

            # Cross-pattern sidewalks at quarter points
            for x_frac in [0.25, 0.75]:
                fig.add_trace(box_mesh_global(
                    site_w * x_frac - 0.4, 1, 0.025,
                    0.8, site_h - 2, 0.03,
                    sidewalk_color, 0.9, f'Walkway V {x_frac}'))

            # ─── Bicycle path (green for eco theme) ───
            fig.add_trace(box_mesh_global(
                site_w + 4, 0, 0.025,
                1.0, site_h, 0.04,
                '#558b2f', 0.85, 'Bicycle Path'))
            # Bike symbol markers
            for ly in np.arange(2, site_h - 2, 3.0):
                fig.add_trace(box_mesh_global(
                    site_w + 4.3, ly, 0.07,
                    0.4, 0.4, 0.02,
                    '#fff', 1.0, 'Bike Symbol'))

            # ─── Crosswalks at intersections (zebra stripes) ───
            # Crosswalk at drop-off entrance
            for cw_x in np.arange(site_w/2 - 2.5, site_w/2 + 2.5, 0.5):
                fig.add_trace(box_mesh_global(
                    cw_x, -2.5, 0.07,
                    0.3, 0.8, 0.02,
                    '#ffffff', 1.0, 'Crosswalk'))

            # ─── Direction arrows on main loop ───
            # Arrow markings on access road (going up = direction)
            for ay in [-2, 1, 4]:
                # Arrow shaft
                fig.add_trace(box_mesh_global(
                    site_w + 1.5, ay, 0.07,
                    0.2, 0.4, 0.02,
                    line_color, 1.0, 'Arrow Shaft'))
                # Arrow head (triangular - simulated with small box)
                fig.add_trace(box_mesh_global(
                    site_w + 1.4, ay + 0.4, 0.07,
                    0.4, 0.15, 0.02,
                    line_color, 1.0, 'Arrow Head'))

        # ═══════════════════════════════════════════════════════
        # OUTDOOR SEATING (benches in garden + around perimeter)
        # ═══════════════════════════════════════════════════════
        if h_outdoor_seating:
            bench_acc = MeshAccumulator('#5d4037', 'Outdoor Benches', 1.0)
            bench_back_acc = MeshAccumulator('#6d4c41', 'Bench Backs', 1.0)
            # Benches around the perimeter (front of site, near entrance)
            for bx in [4, site_w/2 - 8, site_w/2 + 8, site_w - 4]:
                # Seat
                bench_acc.add_box(bx - 0.7, -0.3, 0.4, 1.4, 0.4, 0.08)
                # Backrest
                bench_back_acc.add_box(bx - 0.7, 0.05, 0.48, 1.4, 0.05, 0.45)
                # Legs (4)
                for lx, ly in [(-0.65, -0.25), (0.55, -0.25),
                                (-0.65, 0.0), (0.55, 0.0)]:
                    bench_acc.add_box(bx + lx, ly, 0, 0.06, 0.06, 0.4)

            # Garden benches (in the healing garden, around fountain)
            if _fountain_pos is not None:
                fcx, fcy, _ = _fountain_pos
                garden_bench_radius = 5
                # 4 benches around the fountain (cardinal directions)
                for ang_deg in [45, 135, 225, 315]:
                    ang = np.radians(ang_deg)
                    bx = fcx + garden_bench_radius * np.cos(ang)
                    by = fcy + garden_bench_radius * np.sin(ang)
                    # Seat (oriented loosely toward fountain — for simplicity use axis-aligned)
                    bench_acc.add_box(bx - 0.6, by - 0.25, 0.4, 1.2, 0.5, 0.08)
                    bench_back_acc.add_box(bx - 0.6, by + 0.2, 0.48, 1.2, 0.05, 0.4)
                    # Legs
                    for lx, ly in [(-0.55, -0.2), (0.49, -0.2),
                                    (-0.55, 0.15), (0.49, 0.15)]:
                        bench_acc.add_box(bx + lx, by + ly, 0, 0.06, 0.06, 0.4)

            fig.add_trace(bench_acc.to_mesh3d())
            fig.add_trace(bench_back_acc.to_mesh3d())

        # ═══════════════════════════════════════════════════════
        # Building labels as a single Scatter3d trace (much faster)
        # ═══════════════════════════════════════════════════════
        if building_labels:
            label_x = [l[0] for l in building_labels]
            label_y = [l[1] for l in building_labels]
            label_z = [l[2] for l in building_labels]
            label_text = [l[3] for l in building_labels]
            fig.add_trace(go.Scatter3d(
                x=label_x, y=label_y, z=label_z,
                mode='text', text=label_text,
                textfont=dict(size=9, color='#1b5e20', family='Arial Black'),
                hoverinfo='text', hovertext=label_text,
                showlegend=False, name='Labels'))

        # ═══════════════════════════════════════════════════════
        # ANIMATED FOUNTAIN — Water spout pulses + droplets fly
        # ═══════════════════════════════════════════════════════
        if _fountain_pos is not None:
            fcx, fcy, fr = _fountain_pos
            # Initial water spout (will be animated)
            initial_spout_h = 1.2
            initial_spout = cylinder_mesh(fcx, fcy, 0.5,
                                          fr*0.18, initial_spout_h, 'z', 10,
                                          '#81d4fa', 0.7, 'Water Spout')
            spout_trace_idx = len(fig.data)  # remember index for animation
            fig.add_trace(initial_spout)

            # Initial water cap (sphere-like cone on top of spout)
            initial_cap = cone_mesh(fcx, fcy, 0.5 + initial_spout_h,
                                     fr*0.32, 0.35, 14,
                                     '#b3e5fc', 0.65, 'Water Cap')
            cap_trace_idx = len(fig.data)
            fig.add_trace(initial_cap)

            # Initial water droplets (Scatter3d markers — small white spheres)
            n_droplets = 12
            droplet_angles = np.linspace(0, 2*np.pi, n_droplets, endpoint=False)
            initial_dx = [fcx + fr*0.25*np.cos(a) for a in droplet_angles]
            initial_dy = [fcy + fr*0.25*np.sin(a) for a in droplet_angles]
            initial_dz = [0.8 + (i%3)*0.3 for i in range(n_droplets)]
            droplets_trace = go.Scatter3d(
                x=initial_dx, y=initial_dy, z=initial_dz,
                mode='markers',
                marker=dict(size=8, color='#b3e5fc', opacity=0.85,
                            line=dict(color='#4fc3f7', width=1)),
                name='Water Droplets', hoverinfo='skip', showlegend=False
            )
            droplets_trace_idx = len(fig.data)
            fig.add_trace(droplets_trace)

            # Build animation frames (8 keyframes for smooth oscillation)
            n_frames = 12
            frames = []
            for f_idx in range(n_frames):
                phase = 2 * np.pi * f_idx / n_frames
                # Spout height oscillates between 0.6 and 1.6m (sin wave)
                spout_h = 1.1 + 0.45 * np.sin(phase)

                # New spout mesh at this height
                new_spout = cylinder_mesh(fcx, fcy, 0.5,
                                           fr*0.18, spout_h, 'z', 10,
                                           '#81d4fa', 0.7, 'Water Spout')
                # New cap on top
                cap_radius = fr * (0.28 + 0.08 * np.sin(phase))
                new_cap = cone_mesh(fcx, fcy, 0.5 + spout_h,
                                     cap_radius, 0.35, 14,
                                     '#b3e5fc', 0.65, 'Water Cap')
                # Droplets shoot out — radius expands with spout height
                droplet_radius = fr * (0.3 + 0.6 * (0.5 + 0.5*np.sin(phase)))
                # Droplets also have their own falling motion
                new_dx = [fcx + droplet_radius*np.cos(a + phase*0.5)
                          for a in droplet_angles]
                new_dy = [fcy + droplet_radius*np.sin(a + phase*0.5)
                          for a in droplet_angles]
                # Each droplet has its own ballistic z (offset by index for variety)
                new_dz = []
                for i in range(n_droplets):
                    offset_phase = phase + i * (2*np.pi / n_droplets)
                    # Parabolic motion: rises then falls
                    z_val = 0.5 + spout_h + 0.4 * np.sin(offset_phase)
                    new_dz.append(max(0.5, z_val))
                new_droplets = go.Scatter3d(
                    x=new_dx, y=new_dy, z=new_dz,
                    mode='markers',
                    marker=dict(size=8, color='#b3e5fc', opacity=0.85,
                                line=dict(color='#4fc3f7', width=1)),
                    name='Water Droplets', hoverinfo='skip', showlegend=False
                )

                frames.append(go.Frame(
                    data=[new_spout, new_cap, new_droplets],
                    traces=[spout_trace_idx, cap_trace_idx, droplets_trace_idx],
                    name=f'fountain_{f_idx}'
                ))
            fig.frames = frames

        # ─── Layout & Camera + Animation Controls ───
        max_dim = max(site_w, site_h)
        layout_dict = dict(
            scene=dict(
                xaxis=dict(title='Width (m)', range=[-5, site_w+5],
                    backgroundcolor='#e8f5e9', gridcolor='#c8e6c9',
                    showbackground=True),
                yaxis=dict(title='Depth (m)', range=[-5, site_h+5],
                    backgroundcolor='#e8f5e9', gridcolor='#c8e6c9',
                    showbackground=True),
                zaxis=dict(title='Height (m)', range=[0, 14],
                    backgroundcolor='#e3f2fd', gridcolor='#bbdefb',
                    showbackground=True),
                aspectmode='manual',
                aspectratio=dict(x=site_w/max_dim,
                                  y=site_h/max_dim,
                                  z=0.35),
                camera=dict(eye=dict(x=1.6, y=-1.8, z=1.1),
                    center=dict(x=0, y=0, z=-0.2)),
                bgcolor='#f5fdf7'),
            margin=dict(l=0, r=0, b=0, t=40),
            height=680, showlegend=False,
            title=dict(
                text=f"🏥 VIRIDIS HOSPITAL — 3D Site View ({hospital_scale})",
                font=dict(size=14, color='#1b5e20'),
                x=0.5, xanchor='center'))

        # Fountain animation play/pause buttons
        if _fountain_pos is not None:
            layout_dict['updatemenus'] = [dict(
                type='buttons',
                showactive=False,
                y=0.02, x=0.02,
                xanchor='left', yanchor='bottom',
                pad=dict(t=0, r=10),
                bgcolor='rgba(22, 163, 74, 0.85)',
                bordercolor='#22c55e',
                font=dict(color='white', size=12, family='Arial Black'),
                buttons=[
                    dict(label='⛲ Play Fountain',
                         method='animate',
                         args=[None, dict(
                             frame=dict(duration=180, redraw=True),
                             fromcurrent=True,
                             transition=dict(duration=80, easing='linear'),
                             mode='immediate'
                         )]),
                    dict(label='⏸ Pause',
                         method='animate',
                         args=[[None], dict(
                             frame=dict(duration=0, redraw=False),
                             mode='immediate',
                             transition=dict(duration=0)
                         )])
                ]
            )]

        fig.update_layout(**layout_dict)
        return fig
    # ============================================================
    # RENDER UI — 2D/3D toggle
    # ============================================================
    site_view_mode = st.radio(
        "Site Plan View:",
        ["🗺️ 2D Site Plan", "🌐 3D Hospital View", "📊 Both Side-by-Side"],
        horizontal=True
    )

    if site_view_mode == "🗺️ 2D Site Plan":
        st.subheader("🗺️ Hospital Site Plan (2D)")
        site_fig, dept_positions = draw_site_plan()
        st.pyplot(site_fig)
    elif site_view_mode == "🌐 3D Hospital View":
        st.subheader("🌐 Hospital 3D Render — Bird's-eye View")
        with st.spinner("Building 3D hospital model..."):
            hosp_3d = draw_hospital_3d()
        st.plotly_chart(hosp_3d, use_container_width=True)
        st.caption("💡 Tip: Drag to rotate · Scroll to zoom · "
                   "Solar panels show on Service/Diagnostic roofs · "
                   "Green roofs on Wards · Hover any building to see its name.")
        # Site plan still needed for downstream code references
        site_fig, dept_positions = draw_site_plan()
    else:
        col_2d, col_3d = st.columns(2)
        with col_2d:
            st.subheader("🗺️ 2D Plan")
            site_fig, dept_positions = draw_site_plan()
            st.pyplot(site_fig)
        with col_3d:
            st.subheader("🌐 3D View")
            with st.spinner("Building 3D model..."):
                st.plotly_chart(draw_hospital_3d(), use_container_width=True)


    # Export buttons
    col_ex1, col_ex2, col_ex3 = st.columns(3)
    with col_ex1:
        buf_site = BytesIO()
        site_fig.savefig(buf_site, format='png', dpi=200, bbox_inches='tight')
        st.download_button("⬇️ PNG", buf_site.getvalue(),
            file_name=f"viridis_hospital_{hospital_scale.split('(')[0].strip()}.png",
            mime="image/png", use_container_width=True)
    with col_ex2:
        h_config = {
            "scale": hospital_scale,
            "departments": list(selected_depts.keys()),
            "sustainability": {
                "solar": h_solar, "garden": h_garden,
                "green_roof": h_green_roof, "rainwater": h_rainwater,
                "ev_chargers": h_ev_chargers, "greywater": h_greywater,
                "smart_grid": h_smart_grid
            },
            "metrics": {
                "total_area_m2": hm['total_area'],
                "annual_energy_saved_MWh": hm['energy_saved']*12/1000,
                "annual_money_saved_USD": hm['money_saved']*12,
                "annual_CO2_reduced_tons": hm['co2_saved']*12/1000,
                "trees_equivalent_per_year": hm['trees_equiv']*12
            }
        }
        st.download_button("💾 JSON",
            json.dumps(h_config, indent=2),
            file_name="viridis_hospital.json", mime="application/json",
            use_container_width=True)
    with col_ex3:
        # Build PDF report on demand
        try:
            pdf_h_config = {
                'hospital_scale': hospital_scale,
                'selected_depts': selected_depts,
                'building_mode': building_mode,
                'h_solar': h_solar, 'h_garden': h_garden,
                'h_green_roof': h_green_roof, 'h_rainwater': h_rainwater,
                'h_ev_chargers': h_ev_chargers, 'h_greywater': h_greywater,
                'h_smart_grid': h_smart_grid, 'h_garage': h_garage,
                'h_roads': h_roads, 'h_outdoor_seating': h_outdoor_seating,
            }
            pdf_bytes = generate_pdf_report('hospital', pdf_h_config,
                                             hm, blueprint_fig=site_fig)
            st.download_button(t('pdf_download'), pdf_bytes,
                file_name=f"viridis_hospital_report.pdf",
                mime="application/pdf", use_container_width=True)
        except Exception as e:
            st.caption(f"PDF unavailable: {e}")

    # ============================================================
    # DEPARTMENT BREAKDOWN TABLE
    # ============================================================
    st.markdown("---")
    st.subheader("📊 Department-by-Department Breakdown")

    breakdown = []
    for dname, dinfo in selected_depts.items():
        est_units = max(1, int(dinfo["max_units"] * 0.6))
        dept_area = est_units * dinfo["min_area_per_unit"] * 1.4 * scale
        dept_energy = dinfo["base_power"] * (dept_area / 40) * 720 / 1000  # MWh/mo
        dept_water = dinfo["water_usage"] * est_units * 30 / 1000  # m³/mo
        breakdown.append({
            "Department": f"{dinfo['icon']} {dname}",
            "Category": dinfo['category'],
            "Area (m²)": f"{dept_area:,.0f}",
            "Units": f"{est_units}",
            "Energy (MWh/mo)": f"{dept_energy:,.1f}",
            "Water (m³/mo)": f"{dept_water:,.1f}",
            "ACH": dinfo['ach'],
            "Pressure": dinfo['pressure']
        })
    st.dataframe(breakdown, use_container_width=True, hide_index=True)

    # ============================================================
    # CATEGORY PIE CHART
    # ============================================================
    st.markdown("---")
    st.subheader("📈 Hospital Composition Analysis")

    col_p1, col_p2 = st.columns(2)

    with col_p1:
        # Area by category
        cat_areas = {}
        for dname, dinfo in selected_depts.items():
            cat = dinfo['category']
            est_units = max(1, int(dinfo["max_units"] * 0.6))
            dept_area = est_units * dinfo["min_area_per_unit"] * 1.4 * scale
            cat_areas[cat] = cat_areas.get(cat, 0) + dept_area

        fig_pie, ax_pie = plt.subplots(figsize=(5, 5))
        wedges, texts, autotexts = ax_pie.pie(
            cat_areas.values(),
            labels=cat_areas.keys(),
            colors=[CATEGORY_COLORS[c] for c in cat_areas.keys()],
            autopct='%1.1f%%', startangle=90,
            wedgeprops=dict(edgecolor='white', linewidth=2),
            textprops=dict(fontsize=8, weight='bold')
        )
        for t in autotexts:
            t.set_color('white')
        ax_pie.set_title("Built Area by Category", fontsize=11, weight='bold')
        st.pyplot(fig_pie)

    with col_p2:
        # Energy comparison
        fig_eb, ax_eb = plt.subplots(figsize=(5, 5))
        categories = ['Traditional\nHospital', 'Viridis\nGreen Hospital']
        values = [hm['energy_normal']*12/1000, hm['energy_green']*12/1000]
        bars = ax_eb.bar(categories, values,
            color=['#ef5350', '#66bb6a'], width=0.5,
            edgecolor='black', linewidth=1.2)
        ax_eb.set_ylabel("Energy Consumption (MWh/year)", fontsize=10)
        ax_eb.set_title(f"Annual Energy: {hm['savings_pct']*100:.0f}% Saved",
            fontsize=11, weight='bold')
        for bar, val in zip(bars, values):
            ax_eb.text(bar.get_x() + bar.get_width()/2,
                val + max(values)*0.02,
                f"{val:,.0f} MWh", ha='center',
                fontsize=9, weight='bold')
        ax_eb.grid(axis='y', alpha=0.3)
        st.pyplot(fig_eb)

    # ============================================================
    # SUSTAINABILITY IMPACT SUMMARY
    # ============================================================
    st.markdown("---")
    st.subheader("🌍 Environmental Impact Report")

    impact_report = f"""
### Annual Sustainability Achievements

🌲 **Carbon Footprint Reduction**
- CO₂ emissions avoided: **{hm['co2_saved']*12/1000:,.1f} metric tons/year**
- Equivalent to planting **{int(hm['trees_equiv']*12):,} trees** annually 🌳
- Or removing **{int(hm['co2_saved']*12/4600):,} cars** from the road for a year 🚗

💰 **Financial Returns**
- Annual operational savings: **${hm['money_saved']*12:,.0f} USD**
- 10-year cumulative savings: **${hm['money_saved']*120:,.0f} USD**
- Typical solar PV payback period: **5-7 years**

💧 **Water Conservation**
- Water saved: **{hm['water_saved']*12/1000:,.0f} m³/year**
- Sufficient for **{int(hm['water_saved']*12/180):,} people**'s yearly drinking needs

⚡ **Energy Efficiency**
- Energy savings: **{hm['energy_saved']*12/1000:,.0f} MWh/year** ({hm['savings_pct']*100:.0f}%)
- Powers **{int(hm['energy_saved']*12/3500):,} homes** for a year

### Sustainability Features Active
"""
    if h_solar: impact_report += "- ✅ Rooftop Solar PV Farm\n"
    if h_garden: impact_report += "- ✅ Healing Garden (proven to reduce patient recovery time by 15-20%)\n"
    if h_green_roof: impact_report += "- ✅ Green Roof (insulation + biodiversity)\n"
    if h_rainwater: impact_report += "- ✅ Rainwater Harvesting\n"
    if h_greywater: impact_report += "- ✅ Greywater Recycling\n"
    if h_ev_chargers: impact_report += "- ✅ EV Charging Stations\n"
    if h_smart_grid: impact_report += "- ✅ Smart Grid + Battery Storage\n"

    st.success(impact_report)

    # ============================================================
    # DRILL DOWN HINT
    # ============================================================
    st.info("💡 **Want to design a specific department in detail?** "
            "Switch to **🏠 Single Department** mode in the sidebar "
            "and select any of the departments above for a full 2D + 3D blueprint.")